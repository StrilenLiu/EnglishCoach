#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# English Coach (英语导师) — 翻译 + 朗读 学习助手
# Copyright (C) 2026 Strilen Liu <vfx@strilen.com>  https://www.strilen.com
#
# SPDX-License-Identifier: GPL-3.0-or-later
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
# <https://www.gnu.org/licenses/>
"""
EnglishCoach - 英语助手工具
第一版 (v1.0.0)

功能:
    - 翻译 (默认 Google 免费引擎; DeepL / DeepSeek 可在设置中作为备选)
    - 朗读 (edge-tts 在线 TTS, 多嗓音 / 语速)
    - 版本更新说明与管理
    - 关于
    - Readme / 帮助文档与使用说明
    - 开发者介绍
    - 图标系统 (内置 SVG)

作者: Strilen  (vfx@strilen.com  |  www.strilen.com)
"""

import sys
import os

# ============================================================================
#  Hugging Face 端点与离线模型：必须在【任何】import huggingface_hub /
#  transformers / kokoro 之前完成设置。
#
#  原因：huggingface_hub 在【模块导入时】就把 ENDPOINT 读成常量固定下来，
#  之后再改环境变量完全无效。此前这段逻辑写在函数内部，等执行到时
#  transformers 早已把 huggingface_hub 导入完毕，于是设了也白设 ——
#  大陆用户依然去连 huggingface.co 并失败。
# ============================================================================
def _ec_bootstrap_hf():
    # ---- 1) 本地已有模型则直接离线，完全不联网 ----
    cands = []
    if os.environ.get("ENGLISHCOACH_MODELS"):
        cands.append(os.path.join(os.environ["ENGLISHCOACH_MODELS"], "Kokoro"))
    mei = getattr(sys, "_MEIPASS", None)
    if mei:
        cands.append(os.path.join(mei, "kokoro_model"))
    try:
        here = os.path.dirname(os.path.abspath(
            sys.executable if getattr(sys, "frozen", False) else __file__))
        cands.append(os.path.join(here, "kokoro_model"))
    except Exception:
        pass
    cands.append(os.path.expanduser("~/EnglishCoach Models/Kokoro"))

    for d in cands:
        if not os.path.isdir(d):
            continue
        # HF 缓存结构（hub/models--hexgrad--Kokoro-82M）能让模型与音色
        # 全部离线可用，是首选
        if os.path.isdir(os.path.join(d, "hub")):
            os.environ.setdefault("HF_HOME", d)
            os.environ.setdefault("HF_HUB_OFFLINE", "1")
            os.environ["ENGLISHCOACH_KOKORO_DIR"] = d
            return
        # 扁平目录（snapshot_download(local_dir=...) 的产物）：
        # 记下路径，稍后直接把权重路径喂给 KModel
        if any(f.endswith((".pth", ".safetensors", ".onnx"))
               for _r, _d, fs in os.walk(d) for f in fs):
            os.environ["ENGLISHCOACH_KOKORO_DIR"] = d
            return

    # ---- 2) 无本地模型：大陆自动走镜像，避免直连失败 ----
    if os.environ.get("HF_ENDPOINT"):
        return
    cn = False
    cand_langs = []
    try:
        import locale as _loc
        cand_langs.append(_loc.getdefaultlocale()[0] or "")
    except Exception:
        pass
    for v in ("LC_ALL", "LC_CTYPE", "LANG", "LANGUAGE"):
        cand_langs.append(os.environ.get(v, "") or "")
    if any(c.lower().replace("-", "_").startswith("zh_cn") for c in cand_langs):
        cn = True
    if not cn:
        tz = os.environ.get("TZ", "")
        if not tz:
            try:
                with open("/etc/timezone", encoding="utf-8") as f:
                    tz = f.read().strip()
            except Exception:
                try:
                    tz = os.path.realpath("/etc/localtime")
                except Exception:
                    tz = ""
        if any(k in tz for k in ("Shanghai", "Chongqing", "Harbin",
                                 "Urumqi", "PRC")):
            cn = True
    if cn:
        os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"


_ec_bootstrap_hf()
import json
import re
import asyncio
import tempfile
import traceback


# =============================================================================
#  Argos 离线翻译兼容层（必须在任何 argostranslate 导入之前执行）
#  目的：让 Argos 在无 PyTorch、Big Sur 等受限环境下也能离线翻译。
#    1) 注入假 stanza 模块，骗过顶层 import 与运行时的 Pipeline 调用，甩掉 torch
#    2) 设 stanza_available=True，绕过缺 sbd 包导致语言被过滤为空的问题
#    3) 把模型目录锁定到程序内置 / 用户数据目录（解决"装了读不到"）
# =============================================================================

def _setup_argos_compat():
    import types
    # —— 1) 假 stanza（带断句 Pipeline）——
    if "stanza" not in sys.modules:
        fake = types.ModuleType("stanza")

        class _FakeStanzaPipeline:
            def __init__(self, *a, **k):
                pass

            def __call__(self, text):
                # 简单按整段返回单句，断句交给上层；够 Argos 用
                class _Sentence:
                    def __init__(self, t):
                        self.text = t

                class _Doc:
                    sentences = [_Sentence(text)]
                return _Doc()

        fake.Pipeline = _FakeStanzaPipeline
        fake.download = lambda *a, **k: None
        sys.modules["stanza"] = fake


def _configure_argos_dirs():
    """把 Argos 的模型读取/写入目录都锁定到同一处。返回 (target, bundled)。"""
    from pathlib import Path
    import argostranslate.settings as s
    # 程序内置模型目录（打包后在 _MEIPASS/argos_models）
    base = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    bundled = os.path.join(base, "argos_models")
    # 用户数据目录（运行期实际安装位置，避免只读问题）
    user_dir = os.path.join(
        os.path.expanduser("~"), ".englishcoach", "argos_packages")
    os.makedirs(user_dir, exist_ok=True)
    target = Path(user_dir)
    s.package_dirs = [target]
    s.package_data_dir = target
    s.stanza_available = True   # 关键：别因缺 sbd 包过滤掉语言
    return target, bundled


def _ensure_argos_models(required=("en", "zh")):
    """确保所需语言模型已安装；缺失则从内置目录安装。幂等、可重复调用。
    返回 True 表示所需语言齐备。"""
    try:
        target, bundled = _configure_argos_dirs()
        import argostranslate.package as ap
        import argostranslate.translate as at
        have = {l.code for l in at.get_installed_languages()}
        if all(c in have for c in required):
            return True
        # 缺失 -> 从内置 .argosmodel 安装
        if os.path.isdir(bundled):
            for fn in sorted(os.listdir(bundled)):
                if fn.endswith(".argosmodel"):
                    try:
                        ap.install_from_path(os.path.join(bundled, fn))
                    except Exception:
                        pass
        have = {l.code for l in at.get_installed_languages()}
        return all(c in have for c in required)
    except Exception:
        return False


_setup_argos_compat()

import requests
import edge_tts

from PyQt6.QtCore import (Qt, QThread, pyqtSignal, QSize, QUrl, QSettings,
                          QTimer, QBuffer, QByteArray, QIODevice, QElapsedTimer)
from PyQt6.QtGui import (QIcon, QPixmap, QFont, QAction,
                         QSyntaxHighlighter, QTextCharFormat, QColor)
from PyQt6.QtMultimedia import QMediaPlayer, QAudioOutput
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTextEdit, QPushButton, QComboBox, QLabel, QSlider, QToolBar,
    QStatusBar, QDialog, QDialogButtonBox, QLineEdit, QFormLayout,
    QTextBrowser, QMessageBox, QSplitter, QFrame, QSizePolicy,
    QCheckBox, QScrollArea
)

# =============================================================================
#  应用元信息 / 版本管理
# =============================================================================

APP_NAME = "EnglishCoach"
APP_VERSION = "2.15.12"
APP_AUTHOR = "Strilen"
APP_EMAIL = "vfx@strilen.com"
APP_WEBSITE = "www.strilen.com"

# 统一按钮标准宽度（以"显示/隐藏"按钮为准）
BTN_W = 96

# 数字/符号 -> 名称（用于"单个符号或数字串"特殊翻译模式）
_DIGIT_EN = {"0": "Zero", "1": "One", "2": "Two", "3": "Three", "4": "Four",
             "5": "Five", "6": "Six", "7": "Seven", "8": "Eight", "9": "Nine"}
_DIGIT_ZH = {"0": "零", "1": "一", "2": "二", "3": "三", "4": "四",
             "5": "五", "6": "六", "7": "七", "8": "八", "9": "九"}
_SYM_EN = {",": "comma", "，": "comma", ".": "period", "。": "period",
           "!": "exclamation mark", "！": "exclamation mark",
           "?": "question mark", "？": "question mark", ";": "semicolon",
           "；": "semicolon", ":": "colon", "：": "colon", "+": "plus",
           "-": "minus", "*": "asterisk", "/": "slash", "=": "equals",
           "%": "percent", "$": "dollar", "#": "hash", "@": "at", "&": "ampersand",
           "(": "left parenthesis", ")": "right parenthesis",
           "（": "left parenthesis", "）": "right parenthesis"}
_SYM_ZH = {",": "逗号", "，": "逗号", ".": "句号", "。": "句号",
           "!": "感叹号", "！": "感叹号", "?": "问号", "？": "问号",
           ";": "分号", "；": "分号", ":": "冒号", "：": "冒号", "+": "加号",
           "-": "减号", "*": "星号", "/": "斜杠", "=": "等号", "%": "百分号",
           "$": "美元符", "#": "井号", "@": "艾特", "&": "和号",
           "(": "左括号", ")": "右括号", "（": "左括号", "）": "右括号"}


def _app_data_dir():
    """系统约定的用户级应用数据目录。
    Windows: %APPDATA%\\EnglishCoach ；macOS: ~/Library/Application Support/EnglishCoach；
    其它: ~/.local/share/EnglishCoach。返回路径并确保存在。"""
    import os, sys
    if sys.platform.startswith("win"):
        base = os.environ.get("APPDATA") or os.path.expanduser("~")
    elif sys.platform == "darwin":
        base = os.path.expanduser("~/Library/Application Support")
    else:
        base = os.environ.get("XDG_DATA_HOME") or os.path.expanduser("~/.local/share")
    d = os.path.join(base, APP_NAME)
    try:
        os.makedirs(d, exist_ok=True)
    except Exception:
        d = os.path.expanduser("~")
    return d

HISTORY_FILE = None    # 延迟初始化
LOG_FILE = None

def _history_path():
    global HISTORY_FILE
    if HISTORY_FILE is None:
        import os
        HISTORY_FILE = os.path.join(_app_data_dir(), "翻译历史.txt")
    return HISTORY_FILE

def _log_path():
    global LOG_FILE
    if LOG_FILE is None:
        import os
        LOG_FILE = os.path.join(_app_data_dir(), "运行日志.txt")
    return LOG_FILE

def _network_hint(err):
    """按实际错误给出网络提示。不假设用户用哪款代理软件——代理来自系统环境
    变量(http_proxy/https_proxy)，各人配置不同，因此只描述现象与通用做法。"""
    t = str(err)
    import os as _os
    _pv = next((_os.environ.get(k) for k in
                ("https_proxy", "HTTPS_PROXY", "http_proxy", "HTTP_PROXY",
                 "all_proxy", "ALL_PROXY") if _os.environ.get(k)), None)
    if "ProxyError" in t or "Unable to connect to proxy" in t:
        if _pv:
            return (f"（系统已配置代理 {_pv}，但连接被拒绝：请确认代理软件正在运行，"
                    f"或清除代理环境变量后重试）")
        return "（代理连接失败：请确认代理软件正在运行，或取消代理设置后重试）"
    if "SSLError" in t or "CERTIFICATE" in t.upper():
        return "（TLS 握手失败：可能是代理或防火墙拦截，请检查网络环境）"
    if "Timed out" in t or "timeout" in t.lower():
        return "（连接超时：请检查网络；若该服务在本地区受限，需自行配置网络代理）"
    if _pv:
        return f"（当前经由代理 {_pv} 访问，若异常请检查代理是否正常工作）"
    return "（请检查网络连接；部分线上引擎在部分地区需自行配置网络代理）"


def _log_error(msg):
    """把出错记录追加到日志文件（带时间戳）。"""
    import datetime
    try:
        with open(_log_path(), "a", encoding="utf-8") as f:
            ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            f.write(f"[{ts}] {msg}\n")
    except Exception:
        pass

# 历史记录用 Markdown 文本格式存储：人可读，任意文本/浏览器打开都清晰；
# 每条以 "## " 开头便于解析回结构。
_HIST_SEP = "\n\n---\n\n"

def _load_history():
    import os, re
    p = _history_path()
    if not os.path.exists(p):
        return []
    try:
        with open(p, "r", encoding="utf-8") as f:
            raw = f.read()
    except Exception:
        return []
    items = []
    for block in raw.split(_HIST_SEP):
        block = block.strip()
        if not block.startswith("## "):
            continue
        try:
            header = block.splitlines()[0][3:].strip()
            parts = header.split(" · ")
            ts = parts[0].strip()
            engine = parts[1].strip() if len(parts) > 1 else ""
            body = block.split("\n", 1)[1] if "\n" in block else ""
            src = ""; tgt = ""
            if "【原文】" in body:
                after = body.split("【原文】", 1)[1]
                if "【译文】" in after:
                    src, tgt = after.split("【译文】", 1)
                else:
                    src = after
            items.append({"ts": ts, "engine": engine,
                          "src": src.strip(), "tgt": tgt.strip()})
        except Exception:
            continue
    return items

def _save_history(items):
    blocks = []
    for it in items[-500:]:
        blocks.append(
            f"## {it.get('ts','')} · {it.get('engine','')}\n\n"
            f"【原文】{it.get('src','')}\n\n"
            f"【译文】{it.get('tgt','')}")
    try:
        with open(_history_path(), "w", encoding="utf-8") as f:
            f.write("# EnglishCoach 翻译历史" + _HIST_SEP)
            f.write(_HIST_SEP.join(blocks))
            f.write("\n")
    except Exception as e:
        _log_error(f"保存历史失败: {e}")

def _add_history(src_text, tgt_text, engine):
    """新增一条翻译历史。返回更新后的列表。"""
    import datetime
    if not src_text.strip():
        return _load_history()
    items = _load_history()
    if items and items[-1].get("src") == src_text.strip():
        return items
    items.append({
        "ts": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "src": src_text.strip(), "tgt": tgt_text.strip(), "engine": engine,
    })
    _save_history(items)
    return items

# 版本更新说明 —— 以后每版在最前面追加一条记录即可
CHANGELOG = [
    {
        "version": "2.15.12",
        "date": "2026-07-28",
        "title": "恢复 API Key 输入框样式 · 复选框勾选色恢复为蓝色（不再跟随系统强调色）",
        "notes": [
            "修复 API Key 输入框在 Win10 深色下变白、Win11 出现原生底部亮条：输入框原本与下拉闭合框共用同一条样式规则，改用混合方案时只搬了下拉、把输入框漏了，导致它退回原生渲染。现已恢复，深色 #2d2d30 / 浅色 #ffffff，与下拉闭合框一致，悬停与聚焦为蓝边",
            "修复 Win11 复选框勾选后是黄色：复选框现在走系统原生渲染，勾选色默认取自【系统强调色】，系统若设成黄色勾选就是黄色。现在显式把调色板的强调色固定为程序蓝 #1e88e5，原生渲染与边线保持不变，Win10 同样受益",
            "关于朗读与卡拉OK字幕的原则：经实测验证，卡拉OK（青蓝 #5aa8b0 + 白字）确实覆盖在蓝色主动选区与灰色被动选区之上，三层重叠时最终显示卡拉OK；拖动进度条的『字幕铁律』（该侧音频缓存在且文字在则必有字幕，边界丢失会先从缓存恢复、再按朗读范围重建）也已在代码中实现",
            "本次仅改动非 mac 路径的两处样式；mac 分支未动（经伪装 darwin 验证不调用非 mac 的调色板函数、不重建样式）；6.11 与 6.4.2 两套 Qt 验证：主界面下拉、设置窗下拉、反复切换主题均正常，设置窗与主窗样式一致，无崩溃",
            "新增三平台的安装与卸载脚本：Windows 的 Install.bat / Uninstall.bat（装到用户目录并创建开始菜单与可选桌面快捷方式，无需管理员权限）、macOS 的 Install.command / Uninstall.command（复制到应用程序并自动移除隔离标记，未签名应用首次打开不再被拦截）、Linux 的 Install.sh / Uninstall.sh（装到 ~/.local/opt 或 /opt，安装图标并创建应用菜单入口）。卸载脚本默认保留翻译历史、设置与 API Key，会单独询问是否一并删除",
            "四个构建脚本现在会把对应平台的安装脚本一并打进产物",
            "Linux 构建脚本改进：conda 由必需改为可选，没有 conda 时自动创建虚拟环境（绕开 Debian/Ubuntu 系统 Python 的 PEP 668 限制），可直接在 Docker 容器里编译；旧虚拟环境若是低版本 Python 建的会自动重建；新增 Python 3.10 版本下限检查",
            "Linux 构建脚本钉死 PyQt6 6.9.1：6.10 起的 Linux 轮子要求 glibc 2.34 以上，在 Debian 11 / Ubuntu 20.04 上装不了，pip 会退去编译源码并因缺少 qmake 而失败。6.9.x 是能装在 glibc 2.31 上的最新版本，用它编出的产物可覆盖 Ubuntu 20.04 及以上",
            "Linux 产物现在捆绑四个常缺的 xcb 支持库（约 94KB）：Qt 6.5+ 需要 libxcb-cursor 才能加载 xcb 平台插件，而多数发行版默认不装，用户一启动就是一堆 qt.qpa.plugin 报错。libxcb.so.1、libX11、libc 等仍用系统的，因为它们必须与用户的 X 服务器和显卡驱动匹配",
            "Linux 启动脚本会在启动前检查系统库，缺失时直接给出按发行版适配的安装命令，而不是抛出难以理解的 Qt 插件报错",
            "仓库文档整理：新增 CHANGELOG.md（由 gen_changelog.py 从程序内的更新记录生成，与「关于 → 更新记录」同源）、.gitignore，依赖拆分为 requirements.txt（运行，带平台标记）、requirements-build.txt（打包工具）与 requirements-gpu.txt（CUDA 版 torch）；README 增加下载链接、各平台安装说明与常见问题；移除测试用的 scrollbar_probe.py",
        ],
        "title_en": "Restored the API key field styling · checkbox tick colour is blue again instead of following the system accent",
        "notes_en": [
            "Fixed the API key field turning white in dark mode on Win10 and showing a native bottom accent line on Win11: the field originally shared one style rule with the dropdown closed-box, and when the hybrid approach was introduced only the dropdown was carried over, leaving the field to fall back to native rendering. It is restored now — #2d2d30 in dark, #ffffff in light, matching the dropdown closed-box, with a blue border on hover and focus",
            "Fixed the Win11 checkbox tick appearing yellow: checkboxes are rendered natively, and the tick colour defaults to the system accent, so a yellow system accent produced a yellow tick. The palette's accent role is now pinned to the app blue #1e88e5, leaving native rendering and the indicator border untouched; Win10 benefits from the same change",
            "On the playback and karaoke subtitle principle: verified by test that karaoke (teal #5aa8b0 with white text) does paint above both the blue active selection and the grey passive selection, and wins where all three overlap; the slider's subtitle rule (whenever that side has cached audio and text, dragging must show subtitles, restoring boundaries from cache and otherwise rebuilding them from the spoken range) is already implemented",
            "This release changes only two styling spots on the non-mac path; the mac branch is untouched (a simulated-darwin run confirms it neither calls the non-mac palette helper nor rebuilds the style); verified on Qt 6.11 and 6.4.2: main window dropdowns, settings dropdowns and repeated theme switching all work, the settings dialog matches the main window, and there are no crashes",
            "Added install and uninstall scripts for all three platforms: Install.bat / Uninstall.bat on Windows (installs into the user profile and creates a Start Menu entry plus an optional Desktop shortcut, no administrator rights needed), Install.command / Uninstall.command on macOS (copies into Applications and strips the quarantine flag automatically so the unsigned app is not blocked on first launch), and Install.sh / Uninstall.sh on Linux (installs into ~/.local/opt or /opt, adds the icon and an application-menu entry). The uninstallers keep translation history, settings and API keys by default and ask separately before removing them",
            "All four build scripts now bundle the matching platform's install scripts into the output",
            "Linux build script improvements: conda is now optional rather than required, and without it a virtual environment is created automatically, side-stepping the PEP 668 restriction on Debian and Ubuntu system Python so the build runs inside a Docker container as-is; an existing virtual environment built with an older Python is detected and rebuilt; a Python 3.10 minimum-version check was added",
            "The Linux build pins PyQt6 6.9.1: from 6.10 onwards the Linux wheels require glibc 2.34 or newer, which Debian 11 and Ubuntu 20.04 lack, so pip falls back to building from source and fails for want of qmake. The 6.9.x wheels are the newest installable on glibc 2.31, and building against them keeps the result compatible with Ubuntu 20.04 and later",
            "Linux builds now bundle four commonly-missing xcb support libraries (about 94KB): Qt 6.5+ needs libxcb-cursor to load its xcb platform plugin, yet most distributions do not install it, leaving users with a wall of qt.qpa.plugin errors at startup. libxcb.so.1, libX11 and libc are deliberately left to the system because they must match the user's X server and graphics drivers",
            "The Linux launcher now checks for the required system libraries before starting and prints a ready-to-paste install command matched to the distribution, instead of surfacing an opaque Qt plugin failure",
            "Repository housekeeping: added CHANGELOG.md (generated by gen_changelog.py from the changelog inside the program, so it shares a source with the in-app What's New dialog) and .gitignore; split dependencies into requirements.txt (runtime, with platform markers), requirements-build.txt (packaging tools) and requirements-gpu.txt (CUDA torch); the README gained download links, per-platform installation instructions and troubleshooting notes; removed the diagnostic scrollbar_probe.py",
        ],
    },
    {
        "version": "2.15.11",
        "date": "2026-07-27",
        "title": "修复设置窗语言/主题下拉『整体向上串一行、末尾空白』：弹出列表被卡在滚下一行的位置",
        "notes": [
            "根因(由现象精确定位)：这不是丢项，而是弹出列表被卡在『向下滚了一行』的位置。这两个下拉的弹出高度正好等于项数，本不需要滚动；但选中最后一项时 Qt 会 scrollTo 让当前项可见，把视图滚下一行，而弹出的滚动条是关闭的、滚轮也被忽略，滚下去就卡住回不来 —— 于是显示成 items[1..N-1] 加一行空白，正是『中文/English US』变成『English US/空白』、『深色/浅色/跟随系统』变成『浅色/跟随系统/空白』的现象",
            "修复：只给设置窗的语言与主题这两个下拉，在弹出后把列表滚动位置归零。范围极窄 —— 不改任何弹出机制，不影响主界面下拉，也不影响设置窗里的其它下拉",
            "Linux/Fusion 下项目较矮、内容正好装得下，滚动上限为 0 根本滚不动，这正是此前在测试环境反复复现不出来的原因；windows11 原生样式下项目更高，才会滚出这一行",
            "Mac 分支直接跳过本处理，弹出仍走系统原生；6.11 与 6.4.2 两套 Qt 验证：主界面下拉(14 项)正常，反复选择并弹出 12 次后语言恒 2 项、主题恒 3 项、滚动位置恒为 0，无崩溃",
        ],
        "title_en": "Fixed the settings language/theme dropdowns appearing shifted up by one row with a blank at the end: the popup list was stuck scrolled down by one row",
        "notes_en": [
            "Root cause (pinned down from the exact symptom): nothing was ever lost — the popup list was stuck one row down. These dropdowns size their popup to exactly the number of items so no scrolling should be needed, but selecting the last item makes Qt scroll the view to reveal it, and because the popup's scrollbar is disabled and the wheel is ignored, it stays scrolled and cannot come back. The result is items[1..N-1] plus a blank row, exactly matching 'Chinese / English US' becoming 'English US / blank' and 'Dark / Light / Follow System' becoming 'Light / Follow System / blank'",
            "Fix: for the settings dialog's language and theme dropdowns only, the list scroll position is reset to the top after the popup opens. The change is deliberately narrow — no popup machinery is modified, the main window dropdowns are untouched, and other dropdowns in the settings dialog are unaffected",
            "Under Linux/Fusion the items are shorter and the content fits exactly, so the scroll maximum is zero and nothing can scroll — which is why this never reproduced in the test environment; the native windows11 style renders taller items, which is what produces the extra row",
            "The mac branch skips this handling entirely and popups still use the native system look; verified on both Qt 6.11 and 6.4.2: the main window dropdowns (14 items) work, and after 12 rounds of selecting and reopening the language dropdown stays at 2 items, the theme dropdown at 3, and the scroll position stays at 0, with no crashes",
        ],
    },
    {
        "version": "2.15.10",
        "date": "2026-07-27",
        "title": "回退 2.15.9 对下拉弹出机制的改动（它弄坏了主界面下拉），只保留设置窗语言/主题下拉的丢项自愈",
        "notes": [
            "回退：2.15.9 改了所有下拉共用的弹出机制（每次弹出重算可见项数、开滚动条兜底、滚轮条件放行），本意是修设置窗的丢项，结果把 2.15.8 里已经好用的主界面下拉一起弄坏了。现已完整恢复成 2.15.8 的弹出机制：滚动条关闭、可见项数按创建时项数、滚轮一律忽略、无弹出包装",
            "保留：2.15.8 的下拉弹出背景随主题刷新（深色 #2d2d30 / 浅色 #ffffff）不受影响，仍然生效",
            "保留并收紧：设置窗语言/主题下拉的丢项自愈改为只按身份值判断、且只在项【真的少了】时才重建 —— 项目齐全时完全不碰下拉，零副作用（此前按显示文字比对，会在每次界面重译时误触发重建）",
            "Mac 分支一字未动；6.11 与 6.4.2 两套 Qt 验证：主界面下拉(14 项引擎)与设置窗下拉(语言 2/主题 3)弹出均正常，弹出背景随主题正确，无崩溃",
        ],
        "title_en": "Reverted the 2.15.9 dropdown popup changes (they broke the main window dropdowns), keeping only the settings-dialog language/theme item self-heal",
        "notes_en": [
            "Reverted: 2.15.9 modified the popup machinery shared by every dropdown (recomputing the visible-item count on each popup, a scrollbar fallback, conditional wheel scrolling). It was meant to fix the settings dialog but broke the main window dropdowns that worked fine in 2.15.8. The 2.15.8 popup machinery is fully restored: scrollbar off, visible-item count fixed at creation, wheel always ignored, no popup wrapper",
            "Kept: the 2.15.8 fix that refreshes the dropdown popup background with the theme (#2d2d30 dark / #ffffff light) is unaffected and still active",
            "Kept and tightened: the settings dialog's language/theme item self-heal now checks identity values only and rebuilds only when an item is genuinely missing — when the items are complete it does not touch the dropdown at all, so there are no side effects (it previously compared display text, which triggered a needless rebuild on every UI retranslation)",
            "The mac branch is untouched; verified on both Qt 6.11 and 6.4.2: the main window dropdowns (14-item engine list) and the settings dropdowns (2 language / 3 theme) open correctly, the popup background follows the theme, and there are no crashes",
        ],
    },
    {
        "version": "2.15.9",
        "date": "2026-07-27",
        "title": "修复 Windows 语言/主题下拉反复切换后『丢一项』：弹出高度不够且滚不动，那一项永远够不着",
        "notes": [
            "根因：弹出列表同时做了三件相互冲突的事 —— 关闭垂直滚动条、禁用滚轮、且可见项数只在下拉『创建时』按当时项数算一次。在 windows11 原生样式下项目实际更高(原生内边距叠加样式表 padding)，弹出高度就只装得下 N-1 项；而滚动条和滚轮又都被关掉，那一项便永远够不着，看起来就是『丢了一项』(Linux/Fusion 装得下，所以此前一直复现不出)",
            "修复一：每次弹出前按当前真实项数重算可见项数；弹出后若仍装不下，自动打开滚动条兜底，保证任何一项都能被看到和选到",
            "修复二：滚轮不再一律拦截 —— 确实可滚动时放行(mac 保持原来一律忽略的行为不变)",
            "修复三：语言与主题两个下拉都加上项目完整性自愈 —— 每次界面重译后校验，一旦项数或身份值不符，自动重建为完整项并保留当前选择(此前只有语言下拉有这道保险)",
            "Mac 分支一字未动：尺寸重算与滚动兜底都在 mac 上直接跳过，弹出仍走系统原生(经伪装 darwin 验证)；6.11 与 6.4.2 两套 Qt 反复切换 16 次后语言恒 2 项、主题恒 3 项，无崩溃",
        ],
        "title_en": "Fixed the language/theme dropdowns losing an item after repeated switching on Windows: the popup was too short to fit every item and could not scroll",
        "notes_en": [
            "Root cause: the popup list did three conflicting things at once — the vertical scrollbar was disabled, the mouse wheel was blocked, and the visible-item count was computed only once when the dropdown was created. Under the native windows11 style the items are taller (native padding on top of the stylesheet padding), so the popup only fits N-1 items; with both the scrollbar and the wheel disabled, that last item became permanently unreachable and looked like it had disappeared (everything fits under Linux/Fusion, which is why this never reproduced in testing)",
            "Fix 1: the visible-item count is recomputed from the real item count before every popup, and if items still do not fit the scrollbar is enabled afterwards as a fallback so every item can be seen and selected",
            "Fix 2: the mouse wheel is no longer blocked unconditionally — it now scrolls when the list actually can scroll (mac keeps its previous always-ignore behaviour)",
            "Fix 3: both the language and theme dropdowns now self-heal their item lists — after every UI retranslation the items are validated and rebuilt with the current selection preserved if anything is missing (previously only the language dropdown had this safeguard)",
            "The mac branch is untouched: the resizing and scroll fallback are skipped entirely on mac and popups still use the native system look (verified with a simulated-darwin run); after 16 consecutive switches on both Qt 6.11 and 6.4.2 the language dropdown stays at 2 items and the theme dropdown at 3, with no crashes",
        ],
    },
    {
        "version": "2.15.8",
        "date": "2026-07-27",
        "title": "修复 Windows 切浅色后下拉弹出仍是黑底、以及由此造成的『下拉少一项』（两者是同一个 Bug）",
        "notes": [
            "根因：下拉弹出容器的背景色只在下拉『创建时』设过一次，切换主题时从不更新。所以从深色切到浅色后，弹出容器仍是黑底，而项目文字已随浅色主题变成深色 —— 深字配黑底就完全看不见了",
            "这同时解释了『反复改几次就丢一项』：项目其实一个都没丢（实测反复切换 12 次后语言仍是 2 项、主题仍是 3 项），只是没被选中的那些项深字黑底看不见；被选中那项因为有蓝色高亮加白字所以还看得见，于是看起来就像少了一项",
            "修复：把弹出容器配色抽成与主样式表共用的函数，并在主题热切换时刷新主窗与所有已打开弹窗内的全部下拉 —— 深色 #2d2d30、浅色 #ffffff，与下拉闭合框配色一致",
            "Mac 分支一字未动：mac 的下拉弹出走系统原生、不设任何样式，热切换时也不会调用这个刷新（经伪装 darwin 验证）；6.11 与 6.4.2 两套 Qt 反复切换均无崩溃",
        ],
        "title_en": "Fixed the dropdown popup staying black after switching to the light theme on Windows, and the resulting phantom missing item (both were the same bug)",
        "notes_en": [
            "Root cause: the dropdown popup container's background colour was set only once, when the dropdown was created, and never updated on a theme change. After switching from dark to light the popup container was still black while the item text had followed the light theme and turned dark — dark text on a black background is simply invisible",
            "This also explains the reported item loss after repeated changes: no item was ever removed (after 12 consecutive switches the language dropdown still had 2 items and the theme dropdown 3). The unselected items were merely invisible, while the selected one stayed visible thanks to its blue highlight and white text, making it look like an item had gone missing",
            "Fix: the popup container colours were extracted into a function shared with the main stylesheet, and the theme hot-switch now refreshes every dropdown in the main window and in any open dialog — #2d2d30 in dark, #ffffff in light, matching the dropdown closed-box",
            "The mac branch is untouched: mac dropdown popups use the native system look with no stylesheet, and the hot-switch never calls this refresh on mac (verified with a simulated-darwin run); repeated switching on both Qt 6.11 and 6.4.2 produces no crashes",
        ],
    },
    {
        "version": "2.15.7",
        "date": "2026-07-27",
        "title": "修复 Windows 深色切浅色后按钮图标/文字看不见：切主题时重新生成图标",
        "notes": [
            "找到根因：按钮图标是按当前深浅现场渲染的 SVG（浅色主题用深色 #1f1f22，深色主题用浅色 #e8e8e8）。mac 分支在切主题时一直会重新生成图标，非 mac 分支却漏了这一步——从深色切到浅色后，图标仍是浅灰色，落在浅底按钮上就看不见了",
            "修复：非 mac 的主题热切换现在同样调用图标重生成，并且连已打开的设置窗/弹窗里的按钮图标也一并按新深浅重新着色",
            "已用实际像素颜色验证：深色主题下图标为 #e8e8e8、浅色主题下为 #1f1f22，按钮文字色也随主题正确切换（深色 #dcdcdc / 浅色 #1f1f22）",
            "Mac 分支一字未动：mac 仍走 _mac_hybrid_qss，既不重建样式对象也不调用非 mac 的调色板函数（经伪装 darwin 验证）；6.11 与 6.4.2 两套 Qt 连切 5 次均无崩溃，设置窗与主窗始终一致",
        ],
        "title_en": "Fixed invisible button icons and text on Windows after switching dark to light: icons are now regenerated on theme change",
        "notes_en": [
            "Root cause: button icons are SVGs rendered on the fly for the current theme (dark #1f1f22 for the light theme, light #e8e8e8 for the dark theme). The mac branch has always regenerated icons on a theme change, but the non-mac branch skipped this step — after switching from dark to light the icons stayed light grey and became invisible on light buttons",
            "Fix: the non-mac theme hot-switch now regenerates icons too, including button icons inside any already-open Settings dialog or popup",
            "Verified by sampling actual pixel colors: icons are #e8e8e8 in the dark theme and #1f1f22 in the light theme, and button text follows the theme correctly (#dcdcdc dark / #1f1f22 light)",
            "The mac branch is untouched: mac still uses _mac_hybrid_qss and neither rebuilds the style object nor calls the non-mac palette helper (verified with a simulated-darwin run); five consecutive switches on both Qt 6.11 and 6.4.2 produce no crashes and the Settings dialog stays identical to the main window",
        ],
    },
    {
        "version": "2.15.6",
        "date": "2026-07-27",
        "title": "修复 Windows 切换深浅色后仍错乱：热切换时重建样式对象并全量重绘（Mac 不受影响）",
        "notes": [
            "找到『重启才好』的最后一块拼图：启动时的顺序是 设调色板 -> app.setStyle() 新建样式对象 -> 再设 colorScheme -> 建窗口。原生样式(windows11)在【创建那一刻】确定深浅状态，之后再改 colorScheme，这个已存在的样式对象不会彻底重绘——所以只有重启(重新创建样式对象)才正常",
            "修复：主题热切换时完整复刻启动顺序——重设调色板、重建样式对象、在其后再设一次调色板与 colorScheme(setStyle 会把调色板重置为样式标准值，必须在其后补回)，最后重套混合样式表",
            "新增全量重新 polish：仅调 update() 不足以让已经 polish 过的控件按新样式与调色板重绘，现在对所有控件执行 unpolish/polish，强制彻底刷新",
            "Mac 分支一字未动：mac 仍走 _mac_hybrid_qss，既不重建样式对象也不调用 _apply_win_palette(经伪装 darwin 验证)；6.11 与 6.4.2 两套 Qt 反复切换 5 次均无崩溃，样式引擎保持不变、调色板随主题正确变化、设置窗与主窗始终一致",
        ],
        "title_en": "Fixed Windows still breaking after a light/dark switch: the style object is now rebuilt and everything repolished on switch (Mac unaffected)",
        "notes_en": [
            "Found the last piece of the \"only a restart fixes it\" puzzle: startup does palette -> app.setStyle() creating a fresh style object -> colorScheme -> build window. The native windows11 style fixes its light/dark state at the moment it is created, so changing colorScheme afterwards never fully repaints the existing style object — which is exactly why only a restart looked right",
            "Fix: the theme hot-switch now mirrors the startup sequence exactly — set the palette, rebuild the style object, then set the palette and colorScheme again afterwards (setStyle resets the palette to the style's standard one, so it must be restored), and finally re-apply the hybrid stylesheet",
            "Added a full repolish: calling update() alone does not make already-polished widgets repaint under the new style and palette, so unpolish/polish is now run across all widgets to force a complete refresh",
            "The mac branch is untouched: mac still uses _mac_hybrid_qss and neither rebuilds the style object nor calls _apply_win_palette (verified with a simulated-darwin run); on both Qt 6.11 and 6.4.2, five consecutive switches produce no crashes, the style engine stays the same, the palette tracks the theme, and the Settings dialog stays identical to the main window",
        ],
    },
    {
        "version": "2.15.5",
        "date": "2026-07-26",
        "title": "真正修复 Windows 改深浅色后错乱 + 复选框没边线（Mac 不受影响）",
        "notes": [
            "找到真正根因：启动时非 mac 会给 app 设一套深/浅调色板，但主题热切换(apply_theme)在改用混合方案后漏掉了这一步——原生控件(复选框边线、背景)保留启动时的旧调色板，于是出现『打开好、改主题坏、重启又好』；现在启动与热切换共用同一个 _apply_win_palette，切主题时调色板同步更新",
            "修复复选框小方块没有边线：复选框改为完全走原生(由 setColorScheme+调色板驱动)——移除了窗口级样式表里的 QCheckBox 规则，以及每个复选框的内联样式；只要给复选框设任何样式表，windows11 引擎就会对它整体接管渲染而丢失原生边线(测试程序的混合模式正是不设 QCheckBox 规则才正确)",
            "两条修复合力：调色板同步 + 复选框纯原生，改主题后复选框、按钮、下拉、背景都跟着正确变深浅",
            "Mac 分支一字未动：mac 继续走 _mac_hybrid_qss，且 _apply_win_palette 仅在非 mac 调用(经伪装 darwin 验证 mac 路径不碰它)；6.11 与 6.4.2 两套 Qt 环境反复切换均无崩溃、设置窗与主窗始终一致",
        ],
        "title_en": "Actually fixed Windows corruption after changing light/dark + checkboxes with no border (Mac unaffected)",
        "notes_en": [
            "Found the true root cause: at startup the non-mac path sets a full dark/light palette on the app, but the theme hot-switch (apply_theme) dropped this step when it moved to the hybrid approach — native controls (checkbox borders, backgrounds) kept the startup palette, producing \"fine on open, broken on change, fine after restart\"; startup and hot-switch now share the same _apply_win_palette so the palette updates on every theme change",
            "Fixed checkboxes having no border: checkboxes are now fully native (driven by setColorScheme plus the palette) — the QCheckBox rule was removed from the window stylesheet along with each checkbox's inline style; setting any stylesheet on a checkbox makes the windows11 engine take over its rendering and lose the native border (the test program's hybrid mode was correct precisely because it set no QCheckBox rule)",
            "The two fixes together: palette sync plus fully-native checkboxes mean that after a theme change the checkboxes, buttons, dropdowns and backgrounds all switch light/dark correctly",
            "The mac branch is untouched: mac still uses _mac_hybrid_qss, and _apply_win_palette is called only on non-mac (verified via a simulated-darwin run that the mac path never touches it); repeated switching on both Qt 6.11 and 6.4.2 shows no crashes and the Settings dialog stays identical to the main window",
        ],
    },
    {
        "version": "2.15.4",
        "date": "2026-07-26",
        "title": "修复 Windows 在设置里改深浅色后界面错乱（设置窗改为与主窗完全一致，Mac 不受影响）",
        "notes": [
            "根因找到(靠『刚开好、改主题坏、重启又好』这条线索锁定)：设置窗的下拉/按钮样式函数只有 mac 分支、没有 Windows 分支，改主题时它给设置窗套上一张残缺样式表(只有下拉弹出、没有下拉闭合框和按钮)，覆盖了主窗正确的混合样式，于是整个界面错乱；重启因设置窗未打开而恢复正常",
            "修复：非 mac 时设置窗直接套用主窗的混合样式(_win_hybrid_qss)，与主界面一模一样，且只有一个深浅真相来源，改主题不再错乱",
            "mac 分支一字未动：mac 继续走原有的 _mac_hybrid_qss，效果完全不变(经伪装 darwin 验证 mac 代码路径原样执行)",
            "6.11 与 6.4.2 两套 Qt 环境均验证：反复切换深浅色后，设置窗样式始终与主窗完全一致，下拉正常、无崩溃",
        ],
        "title_en": "Fixed Windows UI corruption after changing light/dark in Settings (Settings dialog now identical to the main window; Mac unaffected)",
        "notes_en": [
            "Root cause (pinned down by the clue \"fine on open, breaks on theme change, fine again after restart\"): the Settings dialog's dropdown/button style function had only a mac branch and no Windows branch, so changing the theme applied an incomplete stylesheet (dropdown popup only, no closed-box or buttons) that overrode the main window's correct hybrid style and corrupted the whole UI; a restart looked fine because the Settings dialog was not open",
            "Fix: on non-mac the Settings dialog now applies the main window's hybrid stylesheet (_win_hybrid_qss) directly, making it identical to the main UI with a single source of truth for light/dark, so changing the theme no longer corrupts it",
            "The mac branch is untouched: mac still uses the existing _mac_hybrid_qss and is completely unchanged (verified via a simulated-darwin run that the mac code path executes as before)",
            "Verified on both Qt 6.11 and 6.4.2: after repeated light/dark switches the Settings dialog stylesheet stays identical to the main window, dropdowns work, no crashes",
        ],
    },
    {
        "version": "2.15.3",
        "date": "2026-07-26",
        "title": "Windows/Linux 改用混合主题方案（照搬 mac 成功模式）：深浅切换正常、复选框原生、按钮下拉自绘保持",
        "notes": [
            "彻底重做 Windows/Linux 的深浅主题：改用与 mac 同构的混合方案——用 Qt 的 setColorScheme 驱动系统原生控件(复选框、滚动条、窗口底色随之正确变深浅)，样式表只绘制按钮、下拉闭合框、下拉弹出(蓝色高亮)、滑杆、状态栏，不再碰复选框指示器与滚动条",
            "解决 Win11 深浅切换后按钮字太浅、复选框没边线/勾选变实心：这些都源于旧方案用样式表强行涂色与 windows11 原生引擎打架；混合方案让原生引擎自己画复选框，样式表只管按钮/下拉，两不相扰",
            "下拉弹出的蓝色高亮(#0e639c)+白字效果完整保留，深浅两主题都在",
            "此改动仅影响非 mac 路径；mac(Intel/Silicon)继续用原有的 _mac_hybrid_qss，效果完全不变",
        ],
        "title_en": "Windows/Linux switched to the hybrid theme approach (mirroring the working mac model): correct dark/light, native checkboxes, self-drawn buttons and dropdowns preserved",
        "notes_en": [
            "Reworked Windows/Linux dark/light theming to mirror mac's hybrid approach: Qt's setColorScheme drives the native controls (checkboxes, scrollbars and window background change light/dark correctly), while the stylesheet only draws buttons, the dropdown closed-box, the dropdown popup (blue highlight), sliders and the status bar — it no longer touches the checkbox indicator or scrollbar",
            "Fixes pale button text and missing checkbox borders / solid-fill checkmarks on Win11 after a theme switch: these came from the old approach fighting the native windows11 engine with stylesheet colors; the hybrid approach lets the native engine draw checkboxes while the stylesheet handles only buttons and dropdowns",
            "The dropdown popup's blue highlight (#0e639c) with white text is fully preserved in both light and dark themes",
            "This change affects only the non-mac path; mac (Intel/Silicon) continues to use the existing _mac_hybrid_qss and is completely unchanged",
        ],
    },
    {
        "version": "2.15.2",
        "date": "2026-07-26",
        "title": "Win11 深浅切换按钮字太浅/复选框实心 · 语言下拉丢项 两个老问题再修（不影响 Mac）",
        "notes": [
            "撤销上一版给复选框指示器加的自绘样式：它在 Win11 上让勾选态变成实心方块(SVG 对勾在 windows11 原生样式下不渲染)。现改为让 Windows 原生样式自己画正确的对勾",
            "修复 Win11 浅色主题下按钮文字太浅看不清：跟随系统且 Qt 返回未知色彩方案时，之前默认按深色处理，导致浅底配浅字；现在未知时保守当作浅色，只有明确深色才用深色配色",
            "语言下拉丢项(中文项消失)彻底根治：改用 userData 存身份值让显示文字与身份解耦(读 currentData 而非文字)，并在每次界面重译后加自愈保险——若因任何原因少了项，自动重建为『中文/English US』两项并保留当前选择",
            "以上仅改动非 Mac 路径与共享样式表，Mac(Intel/Silicon)的原生混合方案完全未动，效果不受影响",
        ],
        "title_en": "Win11 dark/light: pale button text / solid checkbox · language dropdown losing an item — both re-fixed (Mac untouched)",
        "notes_en": [
            "Reverted the custom checkbox-indicator styling added last version: on Win11 it turned the checked state into a solid square (the SVG check does not render under the native windows11 style). The native Windows style now draws the correct checkmark",
            "Fixed pale, unreadable button text in light theme on Win11: when following the system and Qt reported an unknown color scheme, it previously defaulted to dark, producing light text on a light background; unknown is now treated as light, and only an explicit dark scheme uses dark colors",
            "Definitively fixed the language dropdown losing its Chinese item: it now stores identity values in userData (decoupled from display text, read via currentData rather than text), plus a self-healing guard after every UI retranslation that rebuilds the two items (Chinese / English US) and keeps the current selection if any item goes missing",
            "These changes touch only the non-Mac path and the shared stylesheet; the Mac (Intel/Silicon) native hybrid approach is entirely untouched and unaffected",
        ],
    },
    {
        "version": "2.15.1",
        "date": "2026-07-26",
        "title": "多风格换行误判修复 · 选区/卡拉OK统一白字 · 日志写入修复 · Win11深浅复选框 · 被动选区朗读报错",
        "notes": [
            "修复普通翻译时原文含换行被误判为多风格分区：此前用译文第一个空行作直译区/多风格区分界，导致原文如『下载(换行)计算机』译成『Download(空行)computer』时 computer 被误划入多风格区变灰。现在只有多风格模式真正开启时才做空行分界，普通模式绝不分区",
            "文字颜色统一：原文区/译文区只要有选区或卡拉OK效果，被覆盖的字一律显示白色——不论深浅主题、不论直译区黑字还是多风格灰字、不论主动蓝色还是被动灰色背景；未被覆盖的多风格区仍是灰字",
            "修复朗读时到达处的青蓝卡拉OK效果+白字有时不覆盖：三种背景格式(青蓝已读/蓝选区/灰联动)现都显式带白字前景，覆盖一致",
            "修复朗读被动灰色联动区报错『cannot access local variable text』：该分支里 len(text) 用在了 text 赋值之前，现已调整顺序",
            "修复日志文件一直为空(谎报写入)：异常钩子调用的是不存在的 _log 函数，写入在 try 中被静默吞掉，却仍提示已记录；改为正确的 _log_error，日志真正写入(全平台，非仅 Silicon)",
            "修复 Win11 深浅切换后复选框没有边框：给复选框指示器加了显式样式(边框+背景+白色对勾 SVG)，不再依赖 windows11 原生样式在调色板切换时重绘",
            "进度滑杆拖到最左的卡拉OK边界问题与朗读到达覆盖逻辑一并改良",
        ],
        "title_en": "Multi-style newline misdetection fix · unified white text for selection/karaoke · log-write fix · Win11 dark/light checkbox · passive-selection speak error",
        "notes_en": [
            "Fixed normal translations with a newline in the source being misread as a multi-style split: the first blank line in the output was used as the direct/multi-style boundary, so a source like 'download(newline)computer' translated to 'Download(blank line)computer' had computer wrongly grayed into the multi-style area. The blank-line split now only happens when multi-style mode is actually on; normal mode never splits",
            "Unified text color: in both the source and target areas, any text under a selection or karaoke effect now shows white — regardless of light/dark theme, direct-area black text or multi-style gray text, and active blue or passive gray background; uncovered multi-style text stays gray",
            "Fixed the teal karaoke highlight + white text sometimes not covering during playback: all three background formats (teal read / blue selection / gray link) now explicitly carry a white foreground for consistent coverage",
            "Fixed 'cannot access local variable text' when speaking the passive gray link region: that branch used len(text) before text was assigned; the order is now corrected",
            "Fixed the log file always being empty (falsely reporting a write): the exception hook called a non-existent _log function, so the write was silently swallowed in a try while still claiming it logged; switched to the correct _log_error so the log actually writes (all platforms, not just Silicon)",
            "Fixed checkboxes losing their border after a light/dark switch on Win11: the checkbox indicator now has explicit styling (border, background and a white check SVG) instead of relying on the windows11 native style to repaint on a palette change",
            "Improved the karaoke boundary when the progress slider is dragged fully left, alongside the playback coverage logic",
        ],
    },
    {
        "version": "2.15.0",
        "date": "2026-07-25",
        "title": "修复 macOS Apple Silicon：不再弹新窗、深浅色正常 · 按架构分 Qt 版本",
        "notes": [
            "修复 Apple Silicon 打包版生成朗读音频时不断弹出新 App 窗口：torch/Kokoro 用 multiprocessing 起子进程，冻结的 app 里未调用 freeze_support() 会让每个子进程重新启动整个程序；现在入口最前调用 multiprocessing.freeze_support()",
            "修复 Apple Silicon 界面颜色混乱、深色模式按钮全白、切换深浅无反应：mac 深浅之前只靠 pyobjc(AppKit)驱动，Silicon+新系统上失效；现改为优先用 Qt 原生 setColorScheme/colorScheme(6.5+，Apple Silicon 的 6.11 可靠且不依赖 pyobjc)，老 Intel(6.4.2)自动回退 AppKit",
            "按架构分 Qt 版本：Apple Silicon 用 PyQt6 6.11.x(最低 macOS 12.0)，Intel 保持 6.4.2(最低 Big Sur 11.0)；Windows 与 Linux 用 6.11.x。构建脚本按 uname -m 自动选择版本与最低系统",
            "构建脚本两架构都装 pyobjc-framework-Cocoa，作为标题栏等系统装饰深浅同步的辅助",
        ],
        "title_en": "Fixed macOS Apple Silicon: no more popup windows, correct light/dark · Qt version per architecture",
        "notes_en": [
            "Fixed the Apple Silicon packaged build spawning new app windows while generating audio: torch/Kokoro use multiprocessing to start workers, and in a frozen app without freeze_support() each worker relaunches the whole program; multiprocessing.freeze_support() is now called first thing at the entry point",
            "Fixed Apple Silicon UI color chaos, all-white buttons in dark mode, and light/dark switching doing nothing: mac dark/light previously relied solely on pyobjc (AppKit), which fails on Silicon with newer systems; it now prefers Qt's native setColorScheme/colorScheme (6.5+, reliable on Apple Silicon's 6.11 without pyobjc) and falls back to AppKit on older Intel (6.4.2)",
            "Qt version per architecture: Apple Silicon uses PyQt6 6.11.x (min macOS 12.0), Intel stays on 6.4.2 (min Big Sur 11.0); Windows and Linux use 6.11.x. The build script picks the version and minimum system automatically from uname -m",
            "The build script installs pyobjc-framework-Cocoa on both architectures, as an aid for syncing system decorations like the title bar to light/dark",
        ],
    },
    {
        "version": "2.14.9",
        "date": "2026-07-25",
        "title": "滚动条回归系统原生 · 语言下拉丢项/主题错乱/卡拉OK字色/暂停失效/滚轮误改 五项修复 · 新增 Linux 构建脚本",
        "notes": [
            "滚动条不再强加任何自定义样式，各平台一律用系统原生：Win11 原生 Fluent 圆角、Win10 原生 Vista 直角、macOS 原生——与主窗风格统一",
            "修复 Windows 版语言下拉在中英切换后丢失中文项：语言项是身份值(用于判断当前语言)，不应被界面重译改写，现标记为不参与重译",
            "修复深浅主题互换时设置窗下拉/按钮颜色错乱：设置窗缺少主题热切换回调，切主题时保留了旧配色；现补上回调，按新深浅重建自绘样式",
            "卡拉OK字幕字色统一为白色：多风格区原本灰字，被卡拉OK/选区背景覆盖的部分现一律显示白字，与直译区一致；未被覆盖的灰区保持灰字",
            "修复全选文字并有卡拉OK时暂停键失效：全选朗读时选区仍在，暂停点击被误判为新朗读请求而重启；现同段朗读时优先按暂停/继续处理，只有真正不同的新选区才重启",
            "修复设置窗内用滚轮滚动内容时误改下拉选项：下拉框未获焦点时不再响应滚轮，把滚动交给页面",
            "新增 Linux 一键构建脚本 Build Linux.sh：PyInstaller 打包，产物为 dist/Linux-x64/ 下的可执行目录与 tar.gz；建议在较老发行版(如 Ubuntu 20.04)上编译以保 glibc 向下兼容",
        ],
        "title_en": "Scrollbars back to OS-native · five fixes (language dropdown, theme colors, karaoke text, pause, wheel) · new Linux build script",
        "notes_en": [
            "Scrollbars no longer force any custom style; every platform uses its native default: Win11 native Fluent rounded, Win10 native Vista square, macOS native — consistent with the main window",
            "Fixed the Windows language dropdown losing its Chinese item after switching languages: the language items are identity values (used to determine the current language) and must not be rewritten by UI retranslation; they are now marked to skip it",
            "Fixed dropdown and button colors going wrong in the Settings dialog when switching light/dark themes: the dialog lacked a theme-refresh callback and kept its old colors; the callback is now added and rebuilds the styling for the new theme",
            "Karaoke subtitle text is now uniformly white: the multi-style area was gray, and the parts covered by the karaoke/selection background now show white text like the direct-translation area; uncovered gray regions stay gray",
            "Fixed the pause button failing when playing a full-text selection with karaoke: the selection stays active during playback, so a pause click was misread as a new play request and restarted; pausing/resuming now takes priority for the same spoken span, and only a genuinely different selection restarts",
            "Fixed the scroll wheel accidentally changing dropdown values while scrolling the Settings dialog: dropdowns no longer respond to the wheel unless focused, passing the scroll to the page",
            "Added a one-command Linux build script (Build Linux.sh): PyInstaller packaging producing an executable folder and a tar.gz under dist/Linux-x64/; build on an older distro (e.g. Ubuntu 20.04) for forward glibc compatibility",
        ],
    },
    {
        "version": "2.14.8",
        "date": "2026-07-25",
        "title": "自绘滚动条配色改进：Win10 深浅主题下都清晰可见",
        "notes": [
            "修复 Win10 上自绘圆角滚动条看不清：此前滑块用半透明灰，对比依赖底层背景色，Win10 背景不同就与轨道糊在一起",
            "改用不透明实色并主动画一层浅色圆角轨道，滑块与轨道之间保证足够亮度差；按深浅主题自动选配色，浅色主题用中深灰滑块、深色主题用中灰滑块",
            "各类背景(浅色/深色/中灰)下实测滑块与轨道亮度差均 ≥80，清晰可辨",
        ],
        "title_en": "Custom scrollbar colors improved: clearly visible on Win10 in both light and dark themes",
        "notes_en": [
            "Fixed the custom rounded scrollbar being hard to see on Win10: the slider previously used a semi-transparent gray whose contrast depended on the underlying background, so on Win10's different background it blended into the track",
            "Now uses opaque solid colors and draws a light rounded track underneath, guaranteeing enough brightness difference between slider and track; colors are chosen automatically per light/dark theme (a mid-dark gray slider on light themes, a mid gray slider on dark themes)",
            "Measured slider-vs-track brightness difference is >=80 across light, dark and mid-gray backgrounds — clearly distinguishable",
        ],
    },
    {
        "version": "2.14.7",
        "date": "2026-07-25",
        "title": "修复 DeepSeek 模型停用导致翻译失败 · Win10 滚动条改自绘圆角 · 构建脚本消噪",
        "notes": [
            "修复 DeepSeek 翻译报 400：deepseek-chat 模型名已于 2026-07-24 15:59 UTC 停用，改用官方新名 deepseek-v4-flash（原 deepseek-chat 对应的经济档）",
            "翻译任务显式关闭 DeepSeek 的思考模式（v4-flash 默认开启思考，会平添延迟与费用），保持与原来一致的快速非思考行为",
            "Win10 滚动条改用自绘方案（QProxyStyle 直接用 QPainter 画圆角胶囊）：此前给滚动条套 Fusion 样式在实机上仍是直角，因为 Qt 新版的 windows11 引擎会忽略 QSS 圆角、且样式表回退到平台样式；自绘完全绕开样式引擎与 QSS 的层叠，任何 Qt 版本都画出一致圆角",
            "构建脚本消除 en_core_web_sm 的无害报错：该 spaCy 模型不在 PyPI（走 GitHub Releases），镜像返回 0 字节占位导致 Wheel invalid 报错；现改为先检查是否已安装，已装则跳过，未装才从官方地址安装",
        ],
        "title_en": "Fixed DeepSeek translation failure from model retirement · Win10 scrollbars now custom-drawn round · quieter build scripts",
        "notes_en": [
            "Fixed DeepSeek translation returning HTTP 400: the deepseek-chat model name was retired on 2026-07-24 15:59 UTC, replaced with the official new name deepseek-v4-flash (the economical tier the old deepseek-chat mapped to)",
            "Thinking mode is now explicitly disabled for DeepSeek translation (v4-flash enables thinking by default, adding latency and cost), preserving the previous fast non-thinking behaviour",
            "Win10 scrollbars are now custom-drawn (a QProxyStyle painting rounded capsules directly with QPainter): applying the Fusion style still rendered square on real machines because Qt's newer windows11 engine ignores QSS rounding and the stylesheet fell back to the platform style; custom drawing bypasses the style engine and QSS entirely for consistent rounded corners on any Qt version",
            "Build scripts no longer print the harmless en_core_web_sm error: that spaCy model is not on PyPI (it ships via GitHub Releases) and mirrors return a 0-byte placeholder that fails wheel validation; the scripts now check whether it is already installed, skip if so, and otherwise install from the official URL",
        ],
    },
    {
        "version": "2.14.6",
        "date": "2026-07-25",
        "title": "修复 mac 进设置窗闪退（NameError: lv）",
        "notes": [
            "真因找到：v2.14.3 重写下拉弹出函数时，旧函数的尾部代码被遗落在函数体外，成了缩进错误的游离代码块，其中引用的 lv 变量已不在作用域内",
            "这段游离代码恰好是 macOS 专属分支(弹出列表走系统原生透明背景)，所以只在 mac 上触发、Linux 沙盒测不出来——这也是前两版反复没修对的原因",
            "已删除游离代码，并把 macOS 原生弹出处理放回它应在的函数内（lv 有效作用域）",
            "顺带全量扫描了模块结构，确认没有其它同类遗留代码块",
            "上一版加的全局异常钩子发挥了作用：正是它把静默闪退变成了可读的错误提示，才得以一次定位",
        ],
        "title_en": "Fixed macOS crash when opening Settings (NameError: lv)",
        "notes_en": [
            "Real cause found: when the dropdown popup function was rewritten in v2.14.3, the tail of the old function was left outside the function body as a mis-indented orphaned block, referencing the variable lv which was no longer in scope",
            "That orphaned block happened to be the macOS-only branch (native translucent popup), so it fired only on Mac and could not be reproduced in the Linux sandbox — which is why the two previous attempts fixed the wrong thing",
            "The orphaned block has been removed and the macOS native popup handling restored inside the function where lv is actually in scope",
            "Also scanned the whole module structure to confirm no other leftover blocks of this kind exist",
            "The global exception hook added in the previous version did its job: it turned a silent crash into a readable error message, which is what made this diagnosis possible",
        ],
    },
    {
        "version": "2.14.5",
        "date": "2026-07-25",
        "title": "紧急修复：一点设置就闪退（槽函数异常导致 PyQt 直接中止）",
        "notes": [
            "找到闪退真因：崩溃栈显示 pyqt6_err_print → QMessageLogger::fatal → qAbort，且发生在按钮点击(QAbstractButton::mouseReleaseEvent)之后——这是 PyQt6 的行为：槽函数里任何未被捕获的 Python 异常都会让程序直接 abort()，而不是像普通异常那样被忽略",
            "新增全局异常钩子：槽函数抛异常时改为写入日志并弹窗提示，程序继续运行，不再闪退；日志中以 [UNCAUGHT] 标记，便于定位",
            "『保持程序置顶』复选框的响应函数此前完全没有异常保护，正是崩溃栈指向的按钮点击路径，现已加上",
            "置顶收尾函数改为局部导入 Qt 并整体保护，避免作用域问题引发异常",
            "全流程回归通过：设置窗四个按钮逐一点击(含此前必崩的『导出日志』)、置顶勾选与取消、三个下拉弹出、关闭收尾，均无闪退",
        ],
        "title_en": "Critical fix: crash when touching Settings (slot exception aborting PyQt)",
        "notes_en": [
            "Found the real cause: the crash report shows pyqt6_err_print → QMessageLogger::fatal → qAbort right after a button click (QAbstractButton::mouseReleaseEvent). This is PyQt6 behaviour — any uncaught Python exception inside a slot makes the program abort() outright rather than being ignored as a normal exception would be",
            "Added a global exception hook: an exception raised in a slot is now logged and shown in a dialog while the program keeps running, instead of crashing; entries are tagged [UNCAUGHT] in the log for easy tracing",
            "The Keep Window on Top checkbox handler had no exception protection at all — precisely the button-click path the crash stack pointed to — and is now guarded",
            "The always-on-top flush function now imports Qt locally and is fully guarded, removing any scope-related failure",
            "Full regression passed: every Settings button clicked in turn (including Export Log, which previously killed the process), on-top toggled on and off, all three dropdowns opened, and the dialog closed — no crashes",
        ],
    },
    {
        "version": "2.14.4",
        "date": "2026-07-25",
        "title": "紧急修复：进入设置即闪退（滚动条样式对象悬空指针）",
        "notes": [
            "修复上一版一点设置就闪退：QWidget.setStyle() 并不接管样式对象的所有权，上版把 Fusion 样式存成应用的 Python 属性以为能保住它，但 Qt 退出时仍会销毁该对象，而滚动条还指向它——悬空指针导致段错误闪退",
            "改为给样式对象设置真正的 Qt 父对象，生命周期交由 Qt 对象树管理，杜绝悬空",
            "事件过滤器加固：只在窗口 Show 时处理并延后一拍执行，避免控件尚未构造完成或已销毁时被访问；对已销毁控件静默跳过",
            "滚动条样式改为按需设置（已是目标样式则不重复设），减少无谓重绘",
        ],
        "title_en": "Critical fix: crash when opening Settings (dangling scrollbar style pointer)",
        "notes_en": [
            "Fixed the previous version crashing as soon as Settings was touched: QWidget.setStyle() does not take ownership of the style object, and storing the Fusion style as a Python attribute on the application did not keep it alive — Qt still destroyed it at shutdown while scrollbars were pointing at it, producing a dangling pointer and a segmentation fault",
            "The style object is now given a real Qt parent so its lifetime is managed by the Qt object tree, eliminating the dangling pointer",
            "Hardened the event filter: it now acts only on window Show events and defers by one tick, avoiding access to widgets that are not yet fully constructed or have already been destroyed; destroyed widgets are skipped silently",
            "Scrollbar styling is now applied only when it differs from the current state, avoiding redundant repaints",
        ],
    },
    {
        "version": "2.14.3",
        "date": "2026-07-24",
        "title": "下拉弹出根治 · Win10 圆角滚动条(方案B) · 构建目录分平台 · 更新说明全部双语",
        "notes": [
            "根治 Windows 下拉点击无效与文字截断：此前用猴子补丁改写 showPopup 并 setFixedWidth 永久锁死宽度——重复应用会层层包裹导致弹出失败，锁死的宽度又让后续任何重新布局都无法调整，越改越窄；现改为 minimumWidth 由 Qt 自行排版，可反复调用",
            "弹出宽度计算补齐内边距/边框/勾号(此前只算纯文字宽，余量不足才出现 …)",
            "Win10 圆角滚动条改用方案B：只给滚动条套 Fusion 样式(其它控件外观不变)，绕开 Qt 6.7+ 的 windows11 引擎忽略 QSS 圆角的问题；通过全局事件过滤器覆盖主窗与所有对话框",
            "中文朗读引擎下拉在中英切换后不再变窄、弹出项不再显示成 -On…：译后按新文字重算闭合框与弹出宽度，宽度只增不减",
            "构建产物按平台分目录：dist/MacOS-Intel、MacOS-AppleSilicon(自动按 uname -m 判定)、Windows-x64-CPU、Windows-x64-GPU；文件名自描述如 EnglishCoach-2.14.3-Windows-x64-CPU.zip；各脚本只清理自己的目录，不再互相覆盖",
            "程序置顶：首次运行显式写入未勾选状态，确保默认不置顶",
            "历史更新说明翻译全部完成——87 条全部中英双语",
        ],
        "title_en": "Dropdown popup root fix · rounded scrollbars on Win10 (option B) · per-platform build folders · change log fully bilingual",
        "notes_en": [
            "Root fix for Windows dropdowns not opening and text being truncated: the previous code monkey-patched showPopup and used setFixedWidth, so re-applying the styling wrapped the patch in itself until the popup failed to open, while the locked width blocked every later relayout and made the popup progressively narrower; width is now set via minimumWidth and laid out by Qt, safe to re-apply",
            "Popup width calculation now accounts for padding, borders and the checkmark (it previously measured raw text only, leaving too little room and causing ellipses)",
            "Rounded scrollbars on Win10 now use option B: the Fusion style is applied to scrollbars only, leaving every other control's appearance untouched, which sidesteps the Qt 6.7+ windows11 engine ignoring QSS border-radius; a global event filter covers the main window and all dialogs",
            "The Chinese voice dropdown no longer narrows after switching between Chinese and English, and popup entries no longer show as -On…: widths are recalculated from the translated text and only ever grow",
            "Build artifacts are now organized per platform: dist/MacOS-Intel, MacOS-AppleSilicon (chosen automatically from uname -m), Windows-x64-CPU and Windows-x64-GPU, with self-describing names such as EnglishCoach-2.14.3-Windows-x64-CPU.zip; each script cleans only its own folder so platforms no longer overwrite each other",
            "Always-on-top: the unchecked state is written explicitly on first run so the default is genuinely off",
            "Historical change log translation is complete — all 87 entries are now bilingual",
        ],
    },
    {
        "version": "2.14.2",
        "date": "2026-07-24",
        "title": "使用说明补齐分发说明 · 大陆用户免 VPN 下载模型",
        "notes": [
            "使用说明新增『哪些需要联网』与『中国大陆用户须知』两节，并说明无需安装 Python / 依赖 / conda、Linux 暂无预编译版、GPU 版仅适合 NVIDIA 机器",
            "英文版使用说明补齐 System Requirements 整节（此前完全缺失），与中文版内容对齐",
            "中国大陆免 VPN 下载 Kokoro 朗读模型：huggingface.co 大陆无法直连，现按系统区域自动改用 hf-mirror.com 公益镜像；可用 HF_ENDPOINT 环境变量自行覆盖",
        ],
        "title_en": "Distribution details added to the guide · VPN-free model download in mainland China",
        "notes_en": [
            "The user guide gained What Needs a Network Connection and Notes for Users in Mainland China sections, and now states that no Python, dependencies or conda are required, that Linux has no prebuilt binary yet, and that the GPU edition only suits NVIDIA machines",
            "The English guide gained a full System Requirements section (previously missing entirely), matching the Chinese version",
            "VPN-free Kokoro model download in mainland China: huggingface.co is not directly reachable there, so the app now falls back to the hf-mirror.com community mirror based on system locale; override it with the HF_ENDPOINT environment variable",
        ],
    },
    {
        "version": "2.14.1",
        "date": "2026-07-24",
        "title": "弹出窗宽度锚定悬停条 · 置顶失焦根治 · 对话框圆角滚动条",
        "notes": [
            "设置窗下拉弹出列表宽度改在弹出瞬间按最长项内容重设——锚定 mac 悬停蓝条的自然宽度：引擎约占闭合框 27%、语言 18%、样式 16%，Windows 与 mac 一致",
            "根治置顶导致的设置窗异常(v2.14.0 的两个 bug)：改 windowFlags 会销毁重建原生窗口——对对话框做会让它先缩小消失再出现，且 exec() 模态循环失效导致关闭后主界面按钮点了没反应；现在模态窗开着时只记录意图，等它关闭后再应用，全程不重建任何窗口",
            "顺带修掉一个隐藏错误：上版弹出宽度代码引用了尚未定义的变量，整段被异常吞掉从未生效",
            "Win10 圆角滚动条补齐到所有对话框：设置窗/历史窗/文档窗会自设样式表覆盖应用级规则，现在各自带上圆角滚动条样式",
            "继续翻译历史更新说明条目（累计 72/85 条已双语）",
        ],
        "title_en": "Popup width anchored to the hover bar · always-on-top focus fix · rounded scrollbars in dialogs",
        "notes_en": [
            "Settings dropdown popup width is now set at popup time from the longest item — anchored to the natural width of the macOS hover bar: the engine popup is about 27% of the closed box, language 18% and theme 16%, identical on Windows and macOS",
            "Root fix for the Settings dialog problems caused by always-on-top (two bugs from v2.14.0): changing windowFlags destroys and rebuilds the native window — doing it to the dialog made it shrink, vanish and reappear, and broke the exec() modal loop so main-window buttons stopped responding afterwards; the intent is now recorded while a modal dialog is open and applied once it closes, rebuilding no windows at all",
            "Also fixed a hidden error: last version's popup width code referenced a variable before it was defined, so the whole block was swallowed by the exception and never took effect",
            "Rounded scrollbars on Win10 extended to every dialog: Settings, History and document windows set their own stylesheets which override the application-level rule, so each now carries the rounded scrollbar styling itself",
            "Continued translating historical change log entries (72 of 85 now bilingual)",
        ],
    },
    {
        "version": "2.14.0",
        "date": "2026-07-24",
        "title": "下拉弹出窗收窄 · Win10 圆角滚动条 · Windows 置顶失焦修复",
        "notes": [
            "设置窗下拉弹出列表大幅收窄：不再跟随被表单拉伸的闭合框，改为只按最长项内容定宽——引擎约占闭合框 25%，语言约 17%，样式约 14%",
            "Win10 及以下滚动条改为圆角胶囊，与 Win11 / macOS 观感一致；样式提升到应用级，设置窗、历史窗、说明窗等所有对话框一并生效（此前仅主窗有，对话框仍是直角）",
            "修复 Windows 版勾选置顶时设置窗被压到主窗下面：Windows 的窗口层级变更由系统异步处理，同步归还焦点会被随后到达的置顶事件覆盖；改为立即归还后再延后两拍补两次，并让对话框跟随主窗一起置顶",
            "继续翻译历史更新说明条目（累计 58/84 条已双语）",
        ],
        "title_en": "Narrower dropdown popups · rounded scrollbars on Win10 · Windows always-on-top focus fix",
        "notes_en": [
            "Settings dropdown popups made much narrower: instead of following the form-stretched closed box they size to their longest item — the engine popup is about 25% of the closed box, language about 17% and theme about 14%",
            "Scrollbars on Win10 and below are now rounded capsules matching Win11 and macOS; the styling moved to application level so Settings, History, User Guide and every other dialog gets it too (previously only the main window did, leaving dialogs square)",
            "Fixed the Settings dialog being pushed behind the main window when enabling always-on-top on Windows: Windows processes stacking changes asynchronously, so a synchronous focus restore was overridden by the arriving topmost event; focus is now restored immediately and again on two deferred ticks, and the dialog follows the main window's topmost state",
            "Continued translating historical change log entries (58 of 84 now bilingual)",
        ],
    },
    {
        "version": "2.13.9",
        "date": "2026-07-23",
        "title": "提示文字随语言切换 · Windows 下拉弹出窗修复 · 置顶不再抢焦点",
        "notes": [
            "更改语言后，空的原文/译文区提示文字立即互换（通用化根治：遍历重译现在处理所有带 placeholder 的控件，不再只限单行输入框——原文/译文是多行编辑器，此前从未被覆盖）",
            "修复 Windows 版设置窗下拉弹出列表过窄：补齐与主界面同款的按最长项定宽逻辑，弹出宽度与闭合框相当",
            "修复 Windows 版语言/样式下拉弹出列表过高（二三项内容却有十来行）：可见项数=实际项数、按项滚动、关闭滚动条，与 macOS 表现一致",
            "勾选置顶时主界面不再跳到设置窗前面：改用不抢焦点的显示方式，并把层级与焦点归还给正开着的对话框",
            "引擎下拉闭合框再加宽 5 像素，弹出列表加宽 15 像素，其余位置不变",
            "设置窗『保持程序置顶』与日志按钮行之间增加间距",
            "继续翻译历史更新说明条目（累计 45/83 条已双语）",
        ],
        "title_en": "Placeholders follow the UI language · Windows dropdown popup fixes · always-on-top no longer steals focus",
        "notes_en": [
            "Placeholder text in the empty source and target areas now switches instantly with the UI language (generic root fix: the traversal retranslate handles every widget with a placeholder rather than single-line inputs only — the source and target areas are multi-line editors and had never been covered)",
            "Fixed the Settings dropdown popup being too narrow on Windows: it now uses the same longest-item width calculation as the main window, so the popup matches the closed box",
            "Fixed the language and theme dropdown popups being far too tall on Windows (about ten rows for two or three items): visible items now equal the actual item count, with per-item scrolling and scrollbars disabled, matching macOS",
            "Enabling always-on-top no longer makes the main window jump in front of the Settings dialog: the window is shown without stealing focus and the dialog gets its stacking order and focus back",
            "Engine dropdown closed box widened by another 5px and its popup by 15px, with everything else unchanged",
            "Added spacing between the Keep Window on Top checkbox and the log button row in Settings",
            "Continued translating historical change log entries (45 of 83 now bilingual)",
        ],
    },
    {
        "version": "2.13.8",
        "date": "2026-07-23",
        "title": "导出文件名统一 EC 前缀 · 导出文字钮空态灰化 · 文案与宽度微调",
        "notes": [
            "所有导出文件名统一加 EC 前缀：日志 EC LT、翻译历史 EC TH、原文 EC OT 语言、译文 EC TT 语言、翻译后文件 EC TT 语言 … T",
            "原文/译文区无文字时，导出文字按钮呈灰色不可点（与导出音频钮一致），有文字即恢复",
            "引擎下拉框恢复上一版宽度（+10 像素），其余位置不变",
            "英文文案改为标题式大小写：API Keys (Optional)、Multi-Style Translation、Keep Window on Top",
            "继续翻译历史更新说明条目（累计 33/82 条已双语）",
        ],
        "title_en": "Unified EC filename prefix · export buttons dim when empty · wording and width tweaks",
        "notes_en": [
            "All export filenames now carry the EC prefix: log EC LT, translation history EC TH, source EC OT <lang>, target EC TT <lang>, translated file EC TT <lang> ... T",
            "Export text buttons are grayed out when their area has no text (matching the export audio buttons) and re-enable as soon as text appears",
            "Engine dropdown restored to the previous width (+10px) with everything else unchanged",
            "English wording switched to title case: API Keys (Optional), Multi-Style Translation, Keep Window on Top",
            "Continued translating historical change log entries (33 of 82 now bilingual)",
        ],
    },
    {
        "version": "2.13.7",
        "date": "2026-07-22",
        "title": "单词分区 · 导出日志 · 窗口置顶 · Windows 脚本修复",
        "notes": [
            "多风格模式下输入单个词：直译区只保留一个最优译法，其余备选全部归入多风格区",
            "设置窗新增『导出日志』按钮（查看日志右侧、等宽同风格）：可选路径/文件名/格式，支持 .txt .log .md .json，纯标准库实现",
            "设置窗新增『保持程序置顶』复选框（多风格与日志行之间），勾选后窗口始终在其它程序之上，设置持久保存、启动自动应用",
            "修复 Windows 构建脚本乱码报错：脚本内中文注释在 GBK 控制台被误读，导致 setlocal/set 等语句失效、版本号提取失败（产物名丢失变成 .exe 与 -版本-windows.zip）；脚本改为纯 ASCII 并把 chcp 前置",
            "原文/译文语言下拉各再加宽 10 像素（原文向左、译文向右），引擎下拉相应减 10，总宽守恒、交换钮居中不变",
            "Windows 版朗读速度滑杆左右滑槽统一为同色，与 macOS 一致",
        ],
        "title_en": "Word partition · Export log · Always on top · Windows script fix",
        "notes_en": [
            "Single-word input in multi-style mode: the literal zone now keeps only one best translation, with all alternatives moved to the multi-style zone",
            "New Export Log button in Settings (right of View Log, same width and style): choose path, filename and format — .txt, .log, .md, .json — implemented with the standard library only",
            "New Keep window on top checkbox in Settings (between multi-style and the log row): keeps the window above other applications; the setting persists and is applied at startup",
            "Fixed Windows build script failures: Chinese comments inside the script were misread by GBK consoles, breaking setlocal/set statements and version extraction (producing nameless .exe and -version-windows.zip); scripts are now pure ASCII with chcp moved to the top",
            "Source/target language dropdowns widened by another 10px (source expands left, target right) with the engine dropdown narrowed by 10 — total width preserved, swap button stays centered",
            "Windows playback-speed slider grooves unified to the same color on both sides, matching macOS",
        ],
    },
    {
        "version": "2.13.6",
        "date": "2026-07-21",
        "title": "极简交换钮结构性居中 + 字幕/音频/文字生命周期铁律",
        "notes": [
            "极简界面交换钮居中根治：右组不再整体隐藏，改为隐藏其内容并在最右端放与极简钮完全镜像的隐形占位——左右结构对称，居中由布局数学保证，零校准、平台无关",
            "字幕铁律：卡拉OK捆绑音频、依附文字，两者都在拖动进度条必有字幕（边界丢失自动从缓存恢复或按朗读范围重建）；任一不在必无字幕",
            "文字清空钮：清文字同时该侧字幕（边界+高亮）同亡，音频缓存保留可继续播放",
            "音频清空/文字变化作废音频时：该侧内存边界与卡拉OK高亮同步清除，杜绝旧字幕残留",
        ],
        "title_en": "Structural swap centering in minimal mode + karaoke/audio/text lifecycle rules",
        "notes_en": [
            "Minimal-mode swap centering solved structurally: the right group's contents hide while an invisible mirror spacer sits at the far right, so centering is guaranteed by layout symmetry — zero calibration, platform-independent",
            "Karaoke iron rule: subtitles are bound to audio and attached to text — with both present, dragging the progress slider always shows karaoke (lost boundaries auto-restore from cache or rebuild from the spoken span); with either missing, karaoke never shows",
            "Text clear buttons: clearing text also kills that side's karaoke (boundaries + highlight) while the audio cache remains playable",
            "Audio clear / text-change invalidation: the side's in-memory boundaries and karaoke highlight are cleared together, eliminating stale-subtitle leftovers",
        ],
    },
    {
        "version": "2.13.5",
        "date": "2026-07-20",
        "title": "首次朗读无字幕修复 + 卡拉OK绑定朗读范围",
        "notes": [
            "修复首次点朗读经常无卡拉OK：估算字幕需要音频时长，首播时时长常未加载导致估算被静默放弃且不重试；现在自动重试等待直到时长就绪",
            "卡拉OK范围永远与朗读范围一一对应：朗读启动时记录实际朗读范围，估算字幕严格在该范围内铺设",
        ],
        "title_en": "First-play karaoke fix + karaoke bound to spoken range",
        "notes_en": [
            "Fixed karaoke often missing on first play: subtitle estimation needs audio duration, which is often not yet loaded on first playback — the estimator gave up silently without retrying; it now retries until duration is ready",
            "Karaoke range always maps one-to-one to the spoken range: the actual spoken span is recorded at TTS start and estimated subtitles are laid strictly within it",
        ],
    },
    {
        "version": "2.13.4",
        "date": "2026-07-19",
        "title": "灰色多风格区卡拉OK修复",
        "notes": [
            "修复多风格灰色区无卡拉OK：灰字格式此前整体覆盖背景层，重写高亮分层——灰区内的蓝色选区/灰色联动/绿色卡拉OK背景照常显示",
            "极简界面交换钮位置微调",
        ],
        "title_en": "Karaoke fix in the gray multi-style zone",
        "notes_en": [
            "Fixed missing karaoke in the gray multi-style zone: the gray-text format used to overwrite background layers; highlighting is now layered so blue selection, gray link and green karaoke backgrounds all render inside the gray zone",
            "Minor position tweak for the minimal-mode swap button",
        ],
    },
    {
        "version": "2.13.3",
        "date": "2026-07-18",
        "title": "打包版本号自动化 + 多风格朗读分区",
        "notes": [
            "修复打包产物版本号错误：三个构建脚本的版本改为自动从 APP_VERSION 提取，单一事实来源",
            "多风格翻译朗读分区：无选区时点朗读译文只读直译区、卡拉OK只覆盖直译区；有选区仍按选区朗读",
        ],
        "title_en": "Build version automation + multi-style TTS scoping",
        "notes_en": [
            "Fixed wrong packaged version numbers: all three build scripts now auto-extract the version from APP_VERSION — a single source of truth",
            "Multi-style TTS scoping: with no selection, reading the translation speaks only the literal section and karaoke covers it alone; with a selection the selection is read as before",
        ],
    },
    {
        "version": "2.13.2",
        "date": "2026-07-17",
        "title": "使用说明命令行文字样式统一",
        "notes": [
            "使用说明中去除隔离命令改为普通正文样式（去掉代码块底色与等宽字体），与其它文字一致",
        ],
        "title_en": "User guide command text style unified",
        "notes_en": [
            "The quarantine-removal command in the user guide now uses plain body text style (code background and monospace font removed) to match surrounding text",
        ],
    },
    {
        "version": "2.13.1",
        "date": "2026-07-16",
        "title": "跟随系统昼夜切换修复 + 分割条居中吸附",
        "notes": [
            "修复『跟随系统』在系统昼夜切换时样式不完全跟变：Qt 6.4 没有对应信号导致监听从未挂上，macOS 改用轮询兜底",
            "原文/译文分割条新增居中吸附：拖到中央附近自动精准均分左右两栏",
        ],
        "title_en": "Follow-system day/night fix + splitter center snap",
        "notes_en": [
            "Fixed Follow System not fully switching on system day/night change: Qt 6.4 lacks the corresponding signal so the listener never attached; macOS now uses a polling fallback",
            "Source/target splitter now snaps to exact center when dragged near the middle",
        ],
    },
    {
        "version": "2.13.0",
        "date": "2026-07-15",
        "title": "界面精修里程碑：交换钮居中收官 + 三处细节",
        "notes": [
            "原文/译文语言下拉框各加宽 10 像素，文字完整显示，交换钮居中不受影响",
            "英文关于窗标题下补回一行空行",
            "翻译历史窗背景跟随深浅主题：浅色模式下立即使用浅色背景",
        ],
        "title_en": "UI polish milestone: swap centering finale + three details",
        "notes_en": [
            "Source/target language dropdowns widened by 10px each for full text display; the swap button stays centered",
            "Restored a blank line under the English About title",
            "Translation history window background now follows the light/dark theme instantly",
        ],
    },
    {
        "version": "2.12.29",
        "date": "2026-07-08",
        "title": "交换钮精确居中(正常+极简)",
        "notes": [
            "交换钮用右侧对称占位精确居中，正常与极简模式偏差均≈0",
            "极简模式保留交换钮并居中",
        ],
        "title_en": "Swap button precisely centered (normal + minimal)",
        "notes_en": [
            "Swap button precisely centered via a symmetric right-side spacer; near-zero offset in both normal and minimal modes",
            "Minimal mode keeps the swap button centered",
        ],
    },
    {
        "version": "2.12.28",
        "date": "2026-07-08",
        "title": "交换钮真正居中 + 极简模式保留交换钮",
        "notes": [
            "交换钮改用右侧对称占位实现精确居中(不再靠猜margin)，正常/极简模式偏差均≈0",
            "极简模式保留交换钮并始终居中",
        ],
        "title_en": "Swap button truly centered + kept in minimal mode",
        "notes_en": [
            "Swap button now centered via a symmetric right-side spacer (no more guessed margins); near-zero offset in both normal and minimal modes",
            "Minimal mode keeps the swap button, always centered",
        ],
    },
    {
        "version": "2.12.27",
        "date": "2026-07-08",
        "title": "界面细节微调（引擎框/交换钮/滑杆圆球/文档英文化）",
        "notes": [
            "英文环境设置窗『Google 云翻译 Key』显示为 Google Cloud Key",
            "主界面最小宽度缩到880",
            "交换钮组左移微调，趋向居中",
            "两条朗读进度条与速度滑杆的圆球改回正圆、亮白色（去掉描边避免椭圆、灰色改白）",
            "关于窗/更新说明窗标题英文化：About English Coach / Change Log",
            "英文关于窗删除拼音副标题 Ying Yu Dao Shi",
            "更新说明自此支持中英双语（按界面语言切换）",
        ],
        "title_en": "UI detail tweaks (engine box / swap button / slider knob / doc localization)",
        "notes_en": [
            "Settings 'Google Cloud Key' label now shows in English",
            "Main window minimum width reduced to 880",
            "Swap button group nudged left toward center",
            "Playback progress bars and speed slider knobs restored to true circles in bright white (removed border that caused ovals, gray changed to white)",
            "About / Change Log dialog titles localized: About English Coach / Change Log",
            "Removed the pinyin subtitle 'Ying Yu Dao Shi' from the English About page",
            "Change Log now supports bilingual display (follows UI language)",
        ],
    },
    {
        "version": "2.12.10",
        "date": "2026-07-08",
        "title": "删除设置窗多风格说明文字",
        "notes": [
            "删除设置窗『（主译文 + 书面/口语/俚语/美英式等辅助译法）』说明文字",
        ],
        "title_en": "Removed multi-style description text from Settings",
        "notes_en": [
            "Removed the '(main translation + formal/casual/slang/US-UK style variants)' description text from Settings",
        ],
    },
    {
        "version": "2.12.9",
        "date": "2026-07-08",
        "title": "删除设置窗两段引擎说明文字",
        "notes": [
            "删除设置窗顶部『Google 免费、无需 Key…』和底部『提示：Google 免费无需 Key…』两段说明文字",
        ],
        "title_en": "Removed two engine description paragraphs from Settings",
        "notes_en": [
            "Removed the top 'Google is free, no key required...' and bottom 'Tip: Google is free...' description paragraphs from Settings",
        ],
    },
    {
        "version": "2.12.8",
        "date": "2026-07-08",
        "title": "5处修复(含语言切换根治+下拉收口)",
        "notes": [
            "主界面最小宽度加回到970",
            "下拉宽度收口：清理设置窗输入框重复的sizePolicy叠加，主界面下拉恢复2.12.4干净公式(最长内容+52)",
            "语言切换根治：改为遍历整个窗口所有控件按文字查表双向替换(中英)，覆盖按钮/气球提示/占位符/下拉项/两段说明，不再逐条点名遗漏；处理前后空格",
            "设置窗『关闭』钮改蓝色主按钮样式(与关于窗一致)",
            "修复极简界面往返后主窗最小宽度被改成770的bug(现与主窗970一致)",
        ],
        "title_en": "Five fixes (language switching root fix + dropdown width cleanup)",
        "notes_en": [
            "Main window minimum width restored to 970",
            "Dropdown width cleanup: removed duplicated sizePolicy stacking on Settings inputs; main window dropdowns restored to the clean 2.12.4 formula (longest content + 52)",
            "Language switching root fix: now traverses every widget in the window and swaps text via dictionary lookup in both directions (CN/EN), covering buttons, tooltips, placeholders, dropdown items and description paragraphs — no more one-by-one omissions; leading/trailing spaces preserved",
            "Settings Close button restyled as a blue primary button, matching the About dialog",
            "Fixed a bug where the main window minimum width became 770 after toggling minimal mode (now consistent at 970)",
        ],
    },
    {
        "version": "2.12.7",
        "date": "2026-07-08",
        "title": "7处界面细节微调",
        "notes": [
            "主界面最小宽度减到920",
            "设置窗下拉恢复正常宽度(min 200)并随窗宽自适应",
            "语言切换即时生效扩展到设置窗内部(标题/关闭/多风格/两段说明即时重译)",
            "设置窗改为即时保存：删除保存按钮，只留『关闭』(Close)，所有设置项自动即时保存生效",
            "显示密钥按钮：文字改『显示密钥』(Show Key)、宽度与查看日志一致(BTN_W)",
            "极简钮图标改为上下横杠+中间方块(SVG，随主题深浅着色)",
            "英文模式设置窗两段引擎说明改为指定英文文案",
        ],
        "title_en": "Seven UI detail tweaks",
        "notes_en": [
            "Main window minimum width reduced to 920",
            "Settings dropdowns restored to normal width (min 200) and now adapt to window width",
            "Instant language switching extended to the Settings dialog (title, Close, multi-style and both description paragraphs retranslate immediately)",
            "Settings now saves instantly: the Save button was removed leaving only Close; every setting persists and takes effect immediately",
            "Show Key button: text changed to 'Show Key', width matched to the View Log button",
            "Minimal button icon redesigned as top/bottom bars with a center square (SVG, recolored with the light/dark theme)",
            "Specified English wording for the two engine description paragraphs in English mode",
        ],
    },
    {
        "version": "2.12.6",
        "date": "2026-07-08",
        "title": "9处界面细节微调",
        "notes": [
            "主界面最小宽度减到970",
            "设置窗下拉恢复正常宽度、并随窗宽拖动自适应",
            "设置窗所有API Key输入框随窗宽自适应变宽",
            "语言/样式下拉随窗宽自适应",
            "显示密钥按钮：文字由『显示API-Key』改『显示密钥』(Show Key)、宽度与查看日志一致",
            "语言切换更彻底：朗读嗓音下拉、更多tooltip/占位符即时重译",
            "修复极简钮按下后消失：极简钮独立于左组，切换后始终可见；图标恢复▣",
            "原文/译文语言下拉框加宽20%，字完整显示",
            "设置窗两段引擎说明补齐英文翻译",
        ],
        "title_en": "Nine UI detail tweaks",
        "notes_en": [
            "Main window minimum width reduced to 970",
            "Settings dropdowns restored to normal width and adapt as the window is resized",
            "All API key inputs in Settings widen with the window",
            "Language and theme dropdowns adapt to window width",
            "Show Key button: renamed from 'Show API Key' to 'Show Key', width matched to the View Log button",
            "More thorough language switching: TTS voice dropdown and additional tooltips/placeholders retranslate immediately",
            "Fixed the minimal button disappearing after being pressed: it is now independent of the left group and always visible",
            "Source/target language dropdowns widened by 20% for full text display",
        ],
    },
    {
        "version": "2.12.5",
        "date": "2026-07-08",
        "title": "7处界面细节微调",
        "notes": [
            "正方形按钮圆角：正常态8px，青色按下态放大到10px",
            "下拉闭合框高度精确对齐正方形按钮(36)，不再偏高",
            "主界面最小宽度净调整到1020",
            "设置窗API Key输入框/语言/样式下拉随窗宽自适应；显示API-Key按钮改回固定宽但文字完整",
            "语言切换即时生效(下拉一选界面文字立即切换，无需保存或重启)",
            "样式风格切换即时生效(承前版)",
            "交换钮真居中(极简钮并入左组，左右对称)，同时设置/关于那排贴到最右",
        ],
        "title_en": "Seven UI detail tweaks",
        "notes_en": [
            "Square button corners: 8px normally, enlarged to 10px in the cyan pressed state",
            "Dropdown closed-box height aligned exactly with square buttons (36), no longer taller",
            "Main window minimum width adjusted to 1020",
            "Settings API key inputs and language/theme dropdowns adapt to window width; the Show API Key button returns to a fixed width with full text",
            "Instant language switching (selecting in the dropdown switches UI text immediately, no save or restart needed)",
            "Instant theme switching (carried over from the previous version)",
            "Swap button truly centered (minimal button merged into the left group for symmetry) while the settings/about row sits flush right",
        ],
    },
    {
        "version": "2.12.4",
        "date": "2026-07-08",
        "title": "12处界面细节微调(基于v25完美版)",
        "notes": [
            "正方形按钮圆角加大到10px(与青色按下态一致)；下拉闭合框圆角8px、高36与按钮等高",
            "载入下一条改用独立右箭头图标(redo)，主题切换不再丢失方向",
            "朗读速度滑杆左右滑槽与圆球全灰(不蓝)；朗读进度条保持左侧蓝",
            "载入上一条/下一条之间分组缝隙去掉，统一小缝",
            "主界面最小宽度加到970，翻译钮居中；交换钮加右侧等宽占位真居中",
            "左上角极简钮图标换实心方块(线条更清晰)",
            "翻译引擎/原文/译文下拉框各加宽10px",
            "设置窗API Key输入框改圆角长条、随窗宽自适应；语言/样式下拉同样跟随窗宽",
            "显示API-Key按钮文字完整显示(自适应宽+padding)",
            "样式风格下拉即时生效(选中即切换，无需保存或重启)",
        ],
        "title_en": "Twelve UI detail tweaks (based on the v25 reference build)",
        "notes_en": [
            "Square button corners enlarged to 10px (matching the cyan pressed state); dropdown closed box uses 8px corners at height 36, level with the buttons",
            "Load Next now uses a dedicated right-arrow (redo) icon so direction is preserved across theme switches",
            "Playback-speed slider grooves and knob fully gray (not blue); the playback progress bar keeps its blue left side",
            "Removed the grouping gap between Load Previous and Load Next in favor of a uniform small gap",
            "Main window minimum width raised to 970 with the translate button centered; the swap button gets an equal-width right spacer for true centering",
            "Minimal button icon in the top-left changed to a solid square for sharper lines",
            "Engine, source and target dropdowns each widened by 10px",
            "Settings API key inputs restyled as rounded bars that adapt to window width; language and theme dropdowns follow the window width too",
            "Show API Key button text fully visible (adaptive width plus padding)",
        ],
    },
    {
        "version": "2.12.3",
        "date": "2026-07-08",
        "title": "基于v25完美版微调8处界面细节",
        "notes": [
            "正方形按钮恢复正方形(36x36)，圆角与青色按下态一致(8px)",
            "浅色按钮图标/文字统一深黑、深色统一浅白：图标名登记到按钮，主题切换遍历重着色(不再漏清空/复制/粘贴等)",
            "朗读速度滑杆左侧滑槽恢复灰色(不再蓝色)",
            "设置/更新/帮助/关于那排贴最右侧(去掉平衡占位)",
            "复制粘贴等按钮排统一小缝隙不分组，原文贴左、译文贴右",
            "主界面最小宽度加宽~100(770->870)，保证翻译钮居中",
        ],
        "title_en": "Eight UI detail tweaks based on the v25 reference build",
        "notes_en": [
            "Square buttons restored to a true square (36x36) with corners matching the cyan pressed state (8px)",
            "Light-theme button icons and text unified to deep black and dark-theme to off-white: icon names are registered on each button so theme switching recolors them by traversal (no longer missing Clear, Copy, Paste and others)",
            "Playback-speed slider left groove restored to gray (no longer blue)",
            "The Settings/Update/Help/About row sits flush right (balance spacer removed)",
            "Copy, paste and related button rows use a uniform small gap without grouping; source flush left, target flush right",
            "Main window minimum width increased by about 100 (770 to 870) to keep the translate button centered",
        ],
    },
    {
        "version": "2.12.2",
        "date": "2026-07-08",
        "title": "主题终极方案落地：原生+绘制精细混合，深浅完美",
        "notes": [
            "Mac 最终混合方案(真机验证成功)正式落地主程序：",
            "下拉闭合框+方按钮+普通按钮+特殊按钮=绘制(深浅两套)；下拉弹出项+气球+滚动条=系统原生",
            "下拉箭头用V形SVG图片(不再是方块)；弹出项悬停走系统原生蓝条",
            "按钮淡蓝按下反馈；特殊checkable按钮按下保持青色；翻译钮亮蓝",
            "深浅切换由pyobjc(AppKit)驱动，切换时重画绘制部分+重生成图标，即时生效",
        ],
        "title_en": "Final theming solution: native plus drawn hybrid, perfect in light and dark",
        "notes_en": [
            "The macOS hybrid solution (verified on a real machine) is now in the main program",
            "Drawn: dropdown closed box, square buttons, regular buttons and special buttons (two color sets); native: dropdown popup items, tooltips and scrollbars",
            "Dropdown arrow uses a chevron SVG image (no longer a square block); popup items use the native blue hover bar",
            "Light blue press feedback on buttons; checkable special buttons stay cyan while active; the translate button stays bright blue",
            "Light/dark switching is driven by pyobjc (AppKit); drawn parts are repainted and icons regenerated on switch, taking effect immediately",
        ],
    },
    {
        "version": "2.12.1",
        "date": "2026-07-08",
        "title": "Mac原生+pyobjc深浅 最终方案 + 正方形按钮圆角(QSS配合)",
        "notes": [
            "确立最终方案：Mac 控件走系统原生 + pyobjc(AppKit)切深浅/浅/跟随，深浅都完美且即时切换",
            "正方形/工具按钮：加 objectName=toolbtn + border-radius 圆角QSS(只加圆角与hover蓝框，不设颜色，保留原生深浅跟随)",
            "翻译大钮 primary、状态栏、激活区蓝框保持设计标识",
            "非Mac(Windows/Linux)不受影响",
        ],
        "title_en": "macOS native + pyobjc theming, final approach, with square button corners via QSS",
        "notes_en": [
            "Final approach settled: macOS controls use the native appearance with pyobjc (AppKit) driving dark/light/follow-system — perfect in both modes with instant switching",
            "Square and tool buttons: objectName=toolbtn plus border-radius QSS (corners and hover outline only, no colors, preserving native light/dark following)",
            "The primary translate button, status bar and active-area blue outline keep their design identity",
            "Non-macOS platforms (Windows/Linux) are unaffected",
        ],
    },
    {
        "version": "2.12.0",
        "date": "2026-07-08",
        "title": "主题系统重构：Mac 原生深浅(AppKit驱动)，彻底终结深浅打架",
        "notes": [
            "根本方案：Mac 用 pyobjc(AppKit) 的 NSApplication.setAppearance_ 驱动系统原生深浅——这是 Qt6.4.2+BigSur 上唯一可靠方案(setColorScheme 是6.8+才有，本机装不了)",
            "Mac 上不再自涂 QSS 深色皮肤与调色板：所有原生控件(下拉/气球/按钮/滚动条/标题栏)由系统原生深浅驱动，自动跟随，杜绝之前十几处深浅细节bug的总根源(自涂与原生打架)",
            "深色/浅色/跟随系统三选项即时生效、无需重启；跟随系统随昼夜自动变",
            "主题切换后按钮图标按新深浅重新渲染",
            "非 Mac(Windows/Linux)保持原有 QSS+调色板皮肤不变",
            "默认主题改为跟随系统",
            "注意：Mac 需安装 pyobjc-framework-Cocoa(下版加入 requirements 与打包脚本)",
        ],
        "title_en": "Theme system rebuilt: native macOS light/dark driven by AppKit",
        "notes_en": [
            "Core approach: macOS uses NSApplication.setAppearance_ via pyobjc (AppKit) — the only reliable option on Qt 6.4.2 with Big Sur (setColorScheme requires 6.8+, which cannot be installed on this machine)",
            "macOS no longer paints its own dark QSS skin or palette: all native controls (dropdowns, tooltips, buttons, scrollbars, title bar) follow the system appearance, eliminating the root cause of a dozen previous light/dark bugs (self-painting fighting the native appearance)",
            "Dark, Light and Follow System all take effect instantly without restart; Follow System changes automatically with day and night",
            "Button icons are re-rendered for the new appearance after a theme switch",
            "Non-macOS platforms keep the existing QSS and palette skin unchanged",
            "Default theme changed to Follow System",
            "Note: macOS requires pyobjc-framework-Cocoa (to be added to requirements and the build scripts next version)",
        ],
    },
    {
        "version": "2.11.3",
        "date": "2026-07-07",
        "title": "主界面下拉灰底块/气球方框/浅色按钮图标 三处修复",
        "notes": [
            "主界面下拉残留灰背景块：mac 下 listview 与 popup 设透明背景(WA_TranslucentBackground)，彻底走系统原生",
            "气球提示仍是方块：mac 下调色板跳过 ToolTipBase/ToolTipText，交给系统原生尖角圆角气球(apply_theme 与 main 启动两处都改)",
            "浅色模式按钮图标仍白色：根因 Icons.icon 固定 #e8e8e8 浅色绘制，改为按当前主题动态取色(浅色用深色#1f1f22)",
        ],
        "title_en": "Three fixes: dropdown gray block, tooltip square corners, light-theme button icons",
        "notes_en": [
            "Residual gray background block behind main window dropdowns: on macOS the listview and popup are set to a translucent background so they fully use the native appearance",
            "Tooltips still rendering as squares: on macOS the palette now skips ToolTipBase/ToolTipText, leaving the native rounded tooltip in place (fixed in both apply_theme and startup)",
            "Button icons still white in light mode: the root cause was Icons.icon hardcoding the light color #e8e8e8; it now picks the color from the current theme (deep #1f1f22 in light mode)",
        ],
    },
    {
        "version": "2.11.2",
        "date": "2026-07-07",
        "title": "找到下拉方框真凶(自定义combo工厂第1734行)",
        "notes": [
            "根因：下拉是自定义combo(setView+自定义popup)，方框来自工厂函数里 popup.setStyleSheet(border:1px)——此前所有版本都在改无关的 QComboBox QAbstractItemView，从未碰到这行",
            "修复：mac 下该 popup 与 listview 一律清空样式(setStyleSheet(''))，回归系统原生圆角无框下拉；其它平台保持深色边框",
            "tooltip 已确保 mac 下无自定义 QSS",
        ],
        "title_en": "Found the real cause of dropdown square borders (custom combo factory)",
        "notes_en": [
            "Root cause: dropdowns are custom combos (setView plus a custom popup) and the square border came from popup.setStyleSheet(border:1px) inside the factory function — every previous version had been editing the unrelated QComboBox QAbstractItemView rule and never touched this line",
            "Fix: on macOS the popup and listview styles are cleared entirely, returning to the native borderless rounded dropdown; other platforms keep the dark border",
            "Tooltips confirmed free of custom QSS on macOS",
        ],
    },
    {
        "version": "2.11.1",
        "date": "2026-07-07",
        "title": "彻底修复Mac下拉/气球仍有方框(收口全部散落样式定义)",
        "notes": [
            "根因：combo/tooltip 样式散布在主窗/设置窗/历史窗共4处，此前只改了1-2处，其余写死样式在Mac上照样生效画出方框",
            "修复①被错位替换破坏的 _tooltip_css() 函数(内嵌了无效占位符)",
            "修复②主窗 QToolTip 实为写死样式(未走占位符)，改为 %TOOLTIP% 占位并注入",
            "修复③历史窗写死的 QToolTip 改为平台函数注入",
            "四处 combo/tooltip 定义全部收口到 _combo_popup_css()/_tooltip_css()，Mac一律返回空走系统原生(圆角无框下拉、尖角气球)",
        ],
        "title_en": "Fully fixed remaining square borders on macOS dropdowns and tooltips",
        "notes_en": [
            "Root cause: combo and tooltip styles were scattered across four places (main window, Settings, History); earlier fixes only touched one or two, so the remaining hardcoded styles still drew square borders on macOS",
            "Fix 1: repaired the _tooltip_css() function that had been broken by a misplaced replacement (it contained a dead placeholder)",
            "Fix 2: the main window QToolTip was hardcoded rather than using the placeholder; it now uses %TOOLTIP% with proper injection",
            "Fix 3: the hardcoded QToolTip in the History window now uses the platform function",
            "All four combo/tooltip definitions consolidated into _combo_popup_css() and _tooltip_css(), which return empty on macOS so the native appearance is used (borderless rounded dropdowns, pointed tooltips)",
        ],
    },
    {
        "version": "2.11.0",
        "date": "2026-07-06",
        "title": "Mac原生下拉/气球 · 多风格标注净化 · 关于使用说明双语 · 极简状态记忆",
        "notes": [
            "Mac 下拉列表与气球提示改用系统原生(圆角无框下拉、尖角tooltip、深浅自适应)，复刻滚动条方案；Windows保持自定义样式",
            "设置窗 Show API Key 按钮改回固定短宽、文字完整；API Key 输入框随窗口放大而变宽",
            "修复浅色下点查看日志误触多风格复选框(查看日志改独立行容器)",
            "翻译历史『重新载入』英文模式显示 Reload(补 L()与词条)",
            "多风格直译区标注强净化：代码层强删 Part 1/第一部分/直译区/----/【..】等，不依赖模型",
            "朗读范围与卡拉OK字幕范围一致(无选区读全文含多风格区，字幕同步全覆盖)",
            "关于/使用说明 中英双语(按界面语言切换)；更新说明标题双语框架就位(历史条目英文下版补齐)",
            "极简界面状态记忆：重启恢复；API-Key 永远默认隐藏",
        ],
        "title_en": "Native macOS dropdowns/tooltips · multi-style label cleanup · bilingual docs · minimal state memory",
        "notes_en": [
            "macOS dropdowns and tooltips now use the native appearance (borderless rounded lists, pointed tooltips, automatic light/dark), mirroring the scrollbar approach; Windows keeps the custom style",
            "Settings Show API Key button restored to a fixed short width with full text; API key inputs widen with the window",
            "Fixed clicking View Log accidentally toggling the multi-style checkbox in light mode (View Log moved to its own row container)",
            "Translation history Reload button now shows 'Reload' in English mode (L() and dictionary entry added)",
            "Strong cleanup of multi-style literal-zone labels: Part 1, literal-zone markers, ---- and bracketed tags are stripped in code rather than relying on the model",
            "Spoken range and karaoke range kept consistent (with no selection the full text including the multi-style zone is read and subtitles cover it all)",
            "About and User Guide are bilingual (following the UI language); the bilingual framework for the Change Log title is in place (historical entries to be translated next)",
            "Minimal mode state is remembered across restarts; API keys always default to hidden",
        ],
    },
    {
        "version": "2.10.0",
        "date": "2026-07-06",
        "title": "主题切换彻底化 · 两侧字幕独立 · 去边框 · 英文补全 · 多风格空行分隔",
        "notes": [
            "深浅主题切换彻底：热切换时刷新标题栏、下拉弹出列表、已打开设置窗，不再残留深色",
            "浅色模式：所有青色提亮加饱和(#00b3c6)，与蓝色一致；按钮字/图标转黑",
            "修复浅色下点查看日志误触多风格复选框（复选框背景透明化）",
            "下拉列表项外框与气球提示外框隐形（边框设为与背景同色）",
            "朗读两侧独立：拖动某侧进度条只刷新该侧卡拉OK，不再串到另一侧",
            "英文补全：占位符改正常句式(Type or paste…)、历史提示、主界面引擎名(ERNIE/Doubao/Qwen/Hunyuan)、晓贝(Xiaobei)、Key提示段落、Reload",
            "设置窗 Show API Key 按钮文字完整显示；API Key 输入框随窗口放大而变宽",
            "多风格：直译区与多风格区空行分隔，去掉【直译区】【多风格区】及---标注",
        ],
        "title_en": "Thorough theme switching · independent per-side subtitles · borderless · English completion · blank-line separator",
        "notes_en": [
            "Theme switching is now thorough: hot switching refreshes the title bar, dropdown popups and any open Settings dialog, leaving no dark remnants",
            "Light mode: all cyan accents brightened and saturated (#00b3c6) to match the blue; button text and icons turn black",
            "Fixed clicking View Log accidentally toggling the multi-style checkbox in light mode (checkbox background made transparent)",
            "Dropdown item outlines and tooltip outlines made invisible (border color matched to the background)",
            "The two playback sides are independent: dragging one side's progress bar only refreshes that side's karaoke",
            "English completion: placeholders reworded to natural sentences (Type or paste...), history hints, main window engine names (ERNIE/Doubao/Qwen/Hunyuan), Xiaobei, key hint paragraphs and Reload",
            "Settings Show API Key button text fully visible; API key inputs widen with the window",
            "Multi-style: literal and multi-style zones separated by a blank line, with bracketed zone labels and --- markers removed",
        ],
    },
    {
        "version": "2.9.0",
        "date": "2026-07-06",
        "title": "英文全覆盖(引擎/嗓音/Key/文档) · 浅色按钮字色+高饱和蓝 · 状态全记忆 · 多风格空行分隔",
        "notes": [
            "English US 全覆盖：翻译钮 Translate、文心一言 ERNIE、豆包 Doubao、通义千问 Qwen、混元 Hunyuan、晓贝 Xiaobei、Reload；设置窗 Key 标签(Baidu AI Studio / Volcengine / Alibaba Cloud Model Studio / Tencent Cloud Hunyuan)、提示段落、更新说明/使用说明/关于文档均英文",
            "浅色模式：按钮文字与图标转黑；蓝色翻译钮与所有关闭钮提亮到状态栏同款高饱和蓝(#1e88e5)",
            "设置窗 Show API Key 按钮宽度与查看日志等统一(BTN_W)",
            "语言与主题真正即时生效不重启；跟随系统随昼夜自动切换",
            "朗读两排组件统一窄间距紧挨、译文排靠右去空位",
            "拖动进度滑杆时卡拉OK字幕实时跟随(无关选区)",
            "多风格：直译区与多风格区改用空行分隔(去掉----标注)；朗读范围与卡拉OK范围一致",
            "交换钮只交换直译区，多风格灰字区不参与",
            "全状态记忆：引擎/源语言/目标语言/主题/语言/朗读语速/嗓音/多风格开关，重启后恢复",
        ],
        "title_en": "Full English coverage · light-mode button colors · state memory · blank-line separator",
        "notes_en": [
            "Full English US coverage: Translate button, ERNIE, Doubao, Qwen, Hunyuan, Xiaobei, Reload; Settings key labels (Baidu AI Studio / Volcengine / Alibaba Cloud Model Studio / Tencent Cloud Hunyuan), hint paragraphs, and the Change Log, User Guide and About documents",
            "Light mode: button text and icons turn black; the blue translate button and all close buttons brightened to the saturated blue used by the status bar (#1e88e5)",
            "Settings Show API Key button width unified with View Log and others (BTN_W)",
            "Language and theme now take effect immediately without restart; Follow System changes automatically with day and night",
            "The two playback rows use consistent tight spacing with the target row flush right",
            "Karaoke subtitles follow in real time while dragging the progress slider (regardless of selection)",
            "Multi-style: literal and multi-style zones separated by a blank line (---- markers removed); spoken range matches karaoke range",
            "The swap button only swaps the literal zone; the gray multi-style zone is excluded",
        ],
    },
    {
        "version": "2.8.0",
        "date": "2026-07-06",
        "title": "英文全覆盖(含下拉/弹窗)+首字母大写 · 主题立即生效 · 多风格直译区分隔",
        "notes": [
            "English US 模式全面覆盖：下拉列表选项、设置/关于/历史弹窗内所有文字均英文，且统一首字母大写(如 View History)",
            "浅色模式补漏：标题栏随主题变浅、按钮悬停与下拉弹出列表统一浅色配色",
            "语言与深浅主题改为立即生效：切主题热切换、切语言自动重启；跟随系统在系统昼夜切换时自动跟变",
            "设置窗『查看日志』上移到多风格行下方；历史弹窗按钮英文不再截断，『载入』改为 Reload",
            "极简界面图标换为更饱满的 ▣；极简最小窗口高度再压缩至 200",
            "顶部交换钮真正居中(与翻译钮对齐)；朗读两排组件左右边距与上方对齐",
            "已有音频缓存时再次点朗读，无论是否有选区都恢复卡拉OK字幕",
            "多风格翻译分区：上半为逐行直译(黑/白字，参与原文↔译文选区联动)，---- 分隔线下为多风格区(灰字，不参与联动)",
        ],
        "title_en": "Full English coverage with title case · instant theme switching · multi-style zone separator",
        "notes_en": [
            "Comprehensive English US coverage: dropdown items and all text inside the Settings, About and History dialogs are in English with consistent title case (e.g. View History)",
            "Light mode gaps closed: the title bar lightens with the theme, and button hover states plus dropdown popups use unified light colors",
            "Language and theme now take effect immediately: theme switches hot, language triggers an automatic restart; Follow System changes with the system's day/night switch",
            "Settings View Log moved below the multi-style row; History dialog buttons no longer truncate in English and Load was renamed Reload",
            "Minimal mode icon changed to a fuller square; minimal window minimum height reduced to 200",
            "The top swap button is truly centered (aligned with the translate button); the two playback rows share the same left and right margins as the rows above",
            "With an audio cache present, pressing play again restores karaoke subtitles whether or not there is a selection",
            "Multi-style zoning: the upper part is the line-by-line literal translation (black/white text, participating in source-target selection linking); below the ---- separator is the multi-style zone (gray text, excluded from linking)",
        ],
    },
    {
        "version": "2.7.0",
        "date": "2026-07-05",
        "title": "界面语言中英切换 · 浅色/深色/跟随系统主题 · 设置窗改版",
        "notes": [
            "设置新增『语言』：中文 / English US，选英文后全部界面文字、气球提示、弹窗、状态栏均切换为英文（重启后生效）",
            "设置新增『样式风格』：浅色 / 深色 / 跟随系统（跟随系统自动检测系统深浅模式，重启后生效）",
            "设置窗改版：显示API-Key按钮移到Key输入区与多风格选项之间、与输入框左对齐，按下青色显示、弹起灰色隐藏；语言与样式风格在其下方；查看日志移至窗口底部左侧",
            "极简界面按钮图标改为更简洁的 ▢",
            "关于页移除电话号码，仅保留网址与邮箱（隐私保护）",
        ],
        "title_en": "UI language switching · light/dark/follow-system themes · Settings redesign",
        "notes_en": [
            "New Language setting: Chinese or English US — choosing English switches all UI text, tooltips, dialogs and the status bar to English (applies after restart)",
            "New Theme setting: Light, Dark or Follow System (Follow System detects the OS appearance; applies after restart)",
            "Settings redesign: the Show API Key button sits between the key inputs and the multi-style option, left-aligned with the inputs, cyan when pressed and gray when hidden; language and theme sit below it; View Log moved to the bottom left",
            "Minimal mode button icon simplified to an outlined square",
            "Phone number removed from the About page, leaving only the website and email (privacy)",
        ],
    },
    {
        "version": "2.6.0",
        "date": "2026-07-05",
        "title": "极简界面模式 · 跨侧续播根治 · docx导出修复 · 选区字幕自愈 · 对齐分句扩充",
        "notes": [
            "新增极简界面：左上角⛶钮一键切换，只留原文/译文区、翻译钮与状态栏；极简下按钮青色、最小窗口可缩至420x320，再点还原",
            "跨侧续播从头播根因修复：重播路径内部stop会把刚存的续播位置清零——改为先取位再stop，且只有真停止才归零",
            "选区朗读偶发无卡拉OK字幕：引擎返回相对选区的词边界时自动平移为全文绝对位置",
            "修复翻译历史导出docx报错(XML控制字符)：写入前统一净化NULL等非法字符，三处docx写入点全覆盖",
            "主动/从属区对齐分句标点扩充：新增逗号(，,)与全角空格参与分割，对应精度更高",
            "正常界面最小宽度720→770（左侧按钮增多后更合适）",
        ],
        "title_en": "Minimal UI mode · cross-side resume fix · docx export fix · selection subtitle self-healing",
        "notes_en": [
            "New minimal UI: one click on the top-left button leaves only the source/target areas, translate button and status bar; buttons turn cyan and the window can shrink to 420x320; click again to restore",
            "Root fix for cross-side resume restarting from the beginning: the internal stop inside the replay path was zeroing the just-saved resume position — the position is now read before stopping and only a true stop resets it",
            "Occasional missing karaoke when reading a selection: word boundaries returned relative to the selection are now shifted to absolute document positions",
            "Fixed translation history docx export errors (XML control characters): NULL and other illegal characters are sanitized before writing, covering all three docx write points",
            "Expanded sentence-splitting punctuation for active/passive area alignment: commas and full-width spaces now participate, improving accuracy",
            "Normal-mode minimum width raised from 720 to 770 to fit the added buttons",
        ],
    },
    {
        "version": "2.5.1",
        "date": "2026-07-04",
        "title": "跨侧续播回位 · 换嗓保字幕保选区(根治) · 光标处粘贴 · 单实例守护",
        "notes": [
            "跨侧暂停后回来点继续，从暂停位置续播不再从头：暂停位置存入该侧缓存，重播时自动定位；同侧继续/按停止会正确清零",
            "换嗓/引擎重读根治两处：①重读时沿用原选区（此前重新推导误判为读全文并清掉蓝色选区）②为重读而停被误当自然播完、250ms后清绿——preserve期间跳过收尾",
            "粘贴细化：主动区有明确光标位置时粘贴到光标处；从属区且无选区才默认贴到末尾",
            "新增单实例守护（QLockFile）：程序已在运行时再次启动会提示并退出，修复偶发双开",
        ],
        "title_en": "Cross-side resume position · voice switch preserves subtitles and selection · paste at cursor · single instance",
        "notes_en": [
            "Resuming after pausing on the other side no longer restarts from the beginning: the pause position is stored in that side's cache and restored on replay; continuing on the same side or pressing stop correctly resets it",
            "Two root fixes for voice/engine re-reads: the original selection is reused (previously re-derivation misread it as full text and cleared the blue selection), and a stop issued for a re-read is no longer mistaken for natural completion that cleared the green highlight 250ms later",
            "Paste refinement: in the active area with a definite cursor position, paste goes to the cursor; in the passive area with no selection it defaults to the end",
            "Added a single-instance guard (QLockFile): launching again while running shows a notice and exits, fixing occasional double launches",
        ],
    },
    {
        "version": "2.5.0",
        "date": "2026-07-03",
        "title": "缓存直播/真清空 · 换嗓保位保字幕 · 数字翻译修复+小数 · 选区复制粘贴 · 混合文本翻译",
        "notes": [
            "有音频缓存时点朗读直接重播（分侧缓存命中），不再重新生成；清空钮改为真清（同步清掉旧全局缓存，重播不再复活）",
            "换嗓音/引擎重读：卡拉OK绿色与蓝/灰选区随进度条一并保持；播放或暂停中切换嗓音必定触发重新生成（修复偶发不触发）",
            "修复强制中/英数字翻译失效根因：目标判断只比对『英语』而下拉项是『English』永不命中；88 现按目标正确输出中/英两式",
            "原文/译文语言下拉变化即强制重新翻译（等同点翻译按钮）；翻译大按钮无条件重翻",
            "数字支持小数：888.89 → 逐位拼读(含点/point) + 数学读法（八百八十八点八九 / point eight nine）",
            "复制/粘贴选区感知：有蓝/灰选区只复制或只覆盖选中部分；无选区复制全部、粘贴默认到末尾",
            "修复混合中英文本翻译：自动检测含中日韩字符时显式声明源为中文，Google 不再 en→en 原样返回（『中文』二字现在会被翻译）",
            "底部忙碌进度条改自绘胶囊：滑块滑到两端也保持圆角（Qt原生样式两端变方的限制已绕开）",
        ],
        "title_en": "Cache replay and true clear · voice switch keeps position · number translation fix · selection-aware copy/paste",
        "notes_en": [
            "With an audio cache present, pressing play replays directly (per-side cache hit) instead of regenerating; the clear button now truly clears (the old global cache is cleared too so replay cannot revive it)",
            "Voice/engine re-reads keep the green karaoke and blue/gray selection in sync with the progress bar; switching voices while playing or paused always triggers regeneration",
            "Fixed the root cause of forced Chinese/English number translation failing: the target check compared against a Chinese label while the dropdown value was 'English', so it never matched; numbers now output correctly for the chosen target",
            "Changing the source or target language dropdown forces a re-translation (equivalent to pressing Translate); the main translate button always re-translates",
            "Decimal support for numbers: 888.89 reads digit by digit (including the point) as well as mathematically",
            "Selection-aware copy/paste: with a blue/gray selection only that part is copied or replaced; with no selection everything is copied and paste defaults to the end",
            "Fixed mixed Chinese-English translation: when CJK characters are detected the source is declared as Chinese explicitly, so Google no longer returns en-to-en unchanged",
            "The bottom busy progress bar is now custom-drawn as a capsule so the slider keeps rounded ends (working around Qt's square-end limitation)",
        ],
    },
    {
        "version": "2.4.0",
        "date": "2026-07-03",
        "title": "多风格翻译修复+单词多译法 · 进度条冻结 · 载入下一条 · 按钮语义与间距",
        "notes": [
            "修复多风格翻译失效根因：导入过一次文件后残留路径会永久禁用多风格，改为仅在原文与导入内容一致时禁用",
            "多风格新增单词模式：输入单字/词（中文≤4字或英文单词）输出多种准确译法，每行一个，无解释无音标；词组/句子仍按书面/口语/正式等风格分类",
            "朗读钮颜色语义严格化：缓存有音频=青色，无=灰色；播放到头/拖到最右不再误变灰",
            "更改嗓音/引擎重读时，进度条冻结在当前位置，生成完成后就地续播，不再跳回开头",
            "新增『载入下一条原文』按钮（上一条右侧，镜像图标），双向循环历史",
            "设置弹窗保存钮改蓝色，取消保持灰色",
            "按钮组间距统一为10px（对标顶排语言框与交换钮间距）",
        ],
        "title_en": "Multi-style fixes and word mode · frozen progress bar · load next · button semantics and spacing",
        "notes_en": [
            "Root fix for multi-style translation being disabled: a leftover path from a single earlier file import permanently disabled it; it is now disabled only while the source text still matches the imported content",
            "New word mode for multi-style: entering a single character or word (up to 4 Chinese characters, or one English word) returns several accurate translations, one per line, with no explanations or phonetics; phrases and sentences still use the formal/casual/etc. style breakdown",
            "Stricter play button color semantics: cyan when an audio cache exists, gray when not; reaching the end or dragging to the far right no longer turns it gray by mistake",
            "When re-reading after a voice or engine change, the progress bar freezes at the current position and resumes in place once generation completes, instead of jumping back to the start",
            "New 'Load next source' button (right of Load previous, mirrored icon) for cycling through history in both directions",
            "Settings Save button turned blue while Cancel stays gray",
            "Button group spacing unified to 10px, matching the gap between the top language boxes and the swap button",
        ],
    },
    {
        "version": "2.3.1",
        "date": "2026-07-03",
        "title": "修复样式表整表失效(按钮丢样式元凶) · 控件统一加高 · 音频/文字清空分离",
        "notes": [
            "根因修复：样式表为三段拼接，占位符替换只作用于最后一段，残留占位符导致整表解析失败被丢弃——翻译钮丢蓝色、按钮变矮变形的元凶",
            "翻译大按钮恢复蓝色并加圆角；朗读钮青色态加圆角；导入钮变青后不再缩小",
            "所有下拉框/按钮统一加高到34px，图标钮改正方形，交换钮不再偏矮",
            "气球提示外框缩小约20%并改圆角",
            "翻译历史弹窗：检查历史去掉蓝色，关闭钮改蓝色，与关于/说明弹窗统一",
            "朗读语速滑杆两侧统一灰色（不再左蓝右灰），与播放进度条区分",
            "主动/从属切换补修：切换后新从属区正确显示灰色选区（去掉误清联动的旧逻辑）",
            "原生滚动条模式下文本框右内边距 20px 恢复为 8px，滚动条贴右不再有空隙",
            "朗读区新增音频清空钮（原文/译文各一，下载钮右侧）：仅释放该侧音频缓存，朗读钮变灰、下载失效",
            "原文/译文区文字清空钮改为只清文字与导入文件，不再清音频（文字与音频清空分离）",
        ],
        "title_en": "Root fix for stylesheet-wide failure (the cause of buttons losing styles) · uniform control height · audio and text clearing separated",
        "notes_en": [
            "Root cause fixed: the stylesheet is assembled from three parts and placeholder substitution only reached the last one; the leftover placeholders made the whole sheet fail to parse and be discarded — the true cause of the translate button losing its blue and buttons becoming short and misshapen",
            "The large translate button is blue and rounded again; the cyan play state gets rounded corners; the import button no longer shrinks when it turns cyan",
            "All dropdowns and buttons raised to a uniform 34px height, icon buttons made square, and the swap button is no longer shorter than its neighbors",
            "Tooltip frames reduced by about 20% and given rounded corners",
            "Translation history dialog: View History loses its blue and Close becomes blue, matching the About and User Guide dialogs",
            "Playback-speed slider is gray on both sides (no longer blue on the left) to distinguish it from the playback progress bar",
            "Active/passive switching follow-up fix: after switching, the new passive area correctly shows the gray selection (the old logic that wrongly cleared the link was removed)",
            "With native scrollbars the text area's right padding returns from 20px to 8px, so the scrollbar sits flush right without a gap",
        ],
    },
    {
        "version": "2.3.0",
        "date": "2026-07-02",
        "title": "主动区点击切换重构 · 原生胶囊滚动条(Mac/Win11) · 高亮清除根治 · 清空按钮重做",
        "notes": [
            "主动/从属区重构：鼠标按下即切换（点文字或空白都生效），只有一个蓝框；从属区灰色选区点击后变蓝，原主动区蓝选区自动转灰",
            "点击当前主动区：保持主动，仅清掉两侧联动选区",
            "改字清高亮根治：真实变更时掐断卡拉OK定时器、清词边界、停联动防抖、作废对齐表；边界过期自动熔断，绿色不再被画回",
            "改字时若该侧正在朗读，自动停止朗读（读的已是旧内容）",
            "滚动条：Mac 与 Win11(Qt>=6.7) 用系统原生胶囊；Win10及以下/Linux用自定义样式；全局样式解毒（QWidget规则改为调色板，不再污染滚动条）",
            "清空按钮移位：原文清空钮在翻译历史右侧（隔开），译文清空钮在最右侧（隔开）",
            "清空增强：原文清空=清文字+清导入(钮变灰)+译文随清+双侧音频释放(朗读/下载钮变灰)+导出钮隐藏；译文清空=清文字+译文音频释放",
        ],
        "title_en": "Active-area click switching rebuilt · native capsule scrollbars (macOS/Win11) · root fix for highlight clearing · clear buttons redone",
        "notes_en": [
            "Active/passive areas rebuilt: switching happens on mouse press (on text or blank space alike) with only one blue frame; clicking a gray selection in the passive area turns it blue while the previously active blue selection turns gray",
            "Clicking the currently active area keeps it active and only clears the linked selections on both sides",
            "Root fix for highlights surviving edits: on a real change the karaoke timer is cut, word boundaries cleared, link debouncing stopped and the alignment table invalidated; expired boundaries trip a fuse so green can no longer be repainted",
            "Editing text on a side that is currently being read automatically stops playback, since what is being read is already stale",
            "Scrollbars: macOS and Win11 (Qt 6.7+) use native capsules; Win10 and below plus Linux use the custom style; global styling detoxified (the QWidget rule now uses the palette so it no longer contaminates scrollbars)",
            "Clear buttons repositioned: the source clear button sits right of Translation History (separated), the target clear button at the far right (separated)",
            "Clearing enhanced: source clear wipes text, import state (button grays out), the target text as well, releases audio on both sides (play/download gray out) and hides the export button; target clear wipes target text and releases target audio",
        ],
    },
    {
        "version": "2.2.2",
        "date": "2026-07-01",
        "title": "根治选区朗读失效（高亮反馈环彻底消灭）",
        "notes": [
            "根因：停止/清除卡拉OK时 rehighlight 触发 textChanged，被误判为用户改字，导致选区信息被抹掉、自动翻译误触发、离线嗓音卡拉OK从全文扫（看似在读全文）",
            "根治：textChanged 处理器入口做文本比对——内容没变（仅格式重绘）直接跳过，一劳永逸消灭此类误伤",
            "选区朗读现在全程保留选区：蓝色底色、卡拉OK范围、边界估算都只在选区内",
        ],
        "title_en": "Root fix for selection playback failing (highlight feedback loop eliminated)",
        "notes_en": [
            "Root cause: stopping or clearing karaoke triggered rehighlight, which fired textChanged and was mistaken for a user edit — wiping the selection, falsely triggering auto-translation, and making offline-voice karaoke scan from the very beginning (appearing to read the whole text)",
            "Root fix: the textChanged handler now compares text at its entry point and skips immediately when the content is unchanged (a format-only repaint), eliminating this whole class of false positives once and for all",
            "Selection playback now preserves the selection throughout: the blue background, the karaoke range and boundary estimation all stay within the selection",
        ],
    },
    {
        "version": "2.2.1",
        "date": "2026-07-01",
        "title": "修复选区朗读被三态切换吃掉的问题",
        "notes": [
            "修复：选中一段文字再点朗读钮，会被『暂停/继续』三态切换拦截而继续播旧音频（v2.1.2 引入）",
            "现在：只要该区有新的选区（与当前朗读内容不同），点朗读钮即停掉旧音频、只朗读选区",
        ],
        "title_en": "Fixed selection playback being swallowed by the three-state toggle",
        "notes_en": [
            "Fixed: selecting text and pressing play was intercepted by the pause/resume three-state toggle and kept playing the old audio (introduced in v2.1.2)",
            "Now: whenever the area has a new selection differing from what is currently being read, pressing play stops the old audio and reads only the selection",
        ],
    },
    {
        "version": "2.2.0",
        "date": "2026-07-01",
        "title": "主动/从属区单蓝框 · 两侧独立音频缓存 · 高亮清除与导出钮修复 · Mac原生滚动条",
        "notes": [
            "主动区只保留一个蓝框（去掉 :focus 导致的第二个蓝框），从属区灰框",
            "两侧朗读音频各自独立缓存：朗读钮青色=音频在内存可下载，灰色=无音频不可下载",
            "改动某侧文字，该侧音频缓存作废、朗读钮变灰、下载禁用；停止后若音频仍在内存则钮保持青色、进度归零",
            "跨侧朗读：一侧朗读时点另一侧，先暂停该侧（进度保留、青色继续态）再读另一侧",
            "彻底修复文字变更时蓝色选区/绿色卡拉OK未清除（根因：译文区未监听 textChanged + guard 误伤，改用精确 _highlighting 标记）",
            "修复导出翻译后文件按钮显隐（根因：setPlainText 时序竞争，导入状态改为先记录后填文本）",
            "Mac 使用系统原生胶囊滚动条（无灰槽、自动隐藏）；Windows 保留自定义细样式",
            "最小窗口宽度调整为 720（较原缩小约 400px）",
        ],
        "title_en": "Single blue frame for the active area · independent per-side audio cache · highlight clearing and export button fixes · native macOS scrollbars",
        "notes_en": [
            "The active area keeps only one blue frame (the second frame caused by :focus was removed) while the passive area uses a gray frame",
            "Each side caches its playback audio independently: a cyan play button means audio is in memory and downloadable, gray means no audio and no download",
            "Editing text on a side invalidates that side's audio cache, grays out its play button and disables download; after stopping, the button stays cyan with progress reset if audio is still in memory",
            "Cross-side playback: pressing play on the other side while one is reading pauses the first side (keeping its progress and cyan resume state) before reading the other",
            "Fully fixed blue selections and green karaoke not clearing on text changes (root cause: the target area never listened to textChanged, plus an overreaching guard; replaced with a precise _highlighting flag)",
            "Fixed the export-translated-file button appearing and disappearing incorrectly (root cause: a setPlainText timing race; import state is now recorded before the text is filled)",
            "macOS uses native capsule scrollbars (no gray trough, auto-hiding) while Windows keeps the custom slim style",
        ],
    },
    {
        "version": "2.1.3",
        "date": "2026-06-30",
        "title": "修复卡拉OK失效与停止键 · 跨侧朗读暂停 · 最小窗口缩小 · 多项细节",
        "notes": [
            "修复原文区卡拉OK字幕失效（高亮 rehighlight 触发 textChanged 误清高亮的反馈环）",
            "修复文字变更时蓝色选区/绿色字幕未清除（补上清蓝色选区）",
            "修复停止朗读键失效（暂停态下也能停止）",
            "跨侧朗读：一侧朗读时点另一侧，先把该侧暂停（进度保留、按钮青色继续态），再读另一侧",
            "设置面板 HY-MT Key 改为『混元 HY-MT Key』",
            "关于窗『英语导师』与『English Coach』同字号同颜色",
            "最小窗口宽度大幅缩小（960→620）",
            "语速气球拖动时持续显示、跟随滑块",
            "进度条/滚动条进一步圆角处理",
        ],
        "title_en": "Karaoke and stop button fixes · cross-side pause · smaller minimum window · assorted details",
        "notes_en": [
            "Fixed karaoke subtitles failing in the source area (a feedback loop where rehighlight fired textChanged and wrongly cleared the highlight)",
            "Fixed blue selections and green subtitles not clearing on text changes (blue selection clearing added)",
            "Fixed the stop button not working (it now also stops from the paused state)",
            "Cross-side playback: pressing play on the other side while one is reading pauses the first side (progress kept, button in cyan resume state) before reading the other",
            "Settings HY-MT key relabeled as Hunyuan HY-MT Key",
            "In the About dialog the Chinese and English product names use the same size and color",
            "Minimum window width reduced substantially (960 to 620)",
        ],
    },
    {
        "version": "2.1.2",
        "date": "2026-06-30",
        "title": "界面标签改气球提示 · 翻译按钮可靠性 · 朗读按钮三态 · 导出钮智能显隐",
        "notes": [
            "关于窗 English Coach 下方加副标题『英语导师』",
            "原文/译文文字变更时，自动清除蓝色选区高亮与卡拉OK绿色高亮",
            "翻译按钮修复点击无效问题；点翻译会先中断所有进行中的翻译再重新开始",
            "界面文字标签全部去除改为悬停气球提示（引擎/语言/嗓音/文字区/进度条/语速等）",
            "语速气球实时显示『朗读语速 正常 / +20% / -50%』",
            "引擎名缩短：GLM-4-Flash→GLM、HY-MT→混元；约束最小窗口宽度让交换钮居中",
            "导出翻译后文件按钮：平时隐藏，导入成功才出现（灰），翻译完成变青可点；原文与导入不一致则消失",
            "所有滚动条只留胶囊滑块，去掉深灰背景轨道；等待进度条去底槽、圆角胶囊形",
            "朗读按钮悬停三态：朗读原文/译文 → 暂停朗读 → 继续朗读，循环（并修复暂停图标不显示的 bug）",
        ],
        "title_en": "Labels replaced by tooltips · translate button reliability · three-state play button · smart export button",
        "notes_en": [
            "Added a subtitle under English Coach in the About dialog",
            "Blue selection highlights and green karaoke highlights are cleared automatically when source or target text changes",
            "Fixed the translate button not responding; pressing translate now aborts any in-flight translation before starting fresh",
            "All inline text labels removed in favor of hover tooltips (engine, language, voice, text areas, progress bar, speed and more)",
            "The speed tooltip shows the live value: Speech rate Normal / +20% / -50%",
            "Engine names shortened (GLM-4-Flash to GLM, HY-MT to Hunyuan); a minimum window width keeps the swap button centered",
            "Export-translated-file button: hidden normally, appears grayed after a successful import, turns cyan and clickable once translation finishes, and disappears if the source no longer matches the import",
        ],
    },
    {
        "version": "2.1.1",
        "date": "2026-06-30",
        "title": "修复 PDF 导出报错 · 文件名空格化 · 进度条与滚动条精修 · 交换钮居中",
        "notes": [
            "修复导出 PDF 报错（reportlab.pdfbase.pdfmetrics 导入路径）",
            "所有导出文件名分隔符由下划线改空格，如 EC ZH XiaoXiao 2026-06-30 013733.mp3",
            "状态栏进度条改极简药丸形（无外框、更细、半透明槽、青色块）",
            "拖文件到文本框任意位置=导入内容（不再粘贴文件名）",
            "原文/译文朗读进度条彻底分离，拖动互不影响",
            "翻译历史按钮改名：检查历史 / 下载文档",
            "文本框滚动条改 mac 风格细药丸、不再挡字",
            "交换钮用网格强制窗口居中；下拉框留白压到最小；原文语言/译文语言改名",
        ],
        "title_en": "PDF export fix · spaces in filenames · progress bar and scrollbar refinements · centered swap button",
        "notes_en": [
            "Fixed a PDF export error (the reportlab.pdfbase.pdfmetrics import path)",
            "All export filenames now use spaces instead of underscores, e.g. EC ZH XiaoXiao 2026-06-30 013733.mp3",
            "Status bar progress bar restyled as a minimal pill (no outer frame, thinner, translucent trough, cyan block)",
            "Dropping a file anywhere on a text box imports its content (instead of pasting the filename)",
            "Source and target playback progress bars fully separated so dragging one does not affect the other",
            "Translation history buttons renamed to View History and Download Document",
            "Text box scrollbars restyled as slim macOS-style pills that no longer cover text",
        ],
    },
    {
        "version": "2.1.0",
        "date": "2026-06-29",
        "title": "选区主动/从属双向联动 · 界面重排 · 文件导入导出完善",
        "notes": [
            "选区双向联动：原文↔译文互选都能定位（依据保存的句对应关系），译文区选择终于能联动原文区",
            "主动/从属区逻辑：选中区为主动区（蓝框蓝底），另一区为从属区（选区灰底）；点击从属区即切换",
            "状态栏『正在生成音频』改用消息区显示，与『翻译完成』等交替不重叠（修正上次实现方式）",
            "进度条改药丸形：细灰黑轮廓、透明外围、青色滚动块",
            "第一排重排：翻译引擎｜源语言｜交换(居中)｜目标语言｜设置区(右)",
            "操作排重排：原文｜复制粘贴清空｜导出原文+导入文件｜上一条+历史｜翻译｜…｜导出译文+导出文件｜译文",
            "导出文件名规则 OT/TT_语言_日期时间，历史 TH_日期时间；保存路径记忆",
            "PDF 导出改 A4 竖版自动换行（修右侧出画），中文字体 PingFang/思源黑体/雅黑回退",
            "清空原文→重置导入状态；清空译文→重置导出状态（青色恢复灰色）",
            "停止键分侧独立：原文停止只停原文，译文停止只停译文",
            "上一条原文改为循环遍历全部历史原文",
            "拖拽文件到窗口=导入文件内容；原文区右内边距加大，滚动条不再挡字",
            "按钮宽度统一（显示隐藏为准），所有 Close 改『关闭』；标题改 English Coach",
        ],
        "title_en": "Two-way selection linking with active/passive areas · layout reshuffle · file import and export completed",
        "notes_en": [
            "Two-way selection linking: selecting in either the source or target area locates the counterpart (based on the stored sentence mapping), so target-area selections finally link back to the source",
            "Active/passive logic: the area you select in becomes active (blue frame, blue background) and the other becomes passive (gray selection background); clicking the passive area switches roles",
            "The status bar 'Generating audio' notice moved to the message area so it alternates with 'Translation complete' instead of overlapping (correcting the previous implementation)",
            "Progress bar restyled as a pill: thin dark-gray outline, transparent surroundings, cyan scrolling block",
            "First row reshuffled: engine, source language, swap (centered), target language, settings group (right)",
            "Action row reshuffled: source, copy/paste/clear, export source plus import file, previous plus history, translate, export target plus export file, target",
            "Export filename rules OT/TT_language_datetime and history TH_datetime, with the save path remembered",
        ],
    },
    {
        "version": "2.0.0",
        "date": "2026-06-29",
        "title": "2.0 大改版：文件导入导出 · 朗读分原文/译文两组 · 多项修复",
        "notes": [
            "修复 Argos 离线翻译『出出出』重复退化（改为逐句翻译 + 重复压缩兜底）",
            "朗读控制分原文/译文两组，上移到复制粘贴排与嗓音排之间，按钮改纯图标方形",
            "新增导入文件（txt/docx/pdf，支持拖拽）：内容填入原文区并自动翻译，按钮变青色",
            "新增导出当前原文/译文文字（txt/md/docx/json/pdf）",
            "新增导出翻译后文件：docx 保留段落结构，文件名加 T 后缀；pdf 暂只读入不导出",
            "翻译历史新增多格式下载（txt/md/json/docx/pdf）",
            "数字/单符号特殊翻译：78→Seven Eight / seventy-eight（中英各两式）；逗号→comma/逗号",
            "文件导入翻译模式下，多风格翻译与符号翻译自动失效",
            "进度条改青色无外框，状态栏提示移到右侧不再重叠，Mac/Win 右侧留 5% 空隙",
            "设置窗取消/保存按钮统一右对齐",
            "模型目录更名为 EnglishCoach Models/Argos、EnglishCoach Models/Kokoro",
        ],
        "title_en": "Version 2.0 overhaul: file import and export · playback split into source and target groups · assorted fixes",
        "notes_en": [
            "Fixed Argos offline translation degenerating into repeated characters (switched to sentence-by-sentence translation with a repetition-collapsing fallback)",
            "Playback controls split into source and target groups, moved between the copy/paste row and the voice row, with buttons changed to plain square icons",
            "New file import (txt/docx/pdf, drag and drop supported): content fills the source area and is translated automatically, turning the button cyan",
            "New export of the current source or target text (txt/md/docx/json/pdf)",
        ],
    },
    {
        "version": "1.9.6",
        "date": "2026-06-29",
        "title": "选区联动重做（句级对齐）· 朗读交互修复 · 状态栏与下拉修复",
        "notes": [
            "选区联动重新设计：翻译时按句建立原文↔译文对应，选区覆盖所有相关句（选全文即全亮）",
            "合成期间界面不变：已选区的蓝/灰背景在生成音频时保持显示",
            "改朗读引擎/嗓音时进度条不再跳回头、按钮保持青色『继续朗读』不闪",
            "朗读中点交换：立即终止朗读并清空音频缓存，自动重新翻译",
            "状态栏『正在生成音频』改为蓝底白字，与其它提示一致",
            "Mac 合成进度条左移，右侧留约 10% 空隙，构图匀称",
            "Windows 嗓音下拉滚轮不再滚出末尾空行",
        ],
        "title_en": "Selection linking redone (sentence-level alignment) · playback interaction fixes · status bar and dropdown fixes",
        "notes_en": [
            "Selection linking redesigned: translation builds a sentence-by-sentence source-to-target mapping, so a selection highlights every related sentence (selecting everything lights up everything)",
            "The interface stays put during synthesis: blue and gray selection backgrounds remain visible while audio is generated",
            "Changing the playback engine or voice no longer sends the progress bar back to the start, and the button stays cyan in its resume state without flickering",
            "Pressing swap during playback immediately stops playback, clears the audio cache and re-translates automatically",
            "The status bar 'Generating audio' notice now uses white text on blue, matching the other messages",
            "On macOS the synthesis progress bar shifts left, leaving about 10% clearance on the right for a balanced layout",
            "The Windows voice dropdown no longer scrolls past the last item into blank space",
        ],
    },
    {
        "version": "1.9.5",
        "date": "2026-06-28",
        "title": "更新应用图标",
        "notes": [
            "重绘应用图标：线条更圆润饱满，接近原生质感",
            "GPU 版图标补回 A/文 标牌（青绿底 + 闪电，与 CPU 版区分）",
        ],
        "title_en": "Updated application icon",
        "notes_en": [
            "Application icon redrawn with rounder, fuller strokes for a more native feel",
            "The GPU edition icon regains its A/文 badge (teal background with a lightning bolt, distinguishing it from the CPU edition)",
        ],
    },
    {
        "version": "1.9.4",
        "date": "2026-06-28",
        "title": "朗读可中断 · 选区联动不再翻译 · GPU 图标 · 打包命名 · 多项修复",
        "notes": [
            "新增朗读中断：合成中按停止可中断（段边界软中断），并提示『正在停止…』",
            "选区联动改为按位置比例映射，完全不再调用翻译引擎（选译文不再触发自动翻译）",
            "新增 GPU 版专属图标（青绿底+闪电）",
            "Mac 打包为 English Coach.app；Win 打包为含 English Coach 文件夹与 English Coach.exe（GPU 版同理）",
            "选区强制蓝色背景（修复某些系统显示绿色）",
            "中文嗓音下拉加宽（Windows 不再截断），并去除末尾空行与多余滚动",
            "合成中『正在生成音频』文字持续到结束，进度条约占窗口 40% 居中偏右、右侧留白",
            "朗读中点交换内容会停止当前朗读，避免青绿字幕错位到对面",
        ],
        "title_en": "Interruptible playback · selection linking no longer translates · GPU icon · packaging names · assorted fixes",
        "notes_en": [
            "New interruptible playback: pressing stop during synthesis aborts it at a segment boundary and shows a 'Stopping...' notice",
            "Selection linking now maps by positional ratio and never calls the translation engine (selecting target text no longer triggers auto-translation)",
            "New dedicated icon for the GPU edition (teal background with a lightning bolt)",
            "macOS packages as English Coach.app; Windows packages as a folder named English Coach containing English Coach.exe (same for the GPU edition)",
            "Selections are forced to a blue background (fixing green selections on some systems)",
            "Chinese voice dropdown widened (no longer truncated on Windows), with trailing blank rows and excess scrolling removed",
            "During synthesis the 'Generating audio' text persists until completion, with the progress bar occupying about 40% of the window, centered slightly right with clearance on the right",
        ],
    },
    {
        "version": "1.9.3",
        "date": "2026-06-28",
        "title": "修复中文离线朗读截断 · 译文区选区稳定 · 多项界面优化",
        "notes": [
            "修复中文离线朗读长文本被截断（按标点切句分段合成，避免 Kokoro token 上限静默截断）",
            "选区联动改为仅『原文→译文』单向，译文区选择不再被刷新重置",
            "翻译历史文件改为 txt 纯文本",
            "翻译历史按钮图标改为列表横线样式（与更新说明区分）",
            "翻译按钮用网格真正居中，不再因左侧按钮增多而偏右",
            "合成进度：文字并入左侧状态栏，进度条加长放右侧",
        ],
        "title_en": "Fixed truncated Chinese offline playback · stable target-area selections · assorted UI improvements",
        "notes_en": [
            "Fixed long Chinese text being truncated during offline playback (text is now split at punctuation and synthesized in segments, avoiding Kokoro's silent token-limit truncation)",
            "Selection linking changed to source-to-target only, so selections in the target area are no longer reset by refreshes",
            "Translation history files changed to plain text",
            "Translation history button icon changed to a list-lines style, distinguishing it from the change log",
            "The translate button is now truly centered using a grid, instead of drifting right as buttons were added on the left",
            "Synthesis progress: the text merged into the status bar on the left with a longer progress bar on the right",
        ],
    },
    {
        "version": "1.9.2",
        "date": "2026-06-28",
        "title": "修复新环境 ctranslate2 因 setuptools 过新缺 pkg_resources",
        "notes": [
            "构建脚本在装 ctranslate2 前固定 setuptools<81，解决新版移除 pkg_resources 导致 Argos 离线翻译导入失败",
        ],
        "title_en": "Fixed ctranslate2 failing in new environments due to a too-new setuptools missing pkg_resources",
        "notes_en": [
            "The build scripts now pin setuptools below 81 before installing ctranslate2, resolving Argos offline translation import failures caused by newer releases removing pkg_resources",
        ],
    },
    {
        "version": "1.9.1",
        "date": "2026-06-28",
        "title": "修复 Windows GPU 打包脚本无法运行（换行符与编码）",
        "notes": [
            "修复 Build Windows GPU.bat 因换行符为 LF、含中文注释导致命令被截断、无法运行",
            "所有 .bat 脚本统一为纯 ASCII + CRLF 换行（GBK 控制台安全）",
        ],
        "title_en": "Fixed the Windows GPU build script failing to run (line endings and encoding)",
        "notes_en": [
            "Fixed Build Windows GPU.bat failing to run because LF line endings combined with Chinese comments truncated commands",
            "All .bat scripts standardized to pure ASCII with CRLF line endings (safe for GBK consoles)",
        ],
    },
    {
        "version": "1.9.0",
        "date": "2026-06-27",
        "title": "翻译按钮防闪 · 历史改 MD · 进度条优化 · GPU 打包脚本",
        "notes": [
            "翻译按钮文字不再变化（不闪动），状态提示移到底部状态栏",
            "翻译历史改用可读的 Markdown 文本格式（任意文本/浏览器打开都清晰，不再乱码）",
            "合成音频进度提示加文字『正在生成音频…』，进度条加长、左文右条、留白匀称",
            "上一条/历史按钮缩小到与复制粘贴一致，间距统一；撤回图标改为逆时针箭头",
            "按钮气泡提示字号缩小；历史弹窗内悬停气泡字号放大到正文大小",
            "暂停时拖动进度条，青色已读实时跟随（未松手也刷新）",
            "朗读按钮激活时喇叭图标同步反色",
            "选区联动响应更快，降低首次朗读读整段的概率",
            "新增 Windows GPU 加速打包脚本（EnglishCoach-GPU 环境 + CUDA），Kokoro 自动用 GPU",
            "构建脚本更名 Build MacOS.sh / Build Windows.bat / Build Windows GPU.bat",
        ],
        "title_en": "Flicker-free translate button · history in Markdown · progress bar improvements · GPU build script",
        "notes_en": [
            "The translate button's label no longer changes (no flicker); status messages moved to the bottom status bar",
            "Translation history switched to readable Markdown text (clear in any text editor or browser, no more garbled characters)",
            "Audio synthesis progress now shows 'Generating audio...' with a longer bar, text on the left and bar on the right, evenly spaced",
            "Previous and History buttons shrunk to match copy and paste, with unified spacing; the undo icon changed to a counter-clockwise arrow",
            "Button tooltip font size reduced; hover tooltips inside the history dialog enlarged to body text size",
            "Dragging the progress slider while paused updates the cyan read position live (refreshing before you release)",
            "The speaker icon inverts in sync when the play button is active",
        ],
    },
    {
        "version": "1.8.1",
        "date": "2026-06-27",
        "title": "中文离线朗读修复 · 选区联动位置加权 · 合成进度提示 · 多项打磨",
        "notes": [
            "修复中文离线朗读缺 pypinyin 等依赖（已补入打包）",
            "移除会 404 的 Bella+Sarah 混合音（HF 无该现成文件）",
            "选区联动加入位置加权与段落感知：相同词句按位置就近匹配，结果不再离谱",
            "合成语音较慢时状态栏提示，超过约 1 秒显示忙碌进度条",
            "历史弹窗：按钮改为『查看历史文档/重新载入/关闭』三等宽，预览限宽无横向滚动条",
            "上一条原文按钮换更大图标",
            "设置：显示隐藏与查看日志按钮等宽；取消/保存改中文并统一取消在左保存在右",
        ],
        "title_en": "Chinese offline playback fix · position-weighted selection linking · synthesis progress · assorted polish",
        "notes_en": [
            "Fixed Chinese offline playback missing pypinyin and related dependencies (now bundled)",
            "Removed the Bella+Sarah blended voice, which returned 404 (no such prebuilt file on Hugging Face)",
            "Selection linking gained position weighting and paragraph awareness: identical words are matched by proximity, so results are no longer wildly off",
            "The status bar reports slow speech synthesis, showing a busy progress bar after about one second",
            "History dialog: buttons changed to three equal-width actions (View History Document / Reload / Close) with a width-limited preview and no horizontal scrollbar",
            "Larger icon for the previous-source button",
            "Settings: Show/Hide and View Log buttons made equal width; Cancel and Save relabeled in Chinese and consistently ordered with Cancel left, Save right",
        ],
    },
    {
        "version": "1.8.0",
        "date": "2026-06-27",
        "title": "翻译历史 · 日志 · 修复 Kokoro 缺 ordered_set · Windows 图标",
        "notes": [
            "修复 Kokoro 缺少 ordered_set 依赖导致离线朗读不可用（已补入打包）",
            "新增翻译历史：原文区删除键右侧加『上一条原文』与『历史』两个按钮",
            "历史弹窗按天分组，每条显示开头提示，悬停蓝色高亮并气泡显示全文，点选可载入并翻译",
            "新增日志：出错记录写入用户数据目录，设置面板加『查看日志』按钮",
            "历史与日志统一存放在系统用户数据目录（win:%APPDATA% / mac:Application Support）",
            "Windows 应用图标改回铺满尺寸（不再显小）；macOS 图标不变",
            "选区联动匹配更智能：英文吸附到整词，不停在半个单词中间",
        ],
        "title_en": "Translation history · logging · fixed Kokoro missing ordered_set · Windows icon",
        "notes_en": [
            "Fixed Kokoro missing the ordered_set dependency, which made offline playback unavailable (now bundled)",
            "New translation history: Previous Source and History buttons added right of the source area's delete key",
            "The history dialog groups entries by day, shows an opening snippet for each, highlights on hover in blue with a full-text tooltip, and loads plus translates the entry you pick",
            "New logging: errors are written to the user data directory, with a View Log button added to Settings",
            "History and logs are stored together in the system user data directory (%APPDATA% on Windows, Application Support on macOS)",
            "The Windows application icon fills its canvas again (no longer appearing small); the macOS icon is unchanged",
            "Smarter selection linking: English snaps to whole words instead of stopping mid-word",
        ],
    },
    {
        "version": "1.7.1",
        "date": "2026-06-27",
        "title": "修复外置硬盘导致 Kokoro 不可用 · 朗读高亮与选区联动优化",
        "notes": [
            "修复 Kokoro 因 conda 在外置 SSD 上读时区文件被 macOS 拒绝（Operation not permitted）导致离线朗读不可用；自动改用系统时区库，并给出权限设置指引",
            "朗读高亮不再改变字色（黑字变白字），只改背景颜色，更清爽",
            "灰色联动区朗读时保持灰色，读到处变青色，读完恢复灰色（不再变蓝、不消失）",
            "下载音频记住上次选择的格式（wav/mp3）作为默认",
            "交换原文译文时，选区随之继承到新原文区并自动联动新译文区",
        ],
        "title_en": "Fixed Kokoro being unavailable on external drives · playback highlight and selection linking improvements",
        "notes_en": [
            "Fixed Kokoro being unavailable for offline playback when conda lives on an external SSD and macOS denied reading the timezone file (Operation not permitted); the system timezone library is now used, with guidance for the permission setting",
            "Playback highlighting no longer changes the text color (black to white) and only changes the background, for a cleaner look",
            "The gray linked area stays gray while being read, turning cyan at the reading position and returning to gray afterwards (no longer turning blue or disappearing)",
            "Audio download remembers the last chosen format (wav/mp3) as the default",
            "Swapping source and target carries the selection into the new source area and links it to the new target area automatically",
        ],
    },
    {
        "version": "1.7.0",
        "date": "2026-06-26",
        "title": "Kokoro 打包落地 · 音质提升 · 音频格式可选 · 设置面板优化",
        "notes": [
            "Kokoro 离线朗读依赖钉死兼容 Big Sur 的版本（torch 2.2.2 + transformers 4.40.2 + spaCy 英文模型）",
            "英文离线嗓音默认改为 Heart（盲测最像真人），新增 Nova 与 Bella+Sarah 混合音",
            "下载音频可选 wav（无损）或 mp3（压缩），按需转换",
            "Argos 离线模型缓存改放 EnglishCoach-models/argos 子目录（win/mac）",
            "帮助文档新增系统要求（macOS/Windows 版本、空间、未签名解除拦截等），便于分享",
            "设置：Key 输入框限宽右侧留白、显示/隐藏按钮居中、滚动条不再挡字、下拉悬停蓝色统一",
            "Key 标签简化（GPT/Gemini/GLM/豆包/HY-MT）",
            "构建脚本加固：醒目确认并强制校验 conda 环境，杜绝误装到 base",
        ],
        "title_en": "Kokoro bundling completed · better audio quality · selectable audio format · Settings improvements",
        "notes_en": [
            "Kokoro offline playback dependencies pinned to Big Sur-compatible versions (torch 2.2.2, transformers 4.40.2 and the spaCy English model)",
            "Default English offline voice changed to Heart (the most human-sounding in blind testing), with Nova and a Bella+Sarah blend added",
            "Audio downloads can be wav (lossless) or mp3 (compressed), converted on demand",
            "Argos offline model cache moved to an EnglishCoach-models/argos subdirectory (Windows and macOS)",
            "Help documentation gained system requirements (macOS/Windows versions, disk space, unblocking unsigned apps) to make sharing easier",
            "Settings: key inputs width-limited with right clearance, Show/Hide button centered, scrollbars no longer covering text, and unified blue dropdown hover",
            "Key labels simplified (GPT/Gemini/GLM/Doubao/HY-MT)",
        ],
    },
    {
        "version": "1.6.0",
        "date": "2026-06-26",
        "title": "修复闪退与 Kokoro 打包 · 重播缓存 · 多项界面优化",
        "notes": [
            "修复选择文字时的闪退（选区联动跨线程刷新高亮的崩溃，已加锁保护）",
            "修复 Kokoro 打包缺失 language_tags 等数据导致离线朗读不可用",
            "无变化时再次点朗读，直接重播已生成音频，不再重新合成",
            "修复原文区有选区时朗读，青色不覆盖蓝色选区的问题（与译文区一致）",
            "卡拉OK字幕提前量回调到中间值，更贴合声音",
            "选区联动匹配更稳，异常不再影响主流程",
            "引擎名简化：GPT / Gemini / GLM-4-Flash / 豆包 / HY-MT 等",
            "界面标签去掉冒号；『进度』改为『朗读进度』",
            "设置：显示隐藏按钮缩短、窗口变窄并自动换行、多风格翻译默认开启",
        ],
        "title_en": "Fixed crashes and Kokoro bundling · replay cache · assorted UI improvements",
        "notes_en": [
            "Fixed a crash when selecting text (selection linking refreshed highlights across threads; now protected by a lock)",
            "Fixed Kokoro bundling missing language_tags and other data, which made offline playback unavailable",
            "Pressing play again with nothing changed replays the already-generated audio instead of re-synthesizing",
        ],
    },
    {
        "version": "1.5.0",
        "date": "2026-06-26",
        "title": "选区联动 · 新增 Google 官方引擎 · 朗读与界面多项打磨",
        "notes": [
            "新增选区联动：选一侧文字，临时翻译后在另一侧灰色高亮最匹配区间，可直接朗读",
            "新增『Google -API-Key联网』官方云翻译 Basic v2 引擎（免费版保留不变）",
            "帮助文档补充各引擎的模型、收费、稳定性说明，便于甄别",
            "卡拉OK字幕加大提前量，修复再次出现的滞后",
            "原文区选区朗读高亮与译文区一致；读完保留普通蓝色选区（可正常取消/重选）",
            "朗读时按钮变青绿背景，暂停仍保持，停止/读完恢复",
            "只有当前朗读语种的嗓音改动才打断朗读（改另一语种嗓音不打断）",
            "中文嗓音移到左、英文嗓音移到右；窗口最窄时不再与设置按钮重叠",
            "更改翻译引擎后自动触发翻译",
            "Kokoro 离线朗读报错时显示真实原因，便于排查",
        ],
        "title_en": "Selection linking · new official Google engine · playback and UI polish",
        "notes_en": [
            "New selection linking: select text on one side and the closest matching range is highlighted in gray on the other side after a temporary translation, ready to be read aloud",
            "New 'Google -API-Key online' official Cloud Translation Basic v2 engine (the free version is unchanged)",
            "Help documentation now covers each engine's model, pricing and reliability to make choosing easier",
            "Increased karaoke subtitle lead time, fixing lag that had reappeared",
            "Source-area selection playback highlighting now matches the target area; after reading, a normal blue selection remains (cancelable and reselectable as usual)",
            "The play button takes a teal background while reading, keeps it while paused, and returns to normal on stop or completion",
            "Only changing the voice for the language currently being read interrupts playback (changing the other language's voice does not)",
        ],
    },
    {
        "version": "1.4.0",
        "date": "2026-06-26",
        "title": "新增 8 个 LLM 翻译引擎 · 多风格翻译",
        "notes": [
            "新增 OpenAI GPT、Google Gemini、Claude、智谱GLM-4-Flash、文心一言、字节豆包、通义千问、Kimi 八个大模型引擎（均需 API Key）",
            "所有 LLM 引擎统一走 OpenAI 兼容接口，设置中分别填 Key",
            "新增『多风格翻译』开关：选用 LLM 引擎时，主译文不变，下方附书面/口语/俚语/美式英式等多种辅助译法",
            "设置页 Key 较多，改为可滚动",
        ],
        "title_en": "Eight new LLM translation engines · multi-style translation",
        "notes_en": [
            "Added eight large-model engines: OpenAI GPT, Google Gemini, Claude, Zhipu GLM-4-Flash, ERNIE, Doubao, Qwen and Kimi (all require an API key)",
            "All LLM engines use the OpenAI-compatible interface, with separate keys entered in Settings",
            "New multi-style translation toggle: with an LLM engine the main translation is unchanged and formal, casual, slang, US and UK variants are appended below",
            "The Settings page became scrollable to accommodate the many keys",
        ],
    },
    {
        "version": "1.3.0",
        "date": "2026-06-26",
        "title": "新增 Kokoro 本地离线朗读 · 引擎/嗓音标注联网离线 · 多项修复",
        "notes": [
            "新增 Kokoro 本地离线朗读引擎：无需联网、CPU 即可、原生对齐时间戳，卡拉OK更精准",
            "嗓音标注来源：edge-tts 标『-线上联网』，Kokoro 标『-离线本地』",
            "引擎名标注联网方式：Google/DeepL/DeepSeek/HunYuan 联网，Argos 本地离线",
            "修复原文区选区朗读无青蓝覆盖；修复朗读结束后蓝色选区无法取消",
            "朗读失败不再反复弹窗打扰，并提示可改用离线嗓音",
            "字幕加入提前量补偿，跟声音更齐；选区朗读字幕只在选区内推进",
            "保存音频记住上次目录；离线为 wav、在线为 mp3，文件名 EC_语种_嗓音_日期_时间",
            "关于页：联系方式一行、版权用标准英文写法",
        ],
        "title_en": "New Kokoro local offline playback · online/offline labels for engines and voices · assorted fixes",
        "notes_en": [
            "New Kokoro local offline playback engine: no internet required, runs on CPU, with native alignment timestamps for more accurate karaoke",
            "Voices labeled by source: edge-tts marked 'online' and Kokoro marked 'offline local'",
            "Engines labeled by connectivity: Google, DeepL, DeepSeek and Hunyuan online, Argos local offline",
            "Fixed the source area lacking the teal overlay during selection playback; fixed blue selections being uncancelable after playback finished",
            "Playback failures no longer trigger repeated popups and now suggest switching to an offline voice",
            "Subtitles gained lead-time compensation for tighter sync; selection playback advances subtitles only within the selection",
            "Saving audio remembers the last directory; offline saves as wav and online as mp3, named EC_language_voice_date_time",
        ],
    },
    {
        "version": "1.2.2",
        "date": "2026-06-25",
        "title": "字幕提前补偿 · 选区字幕只走选区 · 关于页排版",
        "notes": [
            "卡拉OK字幕加入提前量补偿，跟声音对得更齐（中英文都不再慢半拍）",
            "修复选区朗读时字幕仍从全文头走到尾的问题，现在只在选区内推进",
            "关于页：网址与邮箱一行、两个电话另起一行，版权恢复 © 符号",
            "关于页主标题下方空行高度与其它文档一致",
        ],
        "title_en": "Subtitle lead compensation · selection subtitles stay in the selection · About page layout",
        "notes_en": [
            "Karaoke subtitles gained lead-time compensation for tighter audio sync (no longer half a beat behind in either language)",
            "Fixed subtitles running from the start of the whole text during selection playback; they now advance only within the selection",
            "About page: website and email on one line with phone numbers on the next, and the copyright symbol restored",
            "The blank line under the About page's main title now matches the other documents",
        ],
    },
    {
        "version": "1.2.1",
        "date": "2026-06-25",
        "title": "卡拉OK更精准 · 选区朗读恢复 · 音频命名优化",
        "notes": [
            "卡拉OK高亮改用播放器真实位置同步，更精准（暂停时位置准确）",
            "恢复『选中部分朗读』：选区显示蓝底，读到处覆盖青蓝绿，读完恢复蓝色",
            "朗读高亮色改为偏青蓝、降饱和，更柔和",
            "保存音频记住上次保存目录；文件名改为 EC_语种_嗓音_日期_时间（中文嗓音用拼音）",
            "标题字号略增大（仍克制）；原文/译文区字号再放大",
            "关于页联系方式与版权信息更新",
        ],
        "title_en": "More accurate karaoke · selection playback restored · audio naming improvements",
        "notes_en": [
            "Karaoke highlighting now syncs to the player's real position for better accuracy (correct position while paused)",
            "Restored reading a selection aloud: the selection shows a blue background, turns teal at the reading position and returns to blue when finished",
            "Playback highlight color shifted toward a softer, less saturated teal",
            "Saving audio remembers the last directory; filenames changed to EC_language_voice_date_time (Chinese voices use pinyin)",
            "Title font size increased slightly (still restrained); source and target area fonts enlarged further",
            "About page contact details and copyright information updated",
        ],
    },
    {
        "version": "1.2.0",
        "date": "2026-06-25",
        "title": "卡拉OK高亮独立时钟驱动 · 符号翻译 · 多项界面优化",
        "notes": [
            "卡拉OK高亮改为纯独立时钟驱动，彻底脱离播放器状态，两平台稳定逐词高亮",
            "无词边界时按字符比例估算，保证仍有高亮兜底",
            "新增标点/符号翻译：如 . → dot、（ → 左括号，引擎无能为力时由内置词典兜底",
            "单个英文单词、过短文本也会尽力翻译出含义",
            "Windows 下拉弹窗去掉多余滚动条",
            "原文/译文区字号放大一号，更醒目",
            "文档主标题下方增加空行，排版更匀称",
        ],
        "title_en": "Karaoke highlighting driven by an independent clock · symbol translation · assorted UI improvements",
        "notes_en": [
            "Karaoke highlighting now runs on a fully independent clock, decoupled from player state, for stable word-by-word highlighting on both platforms",
            "When word boundaries are unavailable, timing is estimated by character ratio so highlighting still works",
            "New punctuation and symbol translation: a period becomes 'dot' and a full-width parenthesis becomes 'left parenthesis', with a built-in dictionary covering what engines cannot",
            "Single English words and very short text are now translated as far as possible",
            "Removed the redundant scrollbar from Windows dropdown popups",
        ],
    },
    {
        "version": "1.1.9",
        "date": "2026-06-25",
        "title": "卡拉OK高亮改用底层着色（修复 macOS 不显示）· 下拉排版匀称",
        "notes": [
            "卡拉OK高亮改用 QSyntaxHighlighter 底层着色，修复 macOS/Big Sur 完全不显示高亮的问题",
            "朗读时逐词背景变青绿，随进度递增，读完恢复",
            "下拉列表去掉底部多余空白，外边框单层均匀，行距匀称",
        ],
        "title_en": "Karaoke highlighting moved to low-level formatting (fixing macOS) · balanced dropdown layout",
        "notes_en": [
            "Karaoke highlighting now uses QSyntaxHighlighter-level formatting, fixing highlights not appearing at all on macOS Big Sur",
            "While reading, each word's background turns teal in turn, advancing with progress and restoring when finished",
            "Dropdown lists lost their trailing blank space, with a single even outer border and balanced line spacing",
        ],
    },
    {
        "version": "1.1.8",
        "date": "2026-06-25",
        "title": "下拉项加大行距 · 文档标题确实缩小 · 按钮等宽",
        "notes": [
            "下拉列表选项行距加大，文字不再重叠；弹窗去掉底部缝隙",
            "文档标题改用更可靠的方式渲染，确实缩小到接近正文",
            "设置中 Save 与 Cancel 按钮等宽",
        ],
        "title_en": "Larger dropdown line spacing · document titles genuinely smaller · equal-width buttons",
        "notes_en": [
            "Dropdown item line spacing increased so text no longer overlaps, and the popup's bottom gap removed",
            "Document titles now render through a more reliable method and are genuinely reduced to near body-text size",
            "Save and Cancel buttons in Settings made equal width",
        ],
    },
    {
        "version": "1.1.7",
        "date": "2026-06-25",
        "title": "卡拉OK高亮改用墙上时钟 · 文档标题再缩小 · 弹窗去白边",
        "notes": [
            "卡拉OK逐词高亮改用独立时钟驱动，不再依赖播放器位置，macOS 也能逐词推进",
            "文档（关于/更新/帮助）标题字号改用更可靠的方式设置，确实缩小到接近正文",
            "下拉弹出列表去掉上下白边、深色背景",
            "暂停/继续时高亮位置正确衔接",
        ],
        "title_en": "Karaoke highlighting driven by a wall clock · smaller document titles · popup white edges removed",
        "notes_en": [
            "Word-by-word karaoke highlighting now runs on an independent clock instead of the player position, so it advances on macOS too",
            "Document titles (About, Change Log, Help) use a more reliable font-size mechanism and are genuinely close to body text",
            "Dropdown popups lost their top and bottom white edges and use a dark background",
            "Highlight position resumes correctly when pausing and continuing",
        ],
    },
    {
        "version": "1.1.6",
        "date": "2026-06-25",
        "title": "修复下载报错 · 下拉弹窗加宽 · 缩窄不重叠",
        "notes": [
            "修复下载音频第二次报『只读文件系统』错误：默认保存到「下载」文件夹",
            "下载文件名自动命名：EnglishCoach_日期_编号.mp3（编号自增）",
            "下载按钮换成更直观的下载图标",
            "下拉弹出列表加宽，完整显示选项、两侧留空隙、去掉上下白边",
            "修复窗口缩到最小时顶部控件与设置图标重叠",
        ],
        "title_en": "Download error fix · wider dropdown popups · no overlap when narrowed",
        "notes_en": [
            "Fixed a read-only file system error on the second audio download: files now default to the Downloads folder",
            "Downloads are named automatically as EnglishCoach_date_number.mp3 with an incrementing number",
            "Download button replaced with a more intuitive download icon",
            "Dropdown popups widened to show options in full with clearance on both sides and no top or bottom white edges",
            "Fixed top controls overlapping the settings icon when the window is at its minimum size",
        ],
    },
    {
        "version": "1.1.5",
        "date": "2026-06-25",
        "title": "修复崩溃 · 内存播放提速 · 卡拉OK与下载音频",
        "notes": [
            "修复反复朗读/拖动/退出时的崩溃（朗读线程安全退役与回收）",
            "改为内存直接播放，不再生成临时 mp3，启动播放更快",
            "修复卡拉OK逐词高亮（时间轴换算修正 + 高频刷新）",
            "朗读中改参数时进度条停在当前位置，不再归零",
            "新增「下载音频」按钮：点击才生成 mp3 文件保存",
            "按钮文案精简：朗读原文 / 朗读译文 / 停止朗读，播放时显示 暂停/继续朗读",
            "下拉框去掉选中对号、加蓝色高亮（含 macOS），并尽量收窄",
        ],
        "title_en": "Crash fix · faster in-memory playback · karaoke and audio download",
        "notes_en": [
            "Fixed crashes when repeatedly reading, dragging or quitting (playback threads now retire and are reclaimed safely)",
            "Playback moved to memory instead of generating a temporary mp3, so it starts faster",
            "Fixed word-by-word karaoke highlighting (timeline conversion corrected plus higher refresh rate)",
            "Changing parameters while reading leaves the progress bar at its current position instead of resetting it",
            "New Download Audio button: an mp3 file is generated and saved only when clicked",
            "Button labels simplified to Read Source / Read Target / Stop, showing Pause and Resume during playback",
        ],
    },
    {
        "version": "1.1.4",
        "date": "2026-06-24",
        "title": "修复卡拉OK高亮 · 窗口可缩窄 · 进度条不归零",
        "notes": [
            "修复卡拉OK逐词高亮无效：改用高频定时器驱动，高亮随声音顺滑推进",
            "朗读中更改设置时，进度条停在当前位置，不再跳回开头",
            "修复窗口过宽且无法缩窄：朗读控件分两行排布，窗口可自由调窄",
            "下拉框完整显示且不过度拉伸；悬停项目显示蓝色高亮",
            "文档标题字号确认缩小到接近正文",
        ],
        "title_en": "Karaoke highlighting fix · window can be narrowed · progress bar no longer resets",
        "notes_en": [
            "Fixed word-by-word karaoke highlighting not working: a high-frequency timer now drives it so highlights advance smoothly with the audio",
            "Changing settings while reading leaves the progress bar at its current position instead of jumping back to the start",
            "Fixed the window being too wide and impossible to narrow: playback controls now span two rows so the window resizes freely",
            "Dropdowns display in full without over-stretching, and hovered items show a blue highlight",
            "Document title font size confirmed reduced to near body text",
        ],
    },
    {
        "version": "1.1.3",
        "date": "2026-06-24",
        "title": "卡拉OK式朗读高亮 · 进度条 · 选区朗读",
        "notes": [
            "朗读时按真实词级时间戳逐词高亮已读内容（青绿色），类似 MV 字幕",
            "选中部分文字后点朗读，只朗读选中的部分",
            "新增播放进度条，可拖动实时调整播放位置",
            "朗读按钮改为「开始/暂停/继续朗读」三态切换",
            "朗读中更改嗓音或语速，自动以新设置重读并跳回大致进度",
            "所有下拉框再加宽，鼠标悬停项目显示蓝色高亮",
            "更新/使用/关于文档标题字号再缩小一号",
            "macOS 程序图标四周留白，显示更精致协调",
        ],
        "title_en": "Karaoke-style playback highlighting · progress bar · selection playback",
        "notes_en": [
            "While reading, content already spoken is highlighted word by word in teal using real word-level timestamps, like music video subtitles",
            "Selecting text and pressing play reads only the selection",
            "New playback progress bar that can be dragged to change position in real time",
            "The play button became a three-state toggle: Play, Pause and Resume",
            "Changing voice or speed while reading automatically re-reads with the new settings and returns to roughly the same position",
            "All dropdowns widened further, with a blue highlight on hovered items",
        ],
    },
    {
        "version": "1.1.2",
        "date": "2026-06-24",
        "title": "修复 Argos 偶发翻译失败 · 下拉箭头改尖角号",
        "notes": [
            "修复 Argos 离线翻译时好时坏的问题：翻译前确保模型已就绪并自动重试",
            "下拉框右侧箭头改为扁平尖角号（V 形），无边框",
        ],
        "title_en": "Fixed intermittent Argos translation failures · chevron dropdown arrow",
        "notes_en": [
            "Fixed Argos offline translation working only intermittently: the model is now confirmed ready before translating, with automatic retries",
            "The dropdown arrow on the right became a flat chevron (V shape) without a border",
        ],
    },
    {
        "version": "1.1.1",
        "date": "2026-06-24",
        "title": "中英嗓音分离 · 界面与样式优化",
        "notes": [
            "朗读嗓音拆分为「英文嗓音」「中文嗓音」两个下拉，按文本语种自动取用",
            "英文嗓音不再被迫去读中文，避免乱兜底；两个嗓音选择均会记住",
            "第一排标签改为「目标语言」「翻译引擎」",
            "「使用说明」去掉 Readme 字样",
            "更新说明 / 使用说明 / 关于 的标题字号再缩小一号",
            "下拉框与滑竿的尖角图标改为扁平无边框风格",
            "修复 Windows 状态栏右下角灰色拖拽块",
        ],
        "title_en": "Separate Chinese and English voices · interface and styling improvements",
        "notes_en": [
            "Playback voices split into separate English Voice and Chinese Voice dropdowns, selected automatically by the text's language",
            "English voices are no longer forced to read Chinese, avoiding nonsensical fallbacks; both voice choices are remembered",
            "First row labels changed to Target Language and Translation Engine",
            "Removed the word Readme from the User Guide",
            "Change Log, User Guide and About titles reduced by one more size step",
            "Chevron icons on dropdowns and sliders changed to a flat, borderless style",
        ],
    },
    {
        "version": "1.1.0",
        "date": "2026-06-23",
        "title": "下拉框加宽 · 按钮顺序调整 · 模型仓库复用",
        "notes": [
            "源/目标语言、引擎、嗓音下拉框按最长内容精确加宽，文字完整显示",
            "原文 / 译文操作按钮顺序调整为：复制、粘贴、删除",
            "编译脚本支持公共模型仓库（~/EnglishCoach-models）：模型下载一次，所有版本复用",
            "模型下载改用 HTTP/1.1，规避 argos-net 的 HTTP/2 中断问题",
        ],
        "title_en": "Wider dropdowns · button order adjusted · shared model repository",
        "notes_en": [
            "Source/target language, engine and voice dropdowns widened precisely to their longest content so text displays in full",
            "Source and target action buttons reordered to Copy, Paste, Delete",
            "Build scripts support a shared model repository (~/EnglishCoach-models): models download once and are reused by every build",
            "Model downloads switched to HTTP/1.1 to work around HTTP/2 interruptions from argos-net",
        ],
    },
    {
        "version": "1.0.9",
        "date": "2026-06-23",
        "title": "Argos 离线翻译彻底打通 · 内置中英模型",
        "notes": [
            "Argos 离线翻译现已完全可用：中英互译，无需联网、无需 Key",
            "内置中英离线模型，安装即用，不再需要运行时下载",
            "彻底移除 PyTorch 依赖（通过兼容层让 Argos 在无 torch 环境运行）",
            "兼容 macOS Big Sur 等较老系统的依赖版本组合",
            "断网时也能用 Argos 完成中英翻译",
        ],
        "title_en": "Argos offline translation fully working · bundled Chinese-English models",
        "notes_en": [
            "Argos offline translation is now fully functional for Chinese-English in both directions, with no network and no key required",
            "Chinese-English offline models are bundled, so they work on install with no runtime download",
            "PyTorch dependency removed entirely (a compatibility layer lets Argos run without torch)",
            "Dependency versions chosen for compatibility with older systems such as macOS Big Sur",
            "Argos can complete Chinese-English translation even with no internet connection",
        ],
    },
    {
        "version": "1.0.8",
        "date": "2026-06-23",
        "title": "界面重排 · 朗读自动重试 · Argos 报错优化",
        "notes": [
            "顶部重排：第一排左侧为 源/目标语言与引擎，右侧为 设置/更新/帮助/关于",
            "原文 / 译文标题与 粘贴/复制/删除 按钮下移至「翻译」大按钮同一排",
            "朗读合成失败时静默重试，最多 3 次，仍失败才提示",
            "Argos 加载失败时显示真实原因，便于定位（区分源码运行 / 打包缺失）",
            "编译脚本新增 argostranslate 导入校验，装不上会提前明确报错",
            "下拉框加宽，自动检测等文字显示完整",
            "关于页邮箱与网址并列显示",
        ],
        "title_en": "Interface rearranged · automatic playback retry · clearer Argos errors",
        "notes_en": [
            "Top row rearranged: source/target languages and engine on the left, Settings/Change Log/Help/About on the right",
            "Source and target titles plus the Paste/Copy/Delete buttons moved down to the same row as the large Translate button",
            "Failed speech synthesis retries silently up to three times before showing a message",
            "Argos load failures now show the real cause, distinguishing running from source versus a missing bundled package",
            "Build scripts gained an argostranslate import check that fails early and clearly if installation went wrong",
            "Dropdowns widened so entries like Auto Detect display in full",
        ],
    },
    {
        "version": "1.0.7",
        "date": "2026-06-23",
        "title": "修复中文朗读与离线翻译 · 多项界面优化",
        "notes": [
            "修复中文朗读全部失败：读中文时自动切换到中文嗓音；新增 3 个中文嗓音",
            "修复 Argos 英译中报错：改用正确的翻译路径并处理英语中转",
            "Argos 中英模型随程序预置打包，安装后即可离线翻译（无需再下载）",
            "移除有道朗读引擎与「失败换嗓音」提示（已不再需要）",
            "朗读时更改嗓音或语速即时生效（自动以新设置重读当前栏）",
            "原文 / 译文区新增「粘贴」按钮",
            "源语言 / 目标语言 / 引擎下拉框自适应宽度，文字显示完整",
            "帮助与更新说明的标题字号缩小，更协调",
            "关于页开发者信息追加邮箱 vfx@Strilen.com",
        ],
        "title_en": "Fixed Chinese playback and offline translation · assorted UI improvements",
        "notes_en": [
            "Fixed Chinese playback failing entirely: a Chinese voice is now selected automatically for Chinese text, and three Chinese voices were added",
            "Fixed errors translating English to Chinese with Argos by using the correct translation path and handling English as a pivot",
            "Argos Chinese-English models are bundled with the program, so offline translation works right after install with no download",
            "Removed the Youdao playback engine and the switch-voice-on-failure prompt (no longer needed)",
            "Changing voice or speed during playback takes effect immediately, re-reading the current pane with the new settings",
            "New Paste button in the source and target areas",
        ],
    },
    {
        "version": "1.0.6",
        "date": "2026-06-23",
        "title": "翻译新增 Argos 离线 与 混元 引擎 · 双平台打包",
        "notes": [
            "翻译新增「Argos (离线)」引擎：纯本地、无需 Key、无需联网（首次用会提示下载语言模型）",
            "翻译新增「混元 HY-MT」引擎：腾讯混元在线翻译（需腾讯云 Key）",
            "翻译引擎增至 5 个：Google（默认）、DeepL、DeepSeek、Argos、混元",
            "同时提供 macOS 与 Windows 两套编译脚本",
        ],
        "title_en": "New Argos offline and Hunyuan engines · builds for both platforms",
        "notes_en": [
            "New Argos (offline) engine: entirely local, no key and no network required (language models are downloaded on first use)",
            "New Hunyuan HY-MT engine: Tencent Hunyuan online translation (requires a Tencent Cloud key)",
            "Translation engines increased to five: Google (default), DeepL, DeepSeek, Argos and Hunyuan",
        ],
    },
    {
        "version": "1.0.5",
        "date": "2026-06-23",
        "title": "朗读新增有道引擎 · 嗓音失败可改选",
        "notes": [
            "朗读新增「有道」嗓音（国内免费、无需 Key），作为微软 edge-tts 的备选",
            "微软 edge-tts 仍为默认，音质更佳、嗓音更多",
            "某嗓音合成失败时，弹出列表让你改选其它嗓音并立即重试",
            "朗读改用流式/二进制写入，更稳定",
        ],
        "title_en": "New Youdao playback engine · switch voices after a failure",
        "notes_en": [
            "New Youdao voices for playback (free in China, no key required) as an alternative to Microsoft edge-tts",
            "Microsoft edge-tts remains the default, with better audio quality and more voices",
            "When a voice fails to synthesize, a list appears so you can pick another and retry immediately",
            "Playback switched to streaming/binary writing for better stability",
        ],
    },
    {
        "version": "1.0.4",
        "date": "2026-06-23",
        "title": "兼容旧版 macOS · 界面与交互优化",
        "notes": [
            "兼容 macOS Big Sur：使用 PyQt6 6.4.x，解决 IOKit 符号缺失导致无法启动",
            "新增输入即译：停止输入约 0.8 秒后自动翻译（手动「翻译」按钮仍保留）",
            "左右两栏拉开间距，不再重叠；复制与删除按钮间距收紧",
            "引擎下拉中「Google」不再显示「(免费)」后缀",
            "朗读改用流式合成，更稳定；嗓音失效时给出明确提示",
            "移除「开发者介绍」页，作者信息并入「关于」（开发者：Strilen Liu）",
        ],
        "title_en": "Compatibility with older macOS · interface and interaction improvements",
        "notes_en": [
            "Compatible with macOS Big Sur: using PyQt6 6.4.x resolves a missing IOKit symbol that prevented startup",
            "New translate-as-you-type: translation runs automatically about 0.8 seconds after you stop typing (the manual Translate button remains)",
            "Left and right panes spaced apart so they no longer overlap; the gap between Copy and Delete tightened",
            "The Google entry in the engine dropdown no longer shows a (Free) suffix",
            "Playback switched to streaming synthesis for stability, with a clear message when a voice is unavailable",
            "Removed the Developer page; author information merged into About (developer: Strilen Liu)",
        ],
    },
    {
        "version": "1.0.3",
        "date": "2026-06-22",
        "title": "修复界面不弹出 · 兼容 macOS Big Sur · 全新图标",
        "notes": [
            "修复打包后无报错但窗口不显示的问题（强制窗口前置并抢占焦点）",
            "音频后端初始化失败不再阻止主界面启动（朗读不可用时自动降级）",
            "兼容 macOS 11 Big Sur：打包设置最低系统版本为 11.0",
            "启动异常时弹窗显示具体错误，便于排查",
            "全新应用图标：构图居中、字形圆润、双色配色",
        ],
        "title_en": "Fixed the window not appearing · macOS Big Sur compatibility · new icon",
        "notes_en": [
            "Fixed the packaged app showing no error but no window either (the window is now forced to the front and takes focus)",
            "Audio backend initialization failure no longer blocks the main window from starting (playback degrades gracefully)",
            "Compatible with macOS 11 Big Sur: builds now target a minimum system version of 11.0",
            "Startup errors are shown in a dialog with the specific cause, making diagnosis easier",
            "New application icon: centered composition, rounded letterforms and a two-color scheme",
        ],
    },
    {
        "version": "1.0.2",
        "date": "2026-06-22",
        "title": "改用 PyInstaller 打包并加入应用图标",
        "notes": [
            "macOS 打包改用 PyInstaller，更可靠地处理 PyQt6 的 Qt 插件，解决持续的启动崩溃",
            "新增应用图标（A/文 + 翻开的书），编译后 .app / Dock 显示专属图标",
            "窗口与任务栏同步使用新图标",
        ],
        "title_en": "Switched to PyInstaller packaging and added an application icon",
        "notes_en": [
            "macOS packaging moved to PyInstaller, which handles PyQt6's Qt plugins more reliably and resolves the persistent startup crashes",
            "New application icon (A/文 with an open book), shown for the built .app and in the Dock",
            "The window and taskbar use the new icon as well",
        ],
    },
    {
        "version": "1.0.1",
        "date": "2026-06-22",
        "title": "修复 macOS 打包启动崩溃",
        "notes": [
            "修复 py2app 打包后出现「Launch error」无法启动的问题",
            "打包时完整收录 PyQt6 插件（cocoa 平台、SVG、多媒体），解决 Qt 无法初始化",
            "编译脚本加入 conda 环境自动激活",
        ],
        "title_en": "Fixed the macOS packaged app crashing at startup",
        "notes_en": [
            "Fixed the Launch error that prevented the py2app build from starting",
            "Builds now include the full set of PyQt6 plugins (cocoa platform, SVG, multimedia), resolving Qt initialization failures",
            "Build script now activates the conda environment automatically",
        ],
    },
    {
        "version": "1.0.0",
        "date": "2026-06-22",
        "title": "首个正式版本",
        "notes": [
            "新增「翻译」功能：默认 Google 免费引擎（无需 Key），DeepL / DeepSeek 可作备选",
            "翻译支持中英双向与自动检测，可一键互换",
            "新增「朗读」功能：基于 edge-tts 在线语音合成，多嗓音、可调语速",
            "新增「版本更新说明与管理」面板",
            "新增「关于」「Readme / 帮助」「开发者介绍」页面",
            "内置 SVG 图标系统",
        ],
        "title_en": "First release",
        "notes_en": [
            "New translation feature: Google's free engine by default (no key required), with DeepL and DeepSeek as alternatives",
            "Translation supports Chinese-English in both directions plus auto-detection, with one-click swapping",
            "New playback feature: online speech synthesis via edge-tts with multiple voices and adjustable speed",
            "New change log and version management panel",
            "New About, Readme/Help and Developer pages",
            "Built-in SVG icon system",
        ],
    },
]

# =============================================================================
#  图标系统  (内置 SVG, 不依赖外部图片资源)
# =============================================================================

class Icons:
    """集中管理的内置 SVG 图标。颜色用 currentColor 占位，渲染时替换。"""

    _DEFS = {
        "app": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <path d="M4 19.5A2.5 2.5 0 0 1 6.5 17H20"/>
                  <path d="M6.5 2H20v20H6.5A2.5 2.5 0 0 1 4 19.5v-15A2.5 2.5 0 0 1 6.5 2z"/>
                  <path d="M9 7h6M9 11h6"/></svg>""",
        "translate": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <path d="m5 8 6 6"/><path d="m4 14 6-6 2-3"/><path d="M2 5h12"/>
                  <path d="M7 2h1"/><path d="m22 22-5-10-5 10"/><path d="M14 18h6"/></svg>""",
        "speak": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <polygon points="11 5 6 9 2 9 2 15 6 15 11 19 11 5"/>
                  <path d="M15.54 8.46a5 5 0 0 1 0 7.07"/>
                  <path d="M19.07 4.93a10 10 0 0 1 0 14.14"/></svg>""",
        "swap": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <path d="m16 3 4 4-4 4"/><path d="M20 7H4"/>
                  <path d="m8 21-4-4 4-4"/><path d="M4 17h16"/></svg>""",
        "settings": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <circle cx="12" cy="12" r="3"/>
                  <path d="M19.4 15a1.65 1.65 0 0 0 .33 1.82l.06.06a2 2 0 0 1-2.83 2.83l-.06-.06a1.65 1.65 0 0 0-1.82-.33 1.65 1.65 0 0 0-1 1.51V21a2 2 0 0 1-4 0v-.09A1.65 1.65 0 0 0 9 19.4a1.65 1.65 0 0 0-1.82.33l-.06.06a2 2 0 0 1-2.83-2.83l.06-.06a1.65 1.65 0 0 0 .33-1.82 1.65 1.65 0 0 0-1.51-1H3a2 2 0 0 1 0-4h.09A1.65 1.65 0 0 0 4.6 9a1.65 1.65 0 0 0-.33-1.82l-.06-.06a2 2 0 0 1 2.83-2.83l.06.06a1.65 1.65 0 0 0 1.82.33H9a1.65 1.65 0 0 0 1-1.51V3a2 2 0 0 1 4 0v.09a1.65 1.65 0 0 0 1 1.51 1.65 1.65 0 0 0 1.82-.33l.06-.06a2 2 0 0 1 2.83 2.83l-.06.06a1.65 1.65 0 0 0-.33 1.82V9a1.65 1.65 0 0 0 1.51 1H21a2 2 0 0 1 0 4h-.09a1.65 1.65 0 0 0-1.51 1z"/></svg>""",
        "info": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <circle cx="12" cy="12" r="10"/><path d="M12 16v-4"/><path d="M12 8h.01"/></svg>""",
        "history": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <path d="M3 3v5h5"/><path d="M3.05 13A9 9 0 1 0 6 5.3L3 8"/>
                  <path d="M12 7v5l4 2"/></svg>""",
        "help": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <circle cx="12" cy="12" r="10"/>
                  <path d="M9.09 9a3 3 0 0 1 5.83 1c0 2-3 3-3 3"/><path d="M12 17h.01"/></svg>""",
        "dev": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <polyline points="16 18 22 12 16 6"/><polyline points="8 6 2 12 8 18"/></svg>""",
        "stop": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <rect x="6" y="6" width="12" height="12" rx="2"/></svg>""",
        "copy": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <rect x="9" y="9" width="13" height="13" rx="2" ry="2"/>
                  <path d="M5 15H4a2 2 0 0 1-2-2V4a2 2 0 0 1 2-2h9a2 2 0 0 1 2 2v1"/></svg>""",
        "paste": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <path d="M16 4h2a2 2 0 0 1 2 2v14a2 2 0 0 1-2 2H6a2 2 0 0 1-2-2V6a2 2 0 0 1 2-2h2"/>
                  <rect x="8" y="2" width="8" height="4" rx="1" ry="1"/></svg>""",
        "clear": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <path d="M3 6h18"/><path d="M19 6v14a2 2 0 0 1-2 2H7a2 2 0 0 1-2-2V6"/>
                  <path d="M8 6V4a2 2 0 0 1 2-2h4a2 2 0 0 1 2 2v2"/></svg>""",
        "pause": """<svg viewBox="0 0 24 24" fill="{c}" stroke="none">
                  <rect x="6" y="5" width="4" height="14" rx="1"/>
                  <rect x="14" y="5" width="4" height="14" rx="1"/></svg>""",
        "export": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/>
                  <polyline points="16 6 12 2 8 6"/><line x1="12" y1="2" x2="12" y2="15"/></svg>""",
        "file": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/>
                  <polyline points="14 2 14 8 20 8"/></svg>""",
        "file_down": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/>
                  <polyline points="14 2 14 8 20 8"/><line x1="12" y1="12" x2="12" y2="18"/>
                  <polyline points="9 15 12 18 15 15"/></svg>""",
        "list": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <line x1="8" y1="6" x2="21" y2="6"/><line x1="8" y1="12" x2="21" y2="12"/>
                  <line x1="8" y1="18" x2="21" y2="18"/><line x1="3" y1="6" x2="3.01" y2="6"/>
                  <line x1="3" y1="12" x2="3.01" y2="12"/><line x1="3" y1="18" x2="3.01" y2="18"/></svg>""",
        "undo": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <path d="M9 14 4 9l5-5"/><path d="M4 9h10.5a5.5 5.5 0 0 1 5.5 5.5v0a5.5 5.5 0 0 1-5.5 5.5H11"/></svg>""",
        "redo": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <path d="M15 14 20 9l-5-5"/><path d="M20 9H9.5A5.5 5.5 0 0 0 4 14.5v0A5.5 5.5 0 0 0 9.5 20H13"/></svg>""",
        "download": """<svg viewBox="0 0 24 24" fill="none" stroke="{c}" stroke-width="1.8"
                  stroke-linecap="round" stroke-linejoin="round">
                  <path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/>
                  <polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>""",
    }

    @classmethod
    def icon(cls, name: str, color: str = None) -> QIcon:
        if color is None:
            color = "#1f1f22" if _theme_is_light() else "#e8e8e8"
        svg = cls._DEFS.get(name, "").format(c=color).encode("utf-8")
        pm = QPixmap()
        pm.loadFromData(svg, "SVG")
        return QIcon(pm)

    @classmethod
    def pixmap(cls, name: str, color: str = "#e8e8e8", size: int = 48) -> QPixmap:
        svg = cls._DEFS.get(name, "").format(c=color).encode("utf-8")
        pm = QPixmap()
        pm.loadFromData(svg, "SVG")
        return pm.scaled(size, size, Qt.AspectRatioMode.KeepAspectRatio,
                         Qt.TransformationMode.SmoothTransformation)


# =============================================================================
#  翻译引擎  (DeepSeek API, 后台线程)
# =============================================================================

LANG_OPTIONS = ["自动检测", "中文", "English"]

# 翻译引擎标识
ENGINE_GOOGLE = "Google -线上联网"
ENGINE_GOOGLE_API = "Google -API-Key联网"
ENGINE_DEEPL = "DeepL -API-Key联网"
ENGINE_DEEPSEEK = "DeepSeek -API-Key联网"
ENGINE_ARGOS = "Argos -离线本地"
ENGINE_HUNYUAN = "混元 -API-Key联网"
# LLM 多风格翻译引擎（均为 OpenAI 兼容 chat 接口）
ENGINE_OPENAI = "GPT -API-Key联网"
ENGINE_GEMINI = "Gemini -API-Key联网"
ENGINE_CLAUDE = "Claude -API-Key联网"
ENGINE_GLM = "GLM -API-Key联网"
ENGINE_ERNIE = "文心一言 -API-Key联网"
ENGINE_DOUBAO = "豆包 -API-Key联网"
ENGINE_QWEN = "通义千问 -API-Key联网"
ENGINE_KIMI = "Kimi -API-Key联网"

# 各引擎端点
GOOGLE_ENDPOINT = "https://translate.googleapis.com/translate_a/single"
GOOGLE_API_ENDPOINT = "https://translation.googleapis.com/language/translate/v2"  # 官方 Basic v2
DEEPL_ENDPOINT_FREE = "https://api-free.deepl.com/v2/translate"   # 免费版 Key
DEEPL_ENDPOINT_PRO = "https://api.deepl.com/v2/translate"          # 付费版 Key
DEEPSEEK_ENDPOINT = "https://api.deepseek.com/chat/completions"
# 腾讯混元 OpenAI 兼容端点（需腾讯云 Key）
HUNYUAN_ENDPOINT = "https://api.hunyuan.cloud.tencent.com/v1/chat/completions"

# LLM 引擎统一配置：均为 OpenAI 兼容 /chat/completions 接口
# key: 引擎常量；值: {endpoint, model, key_name(设置中保存的键), auth(bearer/x-api-key)}
LLM_ENGINES = {
    ENGINE_DEEPSEEK: {
        "endpoint": "https://api.deepseek.com/chat/completions",
        # deepseek-chat 已于 2026-07-24 15:59 UTC 停用，须用显式的 V4 名称。
        # deepseek-v4-flash 是原 deepseek-chat 对应的经济档；翻译任务不需要
        # 思考模式，用 extra 显式关闭以省时省钱（v4-flash 默认开启思考）。
        "endpoint_v4": True,
        "model": "deepseek-v4-flash", "key_name": "deepseek", "auth": "bearer",
        "label": "DeepSeek",
        "extra": {"thinking": {"type": "disabled"}},
    },
    ENGINE_OPENAI: {
        "endpoint": "https://api.openai.com/v1/chat/completions",
        "model": "gpt-4o-mini", "key_name": "openai", "auth": "bearer",
        "label": "OpenAI GPT",
    },
    ENGINE_GEMINI: {
        # Gemini 的 OpenAI 兼容端点
        "endpoint": "https://generativelanguage.googleapis.com/v1beta/openai/chat/completions",
        "model": "gemini-2.5-flash", "key_name": "gemini", "auth": "bearer",
        "label": "Google Gemini",
    },
    ENGINE_CLAUDE: {
        "endpoint": "https://api.anthropic.com/v1/messages",
        "model": "claude-sonnet-4-6", "key_name": "claude", "auth": "anthropic",
        "label": "Claude",
    },
    ENGINE_GLM: {
        "endpoint": "https://open.bigmodel.cn/api/paas/v4/chat/completions",
        "model": "glm-4-flash", "key_name": "glm", "auth": "bearer",
        "label": "智谱GLM-4-Flash",
    },
    ENGINE_ERNIE: {
        # 文心一言 OpenAI 兼容端点（千帆）
        "endpoint": "https://qianfan.baidubce.com/v2/chat/completions",
        "model": "ernie-4.5-turbo-128k", "key_name": "ernie", "auth": "bearer",
        "label": "文心一言",
    },
    ENGINE_DOUBAO: {
        "endpoint": "https://ark.cn-beijing.volces.com/api/v3/chat/completions",
        "model": "doubao-pro-32k", "key_name": "doubao", "auth": "bearer",
        "label": "字节豆包",
    },
    ENGINE_QWEN: {
        "endpoint": "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions",
        "model": "qwen-plus", "key_name": "qwen", "auth": "bearer",
        "label": "通义千问",
    },
    ENGINE_KIMI: {
        "endpoint": "https://api.moonshot.cn/v1/chat/completions",
        "model": "moonshot-v1-8k", "key_name": "kimi", "auth": "bearer",
        "label": "Kimi",
    },
    ENGINE_HUNYUAN: {
        "endpoint": HUNYUAN_ENDPOINT,
        "model": "hunyuan-turbo", "key_name": "hunyuan", "auth": "bearer",
        "label": "HunYuan",
    },
    '翻译': 'Translate',
    '文心一言': 'ERNIE',
    '豆包': 'Doubao',
    '字节豆包': 'Doubao',
    '通义千问': 'Qwen',
    '混元': 'Hunyuan',
    '晓贝': 'Xiaobei',
    '查看历史': 'View History',
    '文心一言 Key:': 'ERNIE Key:',
    '豆包 Key:': 'Doubao Key:',
    '通义千问 Key:': 'Qwen Key:',
    '混元 HY-MT Key:': 'Hunyuan HY-MT Key:',
    '百度千帆 Key': 'Baidu AI Studio API Key',
    '火山引擎 Key': 'Volcengine API Key',
    '阿里百炼 sk-...': 'Alibaba Cloud Model Studio API Key sk-...',
    '腾讯云混元 sk-...': 'Tencent Cloud Hunyuan sk-...',
    '显示API-Key': 'Show API Key',
    '显示密钥': 'Show Key',
    '查看日志': 'View Log',
    '语言': 'Language',
    '样式风格': 'Theme',
    '深色': 'Dark',
    '浅色': 'Light',
    '跟随系统': 'Follow System',
    '设置': 'Settings',
    '多风格翻译（LLM 引擎）': 'Multi-Style (LLM Engines)',
    'Google免费、无需Key、即开即用；DeepL/Google云需在下方填Key；Argos为纯离线引擎': 'Google: free, no key, ready to use. DeepL / Google Cloud need keys below. Argos runs fully offline.',
    '（主译文 + 书面/口语/俚语/美英式等多种风格辅助译法，仅大模型引擎有效）': '(Main translation + formal / casual / slang / US-UK style variants; LLM engines only)',
    '提示：Google免费无需Key；Argos纯离线；其余引擎需在上方填入对应 API Key': 'Tip: Google is free and key-less; Argos is fully offline; other engines need their API key above.',
    '免费版Key以...': 'Free-tier key starts with ...',
    'Google云翻译API Key': 'Google Cloud Translation API Key',
    '免费版 Key 以 :fx 结尾': 'Free-tier key ends with :fx',
    'Google 云翻译 Key (AIza...)': 'Google Cloud Translation Key (AIza...)',
    '（主译文 + 书面/口语/俚语/美英式等辅助译法）': '(Main + formal/casual/slang/US-UK style variants)',
    '点选一条记录后『载入并翻译』；悬停可见全文。': 'Click a record to Load & Translate; hover to see full text.',
}

# 所有引擎的显示顺序
ALL_ENGINES = [
    ENGINE_GOOGLE, ENGINE_GOOGLE_API, ENGINE_DEEPL, ENGINE_ARGOS,
    ENGINE_DEEPSEEK, ENGINE_OPENAI, ENGINE_GEMINI, ENGINE_CLAUDE,
    ENGINE_GLM, ENGINE_ERNIE, ENGINE_DOUBAO, ENGINE_QWEN, ENGINE_KIMI,
    ENGINE_HUNYUAN,
]
# 哪些是 LLM（支持多风格翻译）
LLM_ENGINE_SET = {
    ENGINE_DEEPSEEK, ENGINE_OPENAI, ENGINE_GEMINI, ENGINE_CLAUDE,
    ENGINE_GLM, ENGINE_ERNIE, ENGINE_DOUBAO, ENGINE_QWEN, ENGINE_KIMI,
    ENGINE_HUNYUAN,
}

# 语言名 -> 各引擎语言代码
_LANG_GOOGLE = {"中文": "zh-CN", "English": "en", "自动检测": "auto"}
_LANG_DEEPL = {"中文": "ZH", "English": "EN", "自动检测": None}
_LANG_ARGOS = {"中文": "zh", "English": "en"}   # Argos 用 ISO 639-1


class TranslateWorker(QThread):
    """多引擎翻译。默认 Google（免费、无需 Key）；DeepL / DeepSeek 需 Key。"""
    finished_ok = pyqtSignal(str)
    failed = pyqtSignal(str)

    # 常见标点/符号的中英文名称（翻译引擎对单个符号常无能为力，这里内置兜底）
    _SYMBOL_NAMES = {
        ".": ("dot", "点 / 句号"), ",": ("comma", "逗号"),
        "。": ("period (Chinese full stop)", "句号"), "，": ("comma", "逗号"),
        "?": ("question mark", "问号"), "？": ("question mark", "问号"),
        "!": ("exclamation mark", "感叹号"), "！": ("exclamation mark", "感叹号"),
        ";": ("semicolon", "分号"), "；": ("semicolon", "分号"),
        ":": ("colon", "冒号"), "：": ("colon", "冒号"),
        "(": ("left parenthesis", "左括号"), ")": ("right parenthesis", "右括号"),
        "（": ("left parenthesis", "左括号"), "）": ("right parenthesis", "右括号"),
        "[": ("left bracket", "左方括号"), "]": ("right bracket", "右方括号"),
        "{": ("left brace", "左花括号"), "}": ("right brace", "右花括号"),
        "<": ("less-than sign", "小于号"), ">": ("greater-than sign", "大于号"),
        "@": ("at sign", "艾特 / @符号"), "#": ("hash / pound sign", "井号"),
        "$": ("dollar sign", "美元符号"), "%": ("percent sign", "百分号"),
        "&": ("ampersand", "和号"), "*": ("asterisk", "星号"),
        "+": ("plus sign", "加号"), "-": ("hyphen / minus", "连字符 / 减号"),
        "=": ("equals sign", "等号"), "/": ("slash", "斜杠"),
        "\\": ("backslash", "反斜杠"), "|": ("vertical bar", "竖线"),
        "_": ("underscore", "下划线"), "~": ("tilde", "波浪号"),
        "^": ("caret", "脱字符"), "`": ("backtick", "反引号"),
        "'": ("apostrophe", "撇号"), '"': ("quotation mark", "引号"),
        "“": ("left quotation mark", "左引号"), "”": ("right quotation mark", "右引号"),
        "…": ("ellipsis", "省略号"), "—": ("em dash", "破折号"),
        "、": ("Chinese enumeration comma", "顿号"),
    }

    def __init__(self, text, src, tgt, engine, keys: dict, multi_style=False, parent=None):
        super().__init__(parent)
        self.text, self.src, self.tgt = text, src, tgt
        self.engine = engine
        self.keys = keys  # {"deepl": "...", "deepseek": "...", ...}
        self.multi_style = multi_style   # LLM 引擎下是否输出多风格翻译
        self._cancelled = False

    # ---- 目标语言解析：「自动检测」目标 = 中→英 / 其它→中 ----
    def _resolve_target(self):
        if self.tgt != "自动检测":
            return self.tgt
        # 含中文字符则译为英文，否则译为中文
        for ch in self.text:
            if "\u4e00" <= ch <= "\u9fff":
                return "English"
        return "中文"

    def _has_letters_or_cjk(self):
        """文本是否包含字母或中文（即可判断语言的内容）。"""
        for ch in self.text:
            if ch.isalpha() or "\u4e00" <= ch <= "\u9fff":
                return True
        return False

    def _try_symbol_fallback(self):
        """纯符号/标点输入时，返回内置的中/英文名称；无法判断语言时按设置或默认英文。
        返回译文字符串；若不适用则返回 None。"""
        stripped = self.text.strip()
        if not stripped or self._has_letters_or_cjk():
            return None   # 含字母/中文，交给翻译引擎
        # 目标语言：显式设置优先；自动检测时，纯符号默认英文
        if self.tgt == "English":
            want_en = True
        elif self.tgt == "中文":
            want_en = False
        else:  # 自动检测：纯符号无法判断 -> 默认英文
            want_en = True
        # 单个已知符号：直接给名称
        if len(stripped) == 1 and stripped in self._SYMBOL_NAMES:
            en, zh = self._SYMBOL_NAMES[stripped]
            return en if want_en else zh
        # 多个符号：逐个翻译拼接
        names = []
        for ch in stripped:
            if ch in self._SYMBOL_NAMES:
                en, zh = self._SYMBOL_NAMES[ch]
                names.append(en if want_en else zh)
            elif ch.strip():
                names.append(ch)   # 未知符号原样保留
        if names:
            return (", ".join(names) if want_en else "、".join(names))
        return None

    def cancel(self):
        self._cancelled = True

    def run(self):
        try:
            # 纯符号/标点：先用内置词典兜底（引擎对单符号常无能为力）
            sym = self._try_symbol_fallback()
            if sym is not None:
                if not getattr(self, "_cancelled", False):
                    self.finished_ok.emit(sym)
                return
            if self.engine == ENGINE_GOOGLE:
                out = self._run_google()
            elif self.engine == ENGINE_GOOGLE_API:
                out = self._run_google_api()
            elif self.engine == ENGINE_DEEPL:
                out = self._run_deepl()
            elif self.engine == ENGINE_ARGOS:
                out = self._run_argos()
            elif self.engine in LLM_ENGINES:
                # 所有 LLM 引擎走统一 OpenAI 兼容处理（含多风格翻译）
                out = self._run_llm(LLM_ENGINES[self.engine])
            else:
                self.failed.emit(f"未知翻译引擎: {self.engine}")
                return
            # 已被取消（用户重新点了翻译）则丢弃本次结果，不回填
            if getattr(self, "_cancelled", False):
                return
            self.finished_ok.emit(out.strip())
        except requests.exceptions.RequestException as e:
            self.failed.emit(f"网络错误: {e}\n{_network_hint(e)}")
        except RuntimeError as e:
            self.failed.emit(str(e))
        except Exception as e:
            self.failed.emit(f"翻译失败: {e}\n{traceback.format_exc()}")

    # ---- Google 免费端点（非官方，无需 Key）----
    def _run_google(self):
        sl = _LANG_GOOGLE.get(self.src, "auto")
        # 自动检测且含中日韩字符时显式声明源为中文：否则 Google 按多数字符判成
        # 英文，en->en 原样返回，夹杂的中文不被翻译
        if sl == "auto" and _text_is_chinese(self.text):
            sl = "zh-CN"
        tl = _LANG_GOOGLE.get(self._resolve_target(), "en")
        params = {
            "client": "gtx", "sl": sl, "tl": tl,
            "dt": "t", "q": self.text,
        }
        resp = requests.get(GOOGLE_ENDPOINT, params=params, timeout=30)
        if resp.status_code != 200:
            raise RuntimeError(f"Google 返回 {resp.status_code}。该免费端点可能不稳定，"
                               f"可改用设置中的 DeepL / DeepSeek 备选引擎。")
        data = resp.json()
        # 结构: [[["译文","原文",...], ...], ...]
        return "".join(seg[0] for seg in data[0] if seg[0])

    # ---- Google 官方 Cloud Translation Basic v2（需 API Key）----
    def _run_google_api(self):
        key = self.keys.get("google_api", "").strip()
        if not key:
            raise RuntimeError(
                "尚未配置 Google 云翻译 API Key。请到「设置」中填写。\n"
                "（Google Cloud Translation Basic v2，需在 Google Cloud 启用计费）")
        tl = _LANG_GOOGLE.get(self._resolve_target(), "en")
        params = {"key": key, "q": self.text, "target": tl, "format": "text"}
        if self.src != "自动检测":
            params["source"] = _LANG_GOOGLE.get(self.src, "")
        resp = requests.post(GOOGLE_API_ENDPOINT, params=params, timeout=30)
        if resp.status_code != 200:
            raise RuntimeError(f"Google 云翻译返回 {resp.status_code}: {resp.text[:200]}")
        data = resp.json()
        return data["data"]["translations"][0]["translatedText"]

    # ---- DeepL（需 Key，自动识别免费/付费）----
    def _run_deepl(self):
        key = self.keys.get("deepl", "").strip()
        if not key:
            raise RuntimeError("尚未配置 DeepL API Key。请到「设置」中填写。")
        endpoint = DEEPL_ENDPOINT_FREE if key.endswith(":fx") else DEEPL_ENDPOINT_PRO
        payload = {"text": [self.text], "target_lang": _LANG_DEEPL[self._resolve_target()]}
        if self.src != "自动检测":
            payload["source_lang"] = _LANG_DEEPL[self.src]
        resp = requests.post(
            endpoint,
            headers={"Authorization": f"DeepL-Auth-Key {key}",
                     "Content-Type": "application/json"},
            json=payload, timeout=30,
        )
        if resp.status_code != 200:
            raise RuntimeError(f"DeepL 返回 {resp.status_code}: {resp.text[:200]}")
        return resp.json()["translations"][0]["text"]

    # ---- 统一 LLM 翻译（OpenAI 兼容；支持多风格）----
    def _run_llm(self, cfg):
        key = self.keys.get(cfg["key_name"], "").strip()
        if not key:
            raise RuntimeError(
                f"尚未配置 {cfg['label']} API Key。请到「设置」中填写。")
        tgt = self._resolve_target()
        tgt_name = "英文" if tgt == "English" else "简体中文"

        if self.multi_style:
            _t = self.text.strip()
            _no_sep = not re.search(r"[\s，。！？、；：,.!?;:\n]", _t)
            _has_cjk = any("\u4e00" <= c <= "\u9fff" for c in _t)
            # 中文词≤4字，纯英文单词≤24字母（如 screenshot）
            _is_word = _no_sep and ((_has_cjk and len(_t) <= 4)
                                    or (not _has_cjk and len(_t) <= 24))
            if _is_word:
                # 单字/单词：第一部分只给唯一最优译法，其余备选放到多风格区
                system = (
                    "你是一名精通中英互译的词典专家。用户给出一个字或词，"
                    "请给出它在" + tgt_name + "中的译法。\n"
                    "输出严格分两部分（不要任何额外说明）：\n"
                    "第一部分：只有一行，写最常用、最贴切的那一个最优译法，"
                    "只写译文本身，不要编号、解释、音标或任何标注。\n"
                    "然后输出一个空行作为分隔（该空行不含任何文字或符号）\n"
                    "第二部分：其余备选译法，每行一个，格式『风格名：译文』或"
                    "『用法：译文』，按常用程度从高到低排序，给出 2-7 个；"
                    "不要重复第一部分那个译法，不要解释。")
                user = _t
                return (self._call_anthropic(cfg, key, system, user, 1.0)
                        if cfg["auth"] == "anthropic"
                        else self._call_openai_compat(cfg, key, system, user, 1.0))
            system = (
                "你是一名精通中英互译的语言老师。请把用户给的文本翻译成" + tgt_name +
                "。输出严格分两部分（不要任何额外说明）：\n"
                "第一部分：逐行对应的准确直译——原文有几行就输出几行，"
                "每行只有译文本身，不加任何标注，风格与普通翻译引擎一致。\n"
                "然后输出一个空行作为分隔（该空行不含任何文字或符号）\n"
                "第二部分：每行一个，格式『风格名：译文』，"
                "风格如 正式书面、口语随意、美式、英式、俚语 等，"
                "只给适用的 2-5 种，不要硬凑，不要解释。")
            user = f"请翻译为{tgt_name}：\n{self.text}"
            temperature = 1.0
        else:
            system = ("你是一名专业的中英翻译。只输出译文本身，不要添加任何解释、"
                      "标注、引号或前后缀。保持原文的语气与格式。")
            user = (f"把以下文本翻译成{'地道、自然的英文' if tgt=='English' else '自然流畅的简体中文'}。"
                    f"\n\n文本:\n{self.text}")
            temperature = 1.0

        # Claude 用 messages 接口，其余用 OpenAI chat 接口
        if cfg["auth"] == "anthropic":
            return self._call_anthropic(cfg, key, system, user, temperature)
        return self._call_openai_compat(cfg, key, system, user, temperature)

    def _call_openai_compat(self, cfg, key, system, user, temperature):
        _body = {"model": cfg["model"],
                 "messages": [{"role": "system", "content": system},
                              {"role": "user", "content": user}],
                 "temperature": temperature, "stream": False}
        _extra = cfg.get("extra")
        if _extra:
            _body.update(_extra)   # 如 DeepSeek 的 thinking=disabled
        resp = requests.post(
            cfg["endpoint"],
            headers={"Authorization": f"Bearer {key}",
                     "Content-Type": "application/json"},
            json=_body,
            timeout=60,
        )
        if resp.status_code != 200:
            raise RuntimeError(f"{cfg['label']} 返回 {resp.status_code}: {resp.text[:200]}")
        return resp.json()["choices"][0]["message"]["content"]

    def _call_anthropic(self, cfg, key, system, user, temperature):
        resp = requests.post(
            cfg["endpoint"],
            headers={"x-api-key": key,
                     "anthropic-version": "2023-06-01",
                     "Content-Type": "application/json"},
            json={"model": cfg["model"], "max_tokens": 1024,
                  "system": system,
                  "messages": [{"role": "user", "content": user}],
                  "temperature": temperature},
            timeout=60,
        )
        if resp.status_code != 200:
            raise RuntimeError(f"{cfg['label']} 返回 {resp.status_code}: {resp.text[:200]}")
        data = resp.json()
        # Anthropic 返回 content 是块列表
        parts = [b.get("text", "") for b in data.get("content", []) if b.get("type") == "text"]
        return "".join(parts)

    # ---- Argos Translate（纯离线，无需 Key / 联网）----
    def _run_argos(self):
        try:
            _ensure_argos_models(("en", "zh"))   # 锁定目录 + 确保模型已装
            import argostranslate.translate as at
            import argostranslate.package as ap
        except ImportError as e:
            raise RuntimeError(
                "Argos 离线翻译库加载失败。\n\n"
                f"真实原因：{e}\n\n"
                "若为源码运行：请在 EnglishCoach 环境执行 pip install argostranslate；\n"
                "若为打包版：说明编译时未成功安装该库或其依赖（ctranslate2 / sentencepiece），"
                "请重新编译并确认依赖安装无误。")
        tgt = self._resolve_target()
        if self.src == "自动检测":
            has_zh = any("\u4e00" <= c <= "\u9fff" for c in self.text)
            src_code = "zh" if has_zh else "en"
        else:
            src_code = _LANG_ARGOS[self.src]
        tgt_code = _LANG_ARGOS[tgt]
        if src_code == tgt_code:
            return self.text

        installed = at.get_installed_languages()
        by_code = {l.code: l for l in installed}
        if src_code not in by_code or tgt_code not in by_code:
            raise RuntimeError(
                "Argos 离线模型未就绪。\n\n"
                "请确认程序已内置中英模型，或在设置中重新安装离线模型。")

        # get_translation 偶发返回 None（语言对象匹配不稳定）。
        # 重试若干次，每次重新拉取已安装语言，提升稳定性。
        translation = None
        last_err = None
        for _ in range(3):
            try:
                langs = at.get_installed_languages()
                bc = {l.code: l for l in langs}
                if src_code in bc and tgt_code in bc:
                    translation = bc[src_code].get_translation(bc[tgt_code])
                    if translation is not None:
                        break
            except Exception as e:
                last_err = e
            import time
            time.sleep(0.2)

        if translation is None:
            raise RuntimeError(
                f"Argos 暂时无法建立 {src_code}→{tgt_code} 的翻译。\n"
                "请重试一次；若持续失败，请在设置中重新安装离线模型。")
        # Argos 对长文本（尤其多句）易出现"重复退化"（同一个字反复输出，
        # 如满屏"出出出"）。改为按句逐句翻译再拼接，显著缓解该问题，
        # 也让句子更短、更稳定。
        text = self.text
        # 按中英文句末标点切句（保留标点与换行）
        parts = re.split(r'(?<=[。！？；…\.\!\?;\n])', text)
        out = []
        for seg in parts:
            if not seg.strip():
                out.append(seg)        # 保留空白/换行
                continue
            try:
                tr = translation.translate(seg)
            except Exception:
                tr = seg
            # 兜底：若结果出现明显单字重复退化，回退原文该句
            tr = _strip_degenerate_repeat(tr)
            out.append(tr)
        return "".join(out)


def _strip_degenerate_repeat(s, threshold=8):
    """检测并修正 Argos 的重复退化：同一字符连续重复超过 threshold 次，
    压缩为单个，避免满屏重复字。"""
    import re
    if not s:
        return s
    # 把任意单字符连续重复 threshold 次以上压成 1 个
    return re.sub(r'(.)\1{' + str(threshold) + r',}', r'\1', s)


# =============================================================================
#  朗读引擎  (edge-tts -> mp3, 后台线程)
# =============================================================================

# 英文嗓音
# 嗓音定义：每个嗓音标注引擎（edge=线上联网, kokoro=本地离线）与 id。
# 显示名后缀：edge -> " -线上联网"；kokoro -> " -离线本地"
EN_VOICES = {
    "Aria (美音·女) -线上联网":        {"id": "en-US-AriaNeural",        "engine": "edge"},
    "Jenny (美音·女) -线上联网":       {"id": "en-US-JennyNeural",       "engine": "edge"},
    "Guy (美音·男) -线上联网":         {"id": "en-US-GuyNeural",         "engine": "edge"},
    "Christopher (美音·男) -线上联网": {"id": "en-US-ChristopherNeural", "engine": "edge"},
    "Sonia (英音·女) -线上联网":       {"id": "en-GB-SoniaNeural",       "engine": "edge"},
    "Ryan (英音·男) -线上联网":        {"id": "en-GB-RyanNeural",        "engine": "edge"},
    "Natasha (澳音·女) -线上联网":     {"id": "en-AU-NatashaNeural",     "engine": "edge"},
    # Kokoro 本地离线英文嗓音（a=美式, b=英式）。Heart 盲测最像真人，置顶为默认
    "Heart (美音·女) -离线本地":       {"id": "af_heart",   "engine": "kokoro", "lang": "a"},
    "Bella (美音·女) -离线本地":       {"id": "af_bella",   "engine": "kokoro", "lang": "a"},
    "Nova (美音·女) -离线本地":        {"id": "af_nova",    "engine": "kokoro", "lang": "a"},
    "Sarah (美音·女) -离线本地":       {"id": "af_sarah",   "engine": "kokoro", "lang": "a"},
    "Adam (美音·男) -离线本地":        {"id": "am_adam",    "engine": "kokoro", "lang": "a"},
    "Michael (美音·男) -离线本地":     {"id": "am_michael", "engine": "kokoro", "lang": "a"},
    "Emma (英音·女) -离线本地":        {"id": "bf_emma",    "engine": "kokoro", "lang": "b"},
    "George (英音·男) -离线本地":      {"id": "bm_george",  "engine": "kokoro", "lang": "b"},
}

# 中文嗓音
ZH_VOICES = {
    "晓晓 (普通话·女) -线上联网":      {"id": "zh-CN-XiaoxiaoNeural", "engine": "edge"},
    "云希 (普通话·男) -线上联网":      {"id": "zh-CN-YunxiNeural",    "engine": "edge"},
    "晓伊 (普通话·女) -线上联网":      {"id": "zh-CN-XiaoyiNeural",   "engine": "edge"},
    # Kokoro 本地离线中文嗓音（z=普通话）
    "晓贝 (普通话·女) -离线本地":      {"id": "zf_xiaobei",  "engine": "kokoro", "lang": "z"},
    "云健 (普通话·男) -离线本地":      {"id": "zm_yunjian",  "engine": "kokoro", "lang": "z"},
}

# 嗓音 -> 文件名用的简短名（中文用拼音首字母大写）
VOICE_SHORTNAME = {
    "Aria (美音·女) -线上联网": "Aria", "Jenny (美音·女) -线上联网": "Jenny",
    "Guy (美音·男) -线上联网": "Guy", "Christopher (美音·男) -线上联网": "Christopher",
    "Sonia (英音·女) -线上联网": "Sonia", "Ryan (英音·男) -线上联网": "Ryan",
    "Natasha (澳音·女) -线上联网": "Natasha",
    "Heart (美音·女) -离线本地": "Heart", "Bella (美音·女) -离线本地": "Bella",
    "Nova (美音·女) -离线本地": "Nova",
    "Sarah (美音·女) -离线本地": "Sarah", "Adam (美音·男) -离线本地": "Adam",
    "Michael (美音·男) -离线本地": "Michael", "Emma (英音·女) -离线本地": "Emma",
    "George (英音·男) -离线本地": "George",
    "晓晓 (普通话·女) -线上联网": "XiaoXiao", "云希 (普通话·男) -线上联网": "YunXi",
    "晓伊 (普通话·女) -线上联网": "XiaoYi",
    "晓贝 (普通话·女) -离线本地": "XiaoBei", "云健 (普通话·男) -离线本地": "YunJian",
}


def _text_is_chinese(text):
    """文本是否主要为中文。"""
    return any("\u4e00" <= c <= "\u9fff" for c in text)


def _install_combo_wheel_guard(combo):
    """让 QComboBox 未获焦点时不响应滚轮(避免滚动页面时误改选项)。
    做法：装事件过滤器，Wheel 事件在未聚焦时忽略并上抛给父控件。"""
    try:
        from PyQt6.QtCore import QObject, QEvent
        from PyQt6.QtWidgets import QComboBox as _QCB

        class _WheelGuard(QObject):
            def eventFilter(self, obj, ev):
                try:
                    if ev.type() == QEvent.Type.Wheel:
                        if not obj.hasFocus():
                            ev.ignore()
                            return True   # 拦下：不改选项，交给父级滚动
                except Exception:
                    pass
                return False

        g = _WheelGuard(combo)
        combo.installEventFilter(g)
        combo._wheel_guard = g            # 保持引用
        # 同时让它默认不通过滚轮获得焦点(点击/Tab 才聚焦)
        from PyQt6.QtCore import Qt as _Qt
        combo.setFocusPolicy(_Qt.FocusPolicy.StrongFocus)
    except Exception:
        pass


def _pin_popup_to_top(combo):
    """【仅用于设置窗的语言/主题下拉】弹出后把列表滚动位置归零。
    这两个下拉的弹出高度正好等于项数，本来就不需要滚动；但选中最后一项时
    Qt 会 scrollTo 让当前项可见，可能把视图滚下一行。而弹出的滚动条是关闭的、
    滚轮也被忽略，滚下去就卡住回不来 —— 于是显示成 items[1..N-1] 加一行空白，
    正是"整体向上串了一行、最后一项空白"的现象。归零即可复原。
    不改任何弹出机制，也不影响其它任何下拉。"""
    import sys as _sp
    if _sp.platform == "darwin" or combo is None:
        return                      # mac 走系统原生、现状完美，不做任何处理
    try:
        _orig = combo.showPopup

        def _show(_c=combo, _o=_orig):
            _o()
            try:
                v = _c.view()
                if v is not None:
                    sb = v.verticalScrollBar()
                    if sb is not None:
                        sb.setValue(0)
                    v.scrollToTop()
            except Exception:
                pass
        combo.showPopup = _show
    except Exception:
        pass


def _ensure_combo_items(combo, want):
    """仅在项【真的少了】时才重建，保留当前选择。
    只按身份值(userData)判断，不看显示文字——文字会随界面语言变化，
    拿文字比对会在每次重译时误触发重建，反而扰乱下拉。用于语言/主题下拉。"""
    if combo is None:
        return
    try:
        have = [combo.itemData(i) for i in range(combo.count())]
        want_data = [d for _t, d in want]
        if len(have) == len(want_data) and all(d in have for d in want_data):
            return          # 项齐全，什么都不做（不碰下拉，零副作用）
        cur = combo.currentData()
        combo.blockSignals(True)
        combo.clear()
        for _t, _d in want:
            combo.addItem(_t, _d)
        _i = combo.findData(cur) if cur is not None else -1
        combo.setCurrentIndex(_i if _i >= 0 else 0)
        combo.blockSignals(False)
    except Exception:
        try:
            combo.blockSignals(False)
        except Exception:
            pass


def _combo_popup_container_css():
    """下拉弹出容器的背景/边框，按当前深浅。与 _win_hybrid_qss 的下拉配色保持一致。
    工厂函数与主题热切换共用本函数——弹出背景此前只在下拉创建时设一次、切主题从不更新，
    导致切到浅色后弹出仍是黑底，而项目文字已变深色 => 深字黑底看不见 => 看似"丢了一项"。"""
    if _theme_is_light():
        return "background:#ffffff; border:1px solid #c4c4c8;"
    return "background:#2d2d30; border:1px solid #3a3a3a;"


def _refresh_combo_popups(root):
    """按当前深浅刷新 root 下所有下拉的弹出容器配色(非 mac；mac 走系统原生不设样式)。"""
    import sys as _s
    if _s.platform == "darwin" or root is None:
        return
    try:
        from PyQt6.QtWidgets import QComboBox
        css = _combo_popup_container_css()
        for cb in root.findChildren(QComboBox):
            try:
                v = cb.view()
                if v is None:
                    continue
                pop = v.parent()
                if pop is not None:
                    pop.setStyleSheet(css)
            except Exception:
                pass
    except Exception:
        pass


def _apply_combo_popup_style(combo):
    """给下拉应用与主界面同款的弹出view处理，使QSS的item:hover悬停高亮生效。
    mac下用自定义view+半透明，让 QComboBox QAbstractItemView::item:hover 规则作用到弹出列表。

    弹出宽度锚点 = "mac 悬停蓝条"的自然宽度：最长项文字 + QSS内边距(7px 14px→左右28)
    + 边框 + 勾号余量。不跟随被表单拉伸的闭合框。

    v2.14.3 重写要点（修复 Windows 点击无效 / 文字截断 …）：
      * 不再猴子补丁 combo.showPopup —— 重复应用会层层包裹，
        Windows 上导致弹出失败(点击没反应)；改用 view 的 minimumWidth，
        由 Qt 自己在弹出时排版。
      * 不再用 setFixedWidth —— 那是永久锁死，之后任何重新布局
        (语言切换、主题热切换)都无法调整，是"越改越窄、字显示不全"的元凶。
      * 宽度计算补齐 padding/边框/勾号，不再只算纯文字宽。
    """
    from PyQt6.QtWidgets import QFrame, QListView
    from PyQt6.QtCore import Qt
    lv = QListView()
    lv.setFrameShape(QFrame.Shape.NoFrame)
    lv.setUniformItemSizes(True)
    lv.setSpacing(0)
    lv.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
    lv.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
    lv.setAutoScroll(False)
    lv.setVerticalScrollMode(QListView.ScrollMode.ScrollPerItem)
    combo.setView(lv)
    combo.setMaxVisibleItems(max(1, combo.count()))
    lv.wheelEvent = lambda e: e.ignore()
    # 闭合框滚轮拦截：设置窗里用滚轮滚动内容时，鼠标划过下拉框常误改选项。
    # 让下拉框只在"已获得焦点"时才响应滚轮；未聚焦时把滚轮事件让给父级(滚动区)。
    _install_combo_wheel_guard(combo)
    # mac：弹出容器与列表都走系统原生（透明背景），保留悬停蓝条与圆角。
    # 注意这段必须在本函数内 —— 它引用 lv，v2.14.3 重写时被遗落到函数外，
    # 导致 mac 上一进设置窗就抛 NameError: name 'lv' is not defined 而闪退。
    popup = combo.view().parent()
    if popup is not None:
        popup.setContentsMargins(0, 0, 0, 0)
        import sys as _sys
        if _sys.platform == "darwin":
            lv.setStyleSheet("")
            lv.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground, True)
            lv.viewport().setAutoFillBackground(False)
            popup.setStyleSheet("")
            popup.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground, True)
    _fit_combo_popup_width(combo)


def _fit_combo_popup_width(combo):
    """按最长项内容设置弹出列表宽度（可反复调用：语言切换后重算）。
    用 minimumWidth 而非 setFixedWidth —— 后者永久锁死，无法再调整。"""
    try:
        from PyQt6.QtGui import QFontMetrics as _FM
        fm = _FM(combo.font())
        w = 0
        for i in range(combo.count()):
            w = max(w, fm.horizontalAdvance(combo.itemText(i)))
        if w <= 0:
            return
        # 文字 + 左右内边距(QSS: padding 7px 14px) + 边框 + 勾号/余量
        pw = w + 28 + 4 + 12
        v = combo.view()
        v.setMinimumWidth(pw)
        v.setMaximumWidth(16777215)      # 解除任何历史遗留的宽度锁
        cont = v.parentWidget()
        if cont is not None:
            cont.setMinimumWidth(pw)
            cont.setMaximumWidth(16777215)
    except Exception:
        pass


def fit_combo_width(combo, extra=0, popup_extra=0):
    """按最长项设为固定显示宽度。弹出列表比按钮略宽、完整显示、
    行距适中、无多余空白、边框均匀。
    extra=闭合框(第一部分)额外加宽；popup_extra=弹出列表(第二部分)额外加宽。"""
    from PyQt6.QtGui import QFontMetrics
    from PyQt6.QtWidgets import QFrame, QListView
    fm = QFontMetrics(combo.font())
    widest = 0
    for i in range(combo.count()):
        widest = max(widest, fm.horizontalAdvance(combo.itemText(i)))
    # 留足右侧下拉箭头 + 最小内边距（尽量紧凑，给交换钮居中腾空间）
    combo.setFixedWidth(widest + 52 + extra)   # +52基础，extra额外加宽
    combo.setFixedHeight(36)   # 与正方形按钮等高
    lv = QListView()
    lv.setFrameShape(QFrame.Shape.NoFrame)
    lv.setUniformItemSizes(True)
    lv.setSpacing(0)
    lv.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
    lv.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
    lv.setResizeMode(QListView.ResizeMode.Fixed)
    lv.setAutoScroll(False)
    lv.setVerticalScrollMode(QListView.ScrollMode.ScrollPerItem)
    combo.setView(lv)
    combo.view().setMinimumWidth(widest + 56 + popup_extra)
    # 关键：可见项数 = 实际项数，杜绝末尾空行 + 可滚动
    combo.setMaxVisibleItems(max(1, combo.count()))
    # 拦截弹出列表的滚轮事件，避免滚出多余空行
    lv.wheelEvent = lambda e: e.ignore()
    # 弹出容器：mac 走系统原生（圆角无框、深浅自适应）；其它平台深色单层边框
    popup = combo.view().parent()
    if popup is not None:
        popup.setContentsMargins(0, 0, 0, 0)
        import sys
        if sys.platform == "darwin":
            # mac：清空样式并让视图背景透明，回归系统原生下拉（圆角无框、无灰底块）
            lv.setStyleSheet("")
            lv.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground, True)
            lv.viewport().setAutoFillBackground(False)
            popup.setStyleSheet("")
            popup.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground, True)
        else:
            popup.setStyleSheet(_combo_popup_container_css())

def _split_for_tts(text, max_len=120):
    """把长文本按标点切成句子片段，每段不超过 max_len 字符，
    避免 Kokoro 单次推理超 token 上限被静默截断。中英文标点都作为切点。"""
    import re
    if not text:
        return []
    # 在中英文句末/分隔标点后切分（保留标点）
    pieces = re.split(r'(?<=[。！？；…\.\!\?;])\s*', text)
    segments = []
    buf = ""
    for p in pieces:
        if not p:
            continue
        # 单段仍超长，再按逗号/顿号等次级标点切
        if len(p) > max_len:
            subs = re.split(r'(?<=[，,、])\s*', p)
            for s in subs:
                if not s:
                    continue
                if len(buf) + len(s) > max_len and buf:
                    segments.append(buf); buf = ""
                # 仍然超长则硬切
                while len(s) > max_len:
                    segments.append(s[:max_len]); s = s[max_len:]
                buf += s
        else:
            if len(buf) + len(p) > max_len and buf:
                segments.append(buf); buf = ""
            buf += p
    if buf.strip():
        segments.append(buf)
    return segments or [text]


class TTSWorker(QThread):
    # 输出: 音频字节, 词边界列表[(char_start,char_end,offset_ms,dur_ms)]
    finished_ok = pyqtSignal(bytes, list)
    failed = pyqtSignal(str)

    def __init__(self, text, voice_spec, rate, char_offset=0, parent=None):
        super().__init__(parent)
        self.text = text
        # voice_spec: {"id":..., "engine":"edge"|"kokoro", "lang":...}
        self.voice_spec = voice_spec
        self.rate = rate
        self.char_offset = char_offset
        self._cancelled = False

    def cancel(self):
        self._cancelled = True

    def run(self):
        try:
            engine = self.voice_spec.get("engine", "edge")
            if engine == "kokoro":
                self._run_kokoro()
            else:
                self._run_edge()
        except Exception as e:
            if not self._cancelled:
                self.failed.emit(f"朗读失败：{e}")

    # ---- edge-tts（线上联网，音质最好）----
    def _run_edge(self):
        import time
        last_err = None
        rate_str = f"{'+' if self.rate >= 0 else ''}{self.rate}%"
        voice_id = self.voice_spec["id"]
        for attempt in range(3):
            if self._cancelled:
                return
            audio_buf = bytearray()
            boundaries = []
            search_pos = 0
            try:
                async def _synth():
                    nonlocal search_pos
                    comm = edge_tts.Communicate(self.text, voice_id, rate=rate_str)
                    got = False
                    async for chunk in comm.stream():
                        if self._cancelled:
                            return got
                        t = chunk.get("type")
                        if t == "audio" and chunk.get("data"):
                            audio_buf.extend(chunk["data"])
                            got = True
                        elif t == "WordBoundary":
                            word = chunk.get("text", "")
                            off_ms = chunk.get("offset", 0) / 10000.0
                            dur_ms = chunk.get("duration", 0) / 10000.0
                            idx = self.text.find(word, search_pos)
                            if idx < 0:
                                idx = self.text.find(word)
                            if idx >= 0:
                                cs = self.char_offset + idx
                                ce = cs + len(word)
                                search_pos = idx + len(word)
                                boundaries.append((cs, ce, off_ms, dur_ms))
                    return got

                got = asyncio.run(_synth())
                if self._cancelled:
                    return
                if got and len(audio_buf) > 0:
                    self.finished_ok.emit(bytes(audio_buf), boundaries)
                    return
                last_err = "未能合成语音（返回空音频）"
            except Exception as e:
                last_err = str(e)
            if self._cancelled:
                return
            time.sleep(0.5)
        if not self._cancelled:
            self.failed.emit(
                f"线上朗读失败（已重试 3 次）：{last_err}\n"
                "（edge-tts 需联网，大陆请确认网络/代理；或改用『-离线本地』嗓音）")

    # ---- Kokoro（本地离线，无需联网，原生对齐时间戳）----
    def _run_kokoro(self):
        try:
            import io, wave
            import numpy as np
            pipeline = _get_kokoro_pipeline(self.voice_spec.get("lang", "a"))
            if pipeline is None:
                reason = _KOKORO_LAST_ERROR or "未知原因"
                hint = ""
                if "Operation not permitted" in reason or "zoneinfo" in reason:
                    hint = ("\n\n这是 macOS 对『外置硬盘』的访问限制：你的 conda 装在外置 SSD（CineeSD）上，"
                            "系统不允许程序读取其中的时区文件。解决办法：\n"
                            "① 打开『系统设置 → 隐私与安全性 → 完全磁盘访问权限』，"
                            "把本程序（或终端 Terminal）加进去并打勾；\n"
                            "② 或把 EnglishCoach 的 conda 环境装到内置硬盘；\n"
                            "③ 已尝试自动改用系统时区库，若仍失败请用上面①。")
                self.failed.emit(
                    "本地离线朗读不可用：Kokoro 引擎未就绪。\n\n"
                    f"真实原因：{reason}{hint}\n\n"
                    "也可改用『-线上联网』嗓音（edge-tts）。")
                return
            voice_id = self.voice_spec["id"]
            speed = 1.0 + (self.rate / 100.0)   # rate -50..50 -> speed 0.5..1.5
            speed = max(0.5, min(2.0, speed))

            audio_chunks = []
            boundaries = []
            sample_rate = 24000
            search_pos = 0
            cum_samples = 0
            # Kokoro 单次推理有 token 上限（约 510），中文长文本若不切分会被静默截断，
            # 导致"声音没读完但进度跑满"。这里先按标点把长文本切成句子片段，逐段合成。
            segments = _split_for_tts(self.text)
            for seg_text in segments:
                if self._cancelled:
                    return
                if not seg_text.strip():
                    continue
                for result in pipeline(seg_text, voice=voice_id, speed=speed):
                    if self._cancelled:
                        return
                    audio = result.output.audio if hasattr(result, "output") and result.output else None
                    if audio is None:
                        continue
                    audio_np = audio.detach().cpu().numpy() if hasattr(audio, "detach") else np.asarray(audio)
                    chunk_start_ms = cum_samples / sample_rate * 1000.0
                    tokens = getattr(result, "tokens", None) or []
                    for tk in tokens:
                        word = getattr(tk, "text", "") or ""
                        st = getattr(tk, "start_ts", None)
                        et = getattr(tk, "end_ts", None)
                        if not word.strip() or st is None:
                            continue
                        idx = self.text.find(word, search_pos)
                        if idx < 0:
                            idx = self.text.find(word)
                        if idx >= 0:
                            cs = self.char_offset + idx
                            ce = cs + len(word)
                            search_pos = idx + len(word)
                            off_ms = chunk_start_ms + st * 1000.0
                            dur_ms = ((et - st) * 1000.0) if et else 200.0
                            boundaries.append((cs, ce, off_ms, dur_ms))
                    audio_chunks.append(audio_np)
                    cum_samples += len(audio_np)

            if self._cancelled:
                return
            if not audio_chunks:
                self.failed.emit("本地朗读未生成音频。")
                return
            # 拼接并编码为 WAV bytes
            full = np.concatenate(audio_chunks)
            full = np.clip(full, -1.0, 1.0)
            pcm16 = (full * 32767).astype("<i2")
            buf = io.BytesIO()
            with wave.open(buf, "wb") as w:
                w.setnchannels(1)
                w.setsampwidth(2)
                w.setframerate(sample_rate)
                w.writeframes(pcm16.tobytes())
            self.finished_ok.emit(buf.getvalue(), boundaries)
        except Exception as e:
            if not self._cancelled:
                import traceback
                self.failed.emit(f"本地朗读失败：{e}\n{traceback.format_exc()[:300]}")


# Kokoro pipeline 缓存（按语言）。首次使用时惰性创建。
_KOKORO_PIPELINES = {}
_KOKORO_AVAILABLE = None

_KOKORO_LAST_ERROR = ""

def _kokoro_available():
    global _KOKORO_AVAILABLE, _KOKORO_LAST_ERROR
    if _KOKORO_AVAILABLE is None:
        # 规避外置 SSD 上 conda 的 zoneinfo 权限问题（Operation not permitted）：
        # 让 Python 优先用系统时区库，并固定 TZ，避免 import 链路去读外置卷的时区文件。
        import os as _os
        _os.environ.setdefault("TZ", "Asia/Shanghai")
        # 系统时区目录优先（macOS/Linux 通常在 /usr/share/zoneinfo）
        sys_tz = "/usr/share/zoneinfo"
        if _os.path.isdir(sys_tz):
            _os.environ.setdefault("PYTHONTZPATH", sys_tz)
            try:
                import zoneinfo
                zoneinfo.reset_tzpath([sys_tz])
            except Exception:
                pass
        try:
            import time as _time
            _time.tzset()
        except Exception:
            pass
        try:
            import kokoro  # noqa
            _KOKORO_AVAILABLE = True
        except Exception as e:
            _KOKORO_AVAILABLE = False
            _KOKORO_LAST_ERROR = f"{type(e).__name__}: {e}"
    return _KOKORO_AVAILABLE

def _get_kokoro_pipeline(lang_code="a"):
    """惰性创建并缓存 Kokoro pipeline。失败返回 None，原因记入 _KOKORO_LAST_ERROR。"""
    global _KOKORO_LAST_ERROR
    if not _kokoro_available():
        return None
    if lang_code in _KOKORO_PIPELINES:
        return _KOKORO_PIPELINES[lang_code]
    try:
        # 优先使用【本地已有的模型】，避免联网下载。按顺序查找：
        #   1) 打包产物内置目录（PyInstaller 的 _MEIPASS，仅打包版有）
        #   2) 可执行文件/源码同级的 kokoro_model 目录
        #   3) 构建脚本使用的公共模型仓库 ~/EnglishCoach Models/Kokoro
        #      —— 从源码运行时此前完全没有这一步，导致明明本地有模型
        #         仍去联网下载。
        import os as _os
        _cands = []
        _mei = getattr(sys, "_MEIPASS", None)
        if _mei:
            _cands.append(_os.path.join(_mei, "kokoro_model"))
        try:
            _here = _os.path.dirname(_os.path.abspath(
                sys.executable if getattr(sys, "frozen", False) else __file__))
            _cands.append(_os.path.join(_here, "kokoro_model"))
        except Exception:
            pass
        _cands.append(_os.path.expanduser("~/EnglishCoach Models/Kokoro"))
        if _os.environ.get("ENGLISHCOACH_MODELS"):
            _cands.insert(0, _os.path.join(
                _os.environ["ENGLISHCOACH_MODELS"], "Kokoro"))
        for _b in _cands:
            # 认作可用需含实际权重文件，避免空目录让 HF 进入离线模式后报错
            if _os.path.isdir(_b) and any(
                    f.endswith((".pth", ".onnx", ".safetensors", ".bin"))
                    for _r, _d, _fs in _os.walk(_b) for f in _fs):
                _os.environ.setdefault("HF_HUB_OFFLINE", "1")
                _os.environ.setdefault("HF_HOME", _b)
                break
        # 中国大陆：huggingface.co 被封，Kokoro 首次下载模型会失败。
        # 若用户未自行设置 HF_ENDPOINT，且系统区域/语言为中国大陆，则默认走
        # hf-mirror.com 公益镜像，无需 VPN 即可下载（用户可用环境变量覆盖）。
        try:
            if not _os.environ.get("HF_ENDPOINT"):
                _cn = False
                # ① 语言环境。Linux 上除 LANG 外还要看 LC_ALL / LANGUAGE，
                #    getdefaultlocale 在部分环境下会返回 None。
                try:
                    import locale as _loc
                    _cands = [(_loc.getdefaultlocale()[0] or "")]
                except Exception:
                    _cands = []
                for _v in ("LC_ALL", "LC_CTYPE", "LANG", "LANGUAGE"):
                    _cands.append(_os.environ.get(_v, "") or "")
                if any(c.lower().replace("-", "_").startswith("zh_cn")
                       for c in _cands):
                    _cn = True
                # ② 时区。Linux 上 TZ 环境变量通常为空，真正的时区在
                #    /etc/timezone 或 /etc/localtime 的软链接目标里 —— 这正是
                #    大陆检测此前在 Linux 上失效、导致仍去连 huggingface.co 的原因。
                if not _cn:
                    _tzname = _os.environ.get("TZ", "")
                    if not _tzname:
                        try:
                            with open("/etc/timezone", encoding="utf-8") as _f:
                                _tzname = _f.read().strip()
                        except Exception:
                            try:
                                _tzname = _os.path.realpath("/etc/localtime")
                            except Exception:
                                _tzname = ""
                    if not _tzname:
                        try:                      # 兜底：Python 3.9+ 的时区库
                            from tzlocal import get_localzone_name as _gz
                            _tzname = _gz() or ""
                        except Exception:
                            pass
                    if any(k in _tzname for k in
                           ("Shanghai", "Chongqing", "Harbin", "Urumqi",
                            "Asia/Beijing", "PRC", "CST-8")):
                        _cn = True
                if _cn:
                    _os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
        except Exception:
            pass
        from kokoro import KPipeline
        # 有 CUDA GPU 则用 GPU 加速，否则 CPU（CUDA 版 torch 无卡时自动回退）
        device = None
        try:
            import torch
            if torch.cuda.is_available():
                device = "cuda"
        except Exception:
            pass
        # 若本地有扁平结构的模型目录（构建脚本用 local_dir 下载的产物），
        # 直接把权重与配置路径喂给 KModel —— 只设 HF_HOME 对这种结构无效，
        # 因为 hf_hub_download 找的是 hub/models--.../ 缓存布局，
        # 结果仍会去联网下载，正是"本地明明有模型却还要联网"的原因。
        _local_dir = _os.environ.get("ENGLISHCOACH_KOKORO_DIR", "")
        _kmodel = None
        if _local_dir and not _os.path.isdir(_os.path.join(_local_dir, "hub")):
            try:
                _cfg = _wt = None
                for _r, _d, _fs in _os.walk(_local_dir):
                    for _f in _fs:
                        if _f == "config.json" and _cfg is None:
                            _cfg = _os.path.join(_r, _f)
                        elif _f.endswith(".pth") and _wt is None:
                            _wt = _os.path.join(_r, _f)
                if _cfg and _wt:
                    from kokoro import KModel
                    _kmodel = KModel(config=_cfg, model=_wt)
                    if device:
                        _kmodel = _kmodel.to(device)
                    _kmodel = _kmodel.eval()
            except Exception as _e:
                print(f"[Kokoro] 本地模型加载失败，将回退默认方式: {_e}")
                _kmodel = None
        try:
            if _kmodel is not None:
                p = KPipeline(lang_code=lang_code, model=_kmodel, device=device)
            else:
                p = KPipeline(lang_code=lang_code, device=device)
        except TypeError:
            if _kmodel is not None:
                try:
                    p = KPipeline(lang_code=lang_code, model=_kmodel)
                except TypeError:
                    p = KPipeline(lang_code=lang_code)
            else:
                p = KPipeline(lang_code=lang_code)   # 老版本不支持 device 参数
        _KOKORO_PIPELINES[lang_code] = p
        return p
    except Exception as e:
        import traceback
        _KOKORO_LAST_ERROR = f"{type(e).__name__}: {e}"
        traceback.print_exc()
        return None


# =============================================================================
#  对话框: 设置 / 关于 / 更新说明 / 帮助 / 开发者
# =============================================================================

DOC_CSS = """
<style>
  body { font-family: 'Segoe UI','Microsoft YaHei',sans-serif; color:#dcdcdc;
         line-height:1.7; }
  h1 { color:#4ea1ff; font-size:13px; }
  h2 { color:#7bbcff; font-size:12px; margin-top:12px;
       border-bottom:1px solid #333; padding-bottom:2px; }
  h3 { color:#9ad; font-size:12px; }
  code { background:#2a2a2a; padding:1px 5px; border-radius:3px; color:#ffcb6b; }
  a { color:#4ea1ff; }
  .ver { color:#4ea1ff; font-weight:bold; }
  .date { color:#888; font-size:12px; }
  ul { margin-top:4px; }
</style>
"""

# QTextDocument 的默认样式表（setDefaultStyleSheet 用，标题字号明确压到接近正文）
DOC_STYLESHEET = """
body { color:#dcdcdc; line-height:170%; }
.t1 { color:#4ea1ff; font-size:15px; font-weight:bold; margin-bottom:10px; }
.t2 { color:#7bbcff; font-size:13px; font-weight:bold; }
.t3 { color:#9ad; font-size:13px; font-weight:bold; }
code { background:#2a2a2a; color:#ffcb6b; }
a { color:#4ea1ff; }
.ver { color:#4ea1ff; font-weight:bold; }
.date { color:#888; font-size:11px; }
"""


class SettingsDialog(QDialog):
    def __init__(self, settings: QSettings, parent=None):
        super().__init__(parent)
        self.settings = settings
        self.setWindowTitle(L("设置"))
        self.setMinimumWidth(440)
        self.resize(460, 640)
        self._apply_own_combo_style()

        outer = QVBoxLayout(self)
        # 内容较多，放进滚动区
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        content = QWidget()
        layout = QVBoxLayout(content)
        # 右侧多留白，避免滚动条压住文字
        layout.setContentsMargins(4, 4, 18, 4)
        scroll.setWidget(content)
        outer.addWidget(scroll, 1)

        # —— 默认翻译引擎 ——
        eng_label = QLabel(L("默认翻译引擎"))
        eng_label.setStyleSheet("font-weight:bold; color:#7bbcff; margin-top:4px;")
        layout.addWidget(eng_label)

        self.engine_combo = QComboBox()
        _combo_fill(self.engine_combo, ALL_ENGINES)
        _combo_select_data(self.engine_combo, 
            settings.value("engine", ENGINE_GOOGLE))
        self.engine_combo.setFixedHeight(36)   # 与主界面下拉等高
        _apply_combo_popup_style(self.engine_combo)
        layout.addWidget(self.engine_combo)


        # —— 备选引擎 Key ——
        key_label = QLabel(L("备选引擎 API Key（可选）"))
        key_label.setStyleSheet("font-weight:bold; color:#7bbcff; margin-top:6px;")
        layout.addWidget(key_label)

        form = QFormLayout()
        form.setRowWrapPolicy(QFormLayout.RowWrapPolicy.WrapLongRows)
        form.setFieldGrowthPolicy(QFormLayout.FieldGrowthPolicy.AllNonFixedFieldsGrow)
        form.setLabelAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)

        self.deepl_edit = QLineEdit(settings.value("deepl_key", ""))
        self.deepl_edit.setEchoMode(QLineEdit.EchoMode.Password)
        self.deepl_edit.setPlaceholderText(L("免费版 Key 以 :fx 结尾"))
        from PyQt6.QtWidgets import QSizePolicy as _SP
        self.deepl_edit.setSizePolicy(_SP.Policy.Expanding, _SP.Policy.Fixed)
        self.deepl_edit.setMinimumWidth(220)
        form.addRow("DeepL Key:", self.deepl_edit)

        self.google_api_edit = QLineEdit(settings.value("google_api_key", ""))
        self.google_api_edit.setEchoMode(QLineEdit.EchoMode.Password)
        self.google_api_edit.setPlaceholderText(L("Google 云翻译 Key (AIza...)"))
        self.google_api_edit.setSizePolicy(_SP.Policy.Expanding, _SP.Policy.Fixed)
        self.google_api_edit.setMinimumWidth(220)
        form.addRow(L("Google 云翻译 Key") + ":", self.google_api_edit)

        # 各 LLM 引擎的 Key 输入框（动态生成）
        self._key_edits = {}   # key_name -> QLineEdit
        _llm_key_rows = [
            ("deepseek", "DeepSeek Key:", "sk-..."),
            ("openai", "GPT Key:", "sk-..."),
            ("gemini", "Gemini Key:", "AIza..."),
            ("claude", "Claude Key:", "sk-ant-..."),
            ("glm", "GLM Key:", "xxxxxxxx.xxxxxxxx"),
            ("ernie", L("文心一言 Key:"), L("百度千帆 Key")),
            ("doubao", L("豆包 Key:"), L("火山引擎 Key")),
            ("qwen", L("通义千问 Key:"), L("阿里百炼 sk-...")),
            ("kimi", "Kimi Key:", "sk-..."),
            ("hunyuan", L("混元 HY-MT Key:"), L("腾讯云混元 sk-...")),
        ]
        for kn, label, ph in _llm_key_rows:
            e = QLineEdit(settings.value(f"{kn}_key", ""))
            e.setEchoMode(QLineEdit.EchoMode.Password)
            e.setPlaceholderText(ph)
            e.setSizePolicy(_SP.Policy.Expanding, _SP.Policy.Fixed)
            e.setMinimumWidth(220)
            form.addRow(label, e)
            self._key_edits[kn] = e
        # 兼容旧引用
        self.deepseek_edit = self._key_edits["deepseek"]
        self.hunyuan_edit = self._key_edits["hunyuan"]

        # 显示API-Key（与输入框左对齐；按下=显示且青色，弹起=隐藏灰色）
        self.show_keys_btn = QPushButton(L("显示密钥"))
        self.show_keys_btn.setCheckable(True)
        # 固定宽度(不随窗宽变)，但取按内容所需宽度确保文字完整
        self.show_keys_btn.setFixedWidth(BTN_W)
        self.show_keys_btn.toggled.connect(self._on_show_keys)
        form.addRow("", self.show_keys_btn)

        # 界面语言 / 样式风格（重启后生效）
        self.lang_combo = QComboBox()
        # 用 userData 存身份值(中文/English US)，显示文字与身份彻底解耦：
        # 即便界面重译改了显示文字，读 currentData() 仍拿到稳定身份，且构造时
        # 固定两项、不依赖任何文本匹配 -> 从根上杜绝"切换后丢项"。
        self.lang_combo.addItem("中文", "中文")
        self.lang_combo.addItem("English US", "English US")
        self.lang_combo.setProperty("no_retranslate", True)
        _lang_cur = settings.value("ui_lang", "中文")
        _lang_idx = self.lang_combo.findData(_lang_cur)
        self.lang_combo.setCurrentIndex(_lang_idx if _lang_idx >= 0 else 0)
        from PyQt6.QtWidgets import QSizePolicy as _SPl
        self.lang_combo.setFixedHeight(36)
        _apply_combo_popup_style(self.lang_combo)
        _pin_popup_to_top(self.lang_combo)
        self.lang_combo.setSizePolicy(_SPl.Policy.Expanding, _SPl.Policy.Fixed)
        self.lang_combo.setMinimumWidth(200)
        form.addRow(L("语言") + ":", self.lang_combo)

        def _lang_live(_=None):
            v = self.lang_combo.currentData() or self.lang_combo.currentText()
            self.settings.setValue("ui_lang", v)
            global _UI_LANG_CACHE
            _UI_LANG_CACHE = v
            _p = self.parent()
            if _p is not None and hasattr(_p, "retranslate_ui"):
                _p.retranslate_ui()
            if hasattr(self, "_retranslate_self"):
                self._retranslate_self()
        self.lang_combo.currentIndexChanged.connect(_lang_live)

        self.theme_combo = QComboBox()
        _combo_fill(self.theme_combo, ["深色", "浅色", "跟随系统"])
        _combo_select_data(self.theme_combo, settings.value("ui_theme", "跟随系统"))
        self.theme_combo.setFixedHeight(36)
        _apply_combo_popup_style(self.theme_combo)
        _pin_popup_to_top(self.theme_combo)
        self.theme_combo.setSizePolicy(_SPl.Policy.Expanding, _SPl.Policy.Fixed)
        self.theme_combo.setMinimumWidth(200)
        form.addRow(L("样式风格") + ":", self.theme_combo)

        def _theme_live(_=None):
            v = self.theme_combo.currentData() or "深色"
            self.settings.setValue("ui_theme", v)
            _p = self.parent()
            if _p is not None and hasattr(_p, "apply_theme"):
                _p.apply_theme()
            self._apply_own_combo_style()   # 设置窗自己的下拉也即时切换深浅
        self.theme_combo.currentIndexChanged.connect(_theme_live)

        # 多风格翻译开关（默认勾选）
        self.multi_style_chk = QCheckBox(L("LLM 引擎多风格翻译"))
        self.multi_style_chk.toggled.connect(
            lambda v: self.settings.setValue("multi_style", "true" if v else "false"))
        self.multi_style_chk.setChecked(settings.value("multi_style", "true") == "true")
        form.addRow("", self.multi_style_chk)
        # 保持程序置顶（在多风格与日志行之间）
        self.on_top_chk = QCheckBox(L("保持程序置顶"))
        self.on_top_chk.setChecked(settings.value("always_on_top", "false") == "true")

        def _on_top_live(v):
            # 槽函数必须自己吞掉异常：PyQt6 对槽里未捕获的异常直接 abort()，
            # 表现为"一点设置就闪退"(崩溃栈 pyqt6_err_print -> qAbort)。
            try:
                self.settings.setValue("always_on_top", "true" if v else "false")
                _p = self.parent()
                if _p is not None and hasattr(_p, "apply_always_on_top"):
                    _p.apply_always_on_top(bool(v))
            except Exception:
                _log_exc("on_top_live")
        self.on_top_chk.toggled.connect(_on_top_live)
        form.addRow("", self.on_top_chk)
        _gap = QWidget(); _gap.setFixedHeight(10)   # 与日志行隔开一点距离
        form.addRow("", _gap)

        log_btn = QPushButton(L("查看日志"))
        log_btn.setFixedWidth(BTN_W)
        log_btn.clicked.connect(self._open_log)
        exp_log_btn = QPushButton(L("导出日志"))
        exp_log_btn.setFixedWidth(BTN_W)          # 与查看日志等宽同风格
        exp_log_btn.clicked.connect(self._export_log)
        _logrow = QHBoxLayout(); _logrow.setContentsMargins(0, 0, 0, 0)
        _logrow.setSpacing(8)
        _logrow.addWidget(log_btn); _logrow.addWidget(exp_log_btn); _logrow.addStretch(1)
        _logw = QWidget(); _logw.setLayout(_logrow)
        form.addRow("", _logw)


        layout.addLayout(form)


        # 所有设置即时保存生效，只保留"关闭"按钮
        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)
        close_btn = QPushButton(L("关闭"))
        close_btn.setFixedWidth(BTN_W)
        close_btn.setStyleSheet("QPushButton{background:#1e88e5;border:none;border-radius:5px;color:white;}"\
            "QPushButton:hover{background:#2b95ef;}")
        def _close_and_save():
            self._persist_keys()
            self.accept()
        close_btn.clicked.connect(_close_and_save)
        self._close_btn = close_btn
        def _retranslate_self():
            to_lang = _ui_lang()
            retranslate_widget_tree(self, to_lang)
            self.show_keys_btn.setText(
                L("显示密钥") if not self.show_keys_btn.isChecked() else L("隐藏密钥"))
            # 保险：语言下拉必须恒为两项(中文/English US)。若因任何原因丢项，自愈重建，
            # 保留当前选择(按 userData 身份)。彻底防"切换后中文项消失"。
            try:
                lc = self.lang_combo
                _ensure_combo_items(
                    lc, [("中文", "中文"), ("English US", "English US")])
                # 主题下拉同样保证三项齐全(文字随界面语言，身份值不变)
                _tc = getattr(self, "theme_combo", None)
                if _tc is not None:
                    _ensure_combo_items(_tc, [(L("深色"), "深色"),
                                              (L("浅色"), "浅色"),
                                              (L("跟随系统"), "跟随系统")])
            except Exception:
                pass
        self._retranslate_self = _retranslate_self
        btn_row.addStretch()
        btn_row.addWidget(close_btn)
        outer.addLayout(btn_row)

    def _on_show_keys(self, on):
        self._toggle_echo(on)
        if on:
            self.show_keys_btn.setStyleSheet(
                "QPushButton{background:#5aa8b0;color:#0e2024;"
                "border:1px solid #5aa8b0;border-radius:5px;}")
        else:
            self.show_keys_btn.setStyleSheet("")

    def _toggle_echo(self, show):
        mode = QLineEdit.EchoMode.Normal if show else QLineEdit.EchoMode.Password
        self.deepl_edit.setEchoMode(mode)
        self.google_api_edit.setEchoMode(mode)
        for e in self._key_edits.values():
            e.setEchoMode(mode)

    def _export_log(self):
        """导出日志到用户选择的路径/文件名/格式：.txt / .log / .md / .json。
        全部用 Python 标准库实现(无第三方依赖)。"""
        import os, json, datetime
        from PyQt6.QtWidgets import QFileDialog, QMessageBox
        src = _log_path()
        try:
            raw = open(src, "r", encoding="utf-8", errors="replace").read() \
                if os.path.exists(src) else ""
        except Exception as e:
            QMessageBox.warning(self, L("导出日志"), f"{L('读取日志失败')}: {e}")
            return
        if not raw.strip():
            QMessageBox.information(self, L("导出日志"), L("日志为空，无内容可导出"))
            return
        stamp = datetime.datetime.now().strftime("%Y-%m-%d %H%M%S")
        default = os.path.join(os.path.expanduser("~"), f"EC LT {stamp}.txt")
        filt = ("Text (*.txt);;Log (*.log);;Markdown (*.md);;JSON (*.json);;"
                "All Files (*)")
        path, chosen = QFileDialog.getSaveFileName(
            self, L("导出日志"), default, filt)
        if not path:
            return
        ext = os.path.splitext(path)[1].lower()
        if not ext:   # 用户没打后缀：按所选过滤器补
            ext = (".log" if "Log" in (chosen or "") else
                   ".md" if "Markdown" in (chosen or "") else
                   ".json" if "JSON" in (chosen or "") else ".txt")
            path += ext
        try:
            lines = raw.splitlines()
            if ext == ".json":
                data = {
                    "app": "English Coach",
                    "version": APP_VERSION,
                    "exported_at": datetime.datetime.now().isoformat(timespec="seconds"),
                    "line_count": len(lines),
                    "lines": lines,
                }
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
            elif ext == ".md":
                head = (f"# English Coach {L('运行日志')}\n\n"
                        f"- **{L('版本')}**: v{APP_VERSION}\n"
                        f"- **{L('导出时间')}**: "
                        f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
                        "```log\n")
                with open(path, "w", encoding="utf-8") as f:
                    f.write(head + raw.rstrip("\n") + "\n```\n")
            else:   # .txt / .log / 其它：原样文本
                with open(path, "w", encoding="utf-8") as f:
                    f.write(raw)
        except Exception as e:
            QMessageBox.warning(self, L("导出日志"), f"{L('导出失败')}: {e}")
            return
        QMessageBox.information(
            self, L("导出日志"), f"{L('日志已导出到')}:\n{path}")

    def _open_log(self):
        """用系统默认程序打开日志文件。"""
        import os
        from PyQt6.QtGui import QDesktopServices
        from PyQt6.QtCore import QUrl
        p = _log_path()
        if not os.path.exists(p):
            try:
                open(p, "a", encoding="utf-8").close()
            except Exception:
                pass
        QDesktopServices.openUrl(QUrl.fromLocalFile(p))

    def _retheme(self):
        """主题热切换时被主窗 apply_theme 调用：按新深浅重建本窗自绘样式
        (下拉/按钮/输入框颜色)。此前缺此方法 → 设置窗开着切主题时颜色错乱。"""
        try:
            self._apply_own_combo_style()
            # show/hide 密钥按钮若处于按下态，其内联样式也随主题重置
            if hasattr(self, "show_keys_btn") and self.show_keys_btn.isChecked():
                self.show_keys_btn.setStyleSheet(
                    "QPushButton{background:#5aa8b0;color:#0e2024;"
                    "border:1px solid #5aa8b0;border-radius:5px;}")
            else:
                if hasattr(self, "show_keys_btn"):
                    self.show_keys_btn.setStyleSheet("")
            self.update()
        except Exception:
            _log_exc("settings_retheme")

    def _apply_own_combo_style(self):
        """(重)应用设置窗自己的下拉样式，按当前深浅。主题切换时可重调实现即时切换。"""
        import sys as _sysd
        if _sysd.platform != "darwin":
            # 非 mac 混合方案：设置窗直接套用主窗的 _win_hybrid_qss，做到与主界面
            # 一模一样，且只有一个深浅真相来源(setColorScheme)，杜绝改主题时下拉/按钮
            # 因残缺样式表覆盖而错乱(这正是"改深浅色才坏、重启又好"的根因)。
            _p = self.parent()
            if _p is not None and hasattr(_p, "_win_hybrid_qss"):
                self.setStyleSheet(_p._win_hybrid_qss())
            return
        _base_css = (_themed(_combo_popup_css())
            + "QLineEdit{border:1px solid #4a4a4a;border-radius:8px;padding:5px 10px;min-height:22px;}")
        if True:
            _light = _mac_current_is_light()
            if _light:
                _cbg, _ctx, _cbd, _arrow = "#ffffff", "#1f1f22", "#c4c4c8", "#5a5a5a"
            else:
                _cbg, _ctx, _cbd, _arrow = "#2d2d30", "#e8e8e8", "#3a3a3a", "#9aa0a6"
            _p = self.parent()
            _ch = _p._chevron_path(_arrow) if (_p is not None and hasattr(_p, "_chevron_path")) else ""
            _ch_hi = _p._chevron_path("#4ea1ff") if (_p is not None and hasattr(_p, "_chevron_path")) else ""
            _arrow_css = ""
            if _ch:
                _arrow_css = (f"QComboBox::down-arrow {{ image:url('{_ch}'); "
                              f"width:12px; height:8px; margin-right:8px; }}"
                              f"QComboBox::down-arrow:hover {{ image:url('{_ch_hi}'); }}")
            _base_css += (
                f"QComboBox {{ background:{_cbg}; color:{_ctx}; border:1px solid {_cbd};"
                f" border-radius:8px; padding:5px 10px; }}"
                f"QComboBox:hover {{ border:1px solid #4ea1ff; }}"
                f"QComboBox::drop-down {{ border:none; width:22px; }}"
                + _arrow_css
                + f"QComboBox QAbstractItemView {{ background:{_cbg}; outline:none; border:none;"
                  f" selection-background-color:#0e639c; selection-color:white; }}"
                  f"QComboBox QAbstractItemView::item {{ padding:7px 14px; border:none; }}"
                  f"QComboBox QAbstractItemView::item:selected {{ background:#0e639c; color:white; }}"
                  f"QComboBox QAbstractItemView::item:hover {{ background:#0e639c; color:white; }}")
        # 本窗自设样式表会覆盖 app 级规则，故把圆角滚动条一并带上
        # (Win10/Linux 原生滚动条是直角，mac/Win11 返回空串不影响原生)
        if not _native_scrollbar_platform():
            _base_css += "\n" + _rounded_scrollbar_qss()
        self.setStyleSheet(_base_css)

    def _persist_keys(self):
        """保存所有 API Key 与多风格开关(关闭设置窗时调用)。"""
        try:
            self.settings.setValue("deepl_key", self.deepl_edit.text().strip())
            self.settings.setValue("google_api_key", self.google_api_edit.text().strip())
            for kn, e in self._key_edits.items():
                self.settings.setValue(f"{kn}_key", e.text().strip())
            self.settings.setValue(
                "multi_style", "true" if self.multi_style_chk.isChecked() else "false")
        except Exception:
            pass

    def save(self):
        self.settings.setValue("engine", self.engine_combo.currentData())
        self.settings.setValue("deepl_key", self.deepl_edit.text().strip())
        self.settings.setValue("google_api_key", self.google_api_edit.text().strip())
        for kn, e in self._key_edits.items():
            self.settings.setValue(f"{kn}_key", e.text().strip())
        self.settings.setValue(
            "multi_style", "true" if self.multi_style_chk.isChecked() else "false")
        _old_lang = self.settings.value("ui_lang", "中文")
        _old_theme = self.settings.value("ui_theme", "跟随系统")
        _new_lang = self.lang_combo.currentText()
        _new_theme = self.theme_combo.currentData() or "深色"
        self.settings.setValue("ui_lang", _new_lang)
        self.settings.setValue("ui_theme", _new_theme)
        self.accept()
        # 立即生效：主题热切换；语言切换自动重启程序（免手动重启）
        try:
            if _new_theme != _old_theme:
                _p = self.parent()
                if _p is not None and hasattr(_p, "apply_theme"):
                    _p.apply_theme()
            if _new_lang != _old_lang:
                from PyQt6.QtCore import QProcess
                if getattr(sys, "frozen", False):
                    QProcess.startDetached(sys.executable, sys.argv[1:])
                else:
                    QProcess.startDetached(sys.executable, sys.argv)
                QApplication.instance().quit()
        except Exception:
            pass


class DocDialog(QDialog):
    """通用文档展示对话框 (关于 / 更新 / 帮助 / 开发者)。"""
    def __init__(self, title, html, parent=None, width=560, height=480):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.resize(width, height)
        layout = QVBoxLayout(self)
        browser = QTextBrowser()
        browser.setOpenExternalLinks(True)
        from PyQt6.QtGui import QFont
        browser.setFont(QFont("Microsoft YaHei", 10))
        # 用 setDefaultStyleSheet 让 QTextBrowser 可靠地应用字号（比内联 <style> 稳）
        browser.document().setDefaultStyleSheet(_themed(DOC_STYLESHEET))
        if not _native_scrollbar_platform():
            browser.setStyleSheet(_rounded_scrollbar_qss())   # Win10/Linux 圆角滚动条
        browser.setHtml(html)
        layout.addWidget(browser)
        btn_row = QHBoxLayout()
        close_btn = QPushButton(L("关闭"))
        close_btn.setFixedWidth(BTN_W)
        close_btn.setStyleSheet("QPushButton{background:#1e88e5;border:none;border-radius:5px;color:white;}"\
            "QPushButton:hover{background:#2b95ef;}")
        close_btn.clicked.connect(self.accept)
        btn_row.addStretch()
        btn_row.addWidget(close_btn)
        layout.addLayout(btn_row)


def changelog_html():
    _en = _ui_lang() == "English US"
    parts = ['<div class="t1">' + L('版本更新说明') + '</div><br>']
    for entry in CHANGELOG:
        parts.append(f'<div class="t2"><span class="ver">v{entry["version"]}</span>'
                     f' &nbsp;<span class="date">{entry["date"]}</span></div>')
        title = entry.get("title_en", entry["title"]) if _en else entry["title"]
        notes = entry.get("notes_en", entry["notes"]) if _en else entry["notes"]
        parts.append(f'<div class="t3">{title}</div><ul>')
        for n in notes:
            parts.append(f"<li>{n}</li>")
        parts.append("</ul>")
    return "".join(parts)


def about_html_en():
    return f"""
    <div class="t1">English Coach</div>
    <br>
    <p>A concise English assistant integrating <b>Translation</b> and <b>Text-to-Speech</b>.</p>
    <p><b>Version:</b> <span class="ver">v{APP_VERSION}</span></p>
    <p><b>Developer:</b> Strilen Liu</p>
    <p><a href="https://www.Strilen.com">www.Strilen.com</a>
       <a href="mailto:vfx@Strilen.com">vfx@Strilen.com</a></p>
    <div class="t2">Core Features</div>
    <ul>
      <li><b>Translation</b> - multiple engines: Google (free), DeepL, Argos (offline), plus
      LLM engines such as DeepSeek / OpenAI GPT / Gemini / Claude / GLM / ERNIE / Doubao /
      Qwen / Kimi / Hunyuan. LLM engines support multi-style translation (main rendering +
      formal / casual / slang / US-UK variants).</li>
      <li><b>Text-to-Speech</b> - dual engines edge-tts (online, high quality) and Kokoro
      (offline, no network); multiple Chinese/English voices, adjustable speed, karaoke
      word-by-word highlighting.</li>
    </ul>
    <div class="t2">Tech Stack</div>
    <ul>
      <li>UI: PyQt6</li>
      <li>Translation: traditional engines + multiple LLMs (unified via OpenAI-compatible API)</li>
      <li>Speech: edge-tts (online) + Kokoro (offline) + Qt Multimedia</li>
    </ul>
    <p class="date">(C) 2026 Strilen Liu. All rights reserved.</p>
    """


def about_html():
    return about_html_en() if _ui_lang() == "English US" else about_html_zh()


def about_html_zh():
    return f"""
    <div class="t1">English Coach</div>
    <div class="t1" style="margin-top:0;">英语导师</div><br>
    <p>一款简洁的英语助手工具，集成<b>翻译</b>与<b>朗读</b>两大核心功能。</p>
    <p><b>当前版本：</b><span class="ver">v{APP_VERSION}</span></p>
    <p><b>开发者：</b>Strilen Liu</p>
    <p><a href="https://www.Strilen.com">www.Strilen.com</a>
       <a href="mailto:vfx@Strilen.com">vfx@Strilen.com</a></p>
    <div class="t2">核心功能</div>
    <ul>
      <li><b>翻译</b> — 多引擎：Google（免费）、DeepL、Argos（离线）等传统引擎，外加 DeepSeek / OpenAI GPT / Gemini / Claude / 智谱GLM / 文心一言 / 豆包 / 通义千问 / Kimi / HunYuan 等大模型引擎；大模型引擎可开启多风格翻译（主译文 + 书面/口语/俚语/美英式等辅助译法）。</li>
      <li><b>朗读</b> — edge-tts（线上联网，音质佳）与 Kokoro（本地离线，无需联网）双引擎，多种中英嗓音、语速可调、卡拉OK逐词高亮。</li>
    </ul>
    <div class="t2">技术栈</div>
    <ul>
      <li>界面：PyQt6</li>
      <li>翻译：传统引擎 + 多家 LLM（OpenAI 兼容接口统一接入）</li>
      <li>语音：edge-tts（在线）+ Kokoro（离线）+ Qt Multimedia</li>
    </ul>
    <p class="date">© 2026 Strilen Liu. All rights reserved.</p>
    """


def readme_html_en():
    return f"""
    <div class="t1">User Guide</div><br>
    <div class="t2">System Requirements (please read before sharing)</div>
    <ul>
      <li><b>No Python, no dependencies, no conda needed.</b> The interpreter and every
      library are bundled into the app, so just download and run.</li>
      <li><b>macOS</b>: runs natively on Intel; Apple Silicon (M series) runs via Rosetta.
      Requires <b>macOS 11.0 Big Sur or newer</b>. Allow <b>2-3GB</b> of disk space
      (including the offline speech model).</li>
      <li><b>Windows</b>: 64-bit <b>Windows 10 / 11</b>. Allow 2-3GB of disk space.</li>
      <li><b>Linux</b>: no prebuilt binary yet — run from source (Python 3.10+ and the
      packages in requirements.txt).</li>
      <li><b>Unsigned app warning</b>: this app is not code-signed and may be blocked on
      first launch. On macOS open System Settings &rarr; Privacy &amp; Security and click
      Open Anyway, or run
      sudo xattr -rd com.apple.quarantine /Applications/EnglishCoach.app in Terminal.
      On Windows click More info &rarr; Run anyway at the SmartScreen prompt.</li>
      <li><b>GPU edition</b>: only worth installing on machines with an NVIDIA GPU; it
      bundles CUDA components and is considerably larger.</li>
    </ul>
    <div class="t2">What Needs a Network Connection</div>
    <ul>
      <li><b>Bundled, works offline immediately</b>: the program itself and all its
      dependencies.</li>
      <li><b>Downloaded once on first use</b>: the Kokoro offline speech model
      (about 330MB) and Argos offline language packs — fully offline afterwards.</li>
      <li><b>Always needs a network</b>: online translation engines and online
      text-to-speech.</li>
      <li><b>LLM engines</b> (GPT / Claude / Gemini / DeepSeek and others) run in the
      cloud and need your own API key plus a connection; nothing is bundled locally.</li>
    </ul>
    <div class="t2">Notes for Users in Mainland China</div>
    <ul>
      <li><b>Works without a VPN</b>: Chinese LLM engines (DeepSeek, ERNIE, Doubao, Qwen,
      Hunyuan, GLM) and Argos offline language packs.</li>
      <li><b>Kokoro speech model</b>: hosted on Hugging Face, which is not directly
      reachable from mainland China. The app automatically falls back to the
      hf-mirror.com community mirror, so the first download normally succeeds without a
      VPN. To choose your own mirror, set the HF_ENDPOINT environment variable.</li>
      <li><b>Needs a VPN</b>: Google and DeepL translation, and online text-to-speech
      (Microsoft Edge voices).</li>
    </ul>
    <div class="t2">Quick Start</div>
    <ul>
      <li>Type or paste text in the left box; the translation appears on the right.</li>
      <li>Click <b>Translate</b> to force a fresh translation with the current engine.</li>
      <li>Pick the engine and source/target languages from the top dropdowns; the swap
      button exchanges the two sides.</li>
    </ul>
    <div class="t2">Translation Engines</div>
    <ul>
      <li><b>Google -Online</b>: free, no key required, ready to use.</li>
      <li><b>DeepL / Google Cloud</b>: fill in the API key in Settings.</li>
      <li><b>Argos -Offline</b>: fully offline, no key or network.</li>
      <li><b>LLM engines</b> (DeepSeek / GPT / Gemini / Claude / GLM / ERNIE / Doubao /
      Qwen / Kimi / Hunyuan): fill your own API key; enable Multi-style for variant renderings.</li>
    </ul>
    <div class="t2">Text-to-Speech</div>
    <ul>
      <li>Each side has its own play / stop / download / clear-audio buttons and progress bar.</li>
      <li>Select text first to read only the selection; otherwise the whole text is read.</li>
      <li>Karaoke highlighting follows the reading; drag the progress bar to seek.</li>
      <li><b>edge-tts</b> needs network; <b>Kokoro</b> runs offline.</li>
    </ul>
    <div class="t2">Import / Export</div>
    <ul>
      <li>Import txt / md / docx / pdf into the source box; export translation to txt / docx.</li>
      <li>Translation history: reopen, reload &amp; translate, or export past records.</li>
    </ul>
    <div class="t2">Settings</div>
    <ul>
      <li>Language (Chinese / English US) and Theme (Dark / Light / Follow System) take effect immediately.</li>
      <li>API keys are stored locally only and hidden by default.</li>
    </ul>
    """


def readme_html():
    return readme_html_en() if _ui_lang() == "English US" else readme_html_zh()


def readme_html_zh():
    return f"""
    <div class="t1">使用说明</div><br>
    <div class="t2">系统要求（分享给他人前请阅读）</div>
    <ul>
      <li><b>macOS</b>：Intel 芯片可直接运行；Apple Silicon（M 系列）需通过 Rosetta 运行。
          系统需 <b>macOS 11.0 Big Sur 或更新</b>。建议预留 <b>2-3GB</b> 空间（含离线朗读模型）。</li>
      <li><b>Windows</b>：64 位 <b>Windows 10 / 11</b>。建议预留 2-3GB 空间。</li>
      <li><b>首次使用离线朗读（Kokoro）</b>：需联网下载语音模型（约 330MB），下载一次后即可完全离线。</li>
      <li><b>未签名应用提示</b>：本程序未签名，首次打开可能被系统拦截。
          macOS 可在「系统设置 → 隐私与安全性」点『仍要打开』，或终端执行
          sudo xattr -rd com.apple.quarantine /Applications/EnglishCoach.app；
          Windows 在 SmartScreen 提示点『更多信息 → 仍要运行』。</li>
      <li><b>无需安装 Python、依赖环境或 conda</b>：解释器与全部依赖已打包进程序，下载即用。</li>
      <li><b>Linux</b>：暂无预编译版本，需从源码运行（Python 3.10+ 与 requirements.txt 中的依赖）。</li>
      <li><b>GPU 版</b>：打包了 CUDA 组件、体积明显更大，仅在有 NVIDIA 显卡的机器上才有意义。</li>
    </ul>
    <div class="t2">哪些需要联网</div>
    <ul>
      <li><b>已打包、立即可用</b>：程序本体与全部依赖库。</li>
      <li><b>首次使用时下载一次</b>：Kokoro 离线朗读模型（约 330MB）、Argos 离线语言包；
      下载后即可完全离线使用。</li>
      <li><b>始终需要联网</b>：在线翻译引擎、在线朗读。</li>
      <li><b>LLM 引擎</b>（GPT / Claude / Gemini / DeepSeek 等）：模型在云端，需自备
      API Key 并联网，本地不打包任何模型。</li>
    </ul>
    <div class="t2">中国大陆用户须知</div>
    <ul>
      <li><b>无需 VPN 即可使用</b>：国内 LLM 引擎（DeepSeek、文心一言、豆包、通义千问、
      混元、智谱 GLM）与 Argos 离线语言包。</li>
      <li><b>Kokoro 朗读模型</b>：托管在 Hugging Face，大陆无法直连。程序会自动改用
      hf-mirror.com 公益镜像，通常无需 VPN 即可完成首次下载；如需指定其它镜像，
      可设置 HF_ENDPOINT 环境变量。</li>
      <li><b>需要 VPN</b>：Google、DeepL 翻译与在线朗读（微软 Edge 嗓音）。</li>
    </ul>
    <div class="t2">快速开始</div>
    <ul>
      <li>① 默认使用 <b>Google 免费引擎</b>，无需任何配置，打开即用。</li>
      <li>② 在左侧输入框粘贴或键入文本。</li>
      <li>③ 选择目标语言，点击 <b>翻译</b>，译文显示在右侧。</li>
      <li>④ 点击任一侧的 <b>朗读</b> 收听该框内文本。</li>
    </ul>
    <div class="t2">翻译引擎对比（请按需甄别）</div>
    <div class="t3">免费 / 离线引擎</div>
    <ul>
      <li><b>Google -线上联网</b>：网页版非官方接口。<b>免费、无需 Key</b>，即开即用。
          模型为 Google 网页翻译。<b>稳定性：不保证</b>——Google 未公开承诺此接口，
          可能随时变动或失效，大陆还需代理。适合日常随手翻。</li>
      <li><b>Google -API-Key联网</b>：官方 <b>Cloud Translation Basic (v2)</b>，NMT 神经翻译模型。
          需 Google Cloud 的 API Key 并<b>启用计费</b>。<b>收费：每月前 50 万字符免费</b>，
          超出按字符计费。<b>稳定性：官方保证，稳定可靠</b>。适合追求稳定的场景。</li>
      <li><b>Argos -离线本地</b>：纯本地离线，<b>无需 Key、无需联网</b>，免费。
          模型为 OpenNMT 离线中英模型（首次随程序内置或下载）。
          质量中等，胜在完全离线、隐私好。</li>
      <li><b>DeepL -API-Key联网</b>：以翻译质量著称。需 DeepL API Key
          （免费版 Key 以 :fx 结尾，每月 50 万字符免费；付费版更高额度）。稳定可靠。</li>
    </ul>
    <div class="t3">大模型（LLM）引擎 —— 支持「多风格翻译」</div>
    <p>以下引擎均为大语言模型，走各自官方 OpenAI 兼容接口，<b>需对应 API Key 且联网</b>。
       在「设置」勾选『多风格翻译』后，主译文下会附书面 / 口语 / 俚语 / 美式英式等多种译法，
       适合英语学习对比。各家收费、免费额度、稳定性以官网为准（政策常变）：</p>
    <ul>
      <li><b>DeepSeek</b>：中英俱佳、价格极低，国内可直接接入，推荐首选。</li>
      <li><b>智谱GLM-4-Flash</b>：中文地道，<b>该型号长期免费</b>，国内接入简单，推荐兜底。</li>
      <li><b>通义千问</b>：阿里，非英语表现强，有新用户免费额度，国内接入简单。</li>
      <li><b>Kimi</b>：月之暗面，不限 Token 仅限频率，适合慢翻长文。</li>
      <li><b>文心一言</b>：百度，中文理解精准，基础模型有免费额度。</li>
      <li><b>字节豆包</b>：火山引擎，有永久免费额度，接口稳定。</li>
      <li><b>HunYuan</b>：腾讯混元，需腾讯云 Key。</li>
      <li><b>OpenAI GPT</b>：能力顶尖，但大陆需海外信用卡 + 纯净海外 IP，封号风险高，门槛大。</li>
      <li><b>Google Gemini</b>：能力强，免费额度高，但大陆需代理，稳定性看网络。</li>
      <li><b>Claude</b>：能力强，大陆同样需海外信用卡 + 代理，门槛高。</li>
    </ul>
    <p style="color:#888;">提示：模型默认版本与端点可能随各家政策调整；若某引擎报错，多为 Key 未填、
       额度用尽、模型名变动或网络问题。Key 仅保存在本机，不上传。</p>
    <div class="t2">朗读引擎</div>
    <ul>
      <li><b>-线上联网（edge-tts）</b>：微软在线语音，音质最好，<b>需联网</b>（大陆需代理）。</li>
      <li><b>-离线本地（Kokoro）</b>：本地离线神经语音，<b>无需联网、CPU 即可</b>，
          英文质量好、原生时间戳让卡拉OK更准；中文为其支持语言但非强项。
          需随程序安装 Kokoro 及模型。</li>
    </ul>
    <div class="t2">选区联动</div>
    <p>在原文或译文一侧选中一段文字，程序会临时翻译该段并在另一侧用<b>灰色</b>自动高亮
       最匹配的区间（近似匹配，可能略有偏差）。灰色联动区间也可直接朗读。</p>
    <div class="t2">常见问题</div>
    <ul>
      <li><b>翻译/朗读报错？</b> 多为网络或 Key 问题；若使用了网络代理请确认其正在运行，
          或改用离线引擎（Argos 翻译 / Kokoro 朗读）。</li>
      <li><b>卡拉OK字幕？</b> 朗读时逐词高亮青蓝绿，跟随进度。</li>
      <li><b>Key 存在哪？</b> 保存在本机 (QSettings)，不上传。</li>
    </ul>
    """


def developer_html():
    return f"""
    <div class="t1">开发者介绍</div>
    <p><b>{APP_AUTHOR}</b> —— VFX 从业者，专注影视特效、合成（Nuke）与创意工具开发。</p>
    <div class="t2">联系方式</div>
    <ul>
      <li>邮箱：<a href="mailto:{APP_EMAIL}">{APP_EMAIL}</a></li>
      <li>个人站点：<a href="https://{APP_WEBSITE}">{APP_WEBSITE}</a></li>
    </ul>
    <div class="t2">关于本工具</div>
    <p>{APP_NAME} 源于日常英语阅读与学习需求，强调「轻量、本地、即用」，
       延续作者一贯的本地工具开发风格。</p>
    """


# =============================================================================
#  主窗口
# =============================================================================

_UI_LANG_CACHE = None

def _ui_lang():
    global _UI_LANG_CACHE
    if _UI_LANG_CACHE is None:
        from PyQt6.QtCore import QSettings as _QS
        _UI_LANG_CACHE = _QS("Strilen", "EnglishCoach").value("ui_lang", "中文")
    return _UI_LANG_CACHE

_EN = {
"极简界面":"Minimal UI","正常界面":"Normal UI","翻译引擎":"Engine","原文语言":"Source language",
"译文语言":"Target language","交换源文译文内容":"Swap source & target","原文文字":"Source text",
"译文文字":"Target text","在此输入或粘贴文本…":"__RAW__Type or paste text here…","译文显示在这里…":"__RAW__Translation appears here…",
"翻译":"Translate","复制":"Copy","粘贴":"Paste","清空":"Clear","导入文件":"Import file",
"导出当前原文":"Export source text","导出当前译文":"Export target text","导出翻译后文件":"Export translated file",
"载入上一条原文":"Load previous source","载入下一条原文":"Load next source","翻译历史":"History",
"中文嗓音":"Chinese voice","英文嗓音":"English voice","原文朗读进度":"Source playback progress",
"译文朗读进度":"Target playback progress","朗读原文":"Read source","朗读译文":"Read target",
"停止朗读原文":"Stop reading source","停止朗读译文":"Stop reading target",
"下载原文朗读音频":"Download source audio","下载译文朗读音频":"Download target audio",
"暂停朗读":"Pause","继续朗读":"Resume","朗读语速":"Speech rate","朗读语速 正常":"Speech rate: normal",
"设置":"Settings","更新说明":"Changelog","使用说明":"Help","关于":"About","关闭":"Close","取消":"Cancel",
"保存":"Save","检查历史":"View history","下载文档":"Download","载入":"Load","查看日志":"View log",
"显示API-Key":"Show API Keys","默认翻译引擎":"Default engine","备选引擎 API Key（可选）":"API Keys (Optional)",
"LLM 引擎多风格翻译":"Multi-Style Translation (LLM)","语言":"Language","样式风格":"Theme",
"浅色":"Light","深色":"Dark","跟随系统":"Follow system","翻译完成":"Translation done","播放中…":"Playing…",
"已停止":"Stopped","就绪":"Ready","请输入要翻译的文本":"Please enter text to translate",
"正在生成音频…":"Generating audio…","剪贴板为空":"Clipboard is empty","已复制选中部分":"Copied selection",
"已复制全部文字":"Copied all text","暂无历史记录":"No history yet","暂无历史原文":"No history sources",
"原文朗读音频已清空":"Source audio cache cleared","译文朗读音频已清空":"Target audio cache cleared",
"导出原文文字":"Export source text","导出译文文字":"Export target text","下载历史文档":"Download history",
"语言/样式风格更改将在重启后生效":"Language/theme changes take effect after restart",
}

_EN_SUB = [
    (" -线上联网", " -Online"), ("-线上联网", " -Online"),
    (" -API-Key联网", " -API Key"), ("-API-Key联网", " -API Key"),
    (" -纯离线", " -Offline"), ("-纯离线", " -Offline"),
    ("离线本地", "Offline Local"), ("自动检测", "Auto Detect"),
    ("(普通话·女)", "(Mandarin·F)"), ("(普通话·男)", "(Mandarin·M)"),
    ("(美音·女)", "(US·F)"), ("(美音·男)", "(US·M)"),
    ("(英音·女)", "(UK·F)"), ("(英音·男)", "(UK·M)"),
    ("(澳音·女)", "(AU·F)"), ("(澳音·男)", "(AU·M)"),
    ("晓晓", "Xiaoxiao"), ("云希", "Yunxi"), ("云扬", "Yunyang"),
    ("晓伊", "Xiaoyi"), ("云健", "Yunjian"), ("晓辰", "Xiaochen"),
]


def _tc(v):
    if v.startswith("__RAW__"):
        return v
    """英文词表值统一 Title Case（保留全大写缩写与非字母开头词）。"""
    return " ".join(
        w if (w.isupper() or not w[:1].isalpha()) else w[0].upper() + w[1:]
        for w in v.split(" "))


_EN = {k: _tc(v) for k, v in _EN.items()}
_EN["载入"] = "Reload"
_EN["载入翻译"] = "Reload"
_EN["显示密钥"] = "Show Key"
_EN["隐藏密钥"] = "Hide Key"
_EN['Google 免费、无需 Key、即开即用（推荐）。Argos 纯离线。其余 LLM 引擎需填对应 Key，并可开启多风格翻译。'] = "Google's free, no key needed, just use it (recommended). Argos is totally offline. For other LLM engines, you gotta fill in the key, and you can turn on multi-style translation."
_EN['提示：Google 免费无需 Key；Argos 纯离线无需 Key 与联网；其余 LLM 引擎需各自 API Key。勾选『多风格翻译』后，选用 LLM 引擎时会在主译文下给出多种风格译法。朗读 edge-tts 无需 Key。Key 仅保存在本机，不上传。'] = 'Note: Google is free and doesn\'t require a key; Argos is fully offline with no key or internet needed; other LLM models need their own API keys. Checking "Multi-style translation" will give you multiple style versions under the main translation when using LLM engines. Reading aloud with edge-tts requires no key. Keys are saved locally and not uploaded.'
_EN["重新载入"] = "Reload"
_EN["版本更新说明"] = "Changelog"
_EN["使用说明"] = "User Guide"
_EN["点选一条记录后『载入并翻译』；悬停可见全文。"] = "Click a record to Load & Translate; hover to see full text."
_EN["免费版 Key 以 :fx 结尾"] = "Free-tier key ends with :fx"
_EN["Google 云翻译 Key (AIza...)"] = "Google Cloud Translation Key (AIza...)"
_EN["Google 云翻译 Key"] = "Google Cloud Key"
_EN["版本更新说明"] = "Change Log"
_EN["关于 EnglishCoach"] = "About English Coach"
_EN["保持程序置顶"] = "Keep Window on Top"
_EN["导出日志"] = "Export Log"
_EN["读取日志失败"] = "Failed to read log"
_EN["日志为空，无内容可导出"] = "Log is empty, nothing to export"
_EN["导出失败"] = "Export failed"
_EN["日志已导出到"] = "Log exported to"
_EN["运行日志"] = "Runtime Log"
_EN["版本"] = "Version"
_EN["导出时间"] = "Exported at"
_EN["（主译文 + 书面/口语/俚语/美英式等辅助译法）"] = "(Main + formal/casual/slang/US-UK variants)"
for _k, _v in {"文心一言": "ERNIE", "字节豆包": "Doubao", "豆包": "Doubao",
               "通义千问": "Qwen", "混元": "Hunyuan", "晓贝": "Xiaobei"}.items():
    _EN_SUB.insert(0, (_k, _v))   # 插到最前，优先替换引擎主体名
# 无条件补齐（避免去重误判漏词）
for _k, _v in {
    "翻译": "Translate", "文心一言": "ERNIE", "字节豆包": "Doubao", "豆包": "Doubao",
    "通义千问": "Qwen", "混元": "Hunyuan", "晓贝": "Xiaobei", "查看历史": "View History",
    "文心一言 Key:": "ERNIE Key:", "豆包 Key:": "Doubao Key:", "通义千问 Key:": "Qwen Key:",
    "混元 HY-MT Key:": "Hunyuan HY-MT Key:", "百度千帆 Key": "Baidu AI Studio API Key",
    "火山引擎 Key": "Volcengine API Key",
    "阿里百炼 sk-...": "Alibaba Cloud Model Studio API Key sk-...",
    "腾讯云混元 sk-...": "Tencent Cloud Hunyuan sk-...",
    "显示API-Key": "Show API Key", "查看日志": "View Log",
    "语言": "Language", "样式风格": "Theme",
    "深色": "Dark", "浅色": "Light", "跟随系统": "Follow System", "设置": "Settings",
}.items():
    _EN[_k] = _v


def L(s):
    """界面文案本地化：English US 时查词表；查不到按子串规则翻译后缀。"""
    if _ui_lang() != "English US":
        return s
    if s in _EN:
        v = _EN[s]
        return v[7:] if v.startswith("__RAW__") else v
    out = s
    for _a, _b in _EN_SUB:
        out = out.replace(_a, _b)
    return out


# ==== 遍历式整体语言切换（根治：不逐条点名，遍历所有控件按文字查表替换）====
def _build_reverse_map():
    """英文->中文 反向映射(用于英->中切换)。"""
    rev = {}
    for zh, en in _EN.items():
        v = en[7:] if en.startswith("__RAW__") else en
        rev[v] = zh
    return rev

_ZH = _build_reverse_map()

def _translate_text(txt, to_lang):
    """把一段界面文字翻成目标语言。to_lang: 'English US' / '中文'。
    处理前后空格 + 尾部冒号：strip后查表，命中则拼回原装饰。"""
    if not txt:
        return txt
    # 提取前后空白
    lead = txt[:len(txt) - len(txt.lstrip())]
    trail = txt[len(txt.rstrip()):]
    core = txt.strip()
    if not core:
        return txt
    # 剥离尾部冒号(中/英)，翻译后拼回
    colon = ""
    if core and core[-1] in ("：", ":"):
        colon = core[-1]
        core = core[:-1]
    if to_lang == "English US":
        if core in _EN:
            v = _EN[core]
            v = v[7:] if v.startswith("__RAW__") else v
            return lead + v + colon + trail
        out = core
        for a, b in _EN_SUB:
            out = out.replace(a, b)
        return lead + out + colon + trail
    else:
        if core in _ZH:
            return lead + _ZH[core] + colon + trail
        out = core
        for a, b in _EN_SUB:
            out = out.replace(b, a)
        return lead + out + colon + trail

def retranslate_widget_tree(root, to_lang):
    """遍历 root 下所有控件，按当前显示文字翻成目标语言。
    覆盖 text/toolTip/placeholder/窗口标题/下拉项。不逐条点名，自动全覆盖。"""
    from PyQt6.QtWidgets import (QAbstractButton, QLabel, QLineEdit,
                                 QComboBox, QGroupBox, QWidget)
    widgets = root.findChildren(QWidget)
    widgets.append(root)
    for w in widgets:
        try:
            # 按钮/复选框/标签的文字
            if isinstance(w, (QAbstractButton, QLabel, QGroupBox)):
                t = w.text() if hasattr(w, "text") else ""
                if t:
                    nt = _translate_text(t, to_lang)
                    if nt != t:
                        w.setText(nt)
            # tooltip 气球
            tip = w.toolTip()
            if tip:
                ntip = _translate_text(tip, to_lang)
                if ntip != tip:
                    w.setToolTip(ntip)
            # 占位符：凡是有 placeholderText 的控件都处理(QLineEdit / QTextEdit /
            # QPlainTextEdit 等)，词表里没有的原样保留(不误伤 Key 示例格式)。
            if hasattr(w, "placeholderText") and hasattr(w, "setPlaceholderText"):
                ph = w.placeholderText()
                if ph:
                    nph = _translate_text(ph, to_lang)
                    if nph != ph:
                        w.setPlaceholderText(nph)
                        w.update()   # placeholder 仅在空内容时可见，强制重绘
            # 下拉项（保留当前选择）
            if isinstance(w, QComboBox):
                if w.property("no_retranslate"):
                    continue          # 身份值下拉(如语言选择)不参与重译
                _changed = False
                for i in range(w.count()):
                    it = w.itemText(i)
                    nit = _translate_text(it, to_lang)
                    if nit != it:
                        w.setItemText(i, nit)
                        _changed = True
                if _changed:
                    # 译后文字长度变了，闭合框与弹出列表都要按新文字重算宽度，
                    # 否则中英切换后会变窄、弹出项显示成 "-On…"（v2.14.3 修复）
                    try:
                        _refit_combo_width(w)
                    except Exception:
                        pass
        except Exception:
            pass


def _refit_combo_width(combo):
    """语言切换后按当前(新语言)文字重算下拉闭合框与弹出列表宽度。
    闭合框只增不减，避免中英来回切换时越切越窄。"""
    from PyQt6.QtGui import QFontMetrics as _FM
    fm = _FM(combo.font())
    w = 0
    for i in range(combo.count()):
        w = max(w, fm.horizontalAdvance(combo.itemText(i)))
    if w <= 0:
        return
    need = w + 52                      # 与 fit_combo_width 同款基础余量
    if combo.minimumWidth() < need:    # 只增不减，宽度始终保持一致
        combo.setMinimumWidth(need)
    if combo.width() < need:
        combo.setFixedWidth(need)
    _fit_combo_popup_width(combo)      # 弹出列表同步重算


def _combo_fill(combo, items):
    """下拉填充：显示文字走本地化 L()，userData 保存中文原值（逻辑比较不受语言影响）。"""
    for it in items:
        combo.addItem(L(it), it)


def _combo_select_data(combo, value):
    for i in range(combo.count()):
        if combo.itemData(i) == value or combo.itemText(i) == value \
                or L(combo.itemData(i) or "") == value:
            combo.setCurrentIndex(i)
            return
    if isinstance(value, str) and value:
        combo.setCurrentText(L(value))


def _mac_set_appearance(mode):
    """mac 专用：用 AppKit 设置整个 App 的原生外观（深浅），即时生效、无需重启。
    mode: '深色'/'浅色'/'跟随系统'。返回是否成功。
    这是 6.4.2 + Big Sur 上唯一可靠的原生深浅方案（setColorScheme 是 6.8+ 才有）。"""
    import sys
    if sys.platform != "darwin":
        return False
    try:
        from AppKit import NSApplication, NSAppearance
        if mode in ("深色", "Dark"):
            ap = NSAppearance.appearanceNamed_("NSAppearanceNameDarkAqua")
        elif mode in ("浅色", "Light"):
            ap = NSAppearance.appearanceNamed_("NSAppearanceNameAqua")
        else:
            ap = None   # 跟随系统
        NSApplication.sharedApplication().setAppearance_(ap)
        return True
    except Exception:
        return False


def _mac_current_is_light():
    """mac：读取 App 当前实际生效的外观是否为浅色。
    优先用 Qt 的 colorScheme(6.5+，Apple Silicon 6.11 可靠且不依赖 pyobjc)，
    不可用再回退 AppKit effectiveAppearance(Intel 6.4.2 走这条)。"""
    try:
        from PyQt6.QtWidgets import QApplication
        from PyQt6.QtCore import Qt as _Q
        app = QApplication.instance()
        if app is not None:
            sh = app.styleHints()
            if hasattr(sh, "colorScheme"):
                cs = sh.colorScheme()
                if cs == _Q.ColorScheme.Dark:
                    return False
                if cs == _Q.ColorScheme.Light:
                    return True
                # Unknown(跟随系统)：继续用 AppKit 判断真实系统深浅
    except Exception:
        pass
    try:
        from AppKit import NSApplication
        eff = NSApplication.sharedApplication().effectiveAppearance()
        name = eff.name() if eff else ""
        return "Dark" not in str(name)
    except Exception:
        # AppKit 不可用(未装 pyobjc)：退回读系统外观(Qt)
        try:
            from PyQt6.QtWidgets import QApplication
            from PyQt6.QtCore import Qt as _Q
            app = QApplication.instance()
            if app is not None:
                return (app.styleHints().colorScheme()
                        != _Q.ColorScheme.Dark)
        except Exception:
            pass
        return True


def _theme_is_light():
    """当前是否浅色。mac 下以 AppKit 实际生效外观为准（含跟随系统时的真实深浅）。"""
    import sys
    from PyQt6.QtCore import QSettings as _QS
    v = _QS("Strilen", "EnglishCoach").value("ui_theme", "跟随系统")
    if sys.platform == "darwin":
        if v in ("深色", "Dark"):
            return False
        if v in ("浅色", "Light"):
            return True
        # 跟随系统：读系统实际外观
        return _mac_current_is_light()
    # 非 mac：沿用 Qt colorScheme
    if v in ("跟随系统", "Follow system"):
        try:
            from PyQt6.QtGui import QGuiApplication
            from PyQt6.QtCore import Qt as _Qt
            cs = QGuiApplication.styleHints().colorScheme()
            # Unknown(未知)时保守当作浅色，避免在浅色系统上误用深色配色
            # 导致按钮浅字浅底看不清(#1)。只有明确 Dark 才判深色。
            return cs != _Qt.ColorScheme.Dark
        except Exception:
            return True
    return v in ("浅色", "Light")

# 深色样式 -> 浅色的颜色映射（仅非 mac 平台用；mac 走原生不涂）
_LIGHT_COLORS = {
    "#1e1e1e": "#f2f2f3", "#252526": "#ffffff", "#2d2d30": "#e9e9ea",
    "#3a3a3a": "#c9c9cc", "#4a4a4a": "#b5b5b8", "#2a2a2a": "#e2e2e4",
    "#dcdcdc": "#1f1f22", "#e0e0e0": "#1f1f22", "#e8e8e8": "#1f1f22",
    "#cccccc": "#333333", "#0e2024": "#083038",
    "#37373d": "#dcdce0",
    "#5aa8b0": "#00b3c6",
}

def _log_exc(where=""):
    """把异常写进日志，绝不向外抛——供槽函数兜底使用。"""
    try:
        import traceback
        _log_error(f"[EXC] {where}\n{traceback.format_exc()}")
    except Exception:
        pass


def _install_global_excepthook():
    """PyQt6 对槽函数中未捕获的异常会直接 abort()（表现为闪退，
    崩溃栈是 pyqt6_err_print -> QMessageLogger::fatal -> qAbort）。
    装一个全局钩子：记录日志并弹窗提示，让程序活下来而不是直接死掉。"""
    import sys as _sys
    import traceback as _tb

    def _hook(etype, value, tb):
        try:
            _log_error("[UNCAUGHT] " + "".join(
                _tb.format_exception(etype, value, tb)))
        except Exception:
            pass
        try:
            from PyQt6.QtWidgets import QApplication, QMessageBox
            if QApplication.instance() is not None:
                QMessageBox.warning(
                    None, "English Coach",
                    f"{value}\n\n(已记录到日志，程序将继续运行)")
        except Exception:
            pass
    _sys.excepthook = _hook


def _rounded_scrollbar_qss():
    """（已停用）历史上用于 Win10/Linux 自绘圆角滚动条。现决定各平台一律用
    系统原生滚动条，故返回空串，不再注入任何滚动条 QSS。"""
    return ""



def _safe_fusion(widget):
    """延后调用的安全包装：窗口可能已被销毁(RuntimeError)，静默跳过。"""
    try:
        if widget is not None and widget.isVisible():
            _force_fusion_scrollbars(widget)
    except RuntimeError:
        pass
    except Exception:
        pass


class _RoundScrollBarStyle:
    """方案C：自绘圆角滚动条的 QProxyStyle。

    为什么不用方案B(Fusion+QSS)：Qt 6.7+ 的 windows11 样式引擎原生绘制滚动条并
    忽略 QSS border-radius；即便把滚动条切到 Fusion，QSS 仍会经由 QStyleSheetStyle
    包装，未覆盖的子控件回退到平台样式，圆角时有时无。自绘则完全绕开样式引擎与
    QSS 的层叠，在任何 Qt 版本/平台上都画出一致的圆角胶囊。

    仅在需要自绘的平台(Win10及以下/Linux)启用；mac 与 Win11 原生已是圆角胶囊。
    实现为惰性类工厂：QProxyStyle 必须在 QApplication 创建后才能子类化，
    故用函数在运行时构造。
    """
    _cls = None


def _make_round_scrollbar_style():
    """运行时构造并缓存 QProxyStyle 子类（需 QApplication 已存在）。"""
    if _RoundScrollBarStyle._cls is not None:
        return _RoundScrollBarStyle._cls
    from PyQt6.QtWidgets import QProxyStyle, QStyle
    from PyQt6.QtCore import Qt, QRectF
    from PyQt6.QtGui import QPainter, QColor, QBrush

    class _RoundSB(QProxyStyle):
        def drawComplexControl(self, cc, opt, painter, widget=None):
            if cc == QStyle.ComplexControl.CC_ScrollBar:
                self._draw_scrollbar(opt, painter, widget)
                return
            super().drawComplexControl(cc, opt, painter, widget)

        def _draw_scrollbar(self, opt, painter, widget):
            # 画法：先铺一层浅色圆角轨道，再画一个明确对比的圆角滑块。
            # 关键（v2.14.8 修复 Win10 看不清）：不再用半透明色——半透明依赖底下
            # 背景才有对比，Win10 背景色不同就糊成一片。改用不透明实色，并按深浅
            # 主题选配色，滑块与轨道之间保证足够对比。
            groove = self.subControlRect(
                QStyle.ComplexControl.CC_ScrollBar, opt,
                QStyle.SubControl.SC_ScrollBarGroove, widget)
            slider = self.subControlRect(
                QStyle.ComplexControl.CC_ScrollBar, opt,
                QStyle.SubControl.SC_ScrollBarSlider, widget)

            # 判断深浅主题：看控件背景色亮度
            dark = False
            try:
                from PyQt6.QtGui import QPalette
                base = (widget.palette().color(QPalette.ColorRole.Base)
                        if widget is not None else None)
                if base is not None:
                    # 亮度 < 128 视为深色主题
                    dark = (base.red() * 299 + base.green() * 587
                            + base.blue() * 114) / 1000 < 128
            except Exception:
                pass

            if dark:
                track_col = QColor(255, 255, 255, 28)   # 深色主题：极浅白轨道
                slider_col = QColor(120, 120, 120)      # 不透明中灰滑块
                slider_hover = QColor(160, 160, 160)
            else:
                track_col = QColor(0, 0, 0, 22)          # 浅色主题：极浅黑轨道
                slider_col = QColor(136, 136, 136)       # 不透明中深灰滑块(够醒目)
                slider_hover = QColor(100, 100, 100)

            painter.save()
            painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)
            painter.setPen(Qt.PenStyle.NoPen)

            # 轨道：浅色圆角背景（画在 groove 内，留 1px 边距）
            gr = QRectF(groove).adjusted(1, 1, -1, -1)
            if gr.width() > 0 and gr.height() > 0:
                grad = (gr.width() if gr.width() < gr.height() else gr.height()) / 2.0
                painter.setBrush(QBrush(track_col))
                painter.drawRoundedRect(gr, grad, grad)

            # 滑块：留 2px 边距，不透明圆角胶囊
            if slider.isValid() and slider.width() > 0 and slider.height() > 0:
                horizontal = slider.width() > slider.height()
                r = QRectF(slider).adjusted(2, 2, -2, -2)
                hovered = bool(opt.state & QStyle.StateFlag.State_MouseOver)
                painter.setBrush(QBrush(slider_hover if hovered else slider_col))
                radius = (r.height() if horizontal else r.width()) / 2.0
                painter.drawRoundedRect(r, radius, radius)
            painter.restore()

        def pixelMetric(self, metric, opt=None, widget=None):
            from PyQt6.QtWidgets import QStyle as _QS
            # 隐藏两端箭头按钮（胶囊风格无箭头，和 mac 一致）
            if metric == _QS.PixelMetric.PM_ScrollBarExtent:
                return 12
            return super().pixelMetric(metric, opt, widget)

    _RoundScrollBarStyle._cls = _RoundSB
    return _RoundSB


def _force_fusion_scrollbars(widget):
    """（已停用）历史上给滚动条套自绘/Fusion 圆角样式。现决定各平台一律用
    系统原生滚动条，本函数改为空操作，保留以兼容旧调用点。"""
    return



def _native_scrollbar_platform():
    """当前平台是否自带圆角胶囊滚动条(macOS 全版本 / Win11+)。"""
    import sys
    if sys.platform == "darwin":
        return True
    if sys.platform == "win32":
        try:
            if sys.getwindowsversion().build >= 22000:
                from PyQt6.QtWidgets import QStyleFactory
                if "windows11" in [k.lower() for k in QStyleFactory.keys()]:
                    return True
        except Exception:
            pass
    return False


def _qt_set_color_scheme(app, v):
    """用 Qt 原生 styleHints().setColorScheme() 设深浅。Qt 6.5+ 才有该 API，
    6.9+ 在 macOS/Windows 上可靠生效。返回是否成功。"""
    try:
        from PyQt6.QtCore import Qt as _Q
        sh = app.styleHints()
        if not hasattr(sh, "setColorScheme"):
            return False
        if v in ("跟随系统", "Follow system"):
            sh.setColorScheme(_Q.ColorScheme.Unknown)
        else:
            sh.setColorScheme(_Q.ColorScheme.Light if v in ("浅色", "Light")
                              else _Q.ColorScheme.Dark)
        return True
    except Exception:
        return False


def _apply_win_palette(app):
    """非 mac：按当前主题给 app 设深/浅调色板。启动与主题热切换共用同一套，
    保证两条路径完全一致——这是原生控件(复选框边线/背景)正确显示深浅的关键。
    (此前热切换漏了设调色板，导致"打开好、改主题坏、复选框没边线"。)"""
    from PyQt6.QtGui import QPalette, QColor
    if _theme_is_light():
        _pc = {"Window": "#f2f2f3", "WindowText": "#1f1f22", "Base": "#ffffff",
               "AlternateBase": "#ececec", "Text": "#1f1f22", "Button": "#e9e9ea",
               "ButtonText": "#1f1f22", "Highlight": "#3a6ea5", "HighlightedText": "#ffffff",
               "ToolTipBase": "#f7f7f7", "ToolTipText": "#1f1f22", "PlaceholderText": "#9a9a9a"}
    else:
        _pc = {"Window": "#1e1e1e", "WindowText": "#dcdcdc", "Base": "#252526",
               "AlternateBase": "#2d2d30", "Text": "#dcdcdc", "Button": "#2d2d30",
               "ButtonText": "#dcdcdc", "Highlight": "#3a6ea5", "HighlightedText": "#ffffff",
               "ToolTipBase": "#2d2d30", "ToolTipText": "#e0e0e0", "PlaceholderText": "#777777"}
    pal = app.palette()
    for _rn, _cv in _pc.items():
        pal.setColor(getattr(QPalette.ColorRole, _rn), QColor(_cv))
    # 原生复选框/单选钮的勾选色取自【强调色】。不设的话会跟随系统强调色
    # (用户系统若设成黄色，勾选就是黄色)。这里固定为程序蓝，保持原生渲染不变。
    if hasattr(QPalette.ColorRole, "Accent"):        # Qt 6.6+ 才有
        pal.setColor(QPalette.ColorRole.Accent, QColor("#1e88e5"))
    app.setPalette(pal)


def _apply_color_scheme(app):
    """按用户主题设置 App 外观。
    mac：优先用 Qt 原生 setColorScheme(6.5+；Apple Silicon 上的 6.11 可用，
         不依赖 pyobjc)，不可用再回退 AppKit(Intel+6.4.2 走这条)。
    非 mac：用 Qt colorScheme。"""
    import sys
    from PyQt6.QtCore import QSettings as _QS
    v = _QS("Strilen", "EnglishCoach").value("ui_theme", "跟随系统")
    if sys.platform == "darwin":
        # 先试 Qt 原生(新版 mac/6.11)；不行再用 AppKit(老 Intel/6.4.2)
        if not _qt_set_color_scheme(app, v):
            _mac_set_appearance(v)
        else:
            # Qt 原生成功后也顺带调一次 AppKit(若可用)，让窗口标题栏等
            # 系统装饰同步深浅；AppKit 不可用则忽略。
            _mac_set_appearance(v)
        return
    _qt_set_color_scheme(app, v)


def _combo_popup_css():
    """mac 走系统原生下拉(圆角无框、深浅自适应)，其它平台用自定义蓝色高亮样式。"""
    import sys
    if sys.platform == "darwin":
        return ""
    return ("""
            QComboBox QAbstractItemView { background:#2d2d30; outline:none; border:1px solid #3a3a3a;
                selection-background-color:#0e639c; selection-color:white; }
            QComboBox QAbstractItemView::item { padding:7px 14px; border:none; }
            QComboBox QAbstractItemView::item:selected { background:#0e639c; color:white; }
            QComboBox QAbstractItemView::item:hover { background:#0e639c; color:white; }
""")


def _tooltip_css():
    """mac 走系统原生气球(尖角、深浅自适应)，其它平台自定义。"""
    import sys
    if sys.platform == "darwin":
        return ""
    return ("""
            QToolTip { background:#2d2d30; color:#e0e0e0; border:1px solid #4a4a4a;
                padding:2px 5px; font-size:11px; border-radius:6px; }
""")


def _themed(css):
    """非 mac 平台：浅色时对样式表整体换色。
    mac 平台：不改色——深浅由 AppKit 原生外观驱动，QSS 只提供无关配色的结构，
    避免自涂颜色与原生外观打架（这是之前深浅混乱的根源）。"""
    import sys
    if sys.platform == "darwin":
        return css
    if not _theme_is_light():
        return css
    for _k, _v in _LIGHT_COLORS.items():
        css = css.replace(_k, _v)
    return css


def _xml_safe(s):
    """去掉 docx/XML 不允许的控制字符（NULL、\\x0b 等），修复导出 docx 报错。"""
    import re as _r
    return _r.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', s or '')


def _sentence_spans(text):
    """按中英文句末标点把文本切成句子，返回 [(start,end), ...]（含位置）。"""
    import re
    if not text:
        return []
    spans = []
    start = 0
    for m in re.finditer(r'[。！？；…，\.\!\?;,\u3000\n]+', text):
        end = m.end()
        if text[start:end].strip():
            spans.append((start, end))
        start = end
    if start < len(text) and text[start:].strip():
        spans.append((start, len(text)))
    return spans or [(0, len(text))]


def _proportional_span(src_full, s0, s1, tgt_full):
    """把原文 [s0,s1) 选区按"段落对齐 + 段内比例"映射到译文区间。
    不调用翻译引擎。中英文标点都作为段落切点，先定位选区落在哪个句段，
    再在译文对应句段内按字符比例取区间，英文吸附整词。"""
    import re
    if not src_full or not tgt_full:
        return None
    s0 = max(0, min(s0, len(src_full)))
    s1 = max(s0, min(s1, len(src_full)))
    # 句段切分（保留位置）
    def segments(text):
        segs = []
        start = 0
        for m in re.finditer(r'[。！？；…，\.\!\?;,\u3000\n]+', text):
            end = m.end()
            segs.append((start, end))
            start = end
        if start < len(text):
            segs.append((start, len(text)))
        return segs or [(0, len(text))]
    src_segs = segments(src_full)
    tgt_segs = segments(tgt_full)
    # 选区中心落在哪个源句段
    mid = (s0 + s1) / 2.0
    si = 0
    for i, (a, b) in enumerate(src_segs):
        if a <= mid < b:
            si = i; break
    else:
        si = len(src_segs) - 1
    # 对应译文句段（按段索引比例映射，段数可能不同）
    ti = int(round(si / max(1, len(src_segs) - 1) * (len(tgt_segs) - 1))) if len(src_segs) > 1 else 0
    ta, tb = tgt_segs[min(ti, len(tgt_segs) - 1)]
    sa, sb = src_segs[si]
    seg_len = max(1, sb - sa)
    # 段内比例
    r0 = (s0 - sa) / seg_len
    r1 = (s1 - sa) / seg_len
    r0 = max(0.0, min(1.0, r0)); r1 = max(0.0, min(1.0, r1))
    tlen = tb - ta
    ts = ta + int(r0 * tlen)
    te = ta + int(r1 * tlen)
    if te <= ts:
        te = min(tb, ts + 1)
    # 英文整词吸附
    while ts > ta and tgt_full[ts-1].isalnum() and tgt_full[ts].isalnum():
        ts -= 1
    while te < tb and tgt_full[te-1].isalnum() and tgt_full[te].isalnum():
        te += 1
    while ts < te and tgt_full[ts].isspace():
        ts += 1
    while te > ts and tgt_full[te-1].isspace():
        te -= 1
    return (ts, te) if te > ts else None


def _best_match_span(haystack, needle, pos_ratio=None):
    """在 haystack 中找与 needle 最相似的连续字符区间，返回 (start, end) 或 None。
    用于选区联动：把临时译文 needle 在对方文本里定位到最近似的合理区间。
    pos_ratio 是源选区在源文本的相对位置（0~1），用于位置加权，避免选到
    文中其它相似但位置不对的段落。"""
    import difflib
    if not haystack or not needle:
        return None
    needle = needle.strip()
    n = len(needle)
    H = len(haystack)
    if n == 0 or H == 0:
        return None
    sm = difflib.SequenceMatcher()
    sm.set_seq2(needle)
    best = (-1.0, None)
    min_w = max(1, int(n * 0.6))
    max_w = min(H, int(n * 1.6) + 2)
    step = max(1, n // 8)
    target_center = (pos_ratio * H) if pos_ratio is not None else None
    for start in range(0, H, step):
        for w in (min_w, n, max_w):
            end = min(H, start + w)
            if end <= start:
                continue
            window = haystack[start:end]
            sm.set_seq1(window)
            ratio = sm.quick_ratio()
            score = ratio
            # 位置加权：窗口中心离目标位置越近，加分越多（最多 +0.15）
            if target_center is not None:
                center = (start + end) / 2.0
                dist = abs(center - target_center) / H   # 0~1
                score += 0.15 * (1.0 - dist)
            if score > best[0]:
                best = (score, (start, end), ratio)
        if start + min_w >= H:
            break
    # 相似度太低则放弃（用原始相似度判断，不含位置加权）
    if best[1] is None or best[2] < 0.35:
        return None
    s, e = best[1]
    # 化零为整：英文不要停在半个单词中间，把边界吸附到最近的词边界/空白
    # 向左把 s 移到单词开头
    while s > 0 and haystack[s-1].isalnum() and haystack[s].isalnum():
        s -= 1
    # 向右把 e 移到单词结尾
    while e < len(haystack) and haystack[e-1].isalnum() and (e < len(haystack) and haystack[e].isalnum()):
        e += 1
    # 去掉首尾空白
    while s < e and haystack[s].isspace():
        s += 1
    while e > s and haystack[e-1].isspace():
        e -= 1
    return (s, e)


class KaraokeHighlighter(QSyntaxHighlighter):
    """逐词高亮：朗读已读部分上青蓝绿背景。可选地在选区上铺蓝色底，
    朗读经过处用绿色覆盖，读完恢复蓝色。比 ExtraSelection 跨平台更可靠。"""
    def __init__(self, document):
        super().__init__(document)
        self._hl_start = 0
        self._hl_end = 0
        self._sel_start = 0      # 蓝色选区
        self._sel_end = 0
        self._link_start = 0     # 灰色联动选区（非活跃，自动匹配）
        self._link_end = 0
        # 朗读已读：偏青蓝、饱和度降低（更柔和）。只改背景，不改字色。
        # 青蓝色卡拉OK已读底色 + 白字（#7：朗读到哪就覆盖青蓝背景+白字）
        self._fmt = QTextCharFormat()
        self._fmt.setBackground(QColor("#5aa8b0"))
        self._fmt.setForeground(QColor("white"))
        # 蓝色选区底色 + 白字（#4：主动选区一律白字）
        self._sel_fmt = QTextCharFormat()
        self._sel_fmt.setBackground(QColor("#3a6ea5"))
        self._sel_fmt.setForeground(QColor("white"))
        # 灰色联动选区底色 + 白字（#4：被动灰色选区也一律白字）
        self._link_fmt = QTextCharFormat()
        self._link_fmt.setBackground(QColor("#5a5a5a"))
        self._link_fmt.setForeground(QColor("white"))

    def set_range(self, start, end):
        self._hl_start = max(0, start)
        self._hl_end = max(0, end)
        self.rehighlight()

    def set_selection(self, start, end):
        self._sel_start = max(0, start)
        self._sel_end = max(0, end)
        self.rehighlight()

    def set_link(self, start, end):
        self._link_start = max(0, start)
        self._link_end = max(0, end)
        self.rehighlight()

    def clear_link(self):
        self._link_start = 0
        self._link_end = 0
        self.rehighlight()

    def clear_selection(self):
        self._sel_start = 0
        self._sel_end = 0
        self.rehighlight()

    def clear_range(self):
        self._hl_start = 0
        self._hl_end = 0
        self.rehighlight()

    def clear_all(self):
        self._hl_start = self._hl_end = 0
        self._sel_start = self._sel_end = 0
        self._link_start = self._link_end = 0
        self.rehighlight()

    def set_dim(self, start):
        """多风格灰字区起点（-1=无）。"""
        self._dim_from = start if (start is not None and start >= 0) else -1
        self.rehighlight()

    def highlightBlock(self, text):
        block_start = self.currentBlock().position()
        block_end = block_start + len(text)
        df = getattr(self, "_dim_from", -1)
        # 底层：多风格灰字(只前景)。带背景的段随后用"背景+灰字"组合格式覆盖，
        # 保证灰区内的蓝选区/灰联动/绿色卡拉OK照常显示(修复灰区无卡拉OK)。
        if df >= 0 and block_end > df:
            ds = max(df, block_start)
            _f = QTextCharFormat()
            _f.setForeground(QColor("#8a8a8a"))
            self.setFormat(ds - block_start, block_end - ds, _f)

        _has_bg = False   # 当前 _apply 的段是否带背景(带背景则灰区也用白字)

        def _apply(a0, b0, base_fmt):
            a = max(a0, block_start); b = min(b0, block_end)
            if b <= a:
                return
            if df >= 0 and a < df < b:
                segs = [(a, df, False), (df, b, True)]
            elif df >= 0 and a >= df:
                segs = [(a, b, True)]
            else:
                segs = [(a, b, False)]
            for s1, e1, dim in segs:
                f = QTextCharFormat(base_fmt)
                # 带背景的格式(_fmt/_sel_fmt/_link_fmt)已自带白字前景，直接用即可，
                # 无论是直译区黑字还是多风格灰区，被覆盖处都会显示为白字(#4/#7)。
                # 仅"无背景的纯多风格灰区"(_has_bg=False 且 dim)保持灰字。
                if dim and not _has_bg:
                    f.setForeground(QColor("#8a8a8a"))
                self.setFormat(s1 - block_start, e1 - s1, f)

        # 三种带背景的格式：有背景 -> 字色白（覆盖多风格灰字）
        def _apply_bg(a0, b0, base_fmt):
            nonlocal _has_bg
            _has_bg = True
            _apply(a0, b0, base_fmt)
            _has_bg = False

        _has_bg = False
        _apply_bg(self._link_start, self._link_end, self._link_fmt)  # 灰色联动(最底)
        _apply_bg(self._sel_start, self._sel_end, self._sel_fmt)     # 蓝色选区
        _apply_bg(self._hl_start, self._hl_end, self._fmt)           # 绿色已读(最上)


class HistoryDialog(QDialog):
    """翻译历史弹窗：按天分组，每条显示开头一段提示，悬停蓝色高亮 + 气泡显示全文。
    双击或选中后点『载入翻译』返回该条。"""
    def __init__(self, items, parent=None):
        super().__init__(parent)
        self.setWindowTitle(L("翻译历史"))
        self.setMinimumSize(520, 560)
        self.chosen = None
        self._items = items

        from PyQt6.QtWidgets import QListWidget, QListWidgetItem
        layout = QVBoxLayout(self)
        info = QLabel(L("点选一条记录后『载入并翻译』；悬停可见全文。"))
        info.setStyleSheet("color:#888; font-size:12px;")
        layout.addWidget(info)

        self.listw = QListWidget()
        self.listw.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.listw.setWordWrap(False)
        self.listw.setTextElideMode(Qt.TextElideMode.ElideRight)
        # 按当前深浅动态配色(mac上_themed不换色，需自己判断)
        _hl = _theme_is_light()
        _hbg, _hbd, _hsep, _htx = (("#ffffff", "#c4c4c8", "#e4e4e6", "#1f1f22") if _hl
                                    else ("#252526", "#3a3a3a", "#2f2f2f", "#e8e8e8"))
        self.listw.setStyleSheet(_themed("""
%HISTTOOLTIP%
            QListWidget { background:%BG%; color:%TX%; border:1px solid %BD%; border-radius:6px; }
            QListWidget::item { padding:8px 10px; border-bottom:1px solid %SEP%; }
            QListWidget::item:hover { background:#0e639c; color:white; }
            QListWidget::item:selected { background:#0e639c; color:white; }
        """.replace("%HISTTOOLTIP%", _tooltip_css())
           .replace("%BG%", _hbg).replace("%BD%", _hbd)
           .replace("%SEP%", _hsep).replace("%TX%", _htx)
           + ("" if _native_scrollbar_platform()
              else "\n" + _rounded_scrollbar_qss())))
        layout.addWidget(self.listw, 1)

        # 按天分组（倒序，新的在上）
        import datetime
        cur_day = None
        for it in reversed(items):
            ts = it.get("ts", "")
            day = ts[:10]
            if day != cur_day:
                cur_day = day
                header = QListWidgetItem(f"——— {day} ———")
                header.setFlags(Qt.ItemFlag.NoItemFlags)
                header.setForeground(QColor("#4ea1ff"))
                self.listw.addItem(header)
            src = it.get("src", "").replace("\n", " ")
            preview = src[:28] + ("…" if len(src) > 28 else "")
            li = QListWidgetItem(f"  {preview}")
            full = f"【原文】{it.get('src','')}\n\n【译文】{it.get('tgt','')}\n\n[{it.get('engine','')}]"
            li.setToolTip(full)                    # 悬停气泡显示全文
            li.setData(Qt.ItemDataRole.UserRole, it)
            self.listw.addItem(li)

        self.listw.itemDoubleClicked.connect(self._choose)

        btns = QHBoxLayout()
        view_file_btn = QPushButton(L("检查历史"))
        view_file_btn.clicked.connect(self._open_history_file)
        dl_btn = QPushButton(L("下载文档"))
        dl_btn.clicked.connect(self._download_history)
        load_btn = QPushButton(L("重新载入"))
        load_btn.clicked.connect(self._load_selected)
        close_btn = QPushButton(L("关闭"))
        close_btn.clicked.connect(self.reject)
        for b in (view_file_btn, dl_btn, load_btn):
            b.setAutoDefault(False); b.setDefault(False)   # 去掉『检查历史』等的默认蓝
        close_btn.setStyleSheet("QPushButton{background:#1e88e5;border:none;border-radius:5px;color:white;}"\
            "QPushButton:hover{background:#2b95ef;}")
        for b in (view_file_btn, dl_btn, load_btn, close_btn):
            b.setMinimumWidth(BTN_W)   # 英文更长时按内容自适应，不截断
        btns.setSpacing(6)
        btns.addStretch()
        btns.addWidget(view_file_btn)
        btns.addWidget(dl_btn)
        btns.addWidget(load_btn)
        btns.addWidget(close_btn)
        layout.addLayout(btns)

    def _download_history(self):
        from PyQt6.QtWidgets import QFileDialog, QMessageBox
        import os, datetime
        filters = "文本 (*.txt);;Markdown (*.md);;JSON (*.json);;Word (*.docx);;PDF (*.pdf)"
        # 默认文件名 TH 日期 时间.txt；默认目录记忆（首次下载目录）
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H%M%S")
        last_dir = ""
        try:
            last_dir = self.parent().settings.value("last_dir_export", "")
        except Exception:
            pass
        if not last_dir or not os.path.isdir(last_dir):
            last_dir = os.path.join(os.path.expanduser("~"), "Downloads")
        default = os.path.join(last_dir, f"EC TH {ts}.txt")
        path, sel = QFileDialog.getSaveFileName(self, L("下载历史文档"), default, filters)
        if not path:
            return
        ext = os.path.splitext(path)[1].lower()
        try:
            items = self._items
            if ext == ".json":
                import json as _j
                with open(path, "w", encoding="utf-8") as f:
                    _j.dump(items, f, ensure_ascii=False, indent=2)
            elif ext == ".docx":
                from docx import Document
                doc = Document()
                doc.add_heading(L("翻译历史"), 0)
                for it in items:
                    doc.add_heading(_xml_safe(f"{it.get('ts','')} · {it.get('engine','')}"), level=2)
                    doc.add_paragraph(_xml_safe("【原文】" + it.get("src", "")))
                    doc.add_paragraph(_xml_safe("【译文】" + it.get("tgt", "")))
                doc.save(path)
            elif ext == ".pdf":
                self._history_pdf(path, items)
            elif ext == ".md":
                with open(path, "w", encoding="utf-8") as f:
                    f.write("# 翻译历史\n\n")
                    for it in items:
                        f.write(f"## {it.get('ts','')} · {it.get('engine','')}\n\n")
                        f.write(f"**原文**：{it.get('src','')}\n\n")
                        f.write(f"**译文**：{it.get('tgt','')}\n\n---\n\n")
            else:  # txt
                with open(path, "w", encoding="utf-8") as f:
                    for it in items:
                        f.write(f"[{it.get('ts','')} · {it.get('engine','')}]\n")
                        f.write("原文：" + it.get("src", "") + "\n")
                        f.write("译文：" + it.get("tgt", "") + "\n\n")
            QMessageBox.information(self, "已下载", f"已保存到：\n{path}")
            try:
                self.parent().settings.setValue("last_dir_export", os.path.dirname(path))
            except Exception:
                pass
        except Exception as e:
            QMessageBox.warning(self, "下载失败", str(e))

    def _history_pdf(self, path, items):
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os
        font = "Helvetica"
        for fp in ("/System/Library/Fonts/STHeiti Light.ttc", "/System/Library/Fonts/PingFang.ttc",
                   "C:/Windows/Fonts/msyh.ttc", "C:/Windows/Fonts/simsun.ttc"):
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont("CJK", fp)); font = "CJK"; break
                except Exception:
                    pass
        c = canvas.Canvas(path, pagesize=A4)
        w, h = A4; x, y = 2*cm, h-2*cm
        c.setFont(font, 11)
        for it in items:
            for line in (f"[{it.get('ts','')} · {it.get('engine','')}]",
                         "原文：" + it.get("src",""), "译文：" + it.get("tgt",""), ""):
                s = line
                while True:
                    chunk = s[:42]; s = s[42:]
                    c.drawString(x, y, chunk); y -= 0.6*cm
                    if y < 2*cm:
                        c.showPage(); c.setFont(font, 11); y = h-2*cm
                    if not s:
                        break
        c.save()

    def _open_history_file(self):
        import os
        from PyQt6.QtGui import QDesktopServices
        from PyQt6.QtCore import QUrl
        p = _history_path()
        if os.path.exists(p):
            QDesktopServices.openUrl(QUrl.fromLocalFile(p))

    def _choose(self, item):
        data = item.data(Qt.ItemDataRole.UserRole)
        if data:
            self.chosen = data
            self.accept()

    def _load_selected(self):
        it = self.listw.currentItem()
        if it:
            data = it.data(Qt.ItemDataRole.UserRole)
            if data:
                self.chosen = data
                self.accept()


class PillBusyBar(QWidget):
    """自绘胶囊形忙碌指示条：Qt 原生不确定进度条的滑块到两端会变方角
    （样式引擎限制），改用 QPainter 自绘，滑块任何位置都是完整胶囊。"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedHeight(6)
        self._pos = 0.0
        self._dir = 1
        self._t = QTimer(self)
        self._t.setInterval(16)
        self._t.timeout.connect(self._tick)

    def showEvent(self, e):
        self._t.start(); super().showEvent(e)

    def hideEvent(self, e):
        self._t.stop(); super().hideEvent(e)

    def _tick(self):
        self._pos += 0.018 * self._dir
        if self._pos >= 1.0:
            self._pos, self._dir = 1.0, -1
        elif self._pos <= 0.0:
            self._pos, self._dir = 0.0, 1
        self.update()

    def paintEvent(self, e):
        from PyQt6.QtGui import QPainter, QColor
        p = QPainter(self)
        p.setRenderHint(QPainter.RenderHint.Antialiasing)
        w, h = self.width(), self.height()
        r = h / 2
        p.setPen(Qt.PenStyle.NoPen)
        p.setBrush(QColor(255, 255, 255, 26))      # 极淡胶囊槽
        p.drawRoundedRect(0, 0, w, h, r, r)
        cw = max(int(w * 0.28), h * 2)              # 滑块宽
        x = int((w - cw) * self._pos)
        p.setBrush(QColor("#5aa8b0"))               # 青色胶囊滑块
        p.drawRoundedRect(x, 0, cw, h, r, r)


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.settings = QSettings("Strilen", "EnglishCoach")
        self.translate_worker = None
        self.tts_worker = None

        # 老系统 / 缺音频后端时，多媒体初始化失败不应阻止主界面弹出
        self.player = None
        self.audio_out = None
        try:
            self.player = QMediaPlayer()
            self.audio_out = QAudioOutput()
            self.player.setAudioOutput(self.audio_out)
            self.player.positionChanged.connect(self._on_play_position)
            self.player.durationChanged.connect(self._on_play_duration)
            self.player.playbackStateChanged.connect(self._on_play_state)
        except Exception as e:
            print(f"[warn] 音频初始化失败，朗读功能将不可用: {e}")

        # 朗读状态
        self._speak_editor = None      # 当前朗读的文本框
        self._speak_boundaries = []    # 词边界列表
        # 两侧各自独立的音频缓存（互不干扰）：
        #   bytes=已生成音频, boundaries=卡拉OK边界, duration=时长, position=暂停位置, sig=签名
        self._side_cache = {
            "src": {"bytes": None, "boundaries": [], "duration": 0, "position": 0, "sig": None},
            "tgt": {"bytes": None, "boundaries": [], "duration": 0, "position": 0, "sig": None},
        }
        self._speak_rate = 0           # 当前语速（用于 offset 换算）
        self._seeking = False          # 进度条拖动中
        self._karaoke_timer = QTimer(self)   # 高频驱动逐词高亮
        self._karaoke_timer.setInterval(50)
        self._karaoke_timer.timeout.connect(self._karaoke_tick)

        self.setWindowTitle(f"English Coach  v{APP_VERSION}")
        self.setAcceptDrops(True)   # 支持拖拽文件导入
        self.setWindowIcon(self._load_app_icon())
        self.setMinimumSize(880, 480)   # 再-10
        self.resize(920, 620)

        self._build_central()
        self._build_statusbar()
        self._apply_style()
        self._install_bundled_argos_models()

    def _load_app_icon(self):
        """优先加载随包的 icon_1024.png；找不到则回退内置 SVG 图标。"""
        # PyInstaller 打包后资源在 sys._MEIPASS；开发时在脚本目录
        base = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
        for name in ("icon_1024.png", "icon.png"):
            p = os.path.join(base, name)
            if os.path.exists(p):
                ic = QIcon(p)
                if not ic.isNull():
                    return ic
        return Icons.icon("app", "#4ea1ff")

    def _install_bundled_argos_models(self):
        """启动时确保内置中英模型已安装（幂等）。"""
        try:
            _ensure_argos_models(("en", "zh"))
        except Exception:
            pass

    # ---------- 顶部功能按钮（设置/更新/帮助/关于）----------
    def _make_tool_buttons(self):
        """返回一个含 设置/更新/帮助/关于 的横向布局，用于放在顶排右侧。"""
        box = QHBoxLayout()
        box.setSpacing(4)
        self._tool_btns = []   # 登记，极简模式逐个隐藏(right_w整体保留)

        def mk(icon, tip, slot):
            b = QPushButton(Icons.icon(icon), "")
            b.setProperty("_icn", icon)
            b.setFixedSize(36, 36)
            b.setObjectName("toolbtn")
            b.setToolTip(tip)
            b.clicked.connect(slot)
            self._tool_btns.append(b)
            return b

        box.addWidget(mk("settings", L("设置"), self.open_settings))
        box.addWidget(mk("history", L("更新说明"),
                         lambda: DocDialog(L("版本更新说明"), changelog_html(), self).exec()))
        box.addWidget(mk("help", L("使用说明"),
                         lambda: DocDialog(L("使用说明"), readme_html(), self).exec()))
        box.addWidget(mk("info", L("关于"),
                         lambda: DocDialog(L("关于 EnglishCoach"), about_html(), self).exec()))
        return box

    # ---------- 中央区域 ----------
    def _build_central(self):
        central = QWidget()
        root = QVBoxLayout(central)
        root.setContentsMargins(14, 12, 14, 10)
        root.setSpacing(10)

        # —— 第一排：翻译引擎(左) | 原文语言(贴交换) | 交换(中央) | 译文语言(贴交换) | 设置区(右) ——
        self.src_combo = QComboBox()
        _combo_fill(self.src_combo, LANG_OPTIONS)
        fit_combo_width(self.src_combo, extra=20)   # 再加宽10(累计20)，向左扩展
        _combo_select_data(self.src_combo, self.settings.value("src_lang", "自动检测"))
        self.src_combo.setToolTip(L("原文语言"))
        self.tgt_combo = QComboBox()
        _combo_fill(self.tgt_combo, LANG_OPTIONS)
        fit_combo_width(self.tgt_combo, extra=20)   # 再加宽10(累计20)，向右扩展
        _combo_select_data(self.tgt_combo, self.settings.value("tgt_lang", "自动检测"))
        self.tgt_combo.setToolTip(L("译文语言"))

        self.engine_combo = QComboBox()
        _combo_fill(self.engine_combo, ALL_ENGINES)
        fit_combo_width(self.engine_combo, extra=20, popup_extra=15)   # 闭合框+20(再+5)，弹出列表再+15
        _combo_select_data(self.engine_combo, 
            self.settings.value("engine", ENGINE_GOOGLE))
        self.engine_combo.currentTextChanged.connect(self._on_engine_changed)
        self.engine_combo.currentTextChanged.connect(
            lambda _t: self.settings.setValue("engine", self.engine_combo.currentData()))
        self.engine_combo.setToolTip(L("翻译引擎"))
        # 原文/译文语言变化 -> 无条件强制重新翻译（等同点翻译按钮）
        self.src_combo.currentTextChanged.connect(lambda _t: self._on_lang_changed())
        self.tgt_combo.currentTextChanged.connect(lambda _t: self._on_lang_changed())
        self.src_combo.currentTextChanged.connect(
            lambda _t: self.settings.setValue("src_lang", self.src_combo.currentData()))
        self.tgt_combo.currentTextChanged.connect(
            lambda _t: self.settings.setValue("tgt_lang", self.tgt_combo.currentData()))

        swap_btn = QPushButton(Icons.icon("swap"), "")
        swap_btn.setProperty("_icn", "swap")
        swap_btn.setToolTip(L("交换源文译文内容"))
        swap_btn.setFixedSize(44, 34)
        swap_btn.clicked.connect(self.swap_sides)

        from PyQt6.QtWidgets import QGridLayout, QWidget as _QWt
        # 极简钮独立(不并入left_w，否则极简时随left_w一起隐藏无法退出)
        self.min_btn = QPushButton()
        _mini_svg = ('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24">'
                     '<rect x="4" y="4" width="16" height="2.5" rx="1" fill="{c}"/>'
                     '<rect x="7" y="9.5" width="10" height="5" rx="1.5" fill="{c}"/>'
                     '<rect x="4" y="17.5" width="16" height="2.5" rx="1" fill="{c}"/></svg>')
        def _set_mini_icon():
            from PyQt6.QtGui import QPixmap as _PX, QIcon as _IC
            col = "#1f1f22" if _theme_is_light() else "#e8e8e8"
            pm = _PX(); pm.loadFromData(_mini_svg.format(c=col).encode(), "SVG")
            self.min_btn.setIcon(_IC(pm))
        _set_mini_icon()
        self._set_mini_icon = _set_mini_icon
        self.min_btn.setFixedSize(36, 36)
        self.min_btn.setObjectName("toolbtn")
        self.min_btn.setToolTip(L("极简界面"))
        self.min_btn.setStyleSheet("QPushButton{font-size:16px;}")
        self.min_btn.clicked.connect(self._toggle_minimal_ui)
        left_w = _QWt(); ll = QHBoxLayout(left_w); ll.setContentsMargins(0,0,0,0)
        ll.setSpacing(6)
        ll.addWidget(self.min_btn)          # 极简钮贴最左(极简时只隐藏两个下拉，不隐藏left_w整体)
        ll.addSpacing(4)
        ll.addWidget(self.engine_combo)     # 引擎下拉
        ll.addStretch()
        ll.addWidget(self.src_combo)        # 原文语言贴近交换钮

        right_w = _QWt(); rl = QHBoxLayout(right_w); rl.setContentsMargins(0,0,0,0)
        rl.setSpacing(6)
        rl.addWidget(self.tgt_combo)        # 译文语言贴近交换钮
        rl.addStretch()
        rl.addLayout(self._make_tool_buttons())   # 设置/更新/帮助/关于 贴最右
        # 极简镜像占位：40=极简钮36+其后固定间隔4，与左端极简钮完全镜像对称。
        # 仅极简模式显示(结构保证交换钮居中，零校准)；正常模式隐藏不占空间，工具钮贴边不受影响。
        _mirror = _QWt(); _mirror.setFixedSize(40, 36)
        _mirror.setVisible(False)
        self._mini_mirror = _mirror
        rl.addWidget(_mirror)
        self._minimal_ui = False
        self._restore_minimal = (self.settings.value("minimal_ui", "false") == "true")
        self._lit_end = None
        self._ui_top_left = left_w
        self._ui_swap = swap_btn
        self._ui_top_right = right_w

        top_grid = QGridLayout()
        top_grid.setContentsMargins(0, 0, 0, 0)
        top_grid.setHorizontalSpacing(10)
        top_grid.addWidget(left_w, 0, 0)
        top_grid.addWidget(swap_btn, 0, 1)
        top_grid.addWidget(right_w, 0, 2)
        top_grid.setColumnStretch(0, 1)
        top_grid.setColumnStretch(2, 1)
        root.addLayout(top_grid)

        # —— 双栏文本（仅文本框，标题与操作按钮移到下方按钮排）——
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setHandleWidth(10)

        self.input_edit = QTextEdit()
        self.input_edit.setPlaceholderText(L("在此输入或粘贴文本…"))
        self.input_edit.setToolTip(L("原文文字"))
        self.output_edit = QTextEdit()
        self.output_edit.setPlaceholderText(L("译文显示在这里…"))
        self.output_edit.setToolTip(L("译文文字"))
        # 原文/译文区字号放大，更醒目
        _edit_font = QFont("Microsoft YaHei", 13)
        self.input_edit.setFont(_edit_font)
        self.output_edit.setFont(_edit_font)
        # 强制选区为蓝色（跨平台一致，避免某些系统出现绿色等异色）
        from PyQt6.QtGui import QPalette
        for _ed in (self.input_edit, self.output_edit):
            _pal = _ed.palette()
            _pal.setColor(QPalette.ColorRole.Highlight, QColor("#3a6ea5"))
            _pal.setColor(QPalette.ColorRole.HighlightedText, QColor("#ffffff"))
            _ed.setPalette(_pal)
        # 卡拉OK高亮器（QSyntaxHighlighter，跨平台可靠渲染）
        self._hl_input = KaraokeHighlighter(self.input_edit.document())
        self._hl_output = KaraokeHighlighter(self.output_edit.document())
        # 选区联动：选一边的文字，自动在另一边灰色高亮对应区间
        self._link_guard = False
        self._link_timer = QTimer(self)
        self._link_timer.setSingleShot(True)
        self._link_timer.setInterval(250)
        self._link_timer.timeout.connect(self._do_selection_link)
        self._link_pending_src = None
        self.input_edit.selectionChanged.connect(lambda: self._on_selection_changed(self.input_edit))
        self.output_edit.selectionChanged.connect(lambda: self._on_selection_changed(self.output_edit))
        # 点击/聚焦某区即把它设为主动区（蓝框跳过去，颜色互换）
        # 区域激活改由 eventFilter 的鼠标按下事件处理（点空白也能切换）
        # 让拖到文本框里的文件也走"导入内容"，而非粘贴文件名/URL
        self.input_edit.installEventFilter(self)
        self.output_edit.installEventFilter(self)
        self.input_edit.viewport().installEventFilter(self)
        self.output_edit.viewport().installEventFilter(self)

        splitter.addWidget(self.input_edit)
        splitter.addWidget(self.output_edit)
        splitter.setSizes([460, 460])
        splitter.setChildrenCollapsible(False)
        # 拖到中央附近自动吸附精准居中(左右差≤32px时对半均分)
        self._split_snap_guard = False
        def _snap_center(_pos, _idx):
            if self._split_snap_guard:
                return
            try:
                sz = splitter.sizes()
                if len(sz) == 2 and abs(sz[0] - sz[1]) <= 32:
                    total = sz[0] + sz[1]
                    half = total // 2
                    self._split_snap_guard = True
                    splitter.setSizes([half, total - half])
                    self._split_snap_guard = False
            except Exception:
                self._split_snap_guard = False
        splitter.splitterMoved.connect(_snap_center)
        root.addWidget(splitter, 1)

        # 输入即译：停止输入约 0.8 秒后自动翻译（手动按钮仍保留）
        self._auto_timer = QTimer(self)
        self._auto_timer.setSingleShot(True)
        self._auto_timer.setInterval(800)
        self._auto_timer.timeout.connect(self._auto_translate)
        self.input_edit.textChanged.connect(self._on_input_changed)
        self.output_edit.textChanged.connect(self._on_output_changed)

        # —— 操作排：原文 | 复制粘贴清空 | 导出原文+导入文件 | 上一条+历史 | 翻译 |
        #             译文复制粘贴清空 | 导出译文+导出文件 | 译文 ——
        def _mini_btn(icon, tip, slot):
            b = QPushButton(Icons.icon(icon), "")
            b.setProperty("_icn", icon)
            b.setFixedSize(36, 36)
            b.setObjectName("toolbtn")
            b.setToolTip(tip)
            b.clicked.connect(slot)
            return b

        # 原文侧：导出当前原文 + 导入文件
        self.export_src_btn = _mini_btn("export", L("导出当前原文"), lambda: self._export_text("src"))
        self.import_btn = _mini_btn("file", L("导入文件"), self._import_file)
        src_io = QHBoxLayout(); src_io.setSpacing(4); src_io.setContentsMargins(0, 0, 0, 0)
        src_io.addWidget(self.export_src_btn); src_io.addWidget(self.import_btn)

        # 上一条 + 历史
        hist_box = QHBoxLayout()
        hist_box.setSpacing(4)
        hist_box.setContentsMargins(0, 0, 0, 0)
        self.prev_src_btn = _mini_btn("undo", L("载入上一条原文"), self._load_prev_source)
        self.history_btn = _mini_btn("list", L("翻译历史"), self._open_history_dialog)
        self.next_src_btn = _mini_btn("redo", L("载入下一条原文"), self._load_next_source)
        hist_box.addWidget(self.prev_src_btn)
        hist_box.addWidget(self.next_src_btn)
        hist_box.addWidget(self.history_btn)

        # 译文侧：导出当前译文 + 导出翻译后文件
        self.export_tgt_btn = _mini_btn("export", L("导出当前译文"), lambda: self._export_text("tgt"))
        self.export_file_btn = _mini_btn("file_down", L("导出翻译后文件"), self._export_file)
        self.export_file_btn.setEnabled(False)
        self.export_file_btn.setVisible(False)   # 平时隐藏，导入成功后才出现
        tgt_io = QHBoxLayout(); tgt_io.setSpacing(4); tgt_io.setContentsMargins(0, 0, 0, 0)
        tgt_io.addWidget(self.export_tgt_btn); tgt_io.addWidget(self.export_file_btn)

        self.translate_btn = QPushButton(Icons.icon("translate", "#ffffff"), "  " + L("翻译"))
        self.translate_btn.setObjectName("primary")
        self.translate_btn.setMinimumHeight(40)
        self.translate_btn.setToolTip(L("翻译"))
        self.translate_btn.clicked.connect(self.do_translate)

        from PyQt6.QtWidgets import QGridLayout, QWidget as _QW
        left_w = _QW(); left_l = QHBoxLayout(left_w)
        left_l.setContentsMargins(0, 0, 0, 0)
        left_l.setSpacing(4)   # 统一小缝隙，不分组
        left_l.addLayout(self._panel_buttons(self.input_edit))
        left_l.addLayout(src_io)
        left_l.addLayout(hist_box)
        clear_src_btn = QPushButton(Icons.icon("clear"), "")
        clear_src_btn.setProperty("_icn", "clear")
        clear_src_btn.setFixedSize(36, 36)
        clear_src_btn.setObjectName("toolbtn")
        clear_src_btn.setToolTip(L("清空"))
        clear_src_btn.clicked.connect(lambda: self._clear_editor(self.input_edit))
        left_l.addWidget(clear_src_btn)
        left_l.addStretch()          # 整组左对齐

        right_w = _QW(); right_l = QHBoxLayout(right_w)
        right_l.setContentsMargins(0, 0, 0, 0)
        right_l.setSpacing(4)   # 统一小缝隙，不分组
        right_l.addStretch()         # 整组右对齐
        right_l.addLayout(self._panel_buttons(self.output_edit))
        right_l.addLayout(tgt_io)
        clear_tgt_btn = QPushButton(Icons.icon("clear"), "")
        clear_tgt_btn.setProperty("_icn", "clear")
        clear_tgt_btn.setFixedSize(36, 36)
        clear_tgt_btn.setObjectName("toolbtn")
        clear_tgt_btn.setToolTip(L("清空"))
        clear_tgt_btn.clicked.connect(lambda: self._clear_editor(self.output_edit))
        right_l.addWidget(clear_tgt_btn)

        grid = QGridLayout()
        grid.setContentsMargins(0, 0, 0, 0)
        grid.addWidget(left_w, 0, 0)
        grid.addWidget(self.translate_btn, 0, 1)
        grid.addWidget(right_w, 0, 2)
        self._ui_bot_left = left_w
        self._ui_bot_right = right_w
        grid.setColumnStretch(0, 1)
        grid.setColumnStretch(2, 1)
        root.addLayout(grid)

        # ===== 朗读音频条：原文 / 译文 各一组 =====
        play_groups = QHBoxLayout()
        play_groups.setContentsMargins(0, 0, 0, 0)
        play_groups.setSpacing(10)
        play_groups.addLayout(self._make_play_group("src"), 1)
        play_groups.addLayout(self._make_play_group("tgt"), 1)
        self._ui_play_w = _QW()
        self._ui_play_w.setLayout(play_groups)
        root.addWidget(self._ui_play_w)

        # 朗读控制条（嗓音 + 语速）：标签去掉，改成悬停气球提示
        tts_bar = QHBoxLayout()
        tts_bar.setContentsMargins(0, 0, 0, 0)
        tts_bar.setSpacing(6)
        self.zh_voice_combo = QComboBox()
        _combo_fill(self.zh_voice_combo, ZH_VOICES.keys())
        fit_combo_width(self.zh_voice_combo)
        _combo_select_data(self.zh_voice_combo, 
            self.settings.value("zh_voice", next(iter(ZH_VOICES))))
        self.zh_voice_combo.currentTextChanged.connect(self._on_zh_voice_changed)
        self.zh_voice_combo.setToolTip(L("中文嗓音"))
        tts_bar.addWidget(self.zh_voice_combo)

        self.en_voice_combo = QComboBox()
        _combo_fill(self.en_voice_combo, EN_VOICES.keys())
        fit_combo_width(self.en_voice_combo)
        _combo_select_data(self.en_voice_combo, 
            self.settings.value("en_voice", next(iter(EN_VOICES))))
        self.en_voice_combo.currentTextChanged.connect(self._on_en_voice_changed)
        self.en_voice_combo.setToolTip(L("英文嗓音"))
        tts_bar.addWidget(self.en_voice_combo)
        tts_bar.addSpacing(12)

        self.rate_slider = QSlider(Qt.Orientation.Horizontal)
        self.rate_slider.setObjectName("rateSlider")
        self.rate_slider.setRange(-50, 50)
        self.rate_slider.setValue(int(self.settings.value("tts_rate", 0)))
        self.rate_slider.setMinimumWidth(100)
        self.rate_label = QLabel("0%")
        self._update_rate_tooltip(0)   # 初始化语速气球
        self.rate_slider.valueChanged.connect(
            lambda v: self.rate_label.setText(f"{'+' if v>=0 else ''}{v}%"))
        self.rate_slider.valueChanged.connect(self._update_rate_tooltip)
        self.rate_slider.valueChanged.connect(lambda _v: self._persist_rate())
        # 语速变更防抖：拖动停止 0.4s 后才重新朗读，避免拖动过程频繁触发
        self._rate_timer = QTimer(self)
        self._rate_timer.setSingleShot(True)
        self._rate_timer.setInterval(400)
        self._rate_timer.timeout.connect(self._on_voice_or_rate_changed)
        self.rate_slider.valueChanged.connect(
            lambda _: self._rate_timer.start())
        tts_bar.addWidget(self.rate_slider, 1)
        self.rate_label.setVisible(False)   # 数值改用气球提示，界面不再显示
        tts_bar.addWidget(self.rate_label)
        self._ui_tts_w = _QW()
        self._ui_tts_w.setLayout(tts_bar)
        root.addWidget(self._ui_tts_w)

        self.setCentralWidget(central)
        # 兼容别名：默认指向原文侧；朗读某侧时在 do_speak 里切换到该侧
        self.play_slider = self.play_slider_src
        self.download_audio_btn = self.dl_src_btn
        self._active_side = "src"
        self._active_editor = self.input_edit
        self.input_edit.setProperty("activeRegion", "1")
        self.output_edit.setProperty("activeRegion", "0")

    def retranslate_ui(self):
        """语言即时切换：遍历整个窗口所有控件按文字查表替换，自动全覆盖。"""
        try:
            to_lang = _ui_lang()
            retranslate_widget_tree(self, to_lang)
            # 下拉项用 data 键重译（更准，保留选择）
            for attr in ("engine_combo", "src_combo", "tgt_combo",
                         "zh_voice_combo", "en_voice_combo"):
                cb = getattr(self, attr, None)
                if cb is None:
                    continue
                cb.blockSignals(True)
                for i in range(cb.count()):
                    d = cb.itemData(i)
                    if d:
                        cb.setItemText(i, L(d))
                cb.blockSignals(False)
        except Exception:
            pass

    def _refresh_icons_for_theme(self):
        """主题切换后重生成所有按钮图标(SVG 按当前深浅色重新渲染)。
        遍历所有按钮，读取创建时登记的图标名 property('_icn') 重新着色——不漏任何按钮。"""
        from PyQt6.QtWidgets import QPushButton, QToolButton
        for btn in self.findChildren((QPushButton, QToolButton)):
            name = btn.property("_icn")
            if name:
                try:
                    btn.setIcon(Icons.icon(name))
                except Exception:
                    pass
        return

    def _refresh_icons_legacy(self):
        mapping = {
            "import_btn": "import", "export_file_btn": "export",
        }
        for attr, icon_name in mapping.items():
            btn = getattr(self, attr, None)
            if btn is not None:
                try:
                    btn.setIcon(Icons.icon(icon_name))
                except Exception:
                    pass

    def apply_theme(self):
        """主题切换即时生效。mac：AppKit 原生驱动(不涂调色板/QSS，避免打架)；非 mac：调色板+样式表。"""
        from PyQt6.QtWidgets import QApplication as _QA
        from PyQt6.QtGui import QPalette, QColor
        app = _QA.instance()
        _apply_color_scheme(app)   # mac 走 AppKit 原生；非 mac 走 Qt colorScheme
        if sys.platform == "darwin":
            # mac：AppKit 外观已切换。重设混合QSS(绘制部分按新深浅重画) + 重生成图标。
            self.setStyleSheet(self._mac_hybrid_qss())
            self._refresh_icons_for_theme()
            if hasattr(self, "_set_mini_icon"):
                self._set_mini_icon()   # 极简钮图标按新深浅重新着色
            for w in app.topLevelWidgets():
                w.update()
            return
        # 非 mac 混合方案：完整复刻启动路径的顺序，否则会出现
        # "打开好、改主题坏、重启又好"。启动时是：设调色板 -> app.setStyle(新建样式对象)
        # -> 再设 colorScheme -> 建窗口。原生样式(windows11)在【创建时】确定深浅状态，
        # 已存在的样式对象不会因后来改 colorScheme 而彻底重绘 —— 这正是重启才好的原因。
        # 所以热切换必须重建样式对象，并按启动顺序重设调色板与 colorScheme。
        _apply_win_palette(app)
        try:
            from PyQt6.QtWidgets import QStyleFactory
            _cur_style = app.style().objectName()          # 如 windows11 / windowsvista / fusion
            _match = [k for k in QStyleFactory.keys() if k.lower() == _cur_style.lower()]
            if _match:
                app.setStyle(QStyleFactory.create(_match[0]))   # 重建样式对象(关键)
        except Exception:
            pass
        # setStyle 会把 app 调色板重置成该样式的标准调色板，必须在其后再设一次；
        # colorScheme 也照启动那样在 setStyle 之后再设一次。
        _apply_win_palette(app)
        _apply_color_scheme(app)
        _ss = self._win_hybrid_qss()
        self._base_ss = _ss
        self.setStyleSheet(_ss)
        # 强制所有控件重新 polish：仅 update() 不足以让已 polish 过的控件按新样式/调色板重绘
        from PyQt6.QtWidgets import QDialog
        try:
            _st = app.style()
            for _w in app.allWidgets():
                _st.unpolish(_w)
                _st.polish(_w)
        except Exception:
            pass
        # 关键(修复"深色切浅色后按钮图标/文字看不见")：图标是按当前深浅渲染的 SVG
        # (浅色主题用深色 #1f1f22，深色主题用浅色 #e8e8e8)。mac 分支一直有重生成图标
        # 这一步，非 mac 分支却漏了 —— 切到浅色后图标仍是浅灰色，在浅底按钮上就看不见。
        self._refresh_icons_for_theme()
        if hasattr(self, "_set_mini_icon"):
            self._set_mini_icon()      # 极简钮图标按新深浅重新着色
        # 关键(修复"切浅色后下拉弹出仍是黑底、看似丢一项")：弹出容器配色只在下拉
        # 创建时设过一次，切主题从不更新 -> 深字配黑底就"看不见"了。这里按新深浅刷新。
        _refresh_combo_popups(self)
        for w in app.topLevelWidgets():
            if isinstance(w, QDialog) and hasattr(w, "_retheme"):
                w._retheme()
            # 弹窗内的下拉弹出容器也要按新深浅刷新
            _refresh_combo_popups(w)
            # 弹窗内的按钮图标同样要按新深浅重生成
            if w is not self and hasattr(w, "findChildren"):
                try:
                    from PyQt6.QtWidgets import QPushButton, QToolButton
                    for _b in w.findChildren((QPushButton, QToolButton)):
                        _n = _b.property("_icn")
                        if _n:
                            _b.setIcon(Icons.icon(_n))
                except Exception:
                    pass
            w.update()

    def _toggle_minimal_ui(self):
        """极简界面：只留 极简钮/原文译文区/翻译钮/状态栏；再点切回正常。"""
        self._minimal_ui = not getattr(self, "_minimal_ui", False)
        mini = self._minimal_ui
        try:
            self.settings.setValue("minimal_ui", "true" if mini else "false")
        except Exception:
            pass
        for w in (self._ui_bot_left, self._ui_bot_right,
                  self._ui_play_w, self._ui_tts_w):
            w.setVisible(not mini)
        # left_w/right_w 都不整体隐藏，只隐藏内容——右端镜像占位与左端极简钮
        # 结构对称，交换钮由布局数学保证居中(方案B，零数值零校准)。
        self.engine_combo.setVisible(not mini)
        self.src_combo.setVisible(not mini)
        self.tgt_combo.setVisible(not mini)
        for _b in getattr(self, "_tool_btns", []):
            _b.setVisible(not mini)
        if hasattr(self, "_mini_mirror"):
            self._mini_mirror.setVisible(mini)
        self.min_btn.setVisible(True)   # 极简钮始终可见（否则无法退出极简）
        self._ui_swap.setVisible(True)  # 极简模式保留交换钮
        if mini:
            self.min_btn.setToolTip(L("正常界面"))
            self.min_btn.setStyleSheet(
                "QPushButton{font-size:16px; background:#5aa8b0; color:#0e2024;"
                " border:1px solid #5aa8b0; border-radius:6px;}")
            self.setMinimumSize(420, 200)   # 极简模式最小约束
            self.resize(420, 280)           # 点极简默认缩到较小尺寸
        else:
            self.min_btn.setToolTip(L("极简界面"))
            self.min_btn.setStyleSheet("QPushButton{font-size:16px;}")
            self.setMinimumSize(880, 480)

    def eventFilter(self, obj, event):
        from PyQt6.QtCore import QEvent
        et = event.type()
        if et == QEvent.Type.MouseButtonPress:
            # 鼠标按下任何位置（文字或空白）都切换主动区——比 cursorPositionChanged
            # 可靠（点空白处光标可能不动、信号不发）
            editor = None
            if obj in (self.input_edit, self.input_edit.viewport()):
                editor = self.input_edit
            elif obj in (self.output_edit, self.output_edit.viewport()):
                editor = self.output_edit
            if editor is not None:
                if editor is getattr(self, "_active_editor", None):
                    # 点主动区：保持主动，清掉联动选区对（原生选区由点击本身清）
                    self._clear_selection_pair()
                else:
                    self._activate_region(editor)
            return False   # 不吞事件，光标照常定位
        if et in (QEvent.Type.DragEnter, QEvent.Type.DragMove):
            if event.mimeData().hasUrls():
                event.acceptProposedAction()
                return True
        elif et == QEvent.Type.Drop:
            if event.mimeData().hasUrls():
                urls = event.mimeData().urls()
                if urls:
                    path = urls[0].toLocalFile()
                    if path:
                        event.acceptProposedAction()
                        self._do_import(path)   # 导入内容，而非粘贴文件名
                        return True
        return super().eventFilter(obj, event)

    def dragEnterEvent(self, e):
        if e.mimeData().hasUrls():
            e.acceptProposedAction()

    def dropEvent(self, e):
        urls = e.mimeData().urls()
        if urls:
            path = urls[0].toLocalFile()
            if path:
                self._do_import(path)

    def _make_play_group(self, side):
        """构建一组朗读控制（side='src' 原文 / 'tgt' 译文）：
        音频条 [朗读▶] [停止■] [下载↓]。导入/导出按钮在上方操作排。"""
        editor = self.input_edit if side == "src" else self.output_edit
        name = "原文" if side == "src" else "译文"
        box = QHBoxLayout()
        box.setSpacing(4)

        # 音频进度条
        slider = QSlider(Qt.Orientation.Horizontal)
        slider.setRange(0, 1000)
        slider.setValue(0)
        slider.setMinimumWidth(90)
        slider.setToolTip(L(f"{name}朗读进度"))
        slider.sliderPressed.connect(lambda s=side: self._on_seek_start(s))
        slider.sliderReleased.connect(lambda s=side: self._on_seek_end_side(s))
        slider.sliderMoved.connect(lambda v, s=side: self._on_seek_moved(v, s))
        box.addWidget(slider, 1)

        speak_btn = QPushButton(Icons.icon("speak"), "")
        speak_btn.setProperty("_icn", "speak")
        speak_btn.setFixedSize(36, 36)
        speak_btn.setObjectName("toolbtn")
        speak_btn.setToolTip(L(f"朗读{name}"))
        speak_btn.clicked.connect(lambda _=False, e=editor: self._toggle_speak(e))
        box.addWidget(speak_btn)

        stop_btn = QPushButton(Icons.icon("stop"), "")
        stop_btn.setProperty("_icn", "stop")
        stop_btn.setFixedSize(36, 36)
        stop_btn.setObjectName("toolbtn")
        stop_btn.setToolTip(L(f"停止朗读{name}"))
        stop_btn.clicked.connect(lambda _=False, e=editor: self._stop_side(e))
        box.addWidget(stop_btn)

        dl_btn = QPushButton(Icons.icon("download"), "")
        dl_btn.setProperty("_icn", "download")
        dl_btn.setFixedSize(36, 36)
        dl_btn.setObjectName("toolbtn")
        dl_btn.setToolTip(L(f"下载{name}朗读音频"))
        dl_btn.setEnabled(False)
        dl_btn.clicked.connect(lambda _=False, s=side: self._download_audio(s))
        box.addWidget(dl_btn)

        clr_btn = QPushButton(Icons.icon("clear"), "")
        clr_btn.setProperty("_icn", "clear")
        clr_btn.setFixedSize(36, 36)
        clr_btn.setObjectName("toolbtn")
        clr_btn.setToolTip(L("清空"))
        clr_btn.clicked.connect(lambda _=False, s=side: self._clear_side_audio(s))
        box.addWidget(clr_btn)

        if side == "src":
            self.play_slider_src = slider
            self.speak_src_btn = speak_btn
            self.stop_src_btn = stop_btn
            self.dl_src_btn = dl_btn
        else:
            self.play_slider_tgt = slider
            self.speak_tgt_btn = speak_btn
            self.stop_tgt_btn = stop_btn
            self.dl_tgt_btn = dl_btn
        return box

    def _clear_editor(self, editor):
        """清空文本框及全部关联状态：
        - 原文清空：文字清、导入文件状态清（导入钮变灰）、译文随之清空、
          两侧朗读音频释放（朗读钮/下载钮变灰）、导出文件钮隐藏。
        - 译文清空：文字清、译文侧音频释放（钮变灰）、导出文件钮隐藏。
        若被清一侧正在朗读，先停止朗读。"""
        # v2.3.1：文字清空与音频清空分离——本钮只清文字+导入文件，音频缓存保留
        self._preserve_audio_on_clear = True
        try:
            editor.clear()
            if editor is self.input_edit:
                self._imported_path = None
                self._imported_ext = None
                self._imported_text = None
                self._can_export_file = False
                self.output_edit.clear()   # 译文文字随之清空（音频仍保留）
        finally:
            self._preserve_audio_on_clear = False
        # 字幕铁律：字幕依附文字——文字不在，字幕(边界+高亮)同亡；音频bytes保留可继续播放
        self._drop_karaoke_for(editor)
        if editor is self.input_edit:
            self._drop_karaoke_for(self.output_edit)   # 译文文字随清，译文字幕同亡
        self._update_file_buttons()

    def apply_always_on_top(self, on: bool, defer_if_modal: bool = True):
        """保持程序置顶。

        关键约束（v2.14.1 根治）：改 windowFlags 会销毁并重建原生窗口。
        若此时有模态对话框(设置窗 exec())开着：
          - 对主窗重建 → mac 上对话框跟着闪一下、Windows 上对话框被压到下面；
          - 对对话框本身重建 → 它会先缩小消失再出现，且 exec() 的模态循环失效，
            关闭后主界面按钮点了没反应（v2.14.0 的两个 bug 就是这么来的）。
        因此：模态对话框开着时只记录意图，等它关闭后再真正应用。
        模态对话框本来就显示在父窗之上，不需要自己加置顶标志。
        """
        from PyQt6.QtCore import Qt as _Qt
        try:
            _app = QApplication.instance()
            if defer_if_modal and _app is not None and _app.activeModalWidget() is not None:
                # 有模态窗开着：只记意图，关闭后由 open_settings 收尾应用
                self._pending_on_top = bool(on)
                return
            flags = self.windowFlags()
            if on:
                flags |= _Qt.WindowType.WindowStaysOnTopHint
            else:
                flags &= ~_Qt.WindowType.WindowStaysOnTopHint
            if flags == self.windowFlags():
                return                      # 状态没变就不重建窗口
            was_visible = self.isVisible()
            self.setWindowFlags(flags)
            if was_visible:
                self.setVisible(True)       # 改 flags 后需重新显示
        except Exception:
            pass

    def _flush_pending_on_top(self):
        """模态对话框关闭后，应用期间被推迟的置顶设置。"""
        try:
            pend = getattr(self, "_pending_on_top", None)
            if pend is None:
                return
            self._pending_on_top = None
            from PyQt6.QtCore import Qt as _Qt   # 局部导入，避免作用域意外
            cur = bool(self.windowFlags()
                       & _Qt.WindowType.WindowStaysOnTopHint)
            if cur != pend:
                self.apply_always_on_top(pend, defer_if_modal=False)
        except Exception:
            _log_exc("flush_pending_on_top")

    def _drop_karaoke_for(self, editor):
        """该编辑器的字幕状态清零：内存边界表 + 卡拉OK高亮。不动音频缓存bytes。"""
        try:
            hl = self._hl_input if editor is self.input_edit else self._hl_output
            hl.clear_range()
        except Exception:
            pass
        if getattr(self, "_speak_editor", None) is editor:
            self._speak_boundaries = []
            self._speak_span = None

    def _sel_or_link_range(self, editor):
        """该区当前选区范围：原生蓝选 > 高亮蓝选 > 灰色联动，都无则 None。"""
        cur = editor.textCursor()
        if cur.hasSelection():
            return (cur.selectionStart(), cur.selectionEnd())
        hl = self._hl_input if editor is self.input_edit else self._hl_output
        if hl._sel_end > hl._sel_start:
            return (hl._sel_start, hl._sel_end)
        if hl._link_end > hl._link_start:
            return (hl._link_start, hl._link_end)
        return None

    def _smart_copy(self, editor):
        """有选区(蓝/灰)只复制选中部分；无选区复制全部。"""
        from PyQt6.QtWidgets import QApplication as _QA
        r = self._sel_or_link_range(editor)
        text = editor.toPlainText()
        _QA.clipboard().setText(text[r[0]:r[1]] if r else text)
        self.status.showMessage(L("已复制选中部分") if r else L("已复制全部文字"), 2000)

    def _smart_paste(self, editor):
        """有选区(蓝/灰)只覆盖选中部分；无选区默认粘贴到末尾。"""
        from PyQt6.QtWidgets import QApplication as _QA
        from PyQt6.QtGui import QTextCursor
        clip = _QA.clipboard().text()
        if not clip:
            self.status.showMessage(L("剪贴板为空"), 2000)
            return
        r = self._sel_or_link_range(editor)
        c = editor.textCursor()
        if r:
            c.setPosition(r[0])
            c.setPosition(r[1], QTextCursor.MoveMode.KeepAnchor)
            c.insertText(clip)
        elif editor is getattr(self, "_active_editor", None):
            # 主动区有明确光标位置：粘贴到光标处
            c.insertText(clip)
        else:
            # 从属区且无选区：默认粘贴到末尾
            c.movePosition(QTextCursor.MoveOperation.End)
            editor.setTextCursor(c)
            c.insertText(clip)

    def _panel_buttons(self, editor):
        """返回某个文本框的 粘贴/复制/删除 按钮横排。"""
        box = QHBoxLayout()
        box.setSpacing(3)
        box.setContentsMargins(0, 0, 0, 0)
        paste_btn = QPushButton(Icons.icon("paste"), "")
        paste_btn.setProperty("_icn", "paste")
        paste_btn.setFixedSize(36, 36)
        paste_btn.setObjectName("toolbtn")
        paste_btn.setToolTip(L("粘贴"))
        paste_btn.clicked.connect(lambda: self._smart_paste(editor))
        copy_btn = QPushButton(Icons.icon("copy"), "")
        copy_btn.setProperty("_icn", "copy")
        copy_btn.setFixedSize(36, 36)
        copy_btn.setObjectName("toolbtn")
        copy_btn.setToolTip(L("复制"))
        copy_btn.clicked.connect(lambda: self._smart_copy(editor))
        box.addWidget(copy_btn)
        box.addWidget(paste_btn)
        return box

    def _paste(self, editor):
        text = QApplication.clipboard().text()
        if text:
            editor.insertPlainText(text)
        self.status.showMessage("已粘贴", 2000)

    def _copy(self, editor):
        QApplication.clipboard().setText(editor.toPlainText())
        self.status.showMessage("已复制到剪贴板", 2000)

    # ---------- 状态栏 ----------
    def _build_statusbar(self):
        self.status = QStatusBar()
        self.status.setSizeGripEnabled(False)   # 去掉 Windows 右下角的灰色拖拽块
        self.setStatusBar(self.status)
        self.status.showMessage(L("就绪"))

    # ---------- 样式 ----------
    def _mac_hybrid_qss(self):
        """mac 最终混合方案样式表（真机验证成功）。按当前 AppKit 外观取深浅两套色。
        绘制：下拉闭合框+方按钮+普通按钮+特殊青色；原生：弹出项+气球+滚动条。"""
        light = _mac_current_is_light()
        if light:
            bg, tx, bd, hv = "#e8e8ea", "#1f1f22", "#c4c4c8", "#dcdce0"
            cbg, ctx, cbd = "#ffffff", "#1f1f22", "#c4c4c8"
            arrow = "#5a5a5a"
        else:
            bg, tx, bd, hv = "#3a3a3c", "#f0f0f0", "#4a4a4a", "#48484b"
            cbg, ctx, cbd = "#2d2d30", "#e8e8e8", "#3a3a3a"
            arrow = "#9aa0a6"
        ch = self._chevron_path(arrow)
        ch_hi = self._chevron_path("#4ea1ff")
        arrow_css = ""
        if ch:
            arrow_css = (
                f"QComboBox::down-arrow {{ image:url('{ch}'); "
                f"width:12px; height:8px; margin-right:8px; }}\n"
                f"QComboBox::down-arrow:hover {{ image:url('{ch_hi}'); }}\n")
        return f"""
            QTextEdit {{ border-radius:6px; padding:8px; font-size:14px; }}
            QTextEdit[activeRegion="1"] {{ border:2px solid #4ea1ff; }}

            /* 正方形工具按钮：绘制、圆角、深浅、淡蓝按下 */
            QPushButton#toolbtn {{ background:{bg}; color:{tx};
                border:1px solid {bd}; border-radius:8px; }}
            QPushButton#toolbtn:hover {{ background:{hv}; border:1px solid #4ea1ff; border-radius:8px; }}
            QPushButton#toolbtn:pressed {{ background:rgba(78,161,255,0.25);
                border:1px solid #4ea1ff; border-radius:8px; }}
            QPushButton#toolbtn:checked {{ background:#00b3c6; color:white;
                border:1px solid #00b3c6; border-radius:10px; }}

            /* 普通按钮：按下即回弹，淡蓝反馈 */
            QPushButton {{ background:{bg}; color:{tx};
                border:1px solid {bd}; border-radius:6px; padding:6px 12px; }}
            QPushButton:hover {{ background:{hv}; border:1px solid #4ea1ff; }}
            QPushButton:pressed {{ background:rgba(78,161,255,0.25);
                border:1px solid #4ea1ff; }}

            /* 特殊按钮：checkable 保持按下时青色 */
            QPushButton:checked {{ background:#00b3c6; color:white;
                border:1px solid #00b3c6; }}

            /* 蓝色主按钮 */
            QPushButton#primary {{ background:#1e88e5; border:none; font-size:15px;
                font-weight:bold; padding:8px 40px; border-radius:8px; color:white; }}
            QPushButton#primary:hover {{ background:#2b95ef; }}
            QPushButton#primary:pressed {{ background:#1565c0; }}
            QPushButton#primary:disabled {{ background:#12557f; color:#dfe8f2; }}

            /* 下拉闭合框(第一部分)：绘制 + V形箭头图片 */
            QComboBox {{ background:{cbg}; color:{ctx}; border:1px solid {cbd};
                border-radius:8px; padding:5px 10px; }}
            QComboBox:hover {{ border:1px solid #4ea1ff; }}
            QComboBox::drop-down {{ border:none; width:22px; }}
            {arrow_css}
            /* 下拉弹出项(第二部分)：照搬v25完美版——border:none + 背景 + 蓝色高亮，
               显式声明阻断父窗QComboBox规则渗透(那会加方边框破坏圆角) */
            QComboBox QAbstractItemView {{ background:{cbg}; outline:none; border:none;
                selection-background-color:#0e639c; selection-color:white; }}
            QComboBox QAbstractItemView::item {{ padding:7px 14px; border:none; }}
            QComboBox QAbstractItemView::item:selected {{ background:#0e639c; color:white; }}
            QComboBox QAbstractItemView::item:hover {{ background:#0e639c; color:white; }}

            /* 朗读进度条(普通slider)：左侧蓝色已读 */
            QSlider::groove:horizontal {{ height:4px; background:{bd}; border-radius:2px; }}
            QSlider::sub-page:horizontal {{ background:#4ea1ff; border-radius:2px; }}
            QSlider::add-page:horizontal {{ background:{bd}; border-radius:2px; }}
            QSlider::handle:horizontal {{ background:#ffffff; width:14px; height:14px;
                margin:-5px 0; border-radius:7px; border:none; }}
            /* 朗读速度滑杆(rateSlider)：左右滑槽都灰、圆球灰，不要蓝 */
            QSlider#rateSlider::sub-page:horizontal {{ background:#8a8a8a; border-radius:2px; }}
            QSlider#rateSlider::add-page:horizontal {{ background:#8a8a8a; border-radius:2px; }}
            QSlider#rateSlider::handle:horizontal {{ background:#ffffff;
                width:14px; height:14px; margin:-5px 0; border-radius:7px; border:none; }}

            QStatusBar {{ background:#007acc; color:white; }}
            QStatusBar QLabel {{ background:transparent; color:white; }}
            QStatusBar::item {{ border:none; }}
            QSizeGrip {{ background:transparent; width:0; height:0; }}
            """

    def _win_hybrid_qss(self):
        """Windows/Linux 混合方案样式表(真机测试验证)：与 mac 混合同构，但深浅由
        Qt 的 setColorScheme 驱动原生控件(复选框/滚动条/窗口底色走原生)，本 QSS 只绘制
        按钮/下拉闭合框/下拉弹出(蓝色高亮)/滑杆/状态栏，不碰复选框指示器与滚动条。"""
        light = _theme_is_light()
        if light:
            bg, tx, bd, hv = "#e9eaec", "#1f1f22", "#c9c9cc", "#dcdce0"
            cbg, ctx, cbd = "#ffffff", "#1f1f22", "#c4c4c8"
            edit_bg, win_tx = "#ffffff", "#1f1f1f"
            arrow = "#5a5a5a"
        else:
            bg, tx, bd, hv = "#2d2d30", "#dcdcdc", "#3a3a3a", "#37373d"
            cbg, ctx, cbd = "#2d2d30", "#dcdcdc", "#3a3a3a"
            edit_bg, win_tx = "#252526", "#dcdcdc"
            arrow = "#9aa0a6"
        ch = self._chevron_path(arrow)
        ch_hi = self._chevron_path("#4ea1ff")
        arrow_css = ""
        if ch:
            arrow_css = (
                f"QComboBox::down-arrow {{ image:url('{ch}'); "
                f"width:12px; height:8px; margin-right:8px; }}\n"
                f"QComboBox::down-arrow:hover {{ image:url('{ch_hi}'); }}\n")
        return f"""
            /* 标签/复选框只设文字色，背景与指示器交给原生(setColorScheme 驱动) */
            QLabel {{ background:transparent; color:{win_tx}; }}
            /* 复选框完全交给原生(setColorScheme+调色板驱动)：不设任何 QCheckBox 规则，
               否则 windows11 引擎会对复选框整体接管渲染，导致小方块失去边线。
               文字色由调色板 WindowText 提供。测试程序的混合模式正是这样(无QCheckBox规则)。 */

            QToolBar {{ border:none; padding:4px; spacing:4px; }}
            QToolBar QToolButton {{ color:{win_tx}; padding:6px 10px; border-radius:5px; }}
            QToolBar QToolButton:hover {{ background:{hv}; }}

            QTextEdit {{ background:{edit_bg}; color:{win_tx}; border:1px solid {bd};
                border-radius:6px; padding:8px; font-size:14px; }}
            QTextEdit[activeRegion="1"] {{ border:1px solid #4ea1ff; }}
            QTextEdit[activeRegion="0"] {{ border:1px solid {bd}; }}

            /* 普通按钮：绘制、圆角、深浅、淡蓝反馈 */
            QPushButton {{ background:{bg}; color:{tx}; border:1px solid {bd};
                border-radius:5px; padding:6px 12px; }}
            QPushButton:hover {{ background:{hv}; border:1px solid #4ea1ff; }}
            QPushButton:pressed {{ background:#094771; }}
            QPushButton#primary {{ background:#1e88e5; border:none; font-size:15px;
                border-radius:8px; color:white; font-weight:bold; padding:8px 40px; }}
            QPushButton#primary:hover {{ background:#2b95ef; }}
            QPushButton#primary:disabled {{ background:#12557f; color:#bbccdd; }}

            /* 下拉闭合框 + 单行输入框：绘制加高 + V形箭头
               (QLineEdit 原本与 QComboBox 共用同一条规则，混合方案时漏掉了 ->
                Win10 深色下 API Key 框变白、Win11 出现原生底部亮条。此处恢复) */
            QComboBox, QLineEdit {{ background:{cbg}; color:{ctx}; border:1px solid {cbd};
                border-radius:5px; padding:5px 8px; }}
            QComboBox:hover, QLineEdit:hover {{ border:1px solid #4ea1ff; }}
            QLineEdit:focus {{ border:1px solid #4ea1ff; }}
            QComboBox QAbstractItemView::indicator {{ width:0px; height:0px; }}
            QComboBox::drop-down {{ border:none; background:transparent; width:24px;
                subcontrol-origin:padding; subcontrol-position:center right; }}
            {arrow_css}
            /* 下拉弹出项：蓝色高亮 + 白字(照搬主程序，显式声明阻断父规则渗透) */
            QComboBox QAbstractItemView {{ background:{cbg}; outline:none;
                border:1px solid {cbd};
                selection-background-color:#0e639c; selection-color:white; }}
            QComboBox QAbstractItemView::item {{ padding:7px 14px; border:none; }}
            QComboBox QAbstractItemView::item:selected {{ background:#0e639c; color:white; }}
            QComboBox QAbstractItemView::item:hover {{ background:#0e639c; color:white; }}

            /* 朗读进度条：左侧蓝色已读 */
            QSlider {{ border:none; }}
            QSlider::groove:horizontal {{ height:4px; background:{bd}; border-radius:2px; }}
            QSlider::sub-page:horizontal {{ background:#4ea1ff; border-radius:2px; }}
            QSlider::add-page:horizontal {{ background:{bd}; border-radius:2px; }}
            QSlider::handle:horizontal {{ background:#ffffff; width:14px; height:14px;
                margin:-5px 0; border-radius:7px; border:none; }}
            QSlider#rateSlider::sub-page:horizontal {{ background:#8a8a8a; border-radius:2px; }}
            QSlider#rateSlider::add-page:horizontal {{ background:#8a8a8a; border-radius:2px; }}

            QStatusBar {{ background:#007acc; color:white; }}
            QStatusBar QLabel {{ background:transparent; color:white; }}
            QStatusBar::item {{ border:none; }}
            QSizeGrip {{ background:transparent; width:0; height:0; }}
            """

    def _chevron_path(self, color="#9aa0a6"):
        """生成一个扁平 V 形尖角号 SVG，写入临时文件，供下拉箭头使用。"""
        import tempfile, os as _os
        svg = (f'<svg xmlns="http://www.w3.org/2000/svg" width="12" height="8" '
               f'viewBox="0 0 12 8"><path d="M1 1.5 L6 6.5 L11 1.5" '
               f'fill="none" stroke="{color}" stroke-width="1.8" '
               f'stroke-linecap="round" stroke-linejoin="round"/></svg>')
        p = _os.path.join(tempfile.gettempdir(),
                          f"ec_chevron_{color.strip('#')}.svg")
        try:
            with open(p, "w", encoding="utf-8") as f:
                f.write(svg)
        except Exception:
            return ""
        return p.replace("\\", "/")

    def _apply_style(self):
        chevron = self._chevron_path("#9aa0a6")
        chevron_hi = self._chevron_path("#4ea1ff")
        arrow_css = ""
        if chevron:
            arrow_css = (
                f"QComboBox::down-arrow {{ image:url('{chevron}'); "
                f"width:12px; height:8px; margin-right:8px; }}\n"
                f"QComboBox::down-arrow:hover {{ image:url('{chevron_hi}'); }}\n")
        _ss = ("""
            QMainWindow, QDialog, QMessageBox { background:#1e1e1e; }
            QLabel { background:transparent; color:#dcdcdc; }
            QCheckBox { background:transparent; color:#dcdcdc; }
%TOOLTIP%
            QToolBar { background:#252526; border:none; padding:4px; spacing:4px; }
            QToolBar QToolButton { color:#dcdcdc; padding:6px 10px; border-radius:5px; }
            QToolBar QToolButton:hover { background:#37373d; }
            QTextEdit { background:#252526; border:1px solid #3a3a3a;
                border-radius:6px; padding:%EDITPAD%; font-size:14px; }
            QTextEdit[activeRegion="1"] { border:1px solid #4ea1ff; }
            QTextEdit[activeRegion="0"] { border:1px solid #3a3a3a; }
%SCROLLBAR%
            QComboBox, QLineEdit { background:#2d2d30; border:1px solid #3a3a3a;
                border-radius:5px; padding:5px 8px; }
            QComboBox:hover { border:1px solid #4ea1ff; }
%COMBOPOPUP%
            QComboBox QAbstractItemView::indicator { width:0px; height:0px; }
            QComboBox::drop-down { border:none; background:transparent;
                width:24px; subcontrol-origin:padding; subcontrol-position:center right; }
            """ + arrow_css + """
            QPushButton { background:#2d2d30; color:#dcdcdc; border:1px solid #3a3a3a;
                border-radius:5px; padding:6px 12px; }
            QPushButton:hover { background:#37373d; border:1px solid #4ea1ff; }
            QPushButton:pressed { background:#094771; }
            QPushButton#primary { background:#1e88e5; border:none; font-size:15px;
                border-radius:8px; color:white; font-weight:bold; padding:8px 40px; }
            QPushButton#primary:disabled { background:#12557f; color:#bbccdd; }
            QPushButton#primary:hover { background:#2b95ef; }
            QLabel { color:#cccccc; }
            QSlider { border:none; }
            QSlider#rateSlider::sub-page:horizontal { background:#8a8a8a; border-radius:2px; }
            QSlider#rateSlider::add-page:horizontal { background:#8a8a8a; border-radius:2px; }
            QSlider::groove:horizontal { height:4px; background:#3a3a3a; border-radius:2px;
                border:none; }
            QSlider::handle:horizontal { background:#4ea1ff; width:14px; height:14px;
                margin:-5px 0; border-radius:7px; border:none; }
            QStatusBar { background:#007acc; color:white; }
            QStatusBar QLabel { background:transparent; color:white; }
            QStatusBar::item { border:none; }
            QSizeGrip { background:transparent; width:0; height:0; }
            QSplitter::handle { background:#1e1e1e; width:8px; }
        """)
        # 关键：替换必须作用于整张拼接后的样式表（此前只作用于最后一段，
        # 占位符残留导致整表解析失败被丢弃——按钮丢样式/翻译钮丢蓝色的元凶）
        _sb = self._scrollbar_css()
        _ss = _ss.replace("%SCROLLBAR%", _sb)
        _ss = _ss.replace("%TOOLTIP%", _tooltip_css())
        _ss = _ss.replace("%COMBOPOPUP%", _combo_popup_css())
        _ss = _ss.replace("%EDITPAD%", "8px" if _sb == "" else "8px 20px 8px 8px")
        self._base_ss = _ss              # 保存原始表，供主题热切换重设
        import sys as _sys
        if _sys.platform == "darwin":
            # mac 最终混合方案（真机验证成功）：
            # - 下拉闭合框(第一部分)+方按钮+普通按钮+特殊青色 = 绘制(深浅两套)
            # - 下拉弹出项(第二部分)+气球+滚动条 = 系统原生(不设QSS，悬停蓝条自动)
            # - 深浅由 pyobjc 驱动；切换时本函数按当前外观重生成绘制部分
            _ss = self._mac_hybrid_qss()
            self._base_ss = _ss
            self.setStyleSheet(_ss)
            return
        # 非 mac 混合方案(测试程序验证)：先用 setColorScheme 驱动原生深浅
        # (复选框/滚动条/窗口底色走原生)，再套只绘制按钮/下拉/滑杆的混合 QSS。
        try:
            from PyQt6.QtWidgets import QApplication as _QA
            _apply_color_scheme(_QA.instance())
        except Exception:
            pass
        _ss = self._win_hybrid_qss()
        self._base_ss = _ss
        self.setStyleSheet(_ss)

    def _scrollbar_css(self):
        """Mac / Win11(Qt>=6.7) 用系统原生胶囊滚动条（不设任何样式）；
        Win10 及以下 / Linux 用自绘圆角胶囊(与原生观感一致)。"""
        if _native_scrollbar_platform():
            return ""
        return _rounded_scrollbar_qss()

    # ====================================================================
    #  动作
    # ====================================================================

    def open_settings(self):
        SettingsDialog(self.settings, self).exec()
        self._flush_pending_on_top()   # 应用对话框期间推迟的置顶设置
        # 设置里改了默认引擎，同步到主界面下拉
        _combo_select_data(self.engine_combo, 
            self.settings.value("engine", ENGINE_GOOGLE))

    def swap_sides(self):
        # 多风格模式下只交换直译区，多风格灰字区不参与
        _le = getattr(self, "_lit_end", None)
        # 朗读中交换：把卡拉OK朗读目标也对换到对面窗，避免青绿色显示在错误文字上
        was_speaking = getattr(self, "_is_speaking", False)
        speak_ed = getattr(self, "_speak_editor", None)
        # 记住交换前原文区的选区（若有），交换后继承到新原文区
        in_cur = self.input_edit.textCursor()
        out_cur = self.output_edit.textCursor()
        out_sel = (out_cur.selectionStart(), out_cur.selectionEnd()) if out_cur.hasSelection() else None

        a, b = self.input_edit.toPlainText(), self.output_edit.toPlainText()
        # 多风格模式：译文只取直译区参与交换，多风格灰字区不换过去
        if _le is not None:
            b = b[:_le].rstrip()
        self._lit_end = None
        self.input_edit.setPlainText(b)
        self.output_edit.setPlainText(a)
        s, t = self.src_combo.currentData(), self.tgt_combo.currentData()
        if s in LANG_OPTIONS and t in LANG_OPTIONS:
            _combo_select_data(self.src_combo, t)
            _combo_select_data(self.tgt_combo, s)
        # 清掉旧的高亮（含可能停在错误文字上的青绿/灰条）
        self._link_guard = True
        try:
            self._hl_input.clear_all(); self._hl_output.clear_all()
        finally:
            self._link_guard = False
        # 朗读中交换：立即终止朗读 + 清空音频缓存，避免青绿字幕错位到对面
        if was_speaking:
            self.stop_speak()
        self._last_audio = None
        self._last_audio_sig = None
        self._speak_boundaries = []
        # 交换后内容已对调，触发翻译刷新译文与对齐
        if self.input_edit.toPlainText().strip():
            self._start_translate(auto=True)

    def _kill_repaint_sources(self, side):
        """文字真实变更时，掐断该侧所有高亮重画源（只擦画面会被定时器立刻画回来）：
        1) 清词边界 2) 停卡拉OK定时器 3) 停联动防抖 4) 作废过期对齐表。
        若该侧正在朗读，顺带停止朗读（文字已变，读的是旧内容）。"""
        ed = self.input_edit if side == "src" else self.output_edit
        if getattr(self, "_speak_editor", None) is ed:
            self._speak_boundaries = []
            self._karaoke_running = False
            try:
                self._karaoke_timer.stop()
            except Exception:
                pass
            try:
                self.stop_speak(clear_only=True)
            except Exception:
                pass
        self._link_timer.stop()
        self._link_pending_src = None
        self._align = None   # 文字变了，旧句对齐全部失效

    # ---------- 翻译 ----------
    def _on_input_changed(self):
        self._sync_export_text_buttons()   # 导出文字钮灰化随文字有无即时更新
        # 根治方案：rehighlight 触发的 textChanged 中文本内容并未变化，
        # 用文本比对判断是否"真实编辑"——从根上消灭所有高亮反馈环误伤。
        cur_text = self.input_edit.toPlainText()
        if cur_text == getattr(self, "_last_input_text", ""):
            return                      # 仅格式重绘，非真实变更
        self._last_input_text = cur_text
        if getattr(self, "_highlighting", False) or getattr(self, "_link_guard", False):
            return
        self._kill_repaint_sources("src")   # 掐断卡拉OK/联动重画源
        self._clear_all_highlights()
        self._update_file_buttons()
        if not getattr(self, "_preserve_audio_on_clear", False):
            self._invalidate_side_audio("src")   # 原文变了 -> 原文音频缓存失效
        if not cur_text.strip():
            self.output_edit.clear()
            self._auto_timer.stop()
            return
        self._auto_timer.start()

    def _on_output_changed(self):
        self._sync_export_text_buttons()   # 导出文字钮灰化随文字有无即时更新
        cur_text = self.output_edit.toPlainText()
        if cur_text == getattr(self, "_last_output_text", ""):
            return                      # 仅格式重绘，非真实变更
        self._last_output_text = cur_text
        if getattr(self, "_highlighting", False) or getattr(self, "_link_guard", False):
            return
        if getattr(self, "_filling_output", False):
            return   # 翻译结果回填不算用户编辑
        self._kill_repaint_sources("tgt")   # 掐断卡拉OK/联动重画源
        self._clear_all_highlights()
        self._update_file_buttons()
        if not getattr(self, "_preserve_audio_on_clear", False):
            self._invalidate_side_audio("tgt")   # 译文变了 -> 译文音频缓存失效

    def _invalidate_side_audio(self, side):
        """某侧文字变化 -> 该侧音频缓存作废，朗读钮变灰、下载钮禁用。"""
        # 字幕铁律：字幕捆绑音频——音频作废，该侧内存边界与卡拉OK高亮同亡
        _ed = self.input_edit if side == "src" else self.output_edit
        self._drop_karaoke_for(_ed)
        c = self._side_cache.get(side)
        if not c:
            return
        c["bytes"] = None; c["boundaries"] = []; c["sig"] = None
        self._last_audio = None          # 同步清全局缓存，防止旧重播路径"复活"
        self._last_audio_sig = None
        c["duration"] = 0; c["position"] = 0
        # 若该侧不是正在朗读，才把按钮变灰（正在朗读的青色由播放状态管理）
        if getattr(self, "_active_side", None) != side or not getattr(self, "_is_speaking", False):
            btn = self.speak_src_btn if side == "src" else self.speak_tgt_btn
            dl = self.dl_src_btn if side == "src" else self.dl_tgt_btn
            self._set_speak_btn_active(btn, False)
            dl.setEnabled(False)

    def _clear_all_highlights(self):
        """清掉两区的蓝色选区、灰色联动、绿色卡拉OK全部高亮。"""
        if getattr(self, "_highlighting", False):
            return
        self._highlighting = True
        try:
            for hl in (self._hl_input, self._hl_output):
                hl._sel_start = hl._sel_end = 0
                hl._link_start = hl._link_end = 0
                hl._hl_start = hl._hl_end = 0
                hl._dim_from = -1
                hl.rehighlight()
        finally:
            self._highlighting = False
        self._sel_range = None
        self._sel_editor = None
        self._sel_is_link = False

    def _auto_translate(self):
        # 自动翻译：与手动相同，但失败时不弹窗（仅状态栏提示）
        self._start_translate(auto=True)

    def _on_lang_changed(self):
        """语言下拉变化 -> 强制重新翻译当前内容。"""
        if self.input_edit.toPlainText().strip():
            self.do_translate()

    def do_translate(self):
        # 手动按钮：先中断所有正在进行的翻译（网络/本地推理都先停），再开始新翻译
        self._auto_timer.stop()
        self._abort_translation()
        self._start_translate(auto=False, force=True)

    def _abort_translation(self):
        """中断当前正在进行的翻译 worker（不等待结果）。"""
        w = getattr(self, "translate_worker", None)
        if w is not None and w.isRunning():
            try:
                if hasattr(w, "cancel"):
                    w.cancel()           # 软取消（worker 内部检查点会停）
                w.requestInterruption()
            except Exception:
                pass
            # 退役到后台，避免阻塞 UI；不 join 等待
            if not hasattr(self, "_retired_tworkers"):
                self._retired_tworkers = []
            try:
                w.finished_ok.disconnect()
                w.failed.disconnect()
            except Exception:
                pass
            self._retired_tworkers.append(w)
            w.finished.connect(lambda w=w: self._retired_tworkers.remove(w)
                               if w in self._retired_tworkers else None)
        self.translate_worker = None

    def _in_file_mode(self):
        """是否处于有效的文件导入翻译模式（原文与导入内容strip后一致）。"""
        it = getattr(self, "_imported_text", None)
        return (bool(getattr(self, "_imported_path", None)) and it is not None
                and self.input_edit.toPlainText().strip() == it)

    def _start_translate(self, auto=False, force=False):
        text = self.input_edit.toPlainText().strip()
        if not text:
            if not auto:
                self.status.showMessage(L("请输入要翻译的文本"), 3000)
            self.translate_btn.setEnabled(True)
            return
        # 特殊模式：原文仅为单个/一串符号或数字 -> 本地处理（不走引擎）
        if not self._in_file_mode():
            special = self._maybe_symbol_number(text)
            if special is not None:
                self.output_edit.setPlainText(special)
                self.status.showMessage("已按符号/数字规则转换", 3000)
                self.translate_btn.setEnabled(True)
                return
        # 上一个翻译还在跑：自动模式下稍后再试；手动(force)已在 do_translate 里中断
        if not force and self.translate_worker is not None and self.translate_worker.isRunning():
            if auto:
                self._auto_timer.start()
                return
            self._abort_translation()
        engine = self.engine_combo.currentData()
        keys = {
            "deepl": self.settings.value("deepl_key", ""),
            "google_api": self.settings.value("google_api_key", ""),
            "deepseek": self.settings.value("deepseek_key", ""),
            "hunyuan": self.settings.value("hunyuan_key", ""),
            "openai": self.settings.value("openai_key", ""),
            "gemini": self.settings.value("gemini_key", ""),
            "claude": self.settings.value("claude_key", ""),
            "glm": self.settings.value("glm_key", ""),
            "ernie": self.settings.value("ernie_key", ""),
            "doubao": self.settings.value("doubao_key", ""),
            "qwen": self.settings.value("qwen_key", ""),
            "kimi": self.settings.value("kimi_key", ""),
        }
        # LLM 引擎且开启了多风格开关时，输出多种译法；
        # 仅在"有效的文件导入模式"(原文与导入内容一致)下才禁用多风格——
        # 之前只看 _imported_path 是否为 None，导入过一次后残留路径会永久禁掉多风格(bug)。
        multi = (engine in LLM_ENGINE_SET and
                 self.settings.value("multi_style", "true") == "true" and
                 not self._in_file_mode())
        self._multi_active = multi   # 记录本次是否真的多风格，供 on_translate_ok 判分区
        self._current_auto = auto
        self.translate_btn.setEnabled(False)
        # 按钮文字始终不变（避免闪动），状态提示放到底部状态栏
        self.status.showMessage(f"正在翻译（{engine}）…")

        self.translate_worker = TranslateWorker(
            text, self.src_combo.currentData(), self.tgt_combo.currentData(),
            engine, keys, multi_style=multi)
        self.translate_worker.finished_ok.connect(self.on_translate_ok)
        self.translate_worker.failed.connect(self.on_translate_fail)
        self.translate_worker.start()

    def _maybe_symbol_number(self, text):
        """若原文仅为单个符号/数字，或符号数字混合串，返回本地转换结果；否则 None。
        纯数字给『逐位拼读』+『整数读法』两种；目标语种按 tgt 决定中/英。"""
        import re as _re
        # 仅含数字、常见符号、空白才进入此模式（不含字母/汉字）
        if not text or _re.search(r'[A-Za-z\u4e00-\u9fff]', text):
            return None
        if not _re.fullmatch(r'[\d\s\.\,\-\+\*/=%\$#@!\?\(\)\[\]{}:;~`\^&|<>，。！？；：（）]+', text):
            return None
        tgt = self._resolve_tgt_lang()
        # 修复：下拉项是 "English"，旧代码只比对 "英语" 永不命中，导致强制英文失效
        to_en = tgt in ("英语", "English", "英文")
        # 纯数字（可带空格）
        num = _re.sub(r'\s+', '', text)
        if _re.fullmatch(r'\d+(\.\d+)?', num):
            results = []
            # 逐位拼读（小数点读 point/点）
            per = [("point" if to_en else "点") if d == '.'
                   else (_DIGIT_EN[d] if to_en else _DIGIT_ZH[d]) for d in num]
            results.append(("Per-digit" if to_en else "逐位拼读", " ".join(per)))
            # 数学读法（整数与小数都支持）
            try:
                if '.' in num:
                    if to_en:
                        from num2words import num2words
                        whole = num2words(float(num))
                    else:
                        import cn2an
                        whole = cn2an.an2cn(num)
                else:
                    whole = self._num_to_words(int(num), to_en)
                results.append(("As a number" if to_en else "数学读法", whole))
            except Exception:
                pass
            sep = ": " if to_en else "："
            return "\n".join(f"{label}{sep}{val}" for label, val in results)
        # 符号/混合串：逐字符映射为名称
        mapped = []
        for ch in text:
            if ch.isspace():
                mapped.append(ch); continue
            if ch.isdigit():
                mapped.append((_DIGIT_EN if to_en else _DIGIT_ZH)[ch])
            else:
                mapped.append((_SYM_EN if to_en else _SYM_ZH).get(ch, ch))
        return " ".join(m for m in mapped if m.strip())

    def _resolve_tgt_lang(self):
        t = self.tgt_combo.currentData()
        if t == "自动检测":
            # 原文是数字/符号，默认中文输入 -> 英文；否则中文
            return "英语"
        return t

    def _num_to_words(self, n, to_en):
        if to_en:
            try:
                from num2words import num2words
                return num2words(n)
            except Exception:
                return str(n)
        else:
            try:
                import cn2an
                return cn2an.an2cn(str(n))
            except Exception:
                return str(n)

    def _sanitize_literal(self, out):
        """强删多风格直译区可能混入的结构标注：Part 1/第一部分/直译区/----/【..】/Literal 等。
        代码层兜底，不依赖模型是否遵守 prompt。"""
        import re as _re
        lines = out.split("\n")
        cleaned = []
        for ln in lines:
            t = ln.strip()
            if _re.fullmatch(r'[-=_*·—\s]{3,}', t):   # 纯分隔线
                continue
            if _re.match(r'^\s*(Part\s*\d|第[一二三四五]部分|直译区|多风格区|'
                         r'Literal|Direct translation|Multi-?style)\s*[:：]?\s*$',
                         t, _re.I):
                continue
            t2 = _re.sub(r'^\s*[【\[（(]\s*(直译区|多风格区|Part\s*\d|第[一二三四五]部分)'
                         r'\s*[】\])）]\s*[:：]?\s*', '', ln)
            cleaned.append(t2)
        # 合并因删除产生的连续多空行为单个空行
        res = _re.sub(r'\n{3,}', '\n\n', "\n".join(cleaned)).strip("\n")
        return res

    def on_translate_ok(self, out):
        _multi = getattr(self, "_multi_active", False)
        if _multi or "\n----" in out:
            out = self._sanitize_literal(out)
        # 多风格分区解析：直译区(参与对齐联动) 与 多风格区(灰字、不联动) 的边界。
        # 关键(#3 修复)：只有在多风格模式真正激活时才做分区。否则——例如普通翻译
        # 且原文本身含空行(“下载\n\n计算机”)——译文也会有空行，绝不能拿它当分界，
        # 否则空行后的内容(computer)会被误判成多风格区而变灰。普通模式 _lit_end 恒 None。
        if _multi:
            _m = re.search(r"\n[ \t]*\n", out)
            self._lit_end = _m.start() if _m else None
        else:
            self._lit_end = None
        self._filling_output = True
        try:
            self.output_edit.setPlainText(out)
        finally:
            self._filling_output = False
        self._reset_translate_btn()
        self.status.showMessage(L("翻译完成"), 3000)
        self._update_file_buttons()
        # 建立"原文句 <-> 译文句"对应关系，供选区联动精确定位
        try:
            src = self.input_edit.toPlainText()
            self._build_alignment(
                src, out if self._lit_end is None else out[:self._lit_end])
            self._highlighting = True
            try:
                self._hl_output.set_dim(
                    self._lit_end if self._lit_end is not None else -1)
            finally:
                self._highlighting = False
            _add_history(src, out, self.engine_combo.currentData())
        except Exception as e:
            _log_error(f"记录历史/对齐失败: {e}")

    def _build_alignment(self, src, tgt):
        """把原文、译文各自按句切分，建立句级对应（按顺序一一对应）。
        存成 [(src_a,src_b,tgt_a,tgt_b), ...]，选区联动时据此定位。"""
        src_segs = _sentence_spans(src)
        tgt_segs = _sentence_spans(tgt)
        self._align = []
        n = max(len(src_segs), 1)
        m = len(tgt_segs)
        for i, (sa, sb) in enumerate(src_segs):
            # 按比例把第 i 个原文句映射到对应的译文句索引
            if m == 0:
                break
            j = int(round(i / max(1, len(src_segs) - 1) * (m - 1))) if len(src_segs) > 1 else 0
            ta, tb = tgt_segs[min(j, m - 1)]
            self._align.append((sa, sb, ta, tb))

    def on_translate_fail(self, msg):
        self._reset_translate_btn()
        self.status.showMessage("翻译失败", 3000)
        _log_error(f"翻译失败 [{self.engine_combo.currentData()}]: {msg}")
        # 自动翻译失败不打扰；手动翻译才弹窗
        if not getattr(self, "_current_auto", False):
            QMessageBox.warning(self, "翻译失败", msg)

    def _reset_translate_btn(self):
        self.translate_btn.setEnabled(True)
        # 文字始终保持L("翻译")不变（不再改文案，避免闪动）

    def _load_prev_source(self):
        """循环载入历史中的所有原文（每点一次往前一条，到头回到最新）。"""
        items = _load_history()
        if not items:
            self.status.showMessage(L("暂无历史记录"), 2500)
            return
        # 维护一个游标，在所有历史原文间循环
        srcs = [it.get("src", "") for it in items if it.get("src", "").strip()]
        if not srcs:
            self.status.showMessage(L("暂无历史原文"), 2500)
            return
        idx = getattr(self, "_prev_src_idx", None)
        cur = self.input_edit.toPlainText().strip()
        if idx is None:
            # 首次：从最新一条开始；若当前正是最新，则跳到上一条
            idx = len(srcs) - 1
            if srcs[idx].strip() == cur and len(srcs) > 1:
                idx -= 1
        else:
            idx -= 1
            if idx < 0:
                idx = len(srcs) - 1   # 循环回最新
        self._prev_src_idx = idx
        self.input_edit.setPlainText(srcs[idx])
        self.status.showMessage(f"已载入历史原文 {idx+1}/{len(srcs)}", 2500)

    def _load_next_source(self):
        """与『上一条』相反方向循环载入历史原文。"""
        items = _load_history()
        srcs = [it.get("src", "") for it in items if it.get("src", "").strip()]
        if not srcs:
            self.status.showMessage(L("暂无历史原文"), 2500)
            return
        idx = getattr(self, "_prev_src_idx", None)
        if idx is None:
            idx = 0
        else:
            idx += 1
            if idx >= len(srcs):
                idx = 0   # 循环回最早
        self._prev_src_idx = idx
        self.input_edit.setPlainText(srcs[idx])
        self.status.showMessage(f"已载入历史原文 {idx+1}/{len(srcs)}", 2500)

    def _open_history_dialog(self):
        items = _load_history()
        if not items:
            self.status.showMessage(L("暂无历史记录"), 2500)
            return
        dlg = HistoryDialog(items, self)
        if dlg.exec() and dlg.chosen is not None:
            src = dlg.chosen.get("src", "")
            self.input_edit.setPlainText(src)
            # 载入后开始翻译，并记为新记录
            self._start_translate(auto=False)

    # ---------- 朗读 ----------
    def _toggle_speak(self, editor):
        """开始 / 暂停 / 继续 三态切换。跨侧朗读时，先暂停另一侧（不停止）。
        关键：若该区有"新的选区"（与当前朗读内容不同），视为新的朗读请求，
        优先按选区重新朗读——否则三态切换会把选区朗读吃掉（v2.1.2 引入的 bug）。"""
        if self.player is None:
            self.status.showMessage("音频后端不可用，无法播放", 3000)
            return
        state = self.player.playbackState()
        # 若正在播放/暂停当前这一栏，优先按"暂停/继续"处理，不被选区判断拦截。
        # 关键(v2.14.9 修复全选朗读时暂停失效)：全选朗读时编辑器里的选区仍然
        # 存在，且卡拉OK/结束回填会让 _sel_range 与当前选区不再精确相等，于是
        # 旧逻辑把"暂停点击"误判为新朗读请求而重启，导致点暂停没反应(在重播)。
        # 现在：同栏且正在播放 -> 直接暂停；同栏且已暂停 -> 直接继续。
        same_side_active = (
            self._speak_editor is editor and
            state in (QMediaPlayer.PlaybackState.PlayingState,
                      QMediaPlayer.PlaybackState.PausedState))
        cur = editor.textCursor()
        if cur.hasSelection():
            new_range = (cur.selectionStart(), cur.selectionEnd())
            # 判断这个选区是不是"正在朗读的那一段"：全选朗读时选区仍在，且卡拉OK/
            # 结束回填会让 _sel_range 变化，故不能只比 _sel_range，还要比朗读范围
            # _speak_span / _speak_scope。任一匹配即视为"同一段"，走暂停/继续。
            _spans = [getattr(self, "_sel_range", None),
                      getattr(self, "_speak_span", None),
                      getattr(self, "_speak_scope", None)]
            is_same_span = new_range in [sp for sp in _spans if sp]
            # 只有"新的、不同的选区"才重启朗读；同段(或同栏正在播放)交给三态暂停/继续
            if not (same_side_active and is_same_span):
                if (self._speak_editor is not editor or
                        getattr(self, "_sel_range", None) != new_range):
                    if state != QMediaPlayer.PlaybackState.StoppedState:
                        self.stop_speak(clear_only=True)
                    self.do_speak(editor)
                    return
        # 正在播放当前栏 -> 暂停（把位置记入该侧缓存，供跨侧回来续播）
        if (self._speak_editor is editor and
                state == QMediaPlayer.PlaybackState.PlayingState):
            self._side_cache[self._active_side]["position"] = self.player.position()
            self.player.pause()
            return
        # 已暂停当前栏 -> 继续（同侧续播不需要缓存位置，清掉防止以后误回跳）
        if (self._speak_editor is editor and
                state == QMediaPlayer.PlaybackState.PausedState):
            self._side_cache[self._active_side]["position"] = 0
            self.player.play()
            return
        # 点了另一侧：若当前有一侧在朗读，先把它暂停在当前位置（按钮保持青色→继续朗读）
        if (self._speak_editor is not None and self._speak_editor is not editor and
                state == QMediaPlayer.PlaybackState.PlayingState):
            self._pause_other_side()
        # 开始朗读这一栏
        self.do_speak(editor)

    def _pause_other_side(self):
        """把当前正在朗读的一侧暂停（保留进度与青色按钮，显示继续朗读态）。"""
        other = self._speak_editor
        if other is None or self.player is None:
            return
        # 记住该侧的暂停位置与音频缓存，便于以后继续
        pos = self.player.position()
        _oside = "src" if other is self.input_edit else "tgt"
        self._side_cache[_oside]["position"] = pos   # 存进该侧缓存，重播时从此处续
        self._paused_side_editor = other
        self._paused_side_pos = pos
        # 暂停播放器（会触发 on_play_state 把该侧按钮设为『继续朗读』青色）
        self.player.pause()
        # 该侧按钮强制设成青色继续态（喇叭图标 + 青底）
        btn = self.speak_src_btn if other is self.input_edit else self.speak_tgt_btn
        self._set_speak_btn_active(btn, True, icon="speak")
        btn.setToolTip(L("继续朗读"))

    def do_speak(self, editor, from_pos_ratio=None):
        if from_pos_ratio is None:
            self._freeze_slider = False   # 全新朗读不冻结
        # 切换"当前活动侧"别名（进度条/下载按钮指向该侧）
        if editor is self.input_edit:
            self._active_side = "src"
            self.play_slider = self.play_slider_src
            self.download_audio_btn = self.dl_src_btn
        else:
            self._active_side = "tgt"
            self.play_slider = self.play_slider_tgt
            self.download_audio_btn = self.dl_tgt_btn
        # 若有选中文字，只读选中部分；否则看是否有灰色联动区间；再否则读全部
        cursor = editor.textCursor()
        hl = self._hl_input if editor is self.input_edit else self._hl_output
        link_has = hl._link_end > hl._link_start
        if (from_pos_ratio is not None and
                getattr(self, "_sel_editor", None) is editor and
                getattr(self, "_sel_range", None)):
            # 变更嗓音/引擎重读：沿用原选区。此时原生选区早已转为高亮，
            # 重新推导会误判成"读全文"并把蓝色选区清掉（#2 根因之一）
            _a, _b = self._sel_range
            full = editor.toPlainText()
            text = full[_a:_b].replace("\u2029", "\n").strip()
            char_offset = _a
            self._speak_scope = (_a, _b)
        elif cursor.hasSelection():
            text = cursor.selectedText().replace("\u2029", "\n").strip()
            char_offset = cursor.selectionStart()
            self._speak_span = (char_offset, char_offset + len(text))
            self._sel_range = (cursor.selectionStart(), cursor.selectionEnd())
            self._sel_editor = editor
            self._sel_is_link = False    # 用户主动选区 -> 蓝色
            _c = editor.textCursor()
            _c.clearSelection()
            editor.setTextCursor(_c)
        elif link_has:
            # 灰色联动选区也可朗读（与主动选区同等效果），但保持灰色不变蓝
            full = editor.toPlainText()
            char_offset = hl._link_start
            text = full[hl._link_start:hl._link_end].strip()   # 必须先取 text
            self._speak_span = (char_offset, char_offset + len(text))
            self._sel_range = (hl._link_start, hl._link_end)
            self._sel_editor = editor
            self._sel_is_link = True     # 灰色联动区 -> 保持灰色
        else:
            full = editor.toPlainText()
            text = full.strip()
            char_offset = 0
            # 多风格翻译时：译文区无任何选区默认只朗读直译区(第一个空行之前)，
            # 卡拉OK也只对应直译区；有选区仍按选区(含多风格灰区)朗读。
            if editor is self.output_edit:
                _le = getattr(self, "_lit_end", None)
                if _le and 0 < _le <= len(full):
                    _lit = full[:_le]
                    char_offset = len(_lit) - len(_lit.lstrip())
                    text = _lit.strip()
            self._speak_span = (char_offset, char_offset + len(text))
            self._sel_range = None
            self._sel_editor = None
            self._sel_is_link = False
        if not text:
            self.status.showMessage("没有可朗读的文本", 3000)
            return
        # 立即铺好选区底色（蓝/灰），让合成期间界面不变、选区保持可见
        self._setup_selection_highlight()
        # 计算本次朗读的"签名"：文本+嗓音+语速+选区。若与上次已合成的一致，
        # 且有缓存音频，则直接重播，不重新生成（像播放音频文件一样）。
        if _text_is_chinese(text):
            _vname_sig = self.zh_voice_combo.currentData()
        else:
            _vname_sig = self.en_voice_combo.currentData()
        sig = (id(editor), text, _vname_sig, self.rate_slider.value(), self._sel_range)
        self._last_lang = "ZH" if _text_is_chinese(text) else "EN"   # 重播路径也要设，供嗓音切换判断
        _sc = self._side_cache[self._active_side]
        if (from_pos_ratio is None and _sc.get("bytes")
                and _sc.get("sig") == sig):
            # 该侧缓存命中 -> 直接重播，不重新生成（缓存被清空后此路必不命中）
            self._speak_boundaries = _sc.get("boundaries") or []
            _resume = _sc.get("position") or 0   # 必须在内部 stop_speak 之前取！
            _sc["position"] = 0                  # 否则 stop 会把续播位置清零(#1根因)
            self._speak_editor = editor
            self._last_speak_editor = editor
            self._is_speaking = True
            self._pending_seek_ratio = None
            self.stop_speak(clear_only=True)
            self._setup_selection_highlight()
            QTimer.singleShot(600, lambda: setattr(self, "_freeze_slider", False))
            self.status.showMessage(L("播放中…"), 2000)
            self._play_bytes(_sc["bytes"])
            if not self._speak_boundaries:
                QTimer.singleShot(300, self._build_fallback_boundaries)
            if _resume > 0:
                # 缓存里存过暂停位置 -> 续播而不是从头（媒体加载后再定位）
                QTimer.singleShot(150, lambda pms=_resume: self.player.setPosition(pms))
            return
        self._pending_sig = sig
        # 变更重读时保留进度条位置，不跳回头
        keep_progress = from_pos_ratio is not None
        self.stop_speak(clear_only=True, keep_slider=keep_progress)
        self._retire_worker()   # 安全退役上一个 worker，避免线程未结束就析构（崩溃根因）
        self._speak_editor = editor
        self._last_speak_editor = editor
        self._is_speaking = True
        self._pending_seek_ratio = from_pos_ratio
        # 按文本语种选嗓音
        if _text_is_chinese(text):
            voice_name = self.zh_voice_combo.currentData()
            voice_spec = ZH_VOICES.get(voice_name, next(iter(ZH_VOICES.values())))
            self._last_lang = "ZH"
        else:
            voice_name = self.en_voice_combo.currentData()
            voice_spec = EN_VOICES.get(voice_name, next(iter(EN_VOICES.values())))
            self._last_lang = "EN"
        self._last_voice_name = voice_name
        rate = self.rate_slider.value()
        self._speak_rate = rate
        is_kokoro = voice_spec.get("engine") == "kokoro"
        self._synth_in_progress = True
        # 文字提示立即显示并持续到合成结束（不被其它状态信息冲掉）
        self._show_synth_busy(bar=False)
        # 进度条延迟 1.2s 再出（短合成不闪进度条）
        QTimer.singleShot(1200, self._show_synth_busy)

        self.tts_worker = TTSWorker(text, voice_spec, rate, char_offset)
        self.tts_worker.finished_ok.connect(self.on_tts_ok)
        self.tts_worker.failed.connect(self.on_tts_fail)
        self.tts_worker.start()

    def _show_synth_busy(self, bar=True):
        """合成时：文字『正在生成音频…』走左侧消息区(showMessage)，与『翻译完成』等
        交替占用同一位置、不重叠；进度条单独放右侧永久区（药丸形、细轮廓、透明外围）。"""
        if not getattr(self, "_synth_in_progress", False):
            return
        # 文字：用 showMessage 进入左侧消息区（同一位置，自动与其它消息互相覆盖）
        self.status.showMessage(L("正在生成音频…"))
        from PyQt6.QtWidgets import QProgressBar, QWidget as _QW, QHBoxLayout as _QH
        if getattr(self, "_synth_holder", None) is None:
            holder = _QW()
            hl = _QH(holder)
            hl.setContentsMargins(0, 0, 0, 0)
            hl.setSpacing(0)
            pb = PillBusyBar()
            hl.addWidget(pb)
            spacer = _QW()
            hl.addWidget(spacer)
            self._synth_holder = holder
            self._synth_bar = pb
            self._synth_spacer = spacer
            self.status.addPermanentWidget(holder)
        self._synth_holder.show()
        if not bar:
            self._synth_bar.hide()
            self._synth_spacer.setFixedWidth(max(20, int(self.width() * 0.05)))
            return
        self._synth_bar.setFixedWidth(max(300, int(self.width() * 0.42)))
        self._synth_spacer.setFixedWidth(max(20, int(self.width() * 0.05)))
        self._synth_bar.show()

    def _hide_synth_busy(self):
        self._synth_in_progress = False
        holder = getattr(self, "_synth_holder", None)
        if holder is not None:
            holder.hide()
        # 清掉左侧"正在生成音频"消息（让其它状态正常显示）
        if self.status.currentMessage().startswith(L(L("正在生成音频…"))[:6]):
            self.status.clearMessage()

    def _next_audio_filename(self):
        """生成 EC_<语种>_<嗓音>_<日期>_<时间>.<ext>，例如 EC_ZH_XiaoXiao_2026-06-26_210825.mp3。
        Kokoro 离线为 wav，edge 在线为 mp3。保存目录优先用上次用户选过的目录。"""
        import datetime
        save_dir = self.settings.value("last_audio_dir", "")
        if not save_dir or not os.path.isdir(save_dir):
            save_dir = os.path.join(os.path.expanduser("~"), "Downloads")
            if not os.path.isdir(save_dir):
                save_dir = os.path.expanduser("~")
        lang = getattr(self, "_last_lang", "EN")
        vname = getattr(self, "_last_voice_name", "")
        short = VOICE_SHORTNAME.get(vname, "Voice")
        data = getattr(self, "_last_audio", b"")
        # 默认扩展名优先用上次选择的格式，否则按源音频类型
        last_fmt = self.settings.value("last_audio_fmt", "")
        if last_fmt in ("wav", "mp3"):
            ext = last_fmt
        else:
            ext = "wav" if data[:4] == b"RIFF" else "mp3"
        now = datetime.datetime.now()
        stamp = now.strftime("%Y-%m-%d %H%M%S")
        name = f"EC {lang} {short} {stamp}.{ext}"
        return os.path.join(save_dir, name)

    def _clear_side_audio(self, side):
        """仅清空该侧朗读音频缓存：缓存释放，朗读钮青->灰，下载钮失效。不动文字。"""
        ed = self.input_edit if side == "src" else self.output_edit
        if getattr(self, "_speak_editor", None) is ed:
            try:
                self.stop_speak(clear_only=True)
            except Exception:
                pass
        self._invalidate_side_audio(side)
        self.status.showMessage(L("原文朗读音频已清空") if side == "src" else L("译文朗读音频已清空"), 2500)

    def _download_audio(self, side=None):
        if side is None:
            side = getattr(self, "_active_side", "src")
        # 从该侧独立缓存取音频；无则提示（钮为灰时不应能点，双保险）
        data = self._side_cache.get(side, {}).get("bytes") or getattr(self, "_last_audio", None)
        if not data:
            self.status.showMessage("该侧没有可下载的音频（朗读钮为灰色时不可下载）", 3000)
            return
        from PyQt6.QtWidgets import QFileDialog
        default_path = self._next_audio_filename()
        # 让用户在保存对话框里选 wav 或 mp3；记住上次选择作为默认
        wav_filt = "WAV 无损 (*.wav)"
        mp3_filt = "MP3 压缩 (*.mp3)"
        last_fmt = self.settings.value("last_audio_fmt", "")
        if last_fmt == "mp3":
            filters = f"{mp3_filt};;{wav_filt}"
        elif last_fmt == "wav":
            filters = f"{wav_filt};;{mp3_filt}"
        else:
            is_wav_src = data[:4] == b"RIFF"
            filters = f"{wav_filt};;{mp3_filt}" if is_wav_src else f"{mp3_filt};;{wav_filt}"
        is_wav_src = data[:4] == b"RIFF"
        path, sel_filt = QFileDialog.getSaveFileName(
            self, "保存朗读音频", default_path, filters)
        if not path:
            return
        want_mp3 = "mp3" in sel_filt.lower()
        ext = ".mp3" if want_mp3 else ".wav"
        self.settings.setValue("last_audio_fmt", "mp3" if want_mp3 else "wav")  # 记住选择
        # 去掉已有扩展名再补正确的
        base, _e = os.path.splitext(path)
        path = base + ext
        try:
            out = data
            if want_mp3 and is_wav_src:
                out = self._wav_bytes_to_mp3(data)        # Kokoro wav -> mp3
            elif (not want_mp3) and (not is_wav_src):
                # edge 是 mp3，用户却要 wav：解码为 wav
                out = self._mp3_bytes_to_wav(data)
            with open(path, "wb") as f:
                f.write(out)
            self.settings.setValue("last_audio_dir", os.path.dirname(path))
            self.status.showMessage(f"已保存：{path}", 4000)
        except Exception as e:
            QMessageBox.warning(self, "保存失败",
                                f"{e}\n（格式转换可能缺少依赖，可改存另一种格式）")

    def _wav_bytes_to_mp3(self, wav_bytes):
        """WAV 字节 -> MP3 字节。优先用 lameenc，失败则提示。"""
        import io, wave
        with wave.open(io.BytesIO(wav_bytes), "rb") as w:
            ch = w.getnchannels(); sr = w.getframerate()
            pcm = w.readframes(w.getnframes())
        try:
            import lameenc
        except Exception:
            raise RuntimeError("缺少 mp3 编码库 lameenc，无法转 mp3")
        enc = lameenc.Encoder()
        enc.set_bit_rate(192)
        enc.set_in_sample_rate(sr)
        enc.set_channels(ch)
        enc.set_quality(2)
        mp3 = enc.encode(pcm) + enc.flush()
        return bytes(mp3)

    def _mp3_bytes_to_wav(self, mp3_bytes):
        """MP3 字节 -> WAV 字节。用 soundfile/miniaudio 之一解码。"""
        import io
        try:
            import soundfile as sf
            import numpy as np
            data, sr = sf.read(io.BytesIO(mp3_bytes))
            buf = io.BytesIO()
            sf.write(buf, data, sr, format="WAV", subtype="PCM_16")
            return buf.getvalue()
        except Exception:
            raise RuntimeError("无法将 mp3 解码为 wav（缺少解码库）")

    def _retire_worker(self):
        """安全停止并回收正在运行的 TTSWorker，防止线程未结束就被销毁导致崩溃。"""
        w = getattr(self, "tts_worker", None)
        if w is not None and w.isRunning():
            try:
                w.cancel()
                w.finished_ok.disconnect()
                w.failed.disconnect()
            except Exception:
                pass
            # 保留引用直到线程真正结束，由 finished 信号触发 deleteLater
            if not hasattr(self, "_retired_workers"):
                self._retired_workers = []
            self._retired_workers.append(w)
            w.finished.connect(lambda w=w: self._cleanup_worker(w))
            w.quit()
        self.tts_worker = None

    def _cleanup_worker(self, w):
        try:
            w.wait(2000)
        except Exception:
            pass
        if hasattr(self, "_retired_workers") and w in self._retired_workers:
            self._retired_workers.remove(w)
        w.deleteLater()

    def on_tts_ok(self, audio_bytes, boundaries):
        self._hide_synth_busy()
        # 选区朗读时，若引擎返回"相对选区"(从0起算)的词边界，统一平移成全文绝对位置，
        # 否则卡拉OK画在文首或画不出来（中文选区偶发无字幕的根因）
        if boundaries and getattr(self, "_sel_range", None):
            _a, _b = self._sel_range
            if _a > 0 and boundaries[-1][1] <= (_b - _a) + 2:
                boundaries = [(cs + _a, ce + _a, off, dur)
                              for (cs, ce, off, dur) in boundaries]
        self._speak_boundaries = boundaries
        self._last_audio = audio_bytes
        self._last_audio_sig = getattr(self, "_pending_sig", None)
        # 存入当前侧的独立缓存（该侧从此"音频在内存"，朗读钮青色，可下载）
        side = getattr(self, "_active_side", "src")
        c = self._side_cache[side]
        c["bytes"] = audio_bytes
        c["boundaries"] = boundaries
        c["sig"] = self._last_audio_sig
        self.download_audio_btn.setEnabled(bool(audio_bytes))
        # 该侧朗读钮标记为"有音频"（青色）——即使还没开始播放也代表可下载
        self._mark_side_has_audio(side, bool(audio_bytes))
        if self.player is None:
            self.status.showMessage("音频后端不可用，无法播放", 3000)
            return
        self._setup_selection_highlight()
        self._play_bytes(audio_bytes)
        if not boundaries:
            QTimer.singleShot(300, self._build_fallback_boundaries)
        if getattr(self, "_pending_seek_ratio", None):
            r = self._pending_seek_ratio
            self._pending_seek_ratio = None
            QTimer.singleShot(150, lambda: self._seek_ratio(r))
        QTimer.singleShot(600, lambda: setattr(self, "_freeze_slider", False))
        self.status.showMessage(L("播放中…"), 2000)

    def _mark_side_has_audio(self, side, has):
        """标记某侧是否有音频在内存：有->朗读钮青色(可下载)，无->灰色。"""
        btn = self.speak_src_btn if side == "src" else self.speak_tgt_btn
        dl = self.dl_src_btn if side == "src" else self.dl_tgt_btn
        if has:
            self._set_speak_btn_active(btn, True, icon="speak")
        # 下载钮可用性跟随
        dl.setEnabled(has)

    def _build_fallback_boundaries(self, _retry=0):
        """没有 WordBoundary 时，按字符/词在总时长上均匀估算高亮时间表。
        估算范围永远 = 实际朗读范围(_speak_span)：选区/联动区/直译区/全文都一一对应。
        首次播放时媒体时长可能尚未加载(=0)，重试等待而不是放弃(修复首次无字幕)。"""
        if self._speak_boundaries:   # 已经有真实边界
            return
        editor = self._speak_editor
        if editor is None:
            return
        dur = getattr(self, "_play_duration", 0) or (self.player.duration() if self.player else 0)
        if dur <= 0:
            if _retry < 10 and getattr(self, "_is_speaking", False):
                QTimer.singleShot(300, lambda: self._build_fallback_boundaries(_retry + 1))
            return
        full = editor.toPlainText()
        span = getattr(self, "_speak_span", None)
        if span and 0 <= span[0] <= span[1] <= len(full):
            base_off = span[0]
            text = full[span[0]:span[1]]
        else:
            # 兜底的兜底：老逻辑(选区或全文)
            sel = getattr(self, "_sel_range", None)
            if sel and getattr(self, "_sel_editor", None) is editor:
                base_off = sel[0]
                text = full[sel[0]:sel[1]]
            else:
                base_off = 0
                text = full
        if not text.strip():
            return
        import re
        if _text_is_chinese(text):
            tokens = [(m.start(), m.end()) for m in re.finditer(r'\S', text)]
        else:
            tokens = [(m.start(), m.end()) for m in re.finditer(r'\S+', text)]
        if not tokens:
            return
        n = len(tokens)
        bounds = []
        for i, (s, e) in enumerate(tokens):
            off = dur * i / n
            bounds.append((base_off + s, base_off + e, off, dur / n))
        self._speak_boundaries = bounds
        try:
            self._side_cache[getattr(self, "_active_side", "src")]["boundaries"] = \
                list(self._speak_boundaries)
        except Exception:
            pass

    def _play_bytes(self, audio_bytes):
        # 用 QBuffer 把内存中的音频喂给 QMediaPlayer；按头部判断格式提示
        self._audio_qba = QByteArray(audio_bytes)
        self._audio_buffer = QBuffer(self)
        self._audio_buffer.setData(self._audio_qba)
        self._audio_buffer.open(QIODevice.OpenModeFlag.ReadOnly)
        hint = "audio.wav" if audio_bytes[:4] == b"RIFF" else "audio.mp3"
        self.player.setSourceDevice(self._audio_buffer, QUrl(hint))
        self.player.play()
        # 纯墙上时钟驱动高亮：一旦开始播放，时钟自走，完全不依赖 player 状态/位置
        self._karaoke_clock = QElapsedTimer()
        self._karaoke_clock.start()
        self._karaoke_paused_at = 0
        self._karaoke_running = True
        self._karaoke_timer.start()

    def on_tts_fail(self, msg):
        self._hide_synth_busy()
        self._is_speaking = False
        self.tts_worker = None
        self._reset_speak_buttons()
        self.status.showMessage("朗读失败（可改用『-线上联网』嗓音）", 6000)
        _log_error(f"朗读失败: {msg}")

    # ---- 播放器信号 ----
    def _on_play_duration(self, dur):
        self._play_duration = dur

    def _on_play_position(self, pos):
        # 更新进度条（仅用于进度条显示）
        dur = getattr(self, "_play_duration", 0) or self.player.duration()
        if dur > 0 and not self._seeking and not getattr(self, "_freeze_slider", False):
            self.play_slider.setValue(int(pos / dur * 1000))

    def _karaoke_elapsed_ms(self):
        """当前已播放毫秒。优先用 player.position()（与音频精确同步），
        若它不推进（某些 macOS 情况）则回退到墙上时钟。"""
        # 1) 优先播放器真实位置
        if self.player is not None:
            try:
                pos = self.player.position()
            except Exception:
                pos = 0
            if pos > 0:
                self._pos_seen = True
                return pos
        # 2) 回退：墙上时钟
        base = getattr(self, "_karaoke_paused_at", 0)
        clk = getattr(self, "_karaoke_clock", None)
        if clk is not None and getattr(self, "_karaoke_running", False):
            return base + clk.elapsed()
        return base

    def _karaoke_tick(self):
        """定时驱动逐词高亮——纯靠墙上时钟，不检查 playbackState。"""
        if self._speak_editor is None or not self._speak_boundaries:
            return
        if not getattr(self, "_karaoke_running", False):
            return
        self._update_karaoke(self._karaoke_elapsed_ms())

    def _on_play_state(self, state):
        # 仅负责按钮文案与时钟暂停/恢复；高亮由墙上时钟独立驱动
        editor = self._speak_editor
        if state == QMediaPlayer.PlaybackState.PausedState:
            # 暂停：累计已播放时间，停表
            self._karaoke_paused_at = self._karaoke_elapsed_ms()
            self._karaoke_running = False
        elif state == QMediaPlayer.PlaybackState.PlayingState:
            # 播放/继续：重启表（基准是累计值）
            if getattr(self, "_karaoke_clock", None) is None:
                self._karaoke_clock = QElapsedTimer()
            self._karaoke_clock.start()
            self._karaoke_running = True
            if not self._karaoke_timer.isActive():
                self._karaoke_timer.start()
        else:  # Stopped
            self._karaoke_running = False
            self._karaoke_timer.stop()
        # 按钮状态（纯图标）：播放/暂停切换图标 + 青绿背景 + 悬停提示三态
        self._reset_speak_buttons()
        if editor is not None:
            btn = self.speak_src_btn if editor is self.input_edit else self.speak_tgt_btn
            name = "原文" if editor is self.input_edit else "译文"
            if state == QMediaPlayer.PlaybackState.PlayingState:
                self._set_speak_btn_active(btn, True, icon="pause")
                btn.setToolTip(L("暂停朗读"))
            elif state == QMediaPlayer.PlaybackState.PausedState:
                self._set_speak_btn_active(btn, True, icon="speak")
                btn.setToolTip(L("继续朗读"))
            else:
                # 停止态颜色由 _reset_speak_buttons 按缓存决定，这里只更新提示
                btn.setToolTip(L(f"朗读{name}"))
        if state == QMediaPlayer.PlaybackState.StoppedState:
            # 为重读而主动 stop（preserve 期间）不是自然播完：跳过收尾闪亮与清绿，
            # 否则 250ms 后 _finish_karaoke 会把要保持的卡拉OK/选区清掉（#2 根因之二）
            if getattr(self, "_preserve_play_ui", False):
                return
            # 播放结束：读完所有词都点亮一瞬，然后清除绿色已读；
            # 若之前是朗读选区，恢复成普通的原生蓝色选区（可正常点击取消/重选）
            if self._speak_editor is not None and self._speak_boundaries:
                last_end = self._speak_boundaries[-1][1]
                first_start = self._speak_boundaries[0][0]
                self._apply_karaoke_selection(first_start, last_end)
                QTimer.singleShot(250, self._finish_karaoke)
            else:
                self._finish_karaoke()

    def _set_speak_btn_active(self, btn, active, icon="speak"):
        """active=True 青绿背景；icon 决定显示喇叭(speak)还是暂停(pause)图标。
        图标在最后设置，不再被背景样式覆盖（修复之前暂停图标看不到的 bug）。"""
        if active:
            btn.setStyleSheet(
                "QPushButton{background:#5aa8b0; color:#0e2024; border:1px solid #5aa8b0; border-radius:6px;}"
                "QPushButton:hover{background:#66b8c0;}")
            btn.setIcon(Icons.icon(icon, "#0e2024"))   # 深色图标（按状态：喇叭或暂停）
        else:
            btn.setStyleSheet("")
            btn.setIcon(Icons.icon("speak"))           # 恢复浅色喇叭图标

    def _finish_karaoke(self):
        """朗读彻底结束：清绿色已读高亮。
        - 若读的是用户主动选区：恢复为普通原生蓝色选区（点击即可取消/重选）。
        - 若读的是灰色联动区：恢复灰色联动高亮（不消失，也不变蓝）。"""
        self._hl_input.clear_range()
        self._hl_output.clear_range()
        sel = getattr(self, "_sel_range", None)
        ed = getattr(self, "_sel_editor", None)
        is_link = getattr(self, "_sel_is_link", False)
        if sel and ed is not None:
            hl = self._hl_input if ed is self.input_edit else self._hl_output
            if is_link:
                # 恢复灰色联动高亮
                self._link_guard = True
                try:
                    hl.clear_selection()
                    hl.set_link(sel[0], sel[1])
                finally:
                    self._link_guard = False
            else:
                # 用原生选区还原蓝色（普通选区，可正常取消/重选）
                hl.clear_selection()
                from PyQt6.QtGui import QTextCursor
                c = ed.textCursor()
                c.setPosition(sel[0])
                c.setPosition(sel[1], QTextCursor.MoveMode.KeepAnchor)
                ed.setTextCursor(c)
        self._sel_range = None
        self._sel_editor = None
        self._sel_is_link = False

    def _reset_speak_buttons(self):
        # 颜色语义原则：该侧缓存有音频=青色(可下载)，无=灰。播放结束不改此语义。
        for _side, _btn in (("src", self.speak_src_btn), ("tgt", self.speak_tgt_btn)):
            _has = bool(self._side_cache.get(_side, {}).get("bytes"))
            self._set_speak_btn_active(_btn, _has, icon="speak")

    def _clear_selection_pair(self):
        """点主动区：保持主动权，清掉两侧联动选区对（蓝/灰高亮），不碰卡拉OK绿色。"""
        self._link_guard = True
        try:
            for hl in (self._hl_input, self._hl_output):
                hl._sel_start = hl._sel_end = 0
                hl._link_start = hl._link_end = 0
                hl.rehighlight()
        finally:
            self._link_guard = False
        self._link_timer.stop()
        self._link_pending_src = None

    def _activate_region(self, editor):
        """统一的主动/从属切换：点击从属区任何位置 -> 它变主动区（蓝框），
        其灰色联动选区变蓝色；原主动区变从属（灰框），其蓝色选区变灰色联动。
        选区内容两侧都保留，只是蓝/灰跟着主动权走。"""
        from PyQt6.QtGui import QTextCursor
        self._link_timer.stop()           # 防止挂起的联动稍后覆盖本次切换结果
        self._link_pending_src = None
        old = getattr(self, "_active_editor", None)
        self._set_active_editor(editor)   # 蓝框跳到新主动区，旧区灰框
        new_hl = self._hl_input if editor is self.input_edit else self._hl_output
        old_hl = self._hl_output if editor is self.input_edit else self._hl_input
        self._link_guard = True
        try:
            # 1) 原主动区：蓝色选区（原生或高亮）-> 灰色联动
            if old is not None and old is not editor:
                old_cur = old.textCursor()
                if old_cur.hasSelection():
                    s0, s1 = old_cur.selectionStart(), old_cur.selectionEnd()
                    c = old.textCursor(); c.clearSelection(); old.setTextCursor(c)
                    old_hl.set_link(s0, s1)
                elif old_hl._sel_end > old_hl._sel_start:
                    old_hl.set_link(old_hl._sel_start, old_hl._sel_end)
                    old_hl._sel_start = old_hl._sel_end = 0
                    old_hl.rehighlight()
            # 2) 新主动区：灰色联动 -> 蓝色高亮选区
            if new_hl._link_end > new_hl._link_start:
                a, b = new_hl._link_start, new_hl._link_end
                new_hl._link_start = new_hl._link_end = 0
                new_hl.set_selection(a, b)
        finally:
            self._link_guard = False

    def _on_selection_changed(self, editor):
        # 多风格灰字区不参与联动：译文侧在直译区之外的选择不建立对应
        if (editor is getattr(self, "output_edit", None)
                and getattr(self, "_lit_end", None)):
            _c0 = editor.textCursor()
            if _c0.hasSelection() and _c0.selectionStart() >= self._lit_end:
                return
        """任一区选择文字 -> 该区成为『主动区』(蓝框+蓝底)，另一区为『从属区』，
        用句对应关系把选区映射到从属区并铺灰色底。双向均可。"""
        if self._link_guard:
            return
        cur = editor.textCursor()
        if cur.hasSelection():
            # 该区成为主动区
            self._set_active_editor(editor)
            self._link_pending_src = editor
            self._link_timer.start()
        else:
            # 选区被清掉：不再顺手清两侧联动（会误伤切换刚画好的灰条）；
            # 联动选区对的清除统一由点击主动区(_clear_selection_pair)负责
            pass

    def _set_active_editor(self, editor):
        """切换主动区：蓝色边框给主动区，从属区去框。"""
        if getattr(self, "_active_editor", None) is editor:
            return
        self._active_editor = editor
        # 蓝框：主动区 focus 边框色已是蓝，这里通过属性强调
        for ed in (self.input_edit, self.output_edit):
            is_active = ed is editor
            ed.setProperty("activeRegion", "1" if is_active else "0")
            ed.style().unpolish(ed); ed.style().polish(ed)

    def _do_selection_link(self):
        """把主动区选区映射到从属区，铺灰色底。支持原文↔译文双向。"""
        try:
            editor = self._link_pending_src
            if editor is None:
                return
            cur = editor.textCursor()
            if not cur.hasSelection():
                return
            s0, s1 = cur.selectionStart(), cur.selectionEnd()
            from_src = editor is self.input_edit
            other = self.output_edit if from_src else self.input_edit
            other_full = other.toPlainText()
            if not other_full.strip():
                return
            align = getattr(self, "_align", None)
            ts = te = None
            if align:
                lo = hi = None
                for (sa, sb, ta, tb) in align:
                    # 主动侧区间 / 从属侧区间
                    a0, a1 = (sa, sb) if from_src else (ta, tb)
                    o0, o1 = (ta, tb) if from_src else (sa, sb)
                    if a1 > s0 and a0 < s1:          # 选区与该句有重叠
                        lo = o0 if lo is None else min(lo, o0)
                        hi = o1 if hi is None else max(hi, o1)
                if lo is not None:
                    ts, te = lo, hi
            if ts is None:
                span = _proportional_span(editor.toPlainText(), s0, s1, other_full)
                if span:
                    ts, te = span
            if ts is None:
                return
            while ts < te and other_full[ts].isspace():
                ts += 1
            while te > ts and other_full[te-1].isspace():
                te -= 1
            other_hl = self._hl_output if from_src else self._hl_input
            self._link_guard = True
            try:
                other_hl.set_link(ts, te)
            finally:
                self._link_guard = False
        except Exception:
            self._link_guard = False

    def _apply_link_match(self, other_editor, translated, src_ratio=None):
        """在 other_editor 文本中找与 translated 最相似的连续区间，灰色高亮。
        src_ratio 是源选区相对位置，用于优先在目标文本对应段落/位置附近匹配。"""
        try:
            translated = (translated or "").strip()
            if not translated:
                return
            text = other_editor.toPlainText()
            if not text.strip():
                return
            span = _best_match_span(text, translated, src_ratio)
            if span is None:
                return
            s, e = span
            hl = self._hl_input if other_editor is self.input_edit else self._hl_output
            self._link_guard = True
            try:
                hl.set_link(s, e)
            finally:
                self._link_guard = False
        except Exception:
            self._link_guard = False

    def _collect_keys(self):
        return {
            "deepl": self.settings.value("deepl_key", ""),
            "google_api": self.settings.value("google_api_key", ""),
            "deepseek": self.settings.value("deepseek_key", ""),
            "hunyuan": self.settings.value("hunyuan_key", ""),
            "openai": self.settings.value("openai_key", ""),
            "gemini": self.settings.value("gemini_key", ""),
            "claude": self.settings.value("claude_key", ""),
            "glm": self.settings.value("glm_key", ""),
            "ernie": self.settings.value("ernie_key", ""),
            "doubao": self.settings.value("doubao_key", ""),
            "qwen": self.settings.value("qwen_key", ""),
            "kimi": self.settings.value("kimi_key", ""),
        }

    def _on_selection_changed_END(self):
        pass

    def _update_karaoke(self, pos_ms):
        """根据当前播放位置高亮已读部分（青绿色）。"""
        if self._speak_editor is None:
            return
        bounds = self._speak_boundaries
        if not bounds:
            return
        # 保险丝：边界超出当前文字长度说明已过期（文字被改过），自动清空防错位重画
        try:
            if bounds[-1][1] > len(self._speak_editor.toPlainText()):
                self._speak_boundaries = []
                return
        except Exception:
            pass
        # 进度拖到最左(位置≈0)时不应有任何卡拉OK效果，连第一个词也不点亮。
        # 之前因提前量(lead)会让 off-lead<=0 成立而误亮第一个词(#2)。
        if pos_ms <= 0:
            self._apply_karaoke_selection(bounds[0][0], bounds[0][0])
            return
        # 提前量：取中间值。之前 350ms 偏多导致字幕抢拍，180ms 偏少又滞后。
        LEAD_FIXED = 120   # 固定提前 120ms
        cur_end = None
        cur_start = bounds[0][0]
        for (cs, ce, off, dur) in bounds:
            lead = LEAD_FIXED + (dur * 0.25 if dur else 0)
            # 用 off>0 时才允许提前量，避免第一个词(off≈0)在 pos=0 被提前点亮
            _eff = off - lead if off > lead else off
            if _eff <= pos_ms:
                cur_end = ce
            else:
                break
        if cur_end is None:
            # 还没到第一个词的时间点，清空已读绿色（但保留蓝色选区）
            self._apply_karaoke_selection(cur_start, cur_start)
            return
        self._apply_karaoke_selection(cur_start, cur_end)

    def _apply_karaoke_selection(self, start, end):
        editor = self._speak_editor
        hl = self._hl_input if editor is self.input_edit else self._hl_output
        # rehighlight 会触发 textChanged -> 用 guard 防止 _on_input_changed 清掉本高亮
        self._link_guard = True
        try:
            hl.set_range(start, end)
        finally:
            self._link_guard = False

    def _clear_karaoke(self):
        # 清掉绿色已读；若之前是朗读选区，恢复蓝色选区底色
        self._hl_input.clear_range()
        self._hl_output.clear_range()

    def _setup_selection_highlight(self):
        """朗读选区时铺底色：用户选区=蓝色；灰色联动区=灰色（保持不变蓝）。"""
        sel = getattr(self, "_sel_range", None)
        ed = getattr(self, "_sel_editor", None)
        is_link = getattr(self, "_sel_is_link", False)
        self._link_guard = True
        try:
            self._hl_input.clear_selection()
            self._hl_output.clear_selection()
            if sel and ed is not None:
                hl = self._hl_input if ed is self.input_edit else self._hl_output
                if is_link:
                    hl.set_link(sel[0], sel[1])      # 保持灰色
                else:
                    hl.clear_link()
                    hl.set_selection(sel[0], sel[1]) # 蓝色
        finally:
            self._link_guard = False

    def _restore_selection_highlight(self):
        """朗读结束：清绿色，恢复蓝色选区（若有）。"""
        self._hl_input.clear_range()
        self._hl_output.clear_range()
        # set_selection 已在 _setup 时设好，clear_range 后蓝色仍在；无选区则什么都不显示

    # ---- 进度条拖动 ----
    def _persist_rate(self):
        try:
            self.settings.setValue("tts_rate", self.rate_slider.value())
        except Exception:
            pass

    def _update_rate_tooltip(self, value=None):
        """语速气球：0 显示『朗读语速 正常』，否则带正负百分比。
        拖动时用 QToolTip 在滑块上方持续显示，不消失。"""
        if value is None:
            value = self.rate_slider.value()
        if value == 0:
            txt = L(L("朗读语速 正常"))
        else:
            txt = L("朗读语速") + f" {'+' if value > 0 else ''}{value}%"
        self.rate_slider.setToolTip(txt)
        # 拖动时在滑块手柄上方持续显示气球
        from PyQt6.QtWidgets import QToolTip
        from PyQt6.QtCore import QPoint
        sl = self.rate_slider
        rng = sl.maximum() - sl.minimum()
        if rng > 0:
            frac = (value - sl.minimum()) / rng
            x = int(frac * sl.width())
            gp = sl.mapToGlobal(QPoint(x, -6))
            QToolTip.showText(gp, txt, sl)

    def _on_seek_start(self, side=None):
        # 只有拖动"当前朗读侧"的进度条才生效；拖另一侧无效（互不影响）
        if side is not None and side != getattr(self, "_active_side", "src"):
            self._seeking = False
            return
        self._seeking = True

    def _on_seek_moved(self, value, side=None):
        """拖动进度条时实时刷新青色已读位置。只对当前朗读侧生效。
        字幕铁律：该侧音频缓存在 且 该侧文字在 → 拖动必有字幕(边界丢了就恢复/重建)；
        任一不在 → 必无字幕。字幕与音频捆绑、依附文字，一起出现一起消失。"""
        s = side if side is not None else getattr(self, "_active_side", "src")
        # 两侧独立：拖动的不是当前朗读侧则不刷字幕（修复滑原文区译文出字幕）
        if s != getattr(self, "_active_side", "src") or self.player is None:
            return
        ed = self.input_edit if s == "src" else self.output_edit
        txt = ed.toPlainText()
        if not txt.strip():
            return   # 文字不在 → 无字幕(音频可照常播放)
        c = self._side_cache.get(s, {}) if hasattr(self, "_side_cache") else {}
        if not (c.get("bytes") or getattr(self, "_is_speaking", False)):
            return   # 音频不在(缓存无bytes且当前非在读) → 无字幕
        # 两者都在：边界必须有——内存空则先从缓存恢复(缓存bytes在即文字未变，必匹配)，
        # 缓存也没有再按实际朗读范围重建估算
        if not self._speak_boundaries:
            b = c.get("boundaries") or []
            if b and b[-1][1] <= len(txt):
                self._speak_boundaries = list(b)
            else:
                try:
                    self._build_fallback_boundaries()
                except Exception:
                    pass
        if not self._speak_boundaries:
            return
        dur = getattr(self, "_play_duration", 0) or (self.player.duration() if self.player else 0)
        if dur > 0:
            pos_ms = dur * (value / 1000.0)
            self._update_karaoke(pos_ms)

    def _on_seek_end(self):
        self._seeking = False
        self._seek_ratio(self.play_slider.value() / 1000.0)

    def _seek_ratio(self, ratio):
        if self.player is None:
            return
        dur = getattr(self, "_play_duration", 0) or self.player.duration()
        if dur > 0:
            self.player.setPosition(int(dur * ratio))
            # 立即刷新高亮到目标位置（暂停时也更新）
            if self._speak_boundaries:
                self._update_karaoke(dur * ratio)

    def _on_en_voice_changed(self, name):
        self.settings.setValue("en_voice", name)
        self._on_voice_or_rate_changed(changed_lang="EN")

    def _on_engine_changed(self, engine):
        self.settings.setValue("engine", engine)
        # 改引擎后自动触发翻译（若有输入文本）
        if self.input_edit.toPlainText().strip():
            self._start_translate(auto=True)

    def _on_zh_voice_changed(self, name):
        self.settings.setValue("zh_voice", name)
        self._on_voice_or_rate_changed(changed_lang="ZH")

    def _on_voice_or_rate_changed(self, changed_lang=None):
        # 朗读进行中更改嗓音/语速/引擎 -> 在当前进度处用新设置重读，
        # 但保持进度条位置不跳回、按钮保持L("继续朗读")青色不闪。
        _playing = False
        if self.player is not None:
            _st = self.player.playbackState()
            _playing = _st in (QMediaPlayer.PlaybackState.PlayingState,
                               QMediaPlayer.PlaybackState.PausedState)
        if not ((getattr(self, "_is_speaking", False) or _playing)
                and self._speak_editor is not None):
            return
        _ll = getattr(self, "_last_lang", None)
        if changed_lang is not None and _ll and changed_lang != _ll:
            return   # 改的是另一种语言的嗓音，与当前朗读无关，不打断
        ratio = 0.0
        dur = getattr(self, "_play_duration", 0)
        if dur > 0 and self.player is not None:
            ratio = self.player.position() / dur
        self._preserve_play_ui = True          # 重读期间保持进度条与按钮状态
        self._preserve_slider_value = self.play_slider.value()
        self._freeze_slider = True             # 生成期间冻结进度条在当前位置
        self.do_speak(self._speak_editor, from_pos_ratio=ratio)
        self._preserve_play_ui = False

    def _stop_side(self, editor):
        """停止指定侧的朗读/暂停/音频生成。当前朗读的是另一侧则不动作。"""
        # 该侧是否是当前朗读/暂停/生成的对象
        is_current = getattr(self, "_speak_editor", None) is editor
        worker_running = getattr(self, "tts_worker", None) and self.tts_worker.isRunning()
        player_active = False
        if self.player is not None:
            st = self.player.playbackState()
            player_active = st in (QMediaPlayer.PlaybackState.PlayingState,
                                   QMediaPlayer.PlaybackState.PausedState)
        if is_current and (worker_running or player_active or getattr(self, "_is_speaking", False)):
            self.stop_speak()
        # 否则本侧无正在进行的朗读，停止键无效

    def _on_seek_end_side(self, side):
        if side != getattr(self, "_active_side", "src"):
            self._seeking = False
            return
        self._on_seek_end()

    # ---------- 导出当前文字 ----------
    def _last_dir(self, key="export"):
        """记住的上次保存目录；首次默认下载目录。"""
        import os
        d = self.settings.value(f"last_dir_{key}", "")
        if d and os.path.isdir(d):
            return d
        return os.path.join(os.path.expanduser("~"), "Downloads")

    def _remember_dir(self, path, key="export"):
        import os
        self.settings.setValue(f"last_dir_{key}", os.path.dirname(path))

    def _gen_filename(self, side, ext=".txt"):
        """文件名规则（空格分隔）：EC OT/TT 语言 日期 时间。
        例如 EC OT ZH 2026-06-30 013733.txt。原文=OT，译文=TT。"""
        import datetime
        text = (self.input_edit if side == "src" else self.output_edit).toPlainText()
        prefix = "OT" if side == "src" else "TT"
        lang = "ZH" if _text_is_chinese(text) else "EN"
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H%M%S")
        return f"EC {prefix} {lang} {ts}{ext}"

    def _export_text(self, side):
        """导出当前原文/译文区的文字（只文字，无时间等信息）。"""
        import os
        text = (self.input_edit if side == "src" else self.output_edit).toPlainText()
        if not text.strip():
            self.status.showMessage("没有可导出的文字", 2500)
            return
        name = "原文" if side == "src" else "译文"
        from PyQt6.QtWidgets import QFileDialog
        filters = "文本 (*.txt);;Markdown (*.md);;Word (*.docx);;JSON (*.json);;PDF (*.pdf)"
        default = os.path.join(self._last_dir("export"), self._gen_filename(side, ".txt"))
        path, sel = QFileDialog.getSaveFileName(self, L(f"导出{name}文字"), default, filters)
        if not path:
            return
        try:
            self._write_text_file(path, text, sel)
            self._remember_dir(path, "export")
            self.status.showMessage(f"已导出：{path}", 4000)
        except Exception as e:
            _log_error(f"导出文字失败: {e}")
            QMessageBox.warning(self, "导出失败", str(e))

    def _write_text_file(self, path, text, sel_filter=""):
        """按扩展名写出文字到 txt/md/json/docx/pdf。"""
        import os
        ext = os.path.splitext(path)[1].lower()
        if ext == ".json":
            with open(path, "w", encoding="utf-8") as f:
                json.dump({"text": text}, f, ensure_ascii=False, indent=2)
        elif ext == ".docx":
            from docx import Document
            doc = Document()
            for line in text.split("\n"):
                doc.add_paragraph(_xml_safe(line))
            doc.save(path)
        elif ext == ".pdf":
            self._write_pdf(path, text)
        else:  # txt / md / 其它都按纯文本
            with open(path, "w", encoding="utf-8") as f:
                f.write(text)

    def _cjk_font(self):
        """注册并返回 PDF 用中文字体名；找不到指定字体则回退系统默认中文字体。"""
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os, sys
        if getattr(self, "_pdf_font_name", None):
            return self._pdf_font_name
        # 优先：Mac=PingFangHK-Light / Win=思源黑体HK Light；回退系统默认中文字体
        candidates = []
        if sys.platform == "darwin":
            candidates = ["/System/Library/Fonts/PingFang.ttc",
                          "/Library/Fonts/PingFang.ttc",
                          "/System/Library/Fonts/STHeiti Light.ttc"]
        elif sys.platform.startswith("win"):
            candidates = [r"C:\Windows\Fonts\SourceHanSansHWHK-Light.otf",
                          r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\msyh.ttf",
                          r"C:\Windows\Fonts\simsun.ttc"]
        else:
            candidates = ["/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                          "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"]
        for fp in candidates:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont("CJK", fp))
                    self._pdf_font_name = "CJK"
                    return "CJK"
                except Exception:
                    continue
        self._pdf_font_name = "Helvetica"
        return "Helvetica"

    def _write_pdf(self, path, text):
        """A4 竖版纯文本 PDF，按可用宽度自动换行（修右侧出画）。"""
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
        from reportlab.pdfbase.pdfmetrics import stringWidth
        font = self._cjk_font()
        size = 12
        c = canvas.Canvas(path, pagesize=A4)
        W, H = A4                       # 595 x 842 pt
        ml = mr = 2 * cm
        mt = mb = 2 * cm
        usable = W - ml - mr
        x = ml
        y = H - mt
        line_h = size * 1.5
        for para in text.split("\n"):
            if para == "":
                y -= line_h
                if y < mb:
                    c.showPage(); c.setFont(font, size); y = H - mt
                continue
            # 按字符累积，超出可用宽度就换行（中英混排通用）
            line = ""
            for ch in para:
                if stringWidth(line + ch, font, size) > usable:
                    c.setFont(font, size); c.drawString(x, y, line); y -= line_h
                    if y < mb:
                        c.showPage(); y = H - mt
                    line = ch
                else:
                    line += ch
            if line:
                c.setFont(font, size); c.drawString(x, y, line); y -= line_h
                if y < mb:
                    c.showPage(); y = H - mt
        c.save()

    # ---------- 导入文件 ----------
    def _import_file(self):
        from PyQt6.QtWidgets import QFileDialog
        filters = "支持的文件 (*.txt *.md *.docx *.pdf);;文本 (*.txt *.md);;Word (*.docx);;PDF (*.pdf)"
        path, _ = QFileDialog.getOpenFileName(self, L("导入文件"), self._last_dir("import"), filters)
        if not path:
            return
        self._remember_dir(path, "import")
        self._do_import(path)

    def _do_import(self, path):
        import os
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext in (".txt", ".md"):
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            elif ext == ".docx":
                from docx import Document
                doc = Document(path)
                text = "\n".join(p.text for p in doc.paragraphs)
            elif ext == ".pdf":
                text = self._read_pdf_text(path)
            else:
                QMessageBox.warning(self, "不支持的格式", f"暂不支持导入：{ext}")
                return
        except Exception as e:
            _log_error(f"导入文件失败: {e}")
            QMessageBox.warning(self, "导入失败", str(e))
            return
        # 关键：先记录导入状态，再 setPlainText。
        # 因为 setPlainText 会同步触发 textChanged -> _on_input_changed -> _update_file_buttons，
        # 若此时导入状态还没记录，会先被判为"不一致"产生错误中间态（这是之前反复失败的根因）。
        self._imported_path = path
        self._imported_ext = ext
        self._imported_text = text.strip()
        self._can_export_file = ext in (".txt", ".md", ".docx")
        self.input_edit.setPlainText(text)
        self._update_file_buttons()
        self.status.showMessage(f"已导入：{os.path.basename(path)}", 4000)
        # 触发翻译
        if text.strip():
            self._start_translate(auto=True)

    def _sync_export_text_buttons(self):
        """导出原文/译文文字钮：该区无文字则灰色不可点（与导出音频钮一致）。"""
        try:
            if hasattr(self, "export_src_btn"):
                self.export_src_btn.setEnabled(
                    bool(self.input_edit.toPlainText().strip()))
            if hasattr(self, "export_tgt_btn"):
                self.export_tgt_btn.setEnabled(
                    bool(self.output_edit.toPlainText().strip()))
        except Exception:
            pass

    def _update_file_buttons(self):
        """根据导入/翻译/一致性状态，更新导入钮颜色与导出钮显隐/颜色。
        逻辑：
        - 原文与导入内容一致 -> 导入钮青色；不一致或无导入 -> 导入钮灰、导出钮隐藏。
        - 导入有效时导出钮出现：译文区有译文且格式可导出 -> 青色可点，否则灰色不可点。
        """
        imported = getattr(self, "_imported_path", None)
        imported_text = getattr(self, "_imported_text", None)
        cur_src = self.input_edit.toPlainText().strip()
        consistent = bool(imported) and imported_text is not None and cur_src == imported_text

        # 导出文字钮：该区无文字则灰色不可点（与导出音频钮一致）
        if hasattr(self, "export_src_btn"):
            self.export_src_btn.setEnabled(bool(cur_src))
        if hasattr(self, "export_tgt_btn"):
            self.export_tgt_btn.setEnabled(
                bool(self.output_edit.toPlainText().strip()))

        # 导入按钮颜色
        if hasattr(self, "import_btn"):
            if consistent:
                self.import_btn.setStyleSheet(
                    "QPushButton{background:#5aa8b0; border:1px solid #5aa8b0; border-radius:5px;}")
            else:
                self.import_btn.setStyleSheet("")

        if not hasattr(self, "export_file_btn"):
            return
        if not consistent:
            # 脱离文件模式：导出钮消失
            self.export_file_btn.setVisible(False)
            self.export_file_btn.setEnabled(False)
            self.export_file_btn.setStyleSheet("")
            return
        # 一致：导出钮出现
        self.export_file_btn.setVisible(True)
        has_tgt = bool(self.output_edit.toPlainText().strip())
        if self._can_export_file and has_tgt:
            self.export_file_btn.setEnabled(True)
            self.export_file_btn.setStyleSheet(
                "QPushButton{background:#5aa8b0; border:1px solid #5aa8b0; border-radius:5px;}")
        else:
            self.export_file_btn.setEnabled(False)
            self.export_file_btn.setStyleSheet("")

    def _read_pdf_text(self, path):
        """优先用 pdfplumber 抽取文字；失败再尝试 pypdf。"""
        try:
            import pdfplumber
            out = []
            with pdfplumber.open(path) as pdf:
                for pg in pdf.pages:
                    out.append(pg.extract_text() or "")
            text = "\n".join(out).strip()
            if text:
                return text
        except Exception as e:
            _log_error(f"pdfplumber 抽取失败: {e}")
        try:
            from pypdf import PdfReader
            r = PdfReader(path)
            return "\n".join((pg.extract_text() or "") for pg in r.pages).strip()
        except Exception as e:
            _log_error(f"pypdf 抽取失败: {e}")
            raise RuntimeError("无法从该 PDF 提取文字（可能是扫描件，需要 OCR）。")

    # ---------- 导出翻译后的文件 ----------
    def _export_file(self):
        if not getattr(self, "_can_export_file", False):
            QMessageBox.information(self, "暂不支持", "当前导入的文件类型暂不支持导出（如 PDF）。")
            return
        import os
        src_path = getattr(self, "_imported_path", "")
        if not src_path:
            return
        import datetime
        d = os.path.dirname(src_path)
        base, ext = os.path.splitext(os.path.basename(src_path))
        _lang = "ZH" if _text_is_chinese(self.output_edit.toPlainText()) else "EN"
        _ts = datetime.datetime.now().strftime("%Y-%m-%d %H%M%S")
        default = os.path.join(d, f"EC TT {_lang} {_ts} T{ext}")   # 统一 EC 前缀 + T 后缀
        from PyQt6.QtWidgets import QFileDialog
        path, _ = QFileDialog.getSaveFileName(self, "导出翻译后的文件", default,
                                              f"原格式 (*{ext})")
        if not path:
            return
        try:
            text = self.output_edit.toPlainText()
            if ext == ".docx":
                self._export_docx_keep(src_path, path, text)
            else:
                self._write_text_file(path, text)
            self.status.showMessage(f"已导出：{path}", 4000)
        except Exception as e:
            _log_error(f"导出文件失败: {e}")
            QMessageBox.warning(self, "导出失败", str(e))

    def _export_docx_keep(self, src_path, out_path, translated):
        """导出 docx：尽量沿用原文档段落结构，逐段替换为译文。
        （基础排版继承：段落数对应时按段替换；不完全对应时整体写入。）"""
        from docx import Document
        src = Document(src_path)
        src_paras = [p for p in src.paragraphs]
        tgt_lines = [l for l in translated.split("\n")]
        if len([p for p in src_paras if p.text.strip()]) == len([l for l in tgt_lines if l.strip()]):
            ti = 0
            tgt_nonempty = [l for l in tgt_lines if l.strip()]
            for p in src_paras:
                if p.text.strip():
                    # 保留段落样式，只换文字
                    for run in p.runs:
                        run.text = ""
                    if p.runs:
                        p.runs[0].text = tgt_nonempty[ti]
                    else:
                        p.add_run(tgt_nonempty[ti])
                    ti += 1
            src.save(out_path)
        else:
            # 段落数不匹配，退回纯文本 docx
            doc = Document()
            for line in tgt_lines:
                doc.add_paragraph(_xml_safe(line))
            doc.save(out_path)

    def stop_speak(self, clear_only=False, keep_slider=False):
        self._is_speaking = False
        preserve = getattr(self, "_preserve_play_ui", False)
        # 中断正在进行的合成：worker 在下一个段边界检查点停止（软中断）
        w = getattr(self, "tts_worker", None)
        if w is not None and w.isRunning():
            w.cancel()
            if not clear_only:
                self.status.showMessage("正在停止…", 1500)
        self._hide_synth_busy()
        if hasattr(self, "_karaoke_timer"):
            self._karaoke_timer.stop()
        if self.player is not None and \
           self.player.playbackState() != QMediaPlayer.PlaybackState.StoppedState:
            self.player.stop()
        if not preserve:
            self._clear_karaoke()   # 变更嗓音/引擎重读期间保留卡拉OK绿色与选区
        if not keep_slider and not preserve:
            self.play_slider.setValue(0)   # 进度跳回开始
        if not preserve:
            # 停止后：若该侧音频仍在内存，朗读钮保持青色（喇叭图标+可下载），
            # 否则才灰化。图标恢复喇叭、提示恢复"朗读原文/译文"。
            side = getattr(self, "_active_side", "src")
            if not clear_only:
                self._side_cache.get(side, {})["position"] = 0   # 真·停止才归零续播位置
            has_audio = bool(self._side_cache.get(side, {}).get("bytes"))
            btn = self.speak_src_btn if side == "src" else self.speak_tgt_btn
            name = "原文" if side == "src" else "译文"
            if has_audio:
                self._set_speak_btn_active(btn, True, icon="speak")
            else:
                self._set_speak_btn_active(btn, False)
            btn.setToolTip(L(f"朗读{name}"))
            # 另一侧按钮状态按其缓存独立设置
            other = "tgt" if side == "src" else "src"
            obtn = self.speak_src_btn if other == "src" else self.speak_tgt_btn
            oname = "原文" if other == "src" else "译文"
            if bool(self._side_cache.get(other, {}).get("bytes")):
                self._set_speak_btn_active(obtn, True, icon="speak")
            else:
                self._set_speak_btn_active(obtn, False)
            obtn.setToolTip(f"朗读{oname}")
        if not clear_only:
            self.status.showMessage(L("已停止"), 2000)

    def closeEvent(self, event):
        # 退出前安全结束朗读线程，避免 "QThread destroyed while running" 崩溃
        try:
            if self.player is not None:
                self.player.stop()
            w = getattr(self, "tts_worker", None)
            if w is not None and w.isRunning():
                w.cancel()
                w.quit()
                w.wait(2000)
            for rw in getattr(self, "_retired_workers", []):
                if rw.isRunning():
                    rw.cancel()
                    rw.quit()
                    rw.wait(1000)
        except Exception:
            pass
        super().closeEvent(event)


def main():
    app = QApplication(sys.argv)
    _install_global_excepthook()   # 越早越好：把槽函数异常从闪退变成提示
    app.setApplicationName(APP_NAME)
    app.setQuitOnLastWindowClosed(True)
    # 单实例守护：已有实例运行则提示并退出（修复偶发双开两个程序）
    from PyQt6.QtCore import QLockFile, QDir
    _lock = QLockFile(QDir.temp().absoluteFilePath("EnglishCoach.single.lock"))
    if not _lock.tryLock(100):
        from PyQt6.QtWidgets import QMessageBox
        QMessageBox.information(None, "English Coach",
                                "English Coach 已在运行，请勿重复启动。")
        sys.exit(0)
    app._single_instance_lock = _lock   # 保持引用直到退出
    from PyQt6.QtGui import QPalette, QColor
    if sys.platform == "darwin":
        # mac：先按用户设置切原生外观(AppKit)，不设任何调色板——
        # 让所有原生控件(下拉/气球/按钮/滚动条/标题栏)自动深浅，杜绝自涂与原生打架。
        _apply_color_scheme(app)
        f = app.font(); f.setPixelSize(13); app.setFont(f)
    else:
        # 各平台一律使用系统原生滚动条：Win11=Fluent圆角，Win10=Vista直角，
        # Linux=各发行版原生。不再强加任何自定义滚动条样式(2026-07 决定)。
        # 非 mac：用调色板着色（与 apply_theme 热切换共用 _apply_win_palette，杜绝漂移）
        _apply_win_palette(app)
        f = app.font(); f.setPixelSize(13); app.setFont(f)
    # Win11（构建号>=22000）且 Qt>=6.7 提供 windows11 样式 -> 用原生 Fluent 胶囊滚动条
    if sys.platform == "win32":
        try:
            if sys.getwindowsversion().build >= 22000:
                from PyQt6.QtWidgets import QStyleFactory
                if "windows11" in [k.lower() for k in QStyleFactory.keys()]:
                    app.setStyle("windows11")
        except Exception:
            pass
    try:
        _apply_color_scheme(app)
        win = MainWindow()
        try:
            def _on_sys_scheme(_sch):
                from PyQt6.QtCore import QSettings as _QS
                if _QS("Strilen", "EnglishCoach").value("ui_theme", "跟随系统") == "跟随系统":
                    win.apply_theme()
            app.styleHints().colorSchemeChanged.connect(_on_sys_scheme)
        except Exception:
            pass
        # Qt 6.4 没有 colorSchemeChanged 信号(6.5+才有)，上面的 connect 会静默失败。
        # mac 兜底：轮询当前深浅，变化时重刷主题(跟随系统模式下绘制部分才能跟上)。
        if sys.platform == "darwin":
            try:
                from PyQt6.QtCore import QTimer as _QTm, QSettings as _QS2
                win._last_sys_light = _theme_is_light()
                def _poll_theme():
                    try:
                        cur = _theme_is_light()
                        if cur != win._last_sys_light:
                            win._last_sys_light = cur
                            if _QS2("Strilen", "EnglishCoach").value(
                                    "ui_theme", "跟随系统") == "跟随系统":
                                win.apply_theme()
                    except Exception:
                        pass
                win._theme_poll = _QTm(win)
                win._theme_poll.timeout.connect(_poll_theme)
                win._theme_poll.start(2000)
            except Exception:
                pass
    except Exception as e:
        # 启动期异常：弹窗显示，避免"无报错也无界面"
        import traceback as _tb
        msg = f"{e}\n\n{_tb.format_exc()}"
        try:
            QMessageBox.critical(None, "EnglishCoach 启动失败", msg)
        except Exception:
            print(msg)
        sys.exit(1)
    try:
        # 首次运行：显式写入 false，保证默认非勾选(不置顶)
        if win.settings.value("always_on_top", None) is None:
            win.settings.setValue("always_on_top", "false")
        if win.settings.value("always_on_top", "false") == "true":
            win.apply_always_on_top(True)
    except Exception:
        pass
    try:
        win._sync_export_text_buttons()   # 启动时空文本 -> 导出文字钮初始为灰
    except Exception:
        pass
    win.show()
    win.raise_()              # 提到最前
    win.activateWindow()      # 抢占焦点（老 macOS 上常需要）
    sys.exit(app.exec())


if __name__ == "__main__":
    # 关键(修复 macOS Silicon 打包后生成音频时不断弹出新 App 窗口)：
    # torch/Kokoro 等库会用 multiprocessing 起工作进程。在 PyInstaller 冻结的
    # app 里，若不在入口最前调用 freeze_support()，每个子进程会重新执行整个
    # 程序 -> 又弹出一个 App 窗。freeze_support() 让子进程正确识别身份、直接
    # 干活而不重启 app。必须是 __main__ 里的第一件事。
    import multiprocessing as _mp
    _mp.freeze_support()
    try:
        # 强制用 spawn 起子进程时也走 freeze 逻辑（macOS 默认已是 spawn）
        _mp.set_start_method("spawn", force=False)
    except Exception:
        pass
    main()
